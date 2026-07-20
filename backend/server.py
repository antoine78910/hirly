"""
Tinder for Jobs - FastAPI backend.

Features:
- Supabase Google OAuth exchange (access_token -> session_token cookie)
- CV upload (PDF/DOCX/TXT) -> Claude Sonnet 4.5 extracts profile JSON
- Job feed with AI-computed match score & reasons (Claude)
- Swipe right -> creates Application with tailored CV + cover letter (Claude)
- Application tracker with status updates
"""
import sys
import asyncio

if sys.platform.startswith("win"):
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

from fastapi import FastAPI, APIRouter, HTTPException, Request, Response, UploadFile, File, Depends, Cookie, Header, Query, Body
from fastapi.responses import JSONResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
import os
import io
import json
import logging
import inspect
import random
import uuid
import re
import base64
import time
import hashlib
import unicodedata
from decimal import Decimal
from pathlib import Path
from urllib.parse import quote as url_quote, urlparse
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any, Literal, Tuple
from datetime import datetime, timezone, timedelta
import atexit
import httpx
import stripe
from posthog import Posthog, new_context, identify_context, capture as posthog_capture

try:
    from opentelemetry import trace as _otel_trace
    from opentelemetry.sdk.trace import TracerProvider as _OtelTracerProvider
    from opentelemetry.sdk.resources import Resource as _OtelResource, SERVICE_NAME as _OTEL_SERVICE_NAME
    from posthog.ai.otel import PostHogSpanProcessor as _PostHogSpanProcessor
    from opentelemetry.instrumentation.openai_v2 import OpenAIInstrumentor as _OpenAIInstrumentor
    _OTEL_AVAILABLE = True
except ImportError:
    _OTEL_AVAILABLE = False

# Optional file parsing libs
from pypdf import PdfReader
import docx as docx_lib
from application_documents import DOCX_MIME, build_application_package, cover_letter_to_text, sanitize_docx_text
from cover_letter_quality import build_cover_letter_prompt_section, validate_cover_letter_quality
from cv_tailoring import (
    apply_minimal_resume_tailoring,
    build_cv_tailoring_prompt_section,
    enrich_cover_letter_from_profile,
    enrich_tailored_resume_contact,
    prepare_profile_for_application_generation,
    validate_minimal_tailoring_preserved,
)

APPLICATION_GENERATION_PIPELINE = "minimal_cv_v1"
from cv_quality import attach_cover_letter_quality_report, normalize_application_generation, validate_resume_quality
from application_email_service import send_application_email
from email_addresses import INBOUND_MANAGED_EMAIL_ENABLED, managed_reply_address
from inbound_email_service import process_inbound_resend_email
from notifications_service import (
    create_notification,
    list_notifications,
    mark_all_notifications_read,
    mark_notification_read,
    resolve_user_language,
)
from creator_applications_service import create_creator_application
from application_failure import classify_application_failure
from application_expiry import (
    expire_open_applications_for_job,
    mark_application_offer_expired,
    maybe_auto_expire_application,
)
from svix.webhooks import Webhook, WebhookVerificationError
from apply_agent.models import ApplyAgentError
from apply_agent.browser import effective_headless
from apply_agent.runner import run_apply_attempt
from auto_apply.executor import execute_application as auto_apply_execute_application
from auto_apply.metrics import latest_attempt as auto_apply_latest_attempt
from auto_apply.metrics import persist_execution_report as auto_apply_persist_execution_report
from auto_apply.metrics import status_safe_attempt as auto_apply_status_safe_attempt
from auto_apply.metrics import summary as auto_apply_metrics_summary
from auto_apply import queue as auto_apply_queue
from db import create_database_adapter
from db.supabase_adapter import count_supabase_table, test_supabase_connection
from influencer_store import create_influencer, get_influencer, list_influencers, update_influencer
from creator_social_service import build_dashboard, refresh_all_creators, refresh_creator
from creator_social_config import add_creator as add_tracked_creator
from creator_social_maintenance import (
    get_creator_social_refresh_status,
    record_refresh_summary,
    run_creator_social_refresh,
    run_creator_social_refresh_loop,
)
from creator_invite_store import (
    INVITE_TYPE_CREATOR,
    INVITE_TYPE_DEMO,
    INVITE_TYPE_TRAINING,
    create_demo_invitation,
    create_invitation,
    create_standalone_invitation,
    ensure_dev_test_invites,
    get_invite_by_code,
    list_demo_invites,
    list_invites_for_influencer,
    list_training_invites,
    mark_invite_clicked,
    mark_invite_redeemed,
    enrich_invite_rows,
    migrate_file_invites_to_db,
    bootstrap_invites_from_influencers,
    backfill_invites_from_users,
    materialize_invite_from_link,
    resolve_invite_type,
    validate_invite,
)
from jobs_service import (
    FEED_COVERAGE_EVALUATOR_VERSION,
    build_feed_coverage_snapshot,
    build_profile_job_query,
    hash_feed_coverage_user_id,
    refresh_greenhouse_boards,
    refresh_jobs_for_profile_if_needed,
    refresh_lever_boards,
    seed_greenhouse_company_boards,
    seed_lever_company_boards,
)
import jobs_service as jobs_service_module
from job_providers import (
    get_board_provider,
    get_configured_job_provider,
    get_job_provider,
    is_france_travail_provider,
    is_job_provider_configured,
    is_job_provider_enabled,
    primary_job_provider_name,
)
from job_providers.apply_eligibility import classify_apply_link, is_manual_fulfillment_ready, is_france_travail_offer
from job_providers.base import JobSearchQuery
from job_cache_maintenance import (
    env_bool as job_cache_env_bool,
    expire_stale_jobs,
    job_cache_status,
    job_inventory_analytics,
    purge_invalid_jobs,
    refresh_jobs_for_query_or_filters,
    revalidate_cached_jobs,
    run_job_cache_maintenance,
)
from france_travail_harvest import (
    harvest_enabled as ft_harvest_enabled,
    harvest_france_travail,
    last_harvest_summary as ft_last_harvest_summary,
    run_france_travail_harvest_loop,
)
from jsearch_harvest import (
    AGGRESSIVE_HARVEST_CITIES,
    AGGRESSIVE_HARVEST_ROLES,
    harvest_enabled as jsearch_harvest_enabled,
    harvest_jsearch,
    last_harvest_summary as jsearch_last_harvest_summary,
    run_jsearch_harvest_loop,
)
from rome_profile_service import get_rome_profile, normalize_rome_code, rome_profile_enabled
from ats_source_service import (
    ats_direct_maintenance_loop_enabled,
    discover_ats_sources_from_cached_jobs,
    discover_friendly_company_career_pages,
    refresh_ats_source,
    refresh_known_ats_sources,
    run_ats_direct_maintenance_loop,
)
from company_discovery_crawlers import (
    company_discovery_loop_enabled,
    last_discovery_summary as company_discovery_last_summary,
    run_company_discovery,
    run_company_discovery_loop,
)
from job_validation import cheap_validate_job_applyability
from datafast_attribution import datafast_stripe_metadata, merge_stripe_metadata
from employment_kind import (
    contract_type_to_job_types,
    employment_kind_rank_bonus,
    enrich_job_employment_kind,
    job_matches_job_types,
    resolve_profile_contract_type,
)
from location_intelligence import COUNTRY_NAME_TO_CODE, country_to_jsearch_language, expand_location_radius, normalize_place_name
from role_query_terms import ACADEMIC_LEVEL_STOPWORDS, resolve_role_match_tokens
from location_search import search_locations
from llm_client import LLMProviderNotConfigured, complete_json_text, extract_text_from_image_bytes, set_llm_user_context
from openai import RateLimitError as LLMRateLimitError
from onboarding_suggestions import suggest_categories, suggest_roles
from profile_search_preferences import (
    resolve_profile_target_location_data,
    resolve_profile_target_location_label,
    resolve_profile_target_role,
)
from feedback_routes import register_feedback_routes
from feedback_store import migrate_file_feedback_to_db
from feedback_resend_backfill import backfill_feedback_from_resend
from gmail_sync import (
    GMAIL_READONLY_SCOPE,
    gmail_connected_payload,
    public_email_message,
    store_gmail_tokens,
    sync_gmail_application_emails,
)
from training_routes import register_training_routes, register_training_admin_routes
from record_tools_routes import register_record_tools_routes
from record_tools_access import require_record_tools_user
from training_service import (
    SEED_COURSE_ID,
    admin_training_analytics as compute_training_analytics,
    enroll_user,
    ensure_training_enrollments_for_access_users,
    is_training_creator,
    seed_training_content,
    sync_training_locale_content,
)
from training_access import require_training_access, training_access_payload
from friend_referral_service import (
    claim_friend_referral_reward,
    enroll_friend_referral,
    friend_referral_status_payload,
    has_redeemed_friend_referral_code,
    redeem_friend_referral_code,
    validate_friend_referral_code,
    FRIEND_REFERRAL_REWARD_DAYS,
    FRIEND_REFERRAL_SIGNUP_DISCOUNT_COUPON_ID,
    FRIEND_REFERRAL_SIGNUP_DISCOUNT_PERCENT,
)
from referral_email_service import (
    send_friend_referral_reward_email,
    send_friend_referral_used_email,
)

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')
try:
    from apply_agent.browser_env import load_browser_secrets

    load_browser_secrets(override=True)
except Exception:
    pass


def _normalize_supabase_url(value: str | None) -> str:
    url = (value or "").strip().rstrip("/")
    for suffix in ("/rest/v1", "/auth/v1"):
        if url.endswith(suffix):
            return url[: -len(suffix)]
    return url


def _cors_origins() -> List[str]:
    configured = [
        origin.strip().rstrip("/")
        for origin in os.environ.get("CORS_ORIGINS", "").split(",")
        if origin.strip()
    ]
    defaults = [
        "http://localhost:3000",
        "http://localhost:3001",
        "https://www.tryhirly.com",
        "https://tryhirly.com",
        "https://app.tryhirly.com",
        "https://hirly-two.vercel.app",
    ]
    for env_name in ("FRONTEND_URL", "APP_URL", "REACT_APP_APP_ORIGIN", "REACT_APP_FRONTEND_URL", "VERCEL_PROJECT_PRODUCTION_URL", "VERCEL_URL"):
        value = (os.environ.get(env_name) or "").strip().rstrip("/")
        if not value:
            continue
        defaults.append(value if value.startswith(("http://", "https://")) else f"https://{value}")
    origins = configured or defaults
    return list(dict.fromkeys([origin for origin in [*origins, *defaults] if origin and origin != "*"]))


DATABASE_PROVIDER = "supabase"
db = create_database_adapter()
_application_generation_locks: Dict[str, asyncio.Lock] = {}
_application_generation_tasks: Dict[str, asyncio.Task] = {}

_posthog_client: Optional[Posthog] = None

app = FastAPI()
api_router = APIRouter(prefix="/api")

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Supabase via httpx logs every GET/POST at INFO — floods Railway and hides
# auto-apply / Bright Data lines. Keep warnings+ only for HTTP client noise.
for _noisy in ("httpx", "httpcore", "hpack", "h2", "urllib3"):
    logging.getLogger(_noisy).setLevel(logging.WARNING)

_feed_job_pool_cache: Dict[str, Any] = {"query_key": "", "rows": [], "fetched_at": 0.0}
_FEED_JOB_POOL_TTL_SECONDS = 90.0
_feed_sync_refresh_cooldown_until = 0.0
_feed_sync_refresh_cooldowns: Dict[str, float] = {}

# Background (non-blocking) provider refresh: dedup + cooldown per query signature
# so we never run the same discovery twice concurrently and never hammer providers.
_feed_background_refresh_tasks: Dict[str, asyncio.Task] = {}
_feed_background_refresh_last_run: Dict[str, float] = {}
_FEED_BACKGROUND_REFRESH_COOLDOWN_SECONDS = 120.0


def _clear_feed_job_pool_cache() -> None:
    _feed_job_pool_cache.update({"query_key": "", "rows": [], "fetched_at": 0.0})


def _prune_background_refresh_state(now: float) -> None:
    done_keys = [key for key, task in _feed_background_refresh_tasks.items() if task.done()]
    for key in done_keys:
        _feed_background_refresh_tasks.pop(key, None)
    expired = [
        key
        for key, last_run in _feed_background_refresh_last_run.items()
        if (now - last_run) > (_FEED_BACKGROUND_REFRESH_COOLDOWN_SECONDS * 4)
    ]
    for key in expired:
        _feed_background_refresh_last_run.pop(key, None)


def schedule_feed_background_refresh(
    signature: str,
    profile_for_refresh: Dict[str, Any],
    refresh_locations: List[Any],
    *,
    search_radius: str,
    role_override: Optional[str],
    requested_limit: int,
    max_results: int,
    max_pages: int,
    page_size: int,
    attempts_per_city: int,
    user_id: str = "",
) -> bool:
    """Run provider discovery in the background (non-blocking) with dedup + cooldown.

    Returns True when a new background task was scheduled.
    """
    now = time.monotonic()
    _prune_background_refresh_state(now)
    if signature in _feed_background_refresh_tasks:
        return False
    last_run = _feed_background_refresh_last_run.get(signature, 0.0)
    if (now - last_run) < _FEED_BACKGROUND_REFRESH_COOLDOWN_SECONDS:
        return False
    _feed_background_refresh_last_run[signature] = now

    async def _runner() -> None:
        imported_total = 0
        started = time.perf_counter()
        try:
            for loc_data in refresh_locations:
                loc_label = loc_data.get("location_label") if isinstance(loc_data, dict) else None
                try:
                    result = await refresh_jobs_for_profile_if_needed(
                        db,
                        profile_for_refresh,
                        require_auto_apply=False,
                        target_auto_apply_count=min(max_results, max(requested_limit, 1)),
                        location_override=loc_label,
                        location_data_override=loc_data if isinstance(loc_data, dict) else None,
                        search_radius=search_radius,
                        role_override=role_override,
                        force_provider_refresh=True,
                        query_limit_override=max_results,
                        provider_max_pages=max_pages,
                        provider_page_size=page_size,
                        max_provider_requests_override=attempts_per_city,
                        max_direct_apply_requests_override=0,
                    )
                except Exception as exc:
                    logger.warning(
                        "feed_background_refresh_location_error signature=%s location=%s error=%s",
                        signature, loc_label, str(exc)[:200],
                    )
                    continue
                imported = int(result.get("jobs_imported", result.get("count") or 0) or 0)
                imported_total += imported
                if result.get("provider_rate_limited"):
                    logger.info("feed_background_refresh_rate_limited signature=%s", signature)
                    break
            if imported_total > 0:
                _clear_feed_job_pool_cache()
            logger.info(
                "feed_background_refresh_complete signature=%s user_id=%s imported=%s elapsed_ms=%s",
                signature,
                user_id,
                imported_total,
                int((time.perf_counter() - started) * 1000),
            )
        except Exception as exc:
            logger.warning("feed_background_refresh_failed signature=%s error=%s", signature, str(exc)[:200])
        finally:
            _feed_background_refresh_tasks.pop(signature, None)

    task = asyncio.create_task(_runner())
    _feed_background_refresh_tasks[signature] = task
    return True


async def _get_feed_job_candidates(base_query: Dict[str, Any], limit: int) -> List[Dict[str, Any]]:
    cache_key = json.dumps({"q": base_query, "light": True}, sort_keys=True, default=str)
    now = time.monotonic()
    if (
        _feed_job_pool_cache["query_key"] == cache_key
        and _feed_job_pool_cache["rows"]
        and len(_feed_job_pool_cache["rows"]) >= limit
        and (now - float(_feed_job_pool_cache["fetched_at"])) < _FEED_JOB_POOL_TTL_SECONDS
    ):
        return list(_feed_job_pool_cache["rows"][:limit])
    jobs_col = db.jobs
    if hasattr(jobs_col, "read_with_select"):
        from db.supabase_adapter import JOB_FEED_LIGHT_SELECT

        # Column-only select — never pull full JSONB for the candidate pool.
        rows = await jobs_col.read_with_select(base_query, limit, select=JOB_FEED_LIGHT_SELECT)
    else:
        rows = await jobs_col.find(base_query, {"_id": 0}).limit(limit).to_list(limit)
    _feed_job_pool_cache["query_key"] = cache_key
    _feed_job_pool_cache["rows"] = rows
    _feed_job_pool_cache["fetched_at"] = now
    return rows


async def _hydrate_feed_jobs(jobs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Load full JSONB payloads only for jobs missing description text (final deck)."""
    if not jobs:
        return jobs
    missing_ids = [
        job.get("job_id")
        for job in jobs
        if job.get("job_id") and not (job.get("description") or job.get("clean_description") or job.get("job_description_sections"))
    ]
    if not missing_ids:
        return jobs
    full_rows = await db.jobs.find({"job_id": {"$in": missing_ids}}, {"_id": 0}).to_list(len(missing_ids))
    by_id = {row.get("job_id"): row for row in full_rows if row.get("job_id")}
    hydrated: List[Dict[str, Any]] = []
    for job in jobs:
        full = by_id.get(job.get("job_id"))
        merged = {**(full or {}), **job} if full else dict(job)
        # Prefer light-row match_* / server fields; prefer full-row description body.
        if full:
            for key in (
                "description",
                "clean_description",
                "job_description_sections",
                "requirements",
                "offer_details",
                "summary",
                "tagline",
                "company_logo",
                "seniority",
                "job_type",
                "employment_type",
                "contract_type",
                "industry",
                "sector",
                "tech_stack",
                "rome_label",
            ):
                if full.get(key) is not None:
                    merged[key] = full[key]
        for key, value in job.items():
            if key.startswith("_"):
                merged[key] = value
        hydrated.append(merged)
    return hydrated


def _env_bool(name: str, default: bool = False) -> bool:
    value = os.environ.get(name)
    if value is None:
        return default
    return value.lower() in ("1", "true", "yes", "on")


def _env_int(name: str, default: int) -> int:
    try:
        return int(os.environ.get(name, default))
    except (TypeError, ValueError):
        return default


def _apply_fulfillment_fields(job: Dict[str, Any]) -> Dict[str, Any]:
    if job.get("apply_fulfillment_status") and job.get("manual_fulfillment_ready") is not None:
        return {
            "apply_fulfillment_status": job.get("apply_fulfillment_status"),
            "apply_fulfillment_reason": job.get("apply_fulfillment_reason"),
            "apply_url_source": job.get("apply_url_source"),
            "apply_url_provider": job.get("apply_url_provider"),
            "selected_apply_url": job.get("selected_apply_url") or job.get("external_url"),
            "manual_fulfillment_ready": bool(job.get("manual_fulfillment_ready")),
            "job_board_account_required": bool(job.get("job_board_account_required")),
        }
    return classify_apply_link(
        job.get("external_url") or job.get("apply_url") or job.get("hosted_url"),
        source=job.get("source") or job.get("provider"),
        apply_options=job.get("apply_options") or [],
    )


def _job_snapshot_for_swipe(job: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(job, dict) or not job.get("job_id"):
        return {}
    return {key: value for key, value in job.items() if key != "_id"}


async def _restore_job_from_swipe_snapshot(
    job_id: str,
    swipe_row: Optional[Dict[str, Any]] = None,
) -> Optional[Dict[str, Any]]:
    job = await db.jobs.find_one({"job_id": job_id}, {"_id": 0})
    if job:
        return job
    snapshot = (swipe_row or {}).get("job_snapshot")
    if not isinstance(snapshot, dict) or snapshot.get("job_id") != job_id:
        return None
    await db.jobs.update_one(
        {"job_id": job_id},
        {"$set": snapshot},
        upsert=True,
    )
    return snapshot


def _swipe_insert_doc(
    user_id: str,
    job_id: str,
    direction: str,
    job: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    doc: Dict[str, Any] = {
        "user_id": user_id,
        "job_id": job_id,
        "direction": direction,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    if direction == "left" and job:
        snapshot = _job_snapshot_for_swipe(job)
        if snapshot:
            doc["job_snapshot"] = snapshot
    return doc


def _job_is_applyable(job: Dict[str, Any]) -> bool:
    if is_france_travail_offer(
        provider=str(job.get("provider") or ""),
        source=str(job.get("source") or ""),
        url=str(job.get("external_url") or job.get("selected_apply_url") or ""),
    ):
        validation_status = str(job.get("validation_status") or "").strip().lower()
        applyability_tier = str(job.get("applyability_tier") or "").strip().upper()
        if validation_status == "invalid" and applyability_tier == "E":
            return False
        if applyability_tier == "E":
            return False
        return bool(job.get("selected_apply_url") or job.get("external_url") or job.get("apply_url") or job.get("hosted_url"))
    validation_status = str(job.get("validation_status") or "").strip().lower()
    applyability_tier = str(job.get("applyability_tier") or "").strip().upper()
    if validation_status == "invalid" or applyability_tier in {"D", "E"}:
        return False
    if validation_status or applyability_tier:
        if not (job.get("selected_apply_url") or job.get("external_url") or job.get("apply_url") or job.get("hosted_url")):
            return False
        if job.get("requires_login") is True or job.get("requires_account_creation") is True or job.get("captcha_detected") is True:
            return False
    return is_manual_fulfillment_ready(job)


def _job_url_candidates(job: Dict[str, Any]) -> List[str]:
    urls: List[str] = []
    data = job.get("data") if isinstance(job.get("data"), dict) else {}
    for value in (
        job.get("selected_apply_url"),
        job.get("external_url"),
        job.get("apply_url"),
        job.get("hosted_url"),
        data.get("selected_apply_url"),
        data.get("external_url"),
        data.get("job_apply_link"),
        data.get("job_google_link"),
        data.get("employer_website"),
    ):
        if isinstance(value, str) and value.strip():
            urls.append(value.strip())
    for key in ("job_apply_links", "apply_options", "related_links"):
        items = data.get(key) or job.get(key) or []
        if not isinstance(items, list):
            continue
        for item in items:
            if isinstance(item, str) and item.strip():
                urls.append(item.strip())
            elif isinstance(item, dict):
                for field in ("url", "link", "apply_url", "job_apply_link"):
                    value = item.get(field)
                    if isinstance(value, str) and value.strip():
                        urls.append(value.strip())
    return list(dict.fromkeys(urls))


def _job_has_usable_apply_url(job: Dict[str, Any]) -> bool:
    return bool(_job_url_candidates(job))


def _job_is_blocked_for_feed(job: Dict[str, Any]) -> bool:
    if is_france_travail_offer(
        provider=str(job.get("provider") or ""),
        source=str(job.get("source") or ""),
        url=str(job.get("external_url") or job.get("selected_apply_url") or ""),
    ):
        validation_status = str(job.get("validation_status") or "").strip().lower()
        applyability_tier = str(job.get("applyability_tier") or "").strip().upper()
        return applyability_tier == "E" or job.get("captcha_detected") is True
    validation_status = str(job.get("validation_status") or "").strip().lower()
    applyability_tier = str(job.get("applyability_tier") or "").strip().upper()
    return (
        validation_status == "invalid"
        or applyability_tier in {"D", "E"}
        or job.get("requires_login") is True
        or job.get("requires_account_creation") is True
        or job.get("captcha_detected") is True
    )


def _with_application_capability_fields(job: Dict[str, Any]) -> Dict[str, Any]:
    job = _with_apply_fulfillment_fields(job)
    tier = str(job.get("applyability_tier") or "").strip().upper()
    status = str(job.get("validation_status") or "").strip().lower()
    direct_provider = str(job.get("provider") or "").lower() in {"greenhouse", "lever", "ashby"}
    can_auto_apply = (
        not _job_is_blocked_for_feed(job)
        and _job_has_usable_apply_url(job)
        and _job_is_applyable(job)
        and (
            (status == "valid" and tier in {"A", "B"})
            or (direct_provider and job.get("auto_apply_supported") is True)
        )
    )
    if can_auto_apply:
        mode = "auto_apply"
    elif not _job_is_blocked_for_feed(job) and _job_has_usable_apply_url(job):
        mode = "manual"
    else:
        mode = "blocked"
    return {
        **job,
        "application_mode": mode,
        "can_auto_apply": can_auto_apply,
        "requires_manual_review": mode != "auto_apply",
    }


def _with_apply_fulfillment_fields(job: Dict[str, Any]) -> Dict[str, Any]:
    fields = _apply_fulfillment_fields(job)
    external_url = fields.get("selected_apply_url") or job.get("external_url")
    return {**job, **fields, "external_url": external_url}


async def validate_job_before_application(job: Dict[str, Any]) -> Dict[str, Any]:
    """Revalidate and persist a job before credits or application creation."""
    job_id = job.get("job_id")
    latest_job = job
    if job_id:
        latest_job = await db.jobs.find_one({"job_id": job_id}, {"_id": 0}) or job
    validation = cheap_validate_job_applyability(latest_job)
    updated_job = {**latest_job, **validation}
    if job_id:
        await db.jobs.update_one({"job_id": job_id}, {"$set": validation})

    tier = str(validation.get("applyability_tier") or "").upper()
    status = str(validation.get("validation_status") or "").lower()
    allow_unknown = _env_bool("JOBS_ALLOW_UNKNOWN_TIER_APPLICATION", False)
    selected_apply_url = validation.get("selected_apply_url")
    blocked_reason = None

    if status == "invalid" or tier in {"D", "E"}:
        blocked_reason = validation.get("validation_reason") or validation.get("rejection_reason") or "Job is not applyable."
    elif not selected_apply_url:
        blocked_reason = "No apply URL is available."
    elif validation.get("requires_login") is True or validation.get("requires_account_creation") is True:
        blocked_reason = "This job requires a candidate login or account creation."
    elif validation.get("captcha_detected") is True:
        blocked_reason = "This job appears protected by CAPTCHA or bot protection."
    elif tier == "C":
        if not allow_unknown:
            blocked_reason = "This job needs additional validation before applications can be submitted."
        elif updated_job.get("manual_fulfillment_ready") is not True:
            blocked_reason = "This job is not ready for manual fulfillment."
    elif tier not in {"A", "B"} or status != "valid":
        blocked_reason = validation.get("validation_reason") or "Job validation did not confirm an applyable job."

    allowed = blocked_reason is None
    return {
        "allowed": allowed,
        "reason": "Job passed pre-application validation." if allowed else blocked_reason,
        "validation_status": validation.get("validation_status"),
        "applyability_tier": validation.get("applyability_tier"),
        "selected_apply_url": selected_apply_url,
        "requires_login": bool(validation.get("requires_login")),
        "requires_account_creation": bool(validation.get("requires_account_creation")),
        "captcha_detected": bool(validation.get("captcha_detected")),
        "job": updated_job,
        "validation": validation,
    }


# ===================== Models =====================

class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    user_id: str
    email: str
    name: str
    picture: Optional[str] = None
    demo_account: bool = False
    training_access: bool = False
    is_admin: bool = False
    require_review_before_send: bool = True
    language: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


class Profile(BaseModel):
    model_config = ConfigDict(extra="ignore")
    user_id: str
    cv_text: Optional[str] = None
    cv_filename: Optional[str] = None
    summary: Optional[str] = None
    skills: List[str] = []
    experience: List[Dict[str, Any]] = []
    education: List[Dict[str, Any]] = []
    target_roles: List[str] = []
    target_role: Optional[str] = None
    target_location: Optional[str] = None
    remote_preference: Optional[str] = "any"   # remote | onsite | hybrid | any
    seniority: Optional[str] = None
    application_answers_profile: Dict[str, Any] = Field(default_factory=dict)
    application_defaults: Dict[str, Any] = Field(default_factory=dict)
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


class Job(BaseModel):
    model_config = ConfigDict(extra="ignore")
    job_id: str
    title: str
    company: str
    company_logo: Optional[str] = None
    location: str
    remote: str  # remote | hybrid | onsite
    salary_min: Optional[int] = None
    salary_max: Optional[int] = None
    currency: str = "USD"
    description: str
    requirements: List[str] = []
    tech_stack: List[str] = []
    seniority: Optional[str] = None
    posted_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


class SwipeRequest(BaseModel):
    job_id: str
    direction: Literal["left", "right"]


class AdminAtsLabGenerateRequest(BaseModel):
    application_id: Optional[str] = None
    user_id: Optional[str] = None
    job_id: Optional[str] = None
    persist: bool = False


class AdminJobsRefreshRequest(BaseModel):
    search_role: Optional[str] = None
    location: Optional[str] = None
    country_code: Optional[str] = None
    lat: Optional[float] = None
    lng: Optional[float] = None
    search_radius: Optional[str] = None
    include_cross_border: Optional[bool] = None
    discover_ats_sources: bool = False
    refresh_discovered_ats_sources: bool = False
    ats_refresh_limit: Optional[int] = None
    remote: Optional[bool] = None
    limit: Optional[int] = None
    pages: Optional[int] = None
    dry_run: bool = False


class AdminJobsRevalidateRequest(BaseModel):
    validation_status: Optional[str] = None
    applyability_tier: Optional[str] = None
    older_than_hours: Optional[int] = None
    country_code: Optional[str] = None
    limit: Optional[int] = None
    dry_run: bool = False


class AdminJobsExpireStaleRequest(BaseModel):
    older_than_days: Optional[int] = None
    provider: Optional[str] = None
    country_code: Optional[str] = None
    limit: Optional[int] = None
    dry_run: bool = False
    completeness_run_id: Optional[str] = None


class AdminJobsPurgeInvalidRequest(BaseModel):
    older_than_days: Optional[int] = None
    country_code: Optional[str] = None
    applyability_tiers: Optional[List[str]] = None
    expire_first: bool = True
    limit: Optional[int] = None
    dry_run: bool = False
    completeness_run_id: Optional[str] = None


class AdminJobsMaintenanceRequest(BaseModel):
    dry_run: bool = False
    refresh_popular: Optional[bool] = None


class AdminJobsFeedDiagnosticRequest(BaseModel):
    search_role: Optional[str] = None
    location: Optional[str] = None
    country_code: Optional[str] = None
    search_radius: str = "worldwide"
    limit: int = 5
    dry_run: bool = True


class AdminFeedCoverageAuditRequest(BaseModel):
    user_ids: List[str]
    limit: int = 25
    freshness_window_days: Literal[1, 7, 30] = 30


class AdminAtsDiscoverSourcesRequest(BaseModel):
    provider: Optional[Literal["greenhouse", "lever", "ashby"]] = None
    country_code: Optional[str] = None
    limit: Optional[int] = 500
    dry_run: bool = False


class AdminAtsRefreshSourceRequest(BaseModel):
    ats_provider: Literal["greenhouse", "lever", "ashby"]
    source_key: str
    limit: Optional[int] = None
    dry_run: bool = False


class AdminDiscoverFriendlyCompanyPagesRequest(BaseModel):
    limit: Optional[int] = 200
    concurrency: Optional[int] = 5
    dry_run: bool = False


class AdminAtsRefreshKnownSourcesRequest(BaseModel):
    provider: Optional[Literal["greenhouse", "lever", "ashby"]] = None
    country_code: Optional[str] = None
    limit: Optional[int] = 25
    older_than_hours: Optional[int] = 12
    dry_run: bool = False


class GreenhousePrepareSubmitRequest(BaseModel):
    job_id: str


class AgentApplyRequest(BaseModel):
    job_id: str


class AgentSubmissionBenchmarkRequest(BaseModel):
    job_ids: List[str]
    run_submit: bool = True
    allow_real_submit: bool = False


class SupabaseSessionRequest(BaseModel):
    access_token: str
    provider_token: Optional[str] = None
    provider_refresh_token: Optional[str] = None
    provider_token_expires_at: Optional[Any] = None


class InviteEmailAuthRequest(BaseModel):
    email: str
    password: str
    code: str
    mode: Literal["signup", "login"] = "signup"


class AnalyticsEventRequest(BaseModel):
    event: str
    properties: Dict[str, Any] = Field(default_factory=dict)
    anonymous_id: Optional[str] = None
    page: Optional[str] = None
    source: Optional[str] = None


class BillingCheckoutRequest(BaseModel):
    plan: Literal["basic", "pro", "ultra", "monthly", "quarterly"]
    interval: Optional[Literal["weekly", "monthly", "quarterly"]] = None
    source: Optional[Literal["app", "credits", "onboarding"]] = None
    return_path: Optional[str] = None
    datafast_visitor_id: Optional[str] = None
    datafast_session_id: Optional[str] = None


class BillingConfirmCheckoutRequest(BaseModel):
    session_id: str


class AdminStripeReconcileRequest(BaseModel):
    payment_intent_id: Optional[str] = None
    customer_id: Optional[str] = None
    email: Optional[str] = None


class AdminGrantCreditsRequest(BaseModel):
    credits: int
    reason: str


class BillingUpgradeSessionRequest(BaseModel):
    plan: Literal["basic", "pro", "ultra"]
    interval: Optional[Literal["weekly", "monthly"]] = None
    return_path: Optional[str] = None


class BillingMasterCodeRequest(BaseModel):
    code: str
    plan: Literal["basic", "pro", "ultra", "monthly", "quarterly"] = "ultra"
    interval: Optional[Literal["weekly", "monthly", "quarterly"]] = None
    source: Optional[Literal["app", "credits", "onboarding"]] = None


class ResolveMissingInfoRequest(BaseModel):
    answers: Dict[str, Any] = Field(default_factory=dict)
    save_to_profile: bool = False


class Application(BaseModel):
    model_config = ConfigDict(extra="ignore")
    application_id: str
    user_id: str
    job_id: str
    status: Literal["applied", "viewed", "interview", "rejected", "offer"] = "applied"
    package_status: Literal["not_generated", "generated", "generated_text_only", "failed", "pending_generation", "needs_profile_data", "needs_job_data"] = "not_generated"
    generation_status: Optional[Literal["pending_generation", "generating", "generated", "failed"]] = None
    generation_error: Optional[str] = None
    submission_status: Literal["not_submitted", "ready", "prepared", "submitted", "failed", "blocked", "action_required", "blocked_captcha", "prepare_failed", "unknown"] = "not_submitted"
    submitted_at: Optional[str] = None
    submission_provider: Optional[str] = None
    submission_response_id: Optional[str] = None
    submission_error: Optional[str] = None
    prepared_application_payload: Optional[Dict[str, Any]] = None
    prepared_generated_answers: List[Dict[str, Any]] = Field(default_factory=list)
    prepared_missing_information: List[Any] = Field(default_factory=list)
    prepared_blockers: List[str] = Field(default_factory=list)
    prepared_at: Optional[str] = None
    submission_response_metadata: Optional[Dict[str, Any]] = None
    tailored_resume: Optional[Dict[str, Any]] = None
    cover_letter: Optional[Dict[str, Any]] = None
    match_score: Optional[int] = None
    match_reasons: List[str] = []
    interview_prep: List[str] = []
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


class StatusUpdate(BaseModel):
    status: Literal["applied", "viewed", "interview", "rejected", "offer"]


class AdminNoteCreate(BaseModel):
    note: str


class AdminStatusUpdate(BaseModel):
    status: Literal["submitted", "needs_user_input", "blocked", "escalated"]


class AdminManualStatusUpdate(BaseModel):
    manual_status: Literal["manual_review_needed", "manual_in_progress", "manually_submitted", "manual_blocked", "needs_user_input", "offer_expired"]
    note: Optional[str] = None


class AdminSendApplicationEmail(BaseModel):
    to_email: Optional[str] = None
    subject: Optional[str] = None
    body_text: Optional[str] = None
    mark_manually_submitted: bool = True


class PreferencesUpdate(BaseModel):
    target_role: Optional[str] = None
    target_roles: Optional[List[str]] = None
    target_location: Optional[str] = None
    target_location_data: Optional[Dict[str, Any]] = None
    remote_preference: Optional[str] = None
    seniority: Optional[str] = None
    contract_type: Optional[str] = None


class ContactUpdate(BaseModel):
    name: Optional[str] = None
    first_name: Optional[str] = None
    last_name: Optional[str] = None
    email: Optional[str] = None       # always overridden server-side with authenticated user's email
    phone: Optional[str] = None
    location: Optional[str] = None
    location_data: Optional[Dict[str, Any]] = None
    linkedin: Optional[str] = None
    website: Optional[str] = None


class ApplicationDefaultsUpdate(BaseModel):
    application_defaults: Dict[str, Any] = Field(default_factory=dict)


class AccountSettingsUpdate(BaseModel):
    require_review_before_send: Optional[bool] = None
    language: Optional[str] = None


class CvSourceUpdate(BaseModel):
    source: Literal["tailored", "original"]


class CoverLetterEditRequest(BaseModel):
    body_text: str


class ExperienceEntryEdit(BaseModel):
    role: str = ""
    company: str = ""
    location: str = ""
    duration: str = ""
    highlights: List[str] = []


class EducationEntryEdit(BaseModel):
    degree: str = ""
    school: str = ""
    year: str = ""


class ResumeEditRequest(BaseModel):
    experience: List[ExperienceEntryEdit] = []
    education: List[EducationEntryEdit] = []
    languages: List[str] = []


class StructuredProfileDataUpdate(BaseModel):
    contact: Dict[str, Any] = Field(default_factory=dict)
    education: List[Dict[str, Any]] = Field(default_factory=list)
    experience_summary: Dict[str, Any] = Field(default_factory=dict)
    skills: List[str] = Field(default_factory=list)
    application_defaults: Dict[str, Any] = Field(default_factory=dict)


class OnboardingCategoryItem(BaseModel):
    id: str
    label: str


class OnboardingSuggestCategoriesRequest(BaseModel):
    location: str = ""
    contract_type: str = ""
    location_data: Optional[Dict[str, Any]] = None


class CreatorApplicationRequest(BaseModel):
    email: str
    first_name: str
    last_name: str
    tiktok_handle: Optional[str] = None
    instagram_handle: Optional[str] = None
    has_company: Optional[str] = None
    whatsapp_country: Optional[str] = None
    whatsapp_number: Optional[str] = None
    country: Optional[str] = None
    referred_by: Optional[str] = None
    message: Optional[str] = None


class OnboardingSuggestRolesRequest(BaseModel):
    location: str = ""
    contract_type: str = ""
    categories: List[OnboardingCategoryItem] = Field(default_factory=list)
    location_data: Optional[Dict[str, Any]] = None


# ===================== Auth helpers =====================

async def get_current_user(
    request: Request,
    response: Response,
    session_token: Optional[str] = Cookie(default=None),
    authorization: Optional[str] = Header(default=None),
) -> User:
    bearer_token = None
    if authorization and authorization.lower().startswith("bearer "):
        bearer_token = authorization.split(" ", 1)[1].strip()

    candidates: List[tuple[str, str]] = []
    if bearer_token:
        candidates.append(("bearer", bearer_token))
    if session_token:
        source = "bearer_failed_cookie_fallback" if bearer_token else "cookie"
        candidates.append((source, session_token))

    if not candidates:
        logger.debug("auth_me token_missing path=%s", request.url.path)
        raise HTTPException(status_code=401, detail="Not authenticated")

    last_failure = "Invalid session"
    for token_source, token in candidates:
        logger.debug("auth_me token_received source=%s path=%s", token_source, request.url.path)
        session = await db.user_sessions.find_one({"session_token": token}, {"_id": 0})
        if not session:
            logger.info("auth_me invalid_token reason=session_not_found source=%s", token_source)
            last_failure = "Invalid session"
            continue

        expires_at = session.get("expires_at")
        if isinstance(expires_at, str):
            expires_at = datetime.fromisoformat(expires_at)
        if expires_at and expires_at.tzinfo is None:
            expires_at = expires_at.replace(tzinfo=timezone.utc)
        if expires_at and expires_at < datetime.now(timezone.utc):
            logger.info("auth_me invalid_token reason=session_expired source=%s user_id=%s", token_source, session.get("user_id"))
            last_failure = "Session expired"
            continue

        user_doc = await db.users.find_one({"user_id": session["user_id"]}, {"_id": 0})
        if not user_doc:
            logger.info("auth_me invalid_token reason=user_not_found source=%s user_id=%s", token_source, session.get("user_id"))
            last_failure = "User not found"
            continue
        # Re-issue the cookie on every authenticated request (not just at
        # login) so a session that only ever reached us via the old
        # host-only cookie or a bearer token self-heals onto the shared
        # .tryhirly.com-domain cookie the very next time it's used
        # successfully, without requiring an explicit re-login.
        _set_app_session_cookie(response, token)
        return User(**user_doc)

    raise HTTPException(status_code=401, detail=last_failure)


# ===================== Auth routes =====================

async def _upsert_auth_user(email: str, name: str, picture: Optional[str], extra: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    normalized_email = (email or "").strip().lower()
    existing = await _find_user_by_email(normalized_email)
    if existing:
        user_id = existing["user_id"]
        update = {"name": name, "picture": picture}
        if extra:
            update.update(extra)
        await db.users.update_one(
            {"user_id": user_id},
            {"$set": update},
        )
    else:
        user_id = f"user_{uuid.uuid4().hex[:12]}"
        user_doc = {
            "user_id": user_id,
            "email": normalized_email,
            "name": name,
            "picture": picture,
            "created_at": datetime.now(timezone.utc).isoformat(),
        }
        if extra:
            user_doc.update(extra)
        await db.users.insert_one(user_doc)
    user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    return user_doc


async def _create_app_session(user_id: str, source: str, session_token: Optional[str] = None) -> str:
    token = session_token or f"sess_{uuid.uuid4().hex}"
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)
    await db.user_sessions.insert_one({
        "user_id": user_id,
        "session_token": token,
        "expires_at": expires_at.isoformat(),
        "created_at": datetime.now(timezone.utc).isoformat(),
        "source": source,
    })
    return token


def _set_app_session_cookie(response: Response, session_token: str) -> None:
    is_dev = os.environ.get("ENVIRONMENT", "").strip().lower() == "development"
    # Shared across app.tryhirly.com / tryhirly.com / www.tryhirly.com so the
    # session is valid at the cookie layer itself on every subdomain, not
    # just the one that issued it. Without this, the cookie defaults to
    # host-only (scoped to whichever domain's Vercel /api rewrite proxied
    # the login), and the frontend's client-side localStorage/bearer-token
    # mirror (api.js) only ever refreshes at the moment of an active login
    # -- never during the routine /auth/me check -- so a session that was
    # only ever validated via this cookie on one domain was invisible on
    # the other, causing an app.tryhirly.com <-> signin redirect loop that
    # made the app unusable (confirmed live 2026-07-12).
    cookie_domain = os.environ.get("SESSION_COOKIE_DOMAIN", "").strip() or (None if is_dev else ".tryhirly.com")
    response.set_cookie(
        key="session_token",
        value=session_token,
        max_age=7 * 24 * 60 * 60,
        httponly=True,
        secure=not is_dev,
        samesite="lax" if is_dev else "none",
        domain=cookie_domain,
        path="/",
    )


def _supabase_admin_config() -> tuple[str, str]:
    supabase_url = _normalize_supabase_url(os.environ.get("SUPABASE_URL"))
    supabase_secret = os.environ.get("SUPABASE_SECRET_KEY", "")
    if not supabase_url or not supabase_secret:
        raise HTTPException(status_code=503, detail="Supabase auth is not configured")
    return supabase_url, supabase_secret


async def _supabase_admin_request(
    method: str,
    path: str,
    json_body: Optional[Dict[str, Any]] = None,
) -> httpx.Response:
    supabase_url, supabase_secret = _supabase_admin_config()
    headers = {
        "apikey": supabase_secret,
        "Authorization": f"Bearer {supabase_secret}",
        "Content-Type": "application/json",
    }
    async with httpx.AsyncClient(timeout=15.0) as http:
        return await http.request(
            method,
            f"{supabase_url}{path}",
            headers=headers,
            json=json_body,
        )


async def _supabase_password_access_token(email: str, password: str) -> str:
    response = await _supabase_admin_request(
        "POST",
        "/auth/v1/token?grant_type=password",
        {"email": email.strip().lower(), "password": password},
    )
    if response.status_code != 200:
        payload = response.json() if response.content else {}
        detail = payload.get("error_description") or payload.get("msg") or "Invalid email or password"
        raise HTTPException(status_code=401, detail=detail)
    access_token = (response.json() or {}).get("access_token")
    if not access_token:
        raise HTTPException(status_code=401, detail="Authentication failed")
    return access_token


async def _supabase_admin_find_user_by_email(email: str) -> Optional[Dict[str, Any]]:
    normalized = email.strip().lower()
    response = await _supabase_admin_request(
        "GET",
        f"/auth/v1/admin/users?email={url_quote(normalized)}",
    )
    if response.status_code != 200:
        return None
    users = (response.json() or {}).get("users") or []
    return users[0] if users else None


async def _supabase_admin_ensure_confirmed_user(email: str, password: str) -> None:
    normalized = email.strip().lower()
    create_response = await _supabase_admin_request(
        "POST",
        "/auth/v1/admin/users",
        {
            "email": normalized,
            "password": password,
            "email_confirm": True,
            "user_metadata": {"full_name": normalized.split("@")[0]},
        },
    )
    if create_response.status_code in (200, 201):
        return

    existing = await _supabase_admin_find_user_by_email(normalized)
    if not existing:
        payload = create_response.json() if create_response.content else {}
        detail = payload.get("msg") or payload.get("error_description") or "Could not create account"
        raise HTTPException(status_code=400, detail=detail)

    user_id = existing.get("id")
    if not user_id:
        raise HTTPException(status_code=400, detail="Could not resolve existing account")

    update_response = await _supabase_admin_request(
        "PUT",
        f"/auth/v1/admin/users/{user_id}",
        {"email_confirm": True, "password": password},
    )
    if update_response.status_code not in (200, 201):
        payload = update_response.json() if update_response.content else {}
        detail = payload.get("msg") or payload.get("error_description") or "Could not update account"
        raise HTTPException(status_code=400, detail=detail)


async def _session_payload_from_supabase_token(
    access_token: str,
    response: Response,
    *,
    source: str = "supabase",
    provider_token: Optional[str] = None,
    provider_refresh_token: Optional[str] = None,
    provider_token_expires_at: Optional[Any] = None,
) -> Dict[str, Any]:
    supabase_url, supabase_secret = _supabase_admin_config()
    async with httpx.AsyncClient(timeout=15.0) as http:
        r = await http.get(
            f"{supabase_url}/auth/v1/user",
            headers={
                "apikey": supabase_secret,
                "Authorization": f"Bearer {access_token}",
            },
        )
    if r.status_code != 200:
        logger.info("supabase_session invalid_token status=%s body=%s", r.status_code, r.text[:200])
        raise HTTPException(status_code=401, detail="Invalid Supabase access token")

    data = r.json()
    email = data.get("email")
    if not email:
        raise HTTPException(status_code=401, detail="Supabase user email is missing")
    metadata = data.get("user_metadata") or {}
    identities = data.get("identities") or []
    identity_data = (identities[0].get("identity_data") if identities and isinstance(identities[0], dict) else {}) or {}
    name = (
        metadata.get("full_name")
        or metadata.get("name")
        or identity_data.get("full_name")
        or identity_data.get("name")
        or email.split("@")[0]
    )
    picture = metadata.get("avatar_url") or metadata.get("picture") or identity_data.get("avatar_url") or identity_data.get("picture")
    existing_user = await db.users.find_one({"email": (email or "").strip().lower()}, {"_id": 0, "user_id": 1})
    user_doc = await _upsert_auth_user(
        email=email,
        name=name,
        picture=picture,
        extra={
            "auth_provider": "supabase",
            "supabase_user_id": data.get("id"),
            "last_login_at": datetime.now(timezone.utc).isoformat(),
        },
    )
    session_token = await _create_app_session(user_doc["user_id"], source)
    _set_app_session_cookie(response, session_token)

    if _posthog_client is not None:
        is_new_user = existing_user is None
        event_name = "user_signed_up" if is_new_user else "user_logged_in"
        with new_context():
            identify_context(user_doc["user_id"])
            posthog_capture(
                event_name,
                properties={
                    "auth_source": source,
                    "has_gmail_provider": bool(provider_token or provider_refresh_token),
                },
            )

    gmail_connection = {"connected": False, "email": None, "last_synced_at": None}
    if provider_token or provider_refresh_token:
        try:
            await store_gmail_tokens(
                db,
                user_id=user_doc["user_id"],
                email=email,
                provider_token=provider_token,
                provider_refresh_token=provider_refresh_token,
                expires_at=provider_token_expires_at,
            )
            gmail_connection = gmail_connected_payload(
                await db.gmail_connections.find_one({"user_id": user_doc["user_id"]}, {"_id": 0})
            )
        except Exception as exc:
            logger.warning("gmail_token_store_failed user_id=%s error=%s", user_doc["user_id"], str(exc)[:200])
            gmail_connection = {
                "connected": False,
                "email": email,
                "last_synced_at": None,
                "last_sync_error": str(exc)[:200],
            }

    profile = await db.profiles.find_one({"user_id": user_doc["user_id"]}, {"_id": 0, "cv_text": 1, "target_role": 1})
    creator_flag, training_access = await asyncio.gather(
        is_training_creator(db, user_doc["user_id"]),
        _resolve_training_access_from_user_doc(user_doc),
    )
    logger.info("supabase_session success user_id=%s email=%s", user_doc["user_id"], email)
    return {
        "user": user_doc,
        "has_profile": profile is not None and bool(profile.get("cv_text")),
        "has_preferences": profile is not None and bool(profile.get("target_role")),
        "is_training_creator": creator_flag,
        "has_training_access": training_access,
        "gmail": gmail_connection,
        "session_token": session_token,
    }


async def _resolve_training_access_from_user_doc(user_doc: Dict[str, Any]) -> bool:
    from training_access import training_open_access_enabled

    if training_open_access_enabled():
        return True
    user_id = user_doc.get("user_id")
    if not user_id:
        return False
    if user_doc.get("training_access"):
        return True
    if user_id == TUTORIAL_FILMING_USER_ID:
        return True
    if _is_admin_email(user_doc.get("email")):
        return True
    return await is_training_creator(db, user_id)


@api_router.post("/auth/supabase-session")
async def auth_supabase_session(body: SupabaseSessionRequest, response: Response):
    """Exchange a verified Supabase access token for the app's session_token."""
    access_token = (body.access_token or "").strip()
    if not access_token:
        raise HTTPException(status_code=400, detail="access_token required")
    return await _session_payload_from_supabase_token(
        access_token,
        response,
        source="supabase_google",
        provider_token=(body.provider_token or "").strip() or None,
        provider_refresh_token=(body.provider_refresh_token or "").strip() or None,
        provider_token_expires_at=body.provider_token_expires_at,
    )


@api_router.post("/auth/invite-email")
async def auth_invite_email(body: InviteEmailAuthRequest, response: Response):
    """Sign up or sign in via email for creator invites — no Supabase confirmation email."""
    email = (body.email or "").strip().lower()
    password = body.password or ""
    code = (body.code or "").strip()
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="Valid email required")
    if len(password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters")
    if not re.fullmatch(r"\d{6}", code):
        raise HTTPException(status_code=400, detail="Enter a valid 6-digit invitation code")

    check = await validate_invite(db, code)
    if not check.get("valid"):
        reason = check.get("reason")
        if reason == "revoked":
            raise HTTPException(status_code=410, detail="This invitation is no longer valid")
        raise HTTPException(status_code=400, detail="Invalid invitation code")

    if body.mode == "signup":
        await _supabase_admin_ensure_confirmed_user(email, password)

    access_token = await _supabase_password_access_token(email, password)
    source = "supabase_invite_signup" if body.mode == "signup" else "supabase_invite_login"
    return await _session_payload_from_supabase_token(access_token, response, source=source)


@api_router.post("/dev/login")
async def dev_login(response: Response):
    """Development-only local login that bypasses Emergent OAuth."""
    if not _dev_tools_enabled():
        raise HTTPException(status_code=404, detail="Not found")

    user_doc = await db.users.find_one({}, {"_id": 0})
    if not user_doc:
        user_id = f"user_{uuid.uuid4().hex[:12]}"
        user_doc = {
            "user_id": user_id,
            "email": "dev@swiipr.local",
            "name": "Dev User",
            "picture": None,
            "created_at": datetime.now(timezone.utc).isoformat(),
        }
        await db.users.insert_one(user_doc)
    else:
        user_id = user_doc["user_id"]

    session_token = f"dev_session_{uuid.uuid4().hex}"
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)
    await db.user_sessions.insert_one({
        "user_id": user_id,
        "session_token": session_token,
        "expires_at": expires_at.isoformat(),
        "created_at": datetime.now(timezone.utc).isoformat(),
        "source": "dev_login",
    })

    response.set_cookie(
        key="session_token",
        value=session_token,
        max_age=7 * 24 * 60 * 60,
        httponly=True,
        secure=False,
        samesite="lax",
        path="/",
    )

    profile = await db.profiles.find_one({"user_id": user_id}, {"_id": 0})
    logger.info(
        "dev_login success user_id=%s has_profile=%s token_prefix=%s",
        user_id,
        bool(profile),
        session_token[:12],
    )
    return {
        "session_token": session_token,
        "token": session_token,
        "user": user_doc,
        "profile": profile,
        "has_profile": profile is not None and bool(profile.get("cv_text")),
        "has_preferences": profile is not None and bool(profile.get("target_role")),
    }


TUTORIAL_FILMING_USER_ID = "tutorial_filming"

TUTORIAL_FILMING_PROFILE = {
    "cv_text": """Alex Martin
Senior Software Engineer — Paris, France

EXPERIENCE
Senior Software Engineer, ProductCo (2022–Present)
- React/TypeScript product used by 200k+ users
- Shipped hiring workflows and candidate tooling

Software Engineer, StartupXYZ (2019–2022)
- Node.js APIs, PostgreSQL, CI/CD

SKILLS
React, TypeScript, Node.js, Python, system design

EDUCATION
MSc Computer Science, 2019
""",
    "cv_filename": "alex-martin-cv.pdf",
    "summary": "Senior software engineer focused on React, TypeScript, and product delivery.",
    "target_role": "Software Engineer",
    "target_roles": ["Software Engineer", "Frontend Engineer", "Full Stack Engineer"],
    "target_location": "Paris, France",
    "target_location_data": {
        "location_label": "Paris, France",
        "country": "France",
        "country_code": "FR",
    },
    "remote_preference": "hybrid",
    "seniority": "mid",
    "skills": ["React", "TypeScript", "Node.js", "Python", "PostgreSQL"],
}


def _demo_profile_for_feed(user_id: str, profile: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """Demo creators can swipe without uploading a CV — use a seeded profile for feed scoring."""
    base = dict(profile or {})
    merged = {**TUTORIAL_FILMING_PROFILE, **base, "user_id": user_id}
    if not merged.get("cv_text"):
        merged["cv_text"] = TUTORIAL_FILMING_PROFILE["cv_text"]
    if not merged.get("target_role"):
        merged["target_role"] = TUTORIAL_FILMING_PROFILE["target_role"]
    if not merged.get("target_roles"):
        merged["target_roles"] = TUTORIAL_FILMING_PROFILE["target_roles"]
    if not merged.get("target_location"):
        merged["target_location"] = TUTORIAL_FILMING_PROFILE["target_location"]
    if not merged.get("target_location_data"):
        merged["target_location_data"] = TUTORIAL_FILMING_PROFILE["target_location_data"]
    if not _profile_has_usable_phone(merged):
        contact = dict(merged.get("contact") or {})
        contact["phone"] = "+33 6 12 34 56 78"
        merged["contact"] = contact
    return merged


async def _ensure_demo_feed_profile(user_id: str) -> None:
    existing = await db.profiles.find_one({"user_id": user_id}, {"_id": 0, "cv_text": 1})
    if existing and existing.get("cv_text"):
        return
    now = datetime.now(timezone.utc).isoformat()
    await db.profiles.update_one(
        {"user_id": user_id},
        {"$set": {**TUTORIAL_FILMING_PROFILE, "user_id": user_id, "updated_at": now}},
        upsert=True,
    )


@api_router.post("/tutorial/session")
async def tutorial_session(response: Response):
    """Temporary filming endpoint: demo account + seeded profile for real job feed."""
    now = datetime.now(timezone.utc).isoformat()
    user_doc = await db.users.find_one({"user_id": TUTORIAL_FILMING_USER_ID}, {"_id": 0})
    if not user_doc:
        user_doc = {
            "user_id": TUTORIAL_FILMING_USER_ID,
            "email": "tutorial@hirly.app",
            "name": "Alex Martin",
            "picture": None,
            "demo_account": True,
            "training_access": True,
            "created_at": now,
        }
        await db.users.insert_one(user_doc)
    else:
        await db.users.update_one(
            {"user_id": TUTORIAL_FILMING_USER_ID},
            {"$set": {
                "demo_account": True,
                "training_access": True,
                "name": user_doc.get("name") or "Alex Martin",
            }},
        )
        user_doc["demo_account"] = True
        user_doc["training_access"] = True

    profile_payload = {
        **TUTORIAL_FILMING_PROFILE,
        "user_id": TUTORIAL_FILMING_USER_ID,
        "updated_at": now,
    }
    await db.profiles.update_one(
        {"user_id": TUTORIAL_FILMING_USER_ID},
        {"$set": profile_payload},
        upsert=True,
    )

    session_token = f"tutorial_session_{uuid.uuid4().hex}"
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)
    await db.user_sessions.insert_one({
        "user_id": TUTORIAL_FILMING_USER_ID,
        "session_token": session_token,
        "expires_at": expires_at.isoformat(),
        "created_at": now,
        "source": "tutorial_filming",
    })

    response.set_cookie(
        key="session_token",
        value=session_token,
        max_age=7 * 24 * 60 * 60,
        httponly=True,
        secure=os.environ.get("ENVIRONMENT", "").strip().lower() != "development",
        samesite="lax",
        path="/",
    )

    return {
        "session_token": session_token,
        "token": session_token,
        "user": user_doc,
        "has_profile": True,
        "has_preferences": True,
    }


@api_router.get("/auth/me")
async def auth_me(user: User = Depends(get_current_user)):
    profile, creator, training_access = await asyncio.gather(
        db.profiles.find_one({"user_id": user.user_id}, {"_id": 0}),
        is_training_creator(db, user.user_id),
        _resolve_training_access(user),
    )
    logger.info(
        "auth_me cv_readiness user_id=%s profile_exists=%s has_cv_text=%s has_cv_filename=%s has_preferences=%s",
        user.user_id,
        profile is not None,
        bool((profile or {}).get("cv_text")),
        bool((profile or {}).get("cv_filename")),
        bool((profile or {}).get("target_role")),
    )
    return {
        "user": user.model_dump(),
        "has_profile": profile is not None and bool(profile.get("cv_text")),
        "has_preferences": profile is not None and bool(profile.get("target_role")),
        "is_training_creator": creator,
        "has_training_access": training_access,
        "is_admin": _is_admin_email(user.email) or bool(getattr(user, "is_admin", False)),
    }


async def _optional_current_user(
    request: Request,
    response: Response,
    session_token: Optional[str] = Cookie(default=None),
    authorization: Optional[str] = Header(default=None),
) -> Optional[User]:
    try:
        return await get_current_user(request, response, session_token=session_token, authorization=authorization)
    except HTTPException:
        return None


def _hash_ip(request: Request) -> Optional[str]:
    forwarded = request.headers.get("x-forwarded-for", "")
    raw_ip = forwarded.split(",", 1)[0].strip() if forwarded else (request.client.host if request.client else "")
    if not raw_ip:
        return None
    salt = os.environ.get("ANALYTICS_IP_HASH_SALT", "hirly-analytics")
    return hashlib.sha256(f"{salt}:{raw_ip}".encode("utf-8")).hexdigest()[:24]


@api_router.post("/analytics/event")
async def track_analytics_event(
    body: AnalyticsEventRequest,
    request: Request,
    user: Optional[User] = Depends(_optional_current_user),
):
    event_name = (body.event or "").strip()
    if not event_name:
        raise HTTPException(status_code=400, detail="event required")
    if len(event_name) > 120:
        raise HTTPException(status_code=400, detail="event too long")
    now = datetime.now(timezone.utc).isoformat()
    event_doc = {
        "event_id": f"evt_{uuid.uuid4().hex}",
        "user_id": user.user_id if user else None,
        "anonymous_id": (body.anonymous_id or "").strip()[:160] or None,
        "event": event_name,
        "properties": body.properties or {},
        "page": (body.page or "").strip()[:300] or None,
        "source": (body.source or "").strip()[:120] or None,
        "created_at": now,
        "user_agent": request.headers.get("user-agent"),
        "ip_hash": _hash_ip(request),
    }
    try:
        await db.analytics_events.insert_one(event_doc)
    except Exception as exc:
        logger.warning("analytics_event_store_failed event=%s error=%s", event_name, str(exc)[:200])
        return {"ok": False, "stored": False}
    return {"ok": True, "stored": True, "event_id": event_doc["event_id"]}


async def _store_friend_referral_analytics(
    event: str,
    *,
    user_id: Optional[str] = None,
    properties: Optional[Dict[str, Any]] = None,
) -> None:
    now = datetime.now(timezone.utc).isoformat()
    try:
        await db.analytics_events.insert_one({
            "event_id": f"evt_{uuid.uuid4().hex}",
            "user_id": user_id,
            "event": event,
            "properties": properties or {},
            "page": "friend_referral",
            "source": "backend",
            "created_at": now,
        })
    except Exception as exc:
        logger.warning(
            "friend_referral_analytics_failed event=%s error=%s",
            event,
            str(exc)[:200],
        )


@api_router.post("/auth/logout")
async def auth_logout(
    response: Response,
    session_token: Optional[str] = Cookie(default=None),
    authorization: Optional[str] = Header(default=None),
):
    token = session_token
    if not token and authorization and authorization.lower().startswith("bearer "):
        token = authorization.split(" ", 1)[1].strip()
    if token:
        await db.user_sessions.delete_many({"session_token": token})
    response.delete_cookie("session_token", path="/")
    return {"ok": True}


# ===================== Billing =====================

def _stripe_secret_key() -> str:
    key = os.environ.get("STRIPE_SECRET_KEY", "").strip()
    if not key:
        raise HTTPException(status_code=503, detail="Stripe is not configured")
    stripe.api_key = key
    return key


def _frontend_url() -> str:
    return os.environ.get("FRONTEND_URL", "http://localhost:3000").strip().rstrip("/")


def _app_url() -> str:
    return (
        os.environ.get("APP_URL")
        or os.environ.get("REACT_APP_APP_ORIGIN")
        or _frontend_url()
    ).strip().rstrip("/")


def _marketing_url() -> str:
    return _frontend_url()


def _sanitize_return_path(path: Optional[str], *, default: str = "/swipe") -> str:
    if not path:
        return default
    cleaned = path.strip()
    if not cleaned.startswith("/") or cleaned.startswith("//"):
        return default
    return cleaned[:512]


def _checkout_return_url(frontend_url: str, return_path: str, status: str) -> str:
    separator = "&" if "?" in return_path else "?"
    url = f"{frontend_url}{return_path}{separator}upgrade={status}"
    if status == "success":
        url += "&session_id={CHECKOUT_SESSION_ID}"
    return url


def _canonical_billing_plan(plan: str) -> str:
    aliases = {
        "monthly": "pro",
        "quarterly": "ultra",
    }
    return aliases.get((plan or "").strip().lower(), (plan or "").strip().lower())


def _stripe_price_env_for_plan(plan: str, *, interval: Optional[str] = None, source: Optional[str] = None) -> str:
    raw_plan = (plan or "").strip().lower()
    if source == "onboarding":
        if raw_plan not in {"monthly", "quarterly"}:
            raise HTTPException(status_code=400, detail="Unsupported onboarding billing plan")
        return f"STRIPE_PRICE_ONBOARDING_{raw_plan.upper()}"

    canonical_plan = _canonical_billing_plan(plan)
    if canonical_plan not in {"basic", "pro", "ultra"}:
        raise HTTPException(status_code=400, detail="Unsupported billing plan")
    if interval == "weekly":
        return f"STRIPE_PRICE_{canonical_plan.upper()}_WEEKLY"
    return f"STRIPE_PRICE_{canonical_plan.upper()}"


def _stripe_price_for_plan(plan: str, *, interval: Optional[str] = None, source: Optional[str] = None) -> str:
    env_name = _stripe_price_env_for_plan(plan, interval=interval, source=source)
    price_id = os.environ.get(env_name, "").strip()
    if not price_id:
        raise HTTPException(status_code=503, detail=f"{env_name} is not configured")
    return price_id


def _ensure_friend_referral_signup_discount_coupon() -> str:
    """Create the one-time 25%-off-first-invoice coupon if it doesn't
    already exist in this Stripe account. Idempotent -- safe to call on
    every checkout for an eligible user."""
    try:
        stripe.Coupon.create(
            id=FRIEND_REFERRAL_SIGNUP_DISCOUNT_COUPON_ID,
            percent_off=FRIEND_REFERRAL_SIGNUP_DISCOUNT_PERCENT,
            duration="once",
            name="Friend referral signup discount",
        )
    except stripe.error.InvalidRequestError as exc:
        if "already exists" not in str(exc):
            raise
    return FRIEND_REFERRAL_SIGNUP_DISCOUNT_COUPON_ID


def _master_billing_code() -> str:
    return os.environ.get("MASTER_BILLING_CODE", "424242").strip()


def _is_master_billing_code(code: str) -> bool:
    expected = _master_billing_code()
    return bool(expected) and (code or "").strip().casefold() == expected.casefold()


def _plan_from_price(price_id: Optional[str]) -> Optional[str]:
    for plan in ("basic", "pro", "ultra"):
        for env_name in (f"STRIPE_PRICE_{plan.upper()}", f"STRIPE_PRICE_{plan.upper()}_WEEKLY"):
            if price_id and price_id == os.environ.get(env_name, "").strip():
                return plan
    for plan in ("monthly", "quarterly"):
        if price_id and price_id == os.environ.get(f"STRIPE_PRICE_ONBOARDING_{plan.upper()}", "").strip():
            return plan
    return "unknown" if price_id else None


def _billing_from_user(user_doc: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    return dict((user_doc or {}).get("billing") or {})


def _billing_status_payload(user_doc: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    billing = _billing_from_user(user_doc)
    status = billing.get("subscription_status") or "none"
    is_premium = status in {"active", "trialing"}
    plan = billing.get("plan")
    # Friend-referral bonus credits live in their own fields (referral_bonus_
    # credits_total/_remaining) so they survive Stripe/billing-period syncs
    # untouched -- _merge_billing_credit_state recomputes credits_total/
    # credits_remaining from the plan allowance on every sync and would
    # otherwise silently claw back anything added directly to those two
    # fields (confirmed by tracing through its "same billing cycle" branch).
    bonus_total = int(billing.get("referral_bonus_credits_total") or 0)
    bonus_remaining = int(billing.get("referral_bonus_credits_remaining") or 0)
    return {
        "subscription_status": status,
        "plan": plan,
        "plan_tier": _canonical_billing_plan(plan) if is_premium and plan else None,
        "interval": billing.get("interval"),
        "source": billing.get("source"),
        "current_period_end": billing.get("current_period_end"),
        "stripe_customer_id_exists": bool(billing.get("stripe_customer_id")),
        "is_premium": is_premium,
        "credits_total": int(billing.get("credits_total") or 0) + bonus_total,
        "credits_remaining": int(billing.get("credits_remaining") or 0) + bonus_remaining,
        "friend_referral": friend_referral_status_payload(user_doc, is_premium=is_premium),
    }


_PLAN_CREDIT_LIMITS = {
    "monthly": 200,
    "quarterly": 600,
    "ultra": 600,
    "pro": 200,
    "basic": 80,
    # The friend-referral comp plan only exists to bypass the paywall
    # (is_premium=true) -- the actual credits for that period come entirely
    # from the separate referral_bonus_credits pool (_grant_friend_referral_
    # billing) so repeating rewards stay additive instead of being capped at
    # one flat plan allowance.
    "friend_referral": 0,
}


def _billing_credit_limit(
    plan: Optional[str],
    is_premium: bool,
    *,
    interval: Optional[str] = None,
    source: Optional[str] = None,
) -> int:
    if not is_premium:
        return 0
    normalized_plan = (plan or "").strip().lower()
    normalized_interval = (interval or "").strip().lower()
    # Onboarding plan ids ("monthly"/"quarterly") intentionally match the same
    # credit allowance as their app-side price equivalent (pro=29.99€=200,
    # ultra=69.99€=600) via the shared _PLAN_CREDIT_LIMITS lookup below.
    if normalized_interval == "weekly":
        monthly_limit = _PLAN_CREDIT_LIMITS.get(normalized_plan, 200)
        return max(1, monthly_limit // 4)
    return _PLAN_CREDIT_LIMITS.get(normalized_plan, 200)


def _billing_period_bounds(billing: Dict[str, Any]) -> tuple[datetime, datetime]:
    now = datetime.now(timezone.utc)
    period_end = _parse_dt(billing.get("current_period_end")) or now
    if period_end.tzinfo is None:
        period_end = period_end.replace(tzinfo=timezone.utc)
    period_start = period_end - timedelta(days=7)
    return period_start, period_end


def _billing_credit_period_key(billing: Dict[str, Any]) -> str:
    return "|".join(str(billing.get(key) or "") for key in (
        "stripe_subscription_id",
        "current_period_start",
        "current_period_end",
        "plan",
        "interval",
        "source",
    ))


def _merge_billing_credit_state(existing_billing: Dict[str, Any], updates: Dict[str, Any]) -> Dict[str, Any]:
    merged = {**existing_billing, **updates}
    is_premium = (merged.get("subscription_status") or "none") in {"active", "trialing"}
    if not is_premium:
        merged["credits_total"] = 0
        merged["credits_remaining"] = 0
        merged["credits_period_key"] = _billing_credit_period_key(merged)
        return merged

    allowance = _billing_credit_limit(
        merged.get("plan"),
        True,
        interval=merged.get("interval"),
        source=merged.get("source"),
    )
    period_key = _billing_credit_period_key(merged)

    # Nothing changed since the last sync — just clamp remaining to the allowance.
    if (
        existing_billing.get("credits_period_key") == period_key
        and isinstance(existing_billing.get("credits_remaining"), int)
        and int(existing_billing.get("credits_total") or 0) == allowance
    ):
        merged["credits_total"] = allowance
        merged["credits_remaining"] = max(0, min(int(existing_billing.get("credits_remaining") or 0), allowance))
        merged["credits_period_key"] = period_key
        return merged

    # Same Stripe subscription and same billing-period boundaries as before, but the
    # plan changed (a prorated mid-cycle upgrade/downgrade) — grant only the
    # difference in allowance instead of resetting to the full new-tier amount.
    # e.g. Pro (200) -> Ultra (600) at 0 remaining grants +400, landing at 400/600.
    same_billing_cycle = (
        bool(existing_billing.get("stripe_subscription_id"))
        and existing_billing.get("stripe_subscription_id") == merged.get("stripe_subscription_id")
        and existing_billing.get("current_period_start") == merged.get("current_period_start")
        and existing_billing.get("current_period_end") == merged.get("current_period_end")
        and isinstance(existing_billing.get("credits_total"), int)
    )
    if same_billing_cycle:
        old_allowance = int(existing_billing.get("credits_total") or 0)
        old_remaining = int(existing_billing.get("credits_remaining") or 0)
        delta = allowance - old_allowance
        merged["credits_total"] = allowance
        merged["credits_remaining"] = max(0, min(allowance, old_remaining + delta))
        merged["credits_period_key"] = period_key
        return merged

    # New billing period (renewal) or a brand-new subscription — full reset.
    merged["credits_total"] = allowance
    merged["credits_remaining"] = allowance
    merged["credits_period_key"] = period_key
    return merged


def _count_records_in_period(
    records: List[Dict[str, Any]],
    period_start: datetime,
    period_end: datetime,
    *,
    date_field: str = "created_at",
) -> int:
    count = 0
    for record in records:
        created_at = _parse_dt(record.get(date_field))
        if not created_at:
            continue
        if created_at.tzinfo is None:
            created_at = created_at.replace(tzinfo=timezone.utc)
        if period_start <= created_at <= period_end:
            count += 1
    return count


async def _get_user_doc(user: User) -> Dict[str, Any]:
    user_doc = await db.users.find_one({"user_id": user.user_id}, {"_id": 0})
    if not user_doc:
        raise HTTPException(status_code=404, detail="User not found")
    return user_doc


async def _grant_master_billing_access(
    user_id: str,
    plan: str,
    *,
    interval: Optional[str] = None,
    source: Optional[str] = None,
) -> Dict[str, Any]:
    checkout_source = source or "app"
    billing_plan = plan if checkout_source == "onboarding" else _canonical_billing_plan(plan)
    if checkout_source == "onboarding":
        if billing_plan not in {"monthly", "quarterly"}:
            raise HTTPException(status_code=400, detail="Unsupported onboarding billing plan")
        billing_interval = interval or billing_plan
    else:
        if billing_plan not in {"basic", "pro", "ultra"}:
            raise HTTPException(status_code=400, detail="Unsupported billing plan")
        billing_interval = interval or "monthly"

    now = datetime.now(timezone.utc)
    period_days = 7 if billing_interval == "weekly" else 90 if billing_interval == "quarterly" else 30
    updates = {
        "subscription_status": "active",
        "plan": billing_plan,
        "interval": billing_interval,
        "source": checkout_source,
        "stripe_subscription_id": f"master_code_{user_id}_{uuid.uuid4().hex[:12]}",
        "last_payment_status": "master_code",
        "current_period_start": now.isoformat(),
        "current_period_end": (now + timedelta(days=period_days)).isoformat(),
    }
    await _update_user_billing_by_user_id(user_id, updates)
    user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    return _billing_status_payload(user_doc)


async def _update_user_billing_by_user_id(user_id: str, updates: Dict[str, Any]) -> None:
    now = datetime.now(timezone.utc).isoformat()
    existing_user = await db.users.find_one({"user_id": user_id}, {"_id": 0, "billing": 1}) or {}
    merged_billing = _merge_billing_credit_state(_billing_from_user(existing_user), updates)
    update_fields = {f"billing.{key}": value for key, value in merged_billing.items()}
    update_fields["billing.updated_at"] = now
    result = await db.users.update_one({"user_id": user_id}, {"$set": update_fields})
    if getattr(result, "matched_count", 0) == 0:
        logger.warning("stripe_billing_user_update_no_match user_id=%s keys=%s", user_id, sorted(updates.keys()))


async def _grant_friend_referral_billing(user_id: str, credits: int) -> None:
    """Grant a friend-referral reward batch: `credits` bonus credits, plus a
    free month of paywall access if the referrer isn't already a genuine
    paying subscriber. Called once per newly-earned batch (every 3
    referrals), so this must be safe to call repeatedly for the same user.
    """
    user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0, "billing": 1})
    billing = _billing_from_user(user_doc)
    is_real_paid_plan = (
        (billing.get("subscription_status") or "none") in {"active", "trialing"}
        and billing.get("source") != "friend_referral"
    )
    if not is_real_paid_plan:
        # Never touches a real Stripe-synced subscription -- only applies
        # when the referrer has no genuine paid plan, or is already on a
        # prior friend_referral comp grant (safe to refresh).
        now = datetime.now(timezone.utc)
        updates = {
            "subscription_status": "active",
            "plan": "friend_referral",
            "interval": "monthly",
            "source": "friend_referral",
            "stripe_subscription_id": f"friend_referral_{user_id}_{uuid.uuid4().hex[:12]}",
            "last_payment_status": "friend_referral_reward",
            "current_period_start": now.isoformat(),
            "current_period_end": (now + timedelta(days=FRIEND_REFERRAL_REWARD_DAYS)).isoformat(),
        }
        # _update_user_billing_by_user_id merges through _merge_billing_credit_state,
        # which resets billing.credits_total/credits_remaining to the
        # friend_referral plan's own allowance (0 -- see _PLAN_CREDIT_LIMITS)
        # since this always looks like a brand-new subscription. The real
        # reward amount lives entirely in the bonus pool below, so this
        # zeroing is exactly what we want, not a bug.
        await _update_user_billing_by_user_id(user_id, updates)
        billing = _billing_from_user(await db.users.find_one({"user_id": user_id}, {"_id": 0, "billing": 1}))

    # Bonus credits live in their own fields, decoupled from the plan-reset
    # machinery above, so repeating rewards accumulate instead of being
    # capped at one flat allowance -- see _billing_status_payload and
    # _consume_application_credit for how they're read/spent.
    await db.users.update_one(
        {"user_id": user_id},
        {"$set": {
            "billing.referral_bonus_credits_total": int(billing.get("referral_bonus_credits_total") or 0) + credits,
            "billing.referral_bonus_credits_remaining": int(billing.get("referral_bonus_credits_remaining") or 0) + credits,
            "billing.updated_at": datetime.now(timezone.utc).isoformat(),
        }},
    )


async def _finalize_user_billing(user_id: str) -> Optional[Dict[str, Any]]:
    """Refresh subscription state from Stripe and grant any missing credits."""
    user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    if not user_doc:
        return None
    user_doc, _warning = await _refresh_billing_from_stripe(user_doc)
    return await _repair_premium_credits_if_needed(user_id, user_doc)


async def _resolve_user_id_from_checkout_session(
    session_obj: Dict[str, Any],
    *,
    fallback_user_id: Optional[str] = None,
) -> Optional[str]:
    session_user_id = _checkout_session_user_id(session_obj) or fallback_user_id
    if session_user_id:
        found = await db.users.find_one({"user_id": session_user_id}, {"_id": 0, "user_id": 1})
        if found:
            return session_user_id

    customer_id = session_obj.get("customer")
    if customer_id:
        resolved = await _resolve_user_id_for_stripe_customer(customer_id)
        if resolved:
            return resolved

    customer_details = session_obj.get("customer_details") or {}
    checkout_email = (customer_details.get("email") or "").strip().lower()
    if checkout_email:
        user_doc = await _find_user_by_email(checkout_email)
        if user_doc:
            return user_doc.get("user_id")
    return None


async def _ensure_app_user_for_billing_email(email: str) -> Optional[Dict[str, Any]]:
    """Ensure a MongoDB app user exists for a Stripe billing email (create from Supabase if needed)."""
    normalized = (email or "").strip().lower()
    if not normalized:
        return None
    user_doc = await _find_user_by_email(normalized)
    if user_doc:
        return user_doc
    supabase_user = await _supabase_admin_find_user_by_email(normalized)
    if not supabase_user:
        return None
    metadata = supabase_user.get("user_metadata") or {}
    name = metadata.get("full_name") or metadata.get("name") or normalized.split("@")[0]
    picture = metadata.get("avatar_url") or metadata.get("picture")
    logger.info("stripe_reconcile_creating_app_user email=%s", normalized)
    return await _upsert_auth_user(
        email=normalized,
        name=name,
        picture=picture,
        extra={
            "auth_provider": "supabase",
            "supabase_user_id": supabase_user.get("id"),
        },
    )


def _stripe_object_id(value: Any) -> Optional[str]:
    if isinstance(value, str):
        return value.strip() or None
    if isinstance(value, dict):
        object_id = value.get("id")
        return str(object_id).strip() if object_id else None
    object_id = getattr(value, "id", None)
    return str(object_id).strip() if object_id else None


def _stripe_to_dict(obj: Any) -> Any:
    """Recursively convert a raw Stripe SDK object (StripeObject/Event/list results/
    etc.) into plain dicts/lists. All of our billing code is written against dict
    .get()/isinstance(dict) semantics, but this stripe-python version's StripeObject
    no longer subclasses dict and has no .get() method at all -- calling .get() on a
    live SDK object raises AttributeError. Every object handed to us directly by the
    SDK (construct_event, .retrieve(), .list(), auto_paging_iter()) must be converted
    at the boundary before any of our helpers touch it.
    """
    if obj is None or isinstance(obj, (dict, str, int, float, bool)):
        return obj
    if isinstance(obj, (list, tuple)):
        return [_stripe_to_dict(item) for item in obj]
    to_dict = getattr(obj, "to_dict", None)
    if callable(to_dict):
        return to_dict()
    return obj


def _resolve_stripe_checkout_session_context(session_id: str) -> Dict[str, Any]:
    _stripe_secret_key()
    try:
        session_obj = stripe.checkout.Session.retrieve(
            session_id,
            expand=["customer", "subscription", "payment_intent"],
        )
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Invalid checkout session: {str(exc)[:200]}")
    session_dict = _stripe_to_dict(session_obj)
    customer_details = session_dict.get("customer_details") or {}
    customer_obj = session_dict.get("customer")
    email = (customer_details.get("email") or "").strip().lower() or None
    if not email and isinstance(customer_obj, dict):
        email = (customer_obj.get("email") or "").strip().lower() or None
    return {
        "customer_id": _stripe_object_id(session_dict.get("customer")),
        "email": email,
        "subscription_id": _stripe_object_id(session_dict.get("subscription")),
        "user_id_hint": _checkout_session_user_id(session_dict),
    }


def _resolve_stripe_payment_intent_context(payment_intent_id: str) -> Dict[str, Any]:
    _stripe_secret_key()
    try:
        payment_intent = _stripe_to_dict(stripe.PaymentIntent.retrieve(
            payment_intent_id,
            expand=["customer", "invoice.subscription", "latest_charge"],
        ))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Invalid payment intent: {str(exc)[:200]}")

    context: Dict[str, Any] = {
        "customer_id": _stripe_object_id(payment_intent.get("customer")),
        "email": None,
        "subscription_id": None,
        "user_id_hint": None,
    }

    customer_obj = payment_intent.get("customer")
    if isinstance(customer_obj, dict):
        context["email"] = (customer_obj.get("email") or "").strip().lower() or None

    if not context["email"]:
        context["email"] = (payment_intent.get("receipt_email") or "").strip().lower() or None

    charge = payment_intent.get("latest_charge")
    if isinstance(charge, dict) and not context["email"]:
        billing = charge.get("billing_details") or {}
        context["email"] = (billing.get("email") or "").strip().lower() or None

    invoice = payment_intent.get("invoice")
    if isinstance(invoice, dict):
        context["subscription_id"] = _stripe_object_id(invoice.get("subscription"))
    elif isinstance(invoice, str) and invoice:
        try:
            invoice_obj = _stripe_to_dict(stripe.Invoice.retrieve(invoice))
            context["subscription_id"] = _stripe_object_id(invoice_obj.get("subscription"))
        except Exception:
            pass

    try:
        sessions = stripe.checkout.Session.list(payment_intent=payment_intent_id, limit=1)
        data = _stripe_to_dict(sessions.get("data") if isinstance(sessions, dict) else getattr(sessions, "data", []))
        if data:
            session_dict = data[0]
            context["user_id_hint"] = _checkout_session_user_id(session_dict)
            context["customer_id"] = context["customer_id"] or _stripe_object_id(session_dict.get("customer"))
            if not context["email"]:
                details = session_dict.get("customer_details") or {}
                context["email"] = (details.get("email") or "").strip().lower() or None
            context["subscription_id"] = context["subscription_id"] or _stripe_object_id(session_dict.get("subscription"))
    except Exception as exc:
        logger.warning(
            "stripe_checkout_session_lookup_failed payment_intent=%s error=%s",
            payment_intent_id,
            str(exc)[:160],
        )

    return context


def _resolve_stripe_customer_context(customer_id: str) -> Dict[str, Any]:
    _stripe_secret_key()
    try:
        customer = stripe.Customer.retrieve(customer_id)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Invalid Stripe customer: {str(exc)[:200]}")
    customer_dict = _stripe_to_dict(customer)
    metadata = customer_dict.get("metadata") or {}
    return {
        "customer_id": customer_id,
        "email": (customer_dict.get("email") or "").strip().lower() or None,
        "subscription_id": None,
        "user_id_hint": metadata.get("user_id") if isinstance(metadata, dict) else None,
    }


async def _update_user_billing_by_customer_id(customer_id: str, updates: Dict[str, Any]) -> None:
    if not customer_id:
        logger.warning("stripe_billing_customer_update_missing_customer keys=%s", sorted(updates.keys()))
        return
    user_id = await _resolve_user_id_for_stripe_customer(customer_id)
    if user_id:
        await _update_user_billing_by_user_id(user_id, updates)
        await _finalize_user_billing(user_id)
        return
    logger.warning("stripe_billing_customer_update_no_match customer_id=%s keys=%s", customer_id, sorted(updates.keys()))


async def _resolve_user_id_for_stripe_customer(customer_id: str) -> Optional[str]:
    if not customer_id:
        return None
    existing = await db.users.find_one({"billing.stripe_customer_id": customer_id}, {"_id": 0, "user_id": 1})
    if existing:
        return existing.get("user_id")
    if not _stripe_configured():
        return None
    try:
        _stripe_secret_key()
        customer = _stripe_to_dict(stripe.Customer.retrieve(customer_id))
        metadata_user_id = (customer.get("metadata") or {}).get("user_id")
        if metadata_user_id:
            found = await db.users.find_one({"user_id": metadata_user_id}, {"_id": 0, "user_id": 1})
            if found:
                return metadata_user_id
        email = (customer.get("email") or "").strip()
        if email:
            user_doc = await _find_user_by_email(email)
            if user_doc:
                return user_doc.get("user_id")
    except Exception as exc:
        logger.warning(
            "stripe_resolve_user_for_customer_failed customer_id=%s error=%s",
            customer_id,
            str(exc)[:200],
        )
    return None


async def _discover_stripe_customer_for_user(user_doc: Dict[str, Any]) -> Optional[str]:
    billing = _billing_from_user(user_doc)
    existing = billing.get("stripe_customer_id")
    if existing:
        return existing
    email = (user_doc.get("email") or "").strip()
    if not email or not _stripe_configured():
        return None
    try:
        _stripe_secret_key()
        result = stripe.Customer.list(email=email, limit=10)
        data = result.get("data") if isinstance(result, dict) else getattr(result, "data", [])
        user_id = user_doc.get("user_id")
        for customer in data or []:
            customer_id = customer.get("id") if isinstance(customer, dict) else getattr(customer, "id", None)
            metadata = customer.get("metadata") if isinstance(customer, dict) else getattr(customer, "metadata", {}) or {}
            if isinstance(metadata, dict) and metadata.get("user_id") == user_id and customer_id:
                return customer_id
        # Never attach another user's Stripe customer that merely shares this email.
    except Exception as exc:
        logger.warning(
            "stripe_discover_customer_for_user_failed user_id=%s error=%s",
            user_doc.get("user_id"),
            str(exc)[:200],
        )
    return None


async def _stripe_customer_for_user(user_doc: Dict[str, Any]) -> str:
    _stripe_secret_key()
    billing = _billing_from_user(user_doc)
    existing_customer_id = billing.get("stripe_customer_id")
    user_id = user_doc.get("user_id", "")
    user_email = (user_doc.get("email") or "").strip()
    if existing_customer_id:
        try:
            stripe.Customer.modify(
                existing_customer_id,
                email=user_email or None,
                metadata={"user_id": user_id},
            )
        except Exception as exc:
            logger.warning(
                "stripe_customer_sync_failed customer_id=%s user_id=%s error=%s",
                existing_customer_id,
                user_id,
                str(exc)[:200],
            )
        return existing_customer_id
    customer = stripe.Customer.create(
        email=user_email,
        name=user_doc.get("name") or user_email,
        metadata={"user_id": user_id},
    )
    customer_id = customer["id"]
    await _update_user_billing_by_user_id(user_doc["user_id"], {"stripe_customer_id": customer_id})
    return customer_id


def _subscription_primary_item(subscription: Any) -> Optional[Dict[str, Any]]:
    try:
        items = subscription["items"]["data"] if isinstance(subscription, dict) else subscription.items.data
        if items:
            item = items[0]
            return dict(item) if not isinstance(item, dict) else item
    except Exception:
        return None
    return None


def _subscription_period_timestamp(subscription: Any, field: str) -> Optional[int]:
    item = _subscription_primary_item(subscription)
    if item:
        value = item.get(field)
        if value:
            try:
                return int(value)
            except Exception:
                pass
    try:
        value = subscription.get(field) if isinstance(subscription, dict) else getattr(subscription, field, None)
        return int(value) if value else None
    except Exception:
        return None


def _period_end_iso(subscription: Any) -> Optional[str]:
    value = _subscription_period_timestamp(subscription, "current_period_end")
    if not value:
        return None
    try:
        return datetime.fromtimestamp(value, tz=timezone.utc).isoformat()
    except Exception:
        return None


def _period_start_iso(subscription: Any) -> Optional[str]:
    value = _subscription_period_timestamp(subscription, "current_period_start")
    if not value:
        return None
    try:
        return datetime.fromtimestamp(value, tz=timezone.utc).isoformat()
    except Exception:
        return None


def _subscription_plan(subscription: Any) -> Optional[str]:
    try:
        items = subscription["items"]["data"]
        price_id = items[0]["price"]["id"] if items else None
    except Exception:
        price_id = None
    return _plan_from_price(price_id)


def _subscription_metadata(subscription: Any) -> Dict[str, Any]:
    metadata = subscription.get("metadata", {}) if isinstance(subscription, dict) else getattr(subscription, "metadata", {}) or {}
    return dict(metadata or {})


def _subscription_interval(subscription: Any) -> Optional[str]:
    metadata = _subscription_metadata(subscription)
    if metadata.get("interval"):
        return str(metadata.get("interval"))
    try:
        recurring = subscription["items"]["data"][0]["price"].get("recurring") or {}
        return recurring.get("interval")
    except Exception:
        return None


def _subscription_source(subscription: Any) -> Optional[str]:
    metadata = _subscription_metadata(subscription)
    return str(metadata.get("source") or "") or None


def _subscription_billing_updates(subscription: Any, *, last_payment_status: Optional[str] = None) -> Dict[str, Any]:
    metadata = _subscription_metadata(subscription)
    detected_plan = _subscription_plan(subscription)
    metadata_plan = str(metadata.get("plan") or "").strip() or None
    resolved_plan = detected_plan if detected_plan and detected_plan != "unknown" else metadata_plan or detected_plan or "unknown"
    updates = {
        "stripe_customer_id": subscription.get("customer"),
        "stripe_subscription_id": subscription.get("id"),
        "subscription_status": subscription.get("status"),
        "plan": resolved_plan,
        "interval": _subscription_interval(subscription),
        "source": _subscription_source(subscription),
        "current_period_start": _period_start_iso(subscription),
        "current_period_end": _period_end_iso(subscription),
    }
    if last_payment_status:
        updates["last_payment_status"] = last_payment_status
    return {key: value for key, value in updates.items() if value is not None}


def _stripe_configured() -> bool:
    return bool(os.environ.get("STRIPE_SECRET_KEY", "").strip())


def _discover_stripe_subscription(customer_id: Optional[str]) -> Optional[Any]:
    """Find an active/trialing subscription when checkout sync did not persist the id yet."""
    if not customer_id or not _stripe_configured():
        return None
    try:
        _stripe_secret_key()
        for status in ("active", "trialing", "past_due"):
            result = stripe.Subscription.list(customer=customer_id, status=status, limit=1)
            data = result.get("data") if isinstance(result, dict) else getattr(result, "data", None)
            if data:
                return _stripe_to_dict(data[0])
    except Exception as exc:
        logger.warning(
            "stripe_subscription_discover_failed customer_id=%s error=%s",
            customer_id,
            str(exc)[:200],
        )
    return None


async def _count_right_swipes_in_billing_period(user_id: str, billing: Dict[str, Any]) -> int:
    period_start, period_end = _billing_period_bounds(billing)
    rows = await db.swipes.find(
        {"user_id": user_id, "direction": "right"},
        {"_id": 0, "created_at": 1},
    ).to_list(5000)
    return _count_records_in_period(rows, period_start, period_end)


async def _repair_premium_credits_if_needed(user_id: str, user_doc: Dict[str, Any]) -> Dict[str, Any]:
    """Grant missing credits for paid users when Stripe sync/webhook did not initialize them."""
    billing = _billing_from_user(user_doc)
    status = billing.get("subscription_status") or "none"
    if status not in {"active", "trialing"}:
        return user_doc

    allowance = _billing_credit_limit(
        billing.get("plan"),
        True,
        interval=billing.get("interval"),
        source=billing.get("source"),
    )
    if allowance <= 0:
        return user_doc

    remaining = int(billing.get("credits_remaining") or 0)
    if remaining > 0:
        return user_doc

    total = int(billing.get("credits_total") or 0)
    if total > 0:
        used = await _count_right_swipes_in_billing_period(user_id, billing)
        if used > 0:
            return user_doc

    now = datetime.now(timezone.utc).isoformat()
    period_key = _billing_credit_period_key(billing)
    await db.users.update_one(
        {"user_id": user_id},
        {"$set": {
            "billing.credits_total": allowance,
            "billing.credits_remaining": allowance,
            "billing.credits_period_key": period_key,
            "billing.updated_at": now,
        }},
    )
    logger.info(
        "billing_credits_repaired user_id=%s plan=%s allowance=%s previous_total=%s",
        user_id,
        billing.get("plan"),
        allowance,
        total,
    )
    return await db.users.find_one({"user_id": user_id}, {"_id": 0}) or user_doc


async def _refresh_billing_from_stripe(
    user_doc: Dict[str, Any],
    *,
    discover: bool = True,
) -> tuple[Dict[str, Any], Optional[str]]:
    billing = _billing_from_user(user_doc)
    subscription_id = billing.get("stripe_subscription_id")
    if str(subscription_id or "").startswith("master_code_"):
        return user_doc, None
    customer_id = billing.get("stripe_customer_id")
    if discover and not customer_id:
        customer_id = await _discover_stripe_customer_for_user(user_doc)
    if discover and customer_id and customer_id != billing.get("stripe_customer_id"):
        await _update_user_billing_by_user_id(user_doc["user_id"], {"stripe_customer_id": customer_id})
        billing = {**billing, "stripe_customer_id": customer_id}
        user_doc = {**user_doc, "billing": billing}
    if discover and not subscription_id and customer_id:
        discovered = _discover_stripe_subscription(customer_id)
        if discovered:
            subscription_id = discovered.get("id") if isinstance(discovered, dict) else getattr(discovered, "id", None)
    if not subscription_id or not _stripe_configured():
        return user_doc, None
    try:
        _stripe_secret_key()
        subscription = _stripe_to_dict(stripe.Subscription.retrieve(subscription_id))
        updates = _subscription_billing_updates(subscription)
        await _update_user_billing_by_user_id(user_doc["user_id"], updates)
        refreshed = await db.users.find_one({"user_id": user_doc["user_id"]}, {"_id": 0})
        return refreshed or {**user_doc, "billing": {**billing, **updates}}, None
    except Exception as exc:
        logger.warning("stripe_billing_status_refresh_failed user_id=%s subscription_id=%s error=%s", user_doc.get("user_id"), subscription_id, str(exc)[:200])
        return user_doc, "stripe_refresh_failed"


async def _billing_apply_credit_status(user: User) -> Dict[str, Any]:
    user_doc = await _get_user_doc(user)
    user_doc, _warning = await _refresh_billing_from_stripe(user_doc)
    payload = _billing_status_payload(user_doc)
    return payload


async def _consume_application_credit(user_id: str) -> Dict[str, int]:
    user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0, "billing": 1})
    billing = _billing_from_user(user_doc)
    plan_total = int(billing.get("credits_total") or 0)
    plan_remaining = int(billing.get("credits_remaining") or 0)
    bonus_total = int(billing.get("referral_bonus_credits_total") or 0)
    bonus_remaining = int(billing.get("referral_bonus_credits_remaining") or 0)
    updates: Dict[str, Any] = {"billing.updated_at": datetime.now(timezone.utc).isoformat()}
    # Spend plan credits first, then referral bonus credits -- order doesn't
    # affect durability (bonus credits never expire/reset) but draining plan
    # credits first keeps the bonus pool as a visible buffer for longer.
    if plan_remaining > 0:
        plan_remaining -= 1
        updates["billing.credits_remaining"] = plan_remaining
    elif bonus_remaining > 0:
        bonus_remaining -= 1
        updates["billing.referral_bonus_credits_remaining"] = bonus_remaining
    await db.users.update_one({"user_id": user_id}, {"$set": updates})
    return {
        "credits_remaining": plan_remaining + bonus_remaining,
        "credits_total": plan_total + bonus_total,
    }


async def _refund_application_credit(user_id: str) -> None:
    """Give back the 1 credit an application consumed (e.g. the job offer
    expired before the user could actually be considered). We don't track
    which pool (plan vs. bonus) a given application drained, so refund into
    the plan pool, clamped to the plan allowance -- mirrors the direct-$set
    pattern used for bonus-credit grants in _grant_friend_referral_billing
    rather than routing through _update_user_billing_by_user_id, which would
    incorrectly reset the whole billing state as if newly subscribed."""
    user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0, "billing": 1})
    billing = _billing_from_user(user_doc)
    plan_total = int(billing.get("credits_total") or 0)
    plan_remaining = int(billing.get("credits_remaining") or 0)
    new_remaining = min(plan_total, plan_remaining + 1) if plan_total > 0 else plan_remaining + 1
    await db.users.update_one(
        {"user_id": user_id},
        {"$set": {
            "billing.credits_remaining": new_remaining,
            "billing.updated_at": datetime.now(timezone.utc).isoformat(),
        }},
    )


@api_router.post("/billing/create-checkout-session")
async def create_billing_checkout_session(
    body: BillingCheckoutRequest,
    request: Request,
    user: User = Depends(get_current_user),
):
    _stripe_secret_key()
    checkout_source = body.source or "app"
    billing_plan = body.plan if checkout_source == "onboarding" else _canonical_billing_plan(body.plan)
    billing_interval = body.interval or ("quarterly" if checkout_source == "onboarding" and billing_plan == "quarterly" else "monthly")
    user_doc = await _get_user_doc(user)
    existing_billing = _billing_from_user(user_doc)
    existing_subscription_id = existing_billing.get("stripe_subscription_id")
    if (
        (existing_billing.get("subscription_status") or "none") in {"active", "trialing"}
        and existing_subscription_id
        and not str(existing_subscription_id).startswith("master_code_")
    ):
        raise HTTPException(
            status_code=409,
            detail="You already have an active subscription. Use the upgrade option to change plans.",
        )
    customer_id = await _stripe_customer_for_user(user_doc)
    marketing_url = _marketing_url()
    app_url = _app_url()
    if checkout_source == "onboarding":
        success_url = f"{marketing_url}/onboarding?step=creatorAccessCode&checkout=success&session_id={{CHECKOUT_SESSION_ID}}"
        # Back from Stripe should enter the app (frontend finishes onboarding), not re-show paywall.
        cancel_url = f"{marketing_url}/onboarding?checkout=cancelled"
    elif checkout_source == "app":
        return_path = _sanitize_return_path(body.return_path)
        success_url = _checkout_return_url(app_url, return_path, "success")
        cancel_url = _checkout_return_url(app_url, return_path, "cancelled")
    else:
        success_url = _checkout_return_url(app_url, "/swipe", "success")
        cancel_url = _checkout_return_url(app_url, "/swipe", "cancelled")
    base_metadata = {
        "user_id": user.user_id,
        "plan": billing_plan,
        "interval": billing_interval,
        "source": checkout_source,
    }
    datafast_metadata = datafast_stripe_metadata(cookies=request.cookies, body=body.model_dump())
    checkout_metadata = merge_stripe_metadata(base_metadata, datafast_metadata)
    session_kwargs: Dict[str, Any] = {
        "mode": "subscription",
        "customer": customer_id,
        "line_items": [{"price": _stripe_price_for_plan(billing_plan, interval=billing_interval, source=checkout_source), "quantity": 1}],
        "success_url": success_url,
        "cancel_url": cancel_url,
        "client_reference_id": user.user_id,
        "metadata": checkout_metadata,
        "subscription_data": {"metadata": checkout_metadata},
    }
    # First-ever checkout for someone who signed up with a friend's referral
    # code -- honor the 25%-off-first-invoice promise from the referral
    # campaign. Stripe doesn't allow combining a preset discount with
    # allow_promotion_codes, so only offer the manual promo box otherwise.
    signup_discount_eligible = (
        not existing_subscription_id
        and await has_redeemed_friend_referral_code(db, user.user_id)
    )
    if signup_discount_eligible:
        coupon_id = await asyncio.to_thread(_ensure_friend_referral_signup_discount_coupon)
        session_kwargs["discounts"] = [{"coupon": coupon_id}]
    else:
        session_kwargs["allow_promotion_codes"] = True
    session = stripe.checkout.Session.create(**session_kwargs)
    if _posthog_client is not None:
        with new_context():
            identify_context(user.user_id)
            posthog_capture(
                "checkout_started",
                properties={
                    "billing_plan": billing_plan,
                    "billing_interval": billing_interval,
                    "checkout_source": checkout_source,
                    "has_referral_discount": signup_discount_eligible,
                },
            )
    return {"url": session["url"]}


@api_router.post("/billing/redeem-master-code")
async def redeem_billing_master_code(body: BillingMasterCodeRequest, user: User = Depends(get_current_user)):
    code = (body.code or "").strip()
    if not re.fullmatch(r"\d{6}", code):
        raise HTTPException(status_code=400, detail="Enter a valid 6-digit access code")
    if not _is_master_billing_code(code):
        raise HTTPException(status_code=404, detail="Access code not found")

    billing = await _grant_master_billing_access(
        user.user_id,
        body.plan,
        interval=body.interval,
        source=body.source,
    )
    await _set_user_demo_account(user.user_id, True)
    await _ensure_demo_feed_profile(user.user_id)
    return {
        "ok": True,
        "master_code": True,
        "demo_account": True,
        "billing": billing,
    }


@api_router.get("/billing/status")
async def billing_status(user: User = Depends(get_current_user)):
    """Read billing from the user record. Does not recover foreign Stripe checkouts."""
    user_doc = await _get_user_doc(user)
    billing = _billing_from_user(user_doc)
    warning = None
    if billing.get("stripe_subscription_id") and _stripe_configured():
        user_doc, warning = await _refresh_billing_from_stripe(user_doc, discover=False)
    payload = _billing_status_payload(user_doc)
    if warning:
        payload["warning"] = warning
    return payload


@api_router.post("/billing/sync")
async def sync_billing(user: User = Depends(get_current_user)):
    """Force a Stripe refresh and repair missing credit grants after checkout."""
    user_doc = await _get_user_doc(user)
    user_doc, warning = await _refresh_billing_from_stripe(user_doc)
    user_doc = await _recover_stalled_checkout_billing(user_doc)
    user_doc = await _repair_premium_credits_if_needed(user.user_id, user_doc)
    payload = _billing_status_payload(user_doc)
    if warning:
        payload["warning"] = warning
    payload["synced"] = True
    return payload


def _checkout_session_user_id(session_obj: Dict[str, Any]) -> Optional[str]:
    metadata = session_obj.get("metadata") or {}
    return session_obj.get("client_reference_id") or metadata.get("user_id")


async def _checkout_session_belongs_to_user(session_obj: Dict[str, Any], user_doc: Dict[str, Any]) -> bool:
    user_id = user_doc.get("user_id")
    session_user_id = _checkout_session_user_id(session_obj)
    metadata_user_id = (session_obj.get("metadata") or {}).get("user_id")
    for bound_user_id in (session_user_id, metadata_user_id):
        if bound_user_id and bound_user_id != user_id:
            return False
    if session_user_id and session_user_id == user_id:
        return True
    if metadata_user_id and metadata_user_id == user_id:
        return True
    customer_id = _stripe_object_id(session_obj.get("customer"))
    user_customer_id = _billing_from_user(user_doc).get("stripe_customer_id")
    if customer_id and user_customer_id and customer_id == user_customer_id:
        return True
    return False


async def _ensure_checkout_entitlements_from_session(user_id: str, session_obj: Dict[str, Any]) -> None:
    """Grant plan/credits from checkout metadata when Stripe subscription sync is incomplete."""
    payment_status = session_obj.get("payment_status")
    session_status = session_obj.get("status")
    if payment_status != "paid" and session_status != "complete":
        return
    if session_obj.get("mode") != "subscription":
        return

    user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0}) or {}
    billing = _billing_from_user(user_doc)
    if (billing.get("subscription_status") or "none") in {"active", "trialing"} and int(billing.get("credits_remaining") or 0) > 0:
        return

    metadata = session_obj.get("metadata") or {}
    now = datetime.now(timezone.utc)
    period_end = now + timedelta(days=30)
    updates: Dict[str, Any] = {
        "subscription_status": "active",
        "plan": metadata.get("plan") or billing.get("plan") or "monthly",
        "interval": metadata.get("interval") or billing.get("interval") or "monthly",
        "source": metadata.get("source") or billing.get("source") or "app",
        "last_payment_status": payment_status or "paid",
        "current_period_start": billing.get("current_period_start") or now.isoformat(),
        "current_period_end": billing.get("current_period_end") or period_end.isoformat(),
    }
    customer_id = _stripe_object_id(session_obj.get("customer")) or billing.get("stripe_customer_id")
    subscription_id = _stripe_object_id(session_obj.get("subscription")) or billing.get("stripe_subscription_id")
    if customer_id:
        updates["stripe_customer_id"] = customer_id
    if subscription_id:
        updates["stripe_subscription_id"] = subscription_id

    await _update_user_billing_by_user_id(user_id, updates)
    user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0}) or user_doc
    await _repair_premium_credits_if_needed(user_id, user_doc)
    logger.info(
        "checkout_entitlements_fallback_applied user_id=%s plan=%s source=%s",
        user_id,
        updates.get("plan"),
        updates.get("source"),
    )


async def _recover_stalled_checkout_billing(user_doc: Dict[str, Any]) -> Dict[str, Any]:
    """Repair billing for paid customers whose subscription state never persisted after checkout."""
    user_id = user_doc.get("user_id")
    if not user_id:
        return user_doc

    billing = _billing_from_user(user_doc)
    if (billing.get("subscription_status") or "none") in {"active", "trialing"}:
        return await _repair_premium_credits_if_needed(user_id, user_doc)

    customer_id = billing.get("stripe_customer_id") or await _discover_stripe_customer_for_user(user_doc)
    if not customer_id or not _stripe_configured():
        return user_doc

    try:
        _stripe_secret_key()
        discovered = _discover_stripe_subscription(customer_id)
        if discovered:
            updates = _subscription_billing_updates(discovered, last_payment_status="paid")
            await _update_user_billing_by_user_id(user_id, {key: value for key, value in updates.items() if value})
            user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0}) or user_doc
            return await _repair_premium_credits_if_needed(user_id, user_doc)

        sessions = stripe.checkout.Session.list(customer=customer_id, limit=5, status="complete")
        data = sessions.get("data") if isinstance(sessions, dict) else getattr(sessions, "data", [])
        for session in data or []:
            session_dict = _stripe_to_dict(session)
            if session_dict.get("payment_status") != "paid":
                continue
            if not await _checkout_session_belongs_to_user(session_dict, user_doc):
                continue
            await _apply_checkout_session_billing(session_dict, fallback_user_id=user_id)
            await _ensure_checkout_entitlements_from_session(user_id, session_dict)
            user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0}) or user_doc
            billing = _billing_from_user(user_doc)
            if (billing.get("subscription_status") or "none") in {"active", "trialing"}:
                return await _repair_premium_credits_if_needed(user_id, user_doc)
    except Exception as exc:
        logger.warning(
            "recover_stalled_checkout_billing_failed user_id=%s error=%s",
            user_id,
            str(exc)[:200],
        )
    return user_doc


async def _apply_checkout_session_billing(
    session_obj: Dict[str, Any],
    *,
    event_id: Optional[str] = None,
    fallback_user_id: Optional[str] = None,
) -> Optional[str]:
    customer_id = _stripe_object_id(session_obj.get("customer"))
    subscription_ref = session_obj.get("subscription")
    subscription_id = _stripe_object_id(subscription_ref)
    if not customer_id:
        logger.warning("stripe_checkout_completed_missing_customer event_id=%s", event_id)
    updates = {
        "stripe_customer_id": customer_id,
        "stripe_subscription_id": subscription_id,
        "last_payment_status": session_obj.get("payment_status"),
    }
    if subscription_id:
        try:
            if isinstance(subscription_ref, dict):
                subscription = subscription_ref
            else:
                subscription = _stripe_to_dict(stripe.Subscription.retrieve(subscription_id))
            updates.update(_subscription_billing_updates(subscription, last_payment_status=session_obj.get("payment_status")))
        except Exception as exc:
            logger.warning(
                "stripe_checkout_subscription_retrieve_failed event_id=%s subscription_id=%s error=%s",
                event_id,
                subscription_id,
                str(exc)[:200],
            )
    user_id = await _resolve_user_id_from_checkout_session(session_obj, fallback_user_id=fallback_user_id)
    if user_id:
        await _update_user_billing_by_user_id(user_id, {key: value for key, value in updates.items() if value})
        await _finalize_user_billing(user_id)
        await _ensure_checkout_entitlements_from_session(user_id, session_obj)
        logger.info(
            "stripe_checkout_billing_applied event_id=%s user_id=%s customer_id=%s subscription_id=%s",
            event_id,
            user_id,
            customer_id,
            subscription_id,
        )
        return user_id
    if customer_id:
        await _update_user_billing_by_customer_id(customer_id, {key: value for key, value in updates.items() if value})
    else:
        logger.warning("stripe_checkout_completed_no_user_or_customer event_id=%s", event_id)
    return None


@api_router.post("/billing/confirm-checkout")
async def confirm_billing_checkout(
    body: BillingConfirmCheckoutRequest,
    user: User = Depends(get_current_user),
):
    _stripe_secret_key()
    session_id = (body.session_id or "").strip()
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id is required")
    try:
        session_obj = stripe.checkout.Session.retrieve(
            session_id,
            expand=["subscription", "customer", "line_items"],
        )
    except Exception as exc:
        logger.warning("stripe_checkout_confirm_retrieve_failed session_id=%s error=%s", session_id, str(exc)[:200])
        raise HTTPException(status_code=400, detail="Invalid checkout session")

    session_dict = _stripe_to_dict(session_obj)
    user_doc = await _get_user_doc(user)
    if not await _checkout_session_belongs_to_user(session_dict, user_doc):
        raise HTTPException(status_code=403, detail="Checkout session does not belong to this user")

    payment_status = session_dict.get("payment_status")
    session_status = session_dict.get("status")
    if payment_status != "paid" and session_status != "complete":
        payload = _billing_status_payload(user_doc)
        payload["checkout_pending"] = True
        return payload

    await _apply_checkout_session_billing(session_dict, fallback_user_id=user.user_id)
    user_doc = await _finalize_user_billing(user.user_id) or user_doc
    await _ensure_checkout_entitlements_from_session(user.user_id, session_dict)
    user_doc = await db.users.find_one({"user_id": user.user_id}, {"_id": 0}) or user_doc
    payload = _billing_status_payload(user_doc)
    payload["confirmed"] = True
    return payload


@api_router.get("/billing/usage")
async def billing_usage(user: User = Depends(get_current_user)):
    user_doc = await _get_user_doc(user)
    user_doc, _warning = await _refresh_billing_from_stripe(user_doc)
    billing = _billing_from_user(user_doc)
    status_payload = _billing_status_payload(user_doc)
    is_premium = bool(status_payload.get("is_premium"))
    plan = status_payload.get("plan")
    credits_total = int(status_payload.get("credits_total") or _billing_credit_limit(
        plan,
        is_premium,
        interval=status_payload.get("interval"),
        source=status_payload.get("source"),
    ))
    credits_remaining = int(status_payload.get("credits_remaining") or 0)
    period_start, period_end = _billing_period_bounds(billing)

    usage_rows = await db.swipes.find(
        {"user_id": user.user_id, "direction": "right"},
        {"_id": 0, "created_at": 1},
    ).to_list(5000)
    credits_used = max(0, credits_total - credits_remaining) if credits_total else 0
    usage_pct = round((credits_used / credits_total) * 100) if credits_total else 0

    return {
        "period_start": period_start.isoformat(),
        "period_end": period_end.isoformat(),
        "credits_used": credits_used,
        "credits_total": credits_total,
        "credits_remaining": credits_remaining,
        "usage_percent": usage_pct,
        "is_premium": is_premium,
        "plan": plan,
        "interval": status_payload.get("interval"),
        "source": status_payload.get("source"),
        "daily_usage": {
            "7d": _series_counts(usage_rows, 7, lambda item: item.get("created_at")),
            "14d": _series_counts(usage_rows, 14, lambda item: item.get("created_at")),
            "30d": _series_counts(usage_rows, 30, lambda item: item.get("created_at")),
        },
    }


@api_router.post("/billing/create-portal-session")
async def create_billing_portal_session(user: User = Depends(get_current_user)):
    _stripe_secret_key()
    user_doc = await _get_user_doc(user)
    customer_id = _billing_from_user(user_doc).get("stripe_customer_id")
    if not customer_id:
        raise HTTPException(status_code=400, detail="No Stripe customer found")
    session = stripe.billing_portal.Session.create(
        customer=customer_id,
        return_url=f"{_app_url()}/billing",
    )
    return {"url": session["url"]}


@api_router.post("/referrals/friends/enroll")
async def enroll_friend_referral_route(user: User = Depends(get_current_user)):
    status = await enroll_friend_referral(db, user.user_id)
    await _store_friend_referral_analytics(
        "friend_referral_code_generated",
        user_id=user.user_id,
        properties={
            "code": status.get("code"),
            "uses_count": status.get("uses_count"),
        },
    )
    return status


@api_router.get("/referrals/friends/status")
async def friend_referral_status_route(user: User = Depends(get_current_user)):
    user_doc = await _get_user_doc(user)
    return friend_referral_status_payload(user_doc)


@api_router.get("/referrals/friends/validate/{code}")
async def validate_friend_referral_route(code: str, user: User = Depends(get_current_user)):
    return await validate_friend_referral_code(db, code=code, user_id=user.user_id)


@api_router.post("/referrals/friends/redeem")
async def redeem_friend_referral_route(payload: Dict[str, Any], user: User = Depends(get_current_user)):
    code = (payload.get("code") or "").strip()
    try:
        result = await redeem_friend_referral_code(
            db,
            code=code,
            redeemer_user_id=user.user_id,
            redeemer_email=user.email,
            send_use_email=send_friend_referral_used_email,
            send_reward_email=send_friend_referral_reward_email,
            grant_reward=_grant_friend_referral_billing,
            app_url=_app_url(),
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    await _store_friend_referral_analytics(
        "friend_referral_code_redeemed",
        user_id=user.user_id,
        properties={
            "code": code.upper(),
            "referrer_uses_count": result.get("uses_count"),
            "reward_unlocked": result.get("reward_unlocked"),
        },
    )
    referrer_id = result.get("referrer_user_id")
    if referrer_id:
        await _store_friend_referral_analytics(
            "friend_referral_invite_received",
            user_id=referrer_id,
            properties={
                "code": code.upper(),
                "uses_count": result.get("uses_count"),
                "reward_unlocked": result.get("reward_unlocked"),
            },
        )
        if result.get("reward_unlocked"):
            await _store_friend_referral_analytics(
                "friend_referral_reward_unlocked",
                user_id=referrer_id,
                properties={"uses_count": result.get("uses_count")},
            )
    return result


@api_router.post("/referrals/friends/claim")
async def claim_friend_referral_route(payload: Dict[str, Any], user: User = Depends(get_current_user)):
    token = (payload.get("token") or "").strip() or None
    try:
        status = await claim_friend_referral_reward(db, user.user_id, token)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    await _store_friend_referral_analytics(
        "friend_referral_reward_claimed",
        user_id=user.user_id,
        properties={
            "uses_count": status.get("uses_count"),
            "reward_batches_granted": status.get("reward_batches_granted"),
            "credits_earned_total": status.get("credits_earned_total"),
        },
    )
    user_doc = await _get_user_doc(user)
    return {**status, "billing": _billing_status_payload(user_doc)}


_STRIPE_UPGRADE_PORTAL_CONFIG_CACHE: Dict[str, str] = {}


async def _stripe_upgrade_portal_configuration_id() -> str:
    """Lazily create (and cache) a Stripe Billing Portal configuration that allows
    subscription plan changes with proration, scoped to basic/pro/ultra prices only."""
    cached = _STRIPE_UPGRADE_PORTAL_CONFIG_CACHE.get("id")
    if cached:
        return cached

    settings_doc = await db.app_settings.find_one({"_id": "stripe_upgrade_portal_configuration"})
    if settings_doc and settings_doc.get("configuration_id"):
        _STRIPE_UPGRADE_PORTAL_CONFIG_CACHE["id"] = settings_doc["configuration_id"]
        return settings_doc["configuration_id"]

    price_ids: List[str] = []
    for plan in ("basic", "pro", "ultra"):
        for env_name in (f"STRIPE_PRICE_{plan.upper()}", f"STRIPE_PRICE_{plan.upper()}_WEEKLY"):
            price_id = os.environ.get(env_name, "").strip()
            if price_id:
                price_ids.append(price_id)
    if not price_ids:
        raise HTTPException(status_code=503, detail="Stripe plan prices are not configured")

    products_by_id: Dict[str, List[str]] = {}
    for price_id in price_ids:
        price = stripe.Price.retrieve(price_id)
        product = price["product"]
        product_id = product if isinstance(product, str) else product["id"]
        products_by_id.setdefault(product_id, []).append(price_id)

    configuration = stripe.billing_portal.Configuration.create(
        business_profile={"headline": "Manage your Hirly subscription"},
        features={
            "subscription_update": {
                "enabled": True,
                "default_allowed_updates": ["price"],
                "proration_behavior": "create_prorations",
                "products": [
                    {"product": product_id, "prices": prices}
                    for product_id, prices in products_by_id.items()
                ],
            },
            "subscription_cancel": {"enabled": True},
            "payment_method_update": {"enabled": True},
            "invoice_history": {"enabled": True},
        },
    )
    configuration_id = configuration["id"]
    await db.app_settings.update_one(
        {"_id": "stripe_upgrade_portal_configuration"},
        {"$set": {"configuration_id": configuration_id}},
        upsert=True,
    )
    _STRIPE_UPGRADE_PORTAL_CONFIG_CACHE["id"] = configuration_id
    return configuration_id


@api_router.post("/billing/create-upgrade-session")
async def create_billing_upgrade_session(
    body: BillingUpgradeSessionRequest,
    user: User = Depends(get_current_user),
):
    """Upgrade an existing active subscription to a higher tier with Stripe-managed
    proration. The user only pays the prorated difference immediately, and — once the
    webhook confirms the change — is granted the *difference* in application credits
    rather than a full reset (see _merge_billing_credit_state)."""
    _stripe_secret_key()
    user_doc = await _get_user_doc(user)
    billing = _billing_from_user(user_doc)
    subscription_id = billing.get("stripe_subscription_id")
    is_premium = (billing.get("subscription_status") or "none") in {"active", "trialing"}
    if not is_premium or not subscription_id or str(subscription_id).startswith("master_code_"):
        raise HTTPException(status_code=400, detail="No active subscription to upgrade")

    target_plan = _canonical_billing_plan(body.plan)
    if target_plan not in {"basic", "pro", "ultra"}:
        raise HTTPException(status_code=400, detail="Unsupported billing plan")
    target_interval = body.interval or billing.get("interval") or "monthly"

    current_allowance = _billing_credit_limit(
        billing.get("plan"), True, interval=billing.get("interval"), source=billing.get("source"),
    )
    target_allowance = _billing_credit_limit(target_plan, True, interval=target_interval, source="app")
    if target_allowance <= current_allowance:
        raise HTTPException(status_code=400, detail="Select a higher tier to upgrade")

    try:
        subscription = stripe.Subscription.retrieve(subscription_id)
    except Exception as exc:
        logger.warning("stripe_upgrade_subscription_retrieve_failed user_id=%s error=%s", user.user_id, str(exc)[:200])
        raise HTTPException(status_code=400, detail="Could not load your current subscription")
    items = subscription["items"]["data"]
    if not items:
        raise HTTPException(status_code=400, detail="Subscription has no line items")
    item_id = items[0]["id"]
    new_price_id = _stripe_price_for_plan(target_plan, interval=target_interval, source="app")

    customer_id = await _stripe_customer_for_user(user_doc)
    configuration_id = await _stripe_upgrade_portal_configuration_id()
    return_path = _sanitize_return_path(body.return_path)
    return_url = _checkout_return_url(_app_url(), return_path, "success")

    session = stripe.billing_portal.Session.create(
        customer=customer_id,
        configuration=configuration_id,
        return_url=return_url,
        flow_data={
            "type": "subscription_update_confirm",
            "subscription_update_confirm": {
                "subscription": subscription_id,
                "items": [{"id": item_id, "price": new_price_id, "quantity": 1}],
            },
            "after_completion": {
                "type": "redirect",
                "redirect": {"return_url": return_url},
            },
        },
    )
    return {"url": session["url"]}


async def _resolve_user_id_for_stripe_subscription(subscription: Any) -> Optional[str]:
    """Resolve a subscription even when its Stripe customer is not linked yet."""
    metadata_user_id = str(_subscription_metadata(subscription).get("user_id") or "").strip()
    if metadata_user_id:
        found = await db.users.find_one({"user_id": metadata_user_id}, {"_id": 0, "user_id": 1})
        if found:
            return metadata_user_id

    customer = subscription.get("customer")
    customer_id = _stripe_object_id(customer)
    if customer_id:
        resolved = await _resolve_user_id_for_stripe_customer(customer_id)
        if resolved:
            return resolved

    if isinstance(customer, dict):
        customer_email = str(customer.get("email") or "").strip().lower()
    else:
        customer_email = str(getattr(customer, "email", "") or "").strip().lower()
    if customer_email:
        user_doc = await _ensure_app_user_for_billing_email(customer_email)
        if user_doc:
            return user_doc.get("user_id")
    return None


async def _handle_subscription_event(subscription: Any, *, last_payment_status: Optional[str] = None) -> None:
    customer_id = _stripe_object_id(subscription.get("customer"))
    if not customer_id:
        logger.warning("stripe_subscription_event_missing_customer subscription_id=%s", subscription.get("id"))
        return
    updates = _subscription_billing_updates(subscription, last_payment_status=last_payment_status)
    user_id = await _resolve_user_id_for_stripe_subscription(subscription)
    if user_id:
        await _update_user_billing_by_user_id(user_id, updates)
        await _finalize_user_billing(user_id)
        return
    logger.error(
        "stripe_subscription_event_no_app_user customer_id=%s subscription_id=%s",
        customer_id,
        subscription.get("id"),
    )

def _list_stripe_subscriptions_for_reconcile() -> List[Any]:
    _stripe_secret_key()
    result = stripe.Subscription.list(status="all", limit=100, expand=["data.customer"])
    return [_stripe_to_dict(subscription) for subscription in result.auto_paging_iter()]


def _stripe_subscription_reconcile_priority(subscription: Any) -> tuple[int, int]:
    status = str(subscription.get("status") or "").lower()
    return (2 if status in {"active", "trialing"} else 1), int(subscription.get("created") or 0)


async def _reconcile_stripe_subscriptions_once() -> Dict[str, int]:
    """Repair webhook misses without rewriting already-synced billing records."""
    if not _stripe_configured():
        return {"scanned": 0, "updated": 0, "unmatched": 0}

    subscriptions = await asyncio.to_thread(_list_stripe_subscriptions_for_reconcile)
    current_by_customer: Dict[str, Any] = {}
    for subscription in subscriptions:
        customer_id = _stripe_object_id(subscription.get("customer"))
        if not customer_id:
            continue
        current = current_by_customer.get(customer_id)
        if current is None or _stripe_subscription_reconcile_priority(subscription) > _stripe_subscription_reconcile_priority(current):
            current_by_customer[customer_id] = subscription

    updated = 0
    unmatched = 0
    for customer_id, subscription in current_by_customer.items():
        # Isolate each customer -- one bad/unresolvable record (Stripe lookup
        # failure, missing app user, etc.) must not abort the whole scan and
        # leave every other subscriber behind it un-reconciled until the next
        # pass. Every customer gets an independent chance to be repaired.
        try:
            user_id = await _resolve_user_id_for_stripe_subscription(subscription)
            if not user_id:
                unmatched += 1
                continue

            user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0})
            if not user_doc:
                unmatched += 1
                continue
            existing_billing = _billing_from_user(user_doc)
            updates = _subscription_billing_updates(subscription)
            updates["stripe_customer_id"] = customer_id
            comparable_updates = {key: value for key, value in updates.items() if key != "last_payment_status"}
            if all(existing_billing.get(key) == value for key, value in comparable_updates.items()):
                continue

            await _update_user_billing_by_user_id(user_id, updates)
            refreshed = await db.users.find_one({"user_id": user_id}, {"_id": 0}) or user_doc
            await _repair_premium_credits_if_needed(user_id, refreshed)
            updated += 1
        except Exception as exc:
            logger.warning(
                "stripe_subscription_reconcile_item_failed customer_id=%s subscription_id=%s error=%s",
                customer_id,
                subscription.get("id"),
                str(exc)[:200],
            )

    logger.info(
        "stripe_subscription_reconcile_complete scanned=%s updated=%s unmatched=%s",
        len(current_by_customer), updated, unmatched,
    )
    return {"scanned": len(current_by_customer), "updated": updated, "unmatched": unmatched}


async def _run_stripe_subscription_reconcile_loop() -> None:
    interval_seconds = max(60, int(os.environ.get("STRIPE_RECONCILE_INTERVAL_SECONDS", "300") or 300))
    while True:
        try:
            await _reconcile_stripe_subscriptions_once()
        except asyncio.CancelledError:
            raise
        except Exception as exc:
            logger.warning("stripe_subscription_reconcile_failed error=%s", str(exc)[:200])
        await asyncio.sleep(interval_seconds)

async def _stripe_event_already_processed(event_id: str) -> bool:
    if not event_id:
        return False
    try:
        existing = await db.stripe_events.find_one({"event_id": event_id}, {"_id": 0, "event_id": 1})
    except Exception as exc:
        # A transient DB hiccup here must not crash the whole webhook request
        # (billing processing is idempotent either way) -- fail open and let
        # the event through rather than 500ing every delivery Stripe sends.
        logger.warning("stripe_event_dedupe_check_failed event_id=%s error=%s", event_id, str(exc)[:200])
        return False
    return bool(existing)


async def _record_processed_stripe_event(event: Any) -> None:
    event_id = event.get("id")
    if not event_id:
        return
    created = event.get("created")
    created_at = None
    if created:
        try:
            created_at = datetime.fromtimestamp(int(created), tz=timezone.utc).isoformat()
        except Exception:
            created_at = None
    try:
        await db.stripe_events.insert_one({
            "event_id": event_id,
            "type": event.get("type"),
            "created_at": created_at,
            "processed_at": datetime.now(timezone.utc).isoformat(),
        })
    except Exception as exc:
        logger.info("stripe_event_record_duplicate_or_failed event_id=%s error=%s", event_id, str(exc)[:160])


_POSTHOG_REVENUE_UUID_NAMESPACE = uuid.UUID("f637c7f8-e2ef-5ee7-a976-c3017843ab42")
_POSTHOG_ZERO_DECIMAL_CURRENCIES = {
    "BIF", "CLP", "DJF", "GNF", "JPY", "KMF", "KRW", "MGA", "PYG", "RWF",
    "UGX", "VND", "VUV", "XAF", "XOF", "XPF",
}
_POSTHOG_THREE_DECIMAL_CURRENCIES = {"BHD", "JOD", "KWD", "OMR", "TND"}
_POSTHOG_TWO_DECIMAL_CURRENCIES = {
    "AED", "AFN", "ALL", "AMD", "ANG", "AOA", "ARS", "AUD", "AWG", "AZN",
    "BAM", "BBD", "BDT", "BGN", "BMD", "BND", "BOB", "BRL", "BSD", "BTN",
    "BWP", "BYN", "BZD", "CAD", "CDF", "CHF", "CNY", "COP", "CRC", "CVE",
    "CZK", "DKK", "DOP", "DZD", "EGP", "ETB", "EUR", "FJD", "FKP", "GBP",
    "GEL", "GIP", "GMD", "GTQ", "GYD", "HKD", "HNL", "HRK", "HTG", "HUF",
    "IDR", "ILS", "INR", "ISK", "JMD", "KES", "KGS", "KHR", "KYD", "KZT",
    "LAK", "LBP", "LKR", "LRD", "LSL", "MAD", "MDL", "MKD", "MMK", "MNT",
    "MOP", "MUR", "MVR", "MWK", "MXN", "MYR", "MZN", "NAD", "NGN", "NIO",
    "NOK", "NPR", "NZD", "PAB", "PEN", "PGK", "PHP", "PKR", "PLN", "QAR",
    "RON", "RSD", "SAR", "SBD", "SCR", "SEK", "SGD", "SHP", "SLE", "SOS",
    "SRD", "STD", "SZL", "THB", "TJS", "TOP", "TRY", "TTD", "TWD", "TZS",
    "UAH", "USD", "UYU", "UZS", "WST", "XCD", "YER", "ZAR", "ZMW",
}
_POSTHOG_ANALYTICS_TIMEOUT_SECONDS = 1.0
_POSTHOG_STRIPE_ENRICHMENT_TIMEOUT_SECONDS = 0.20


def _posthog_server_capture_configured() -> bool:
    return bool(
        os.environ.get("POSTHOG_SERVER_API_KEY", "").strip()
        and _posthog_capture_url()
    )


def _posthog_revenue_enabled(event_name: str) -> bool:
    env_name = (
        "POSTHOG_PAYMENT_REVENUE_ENABLED"
        if event_name == "payment_succeeded"
        else "POSTHOG_REFUND_REVENUE_ENABLED"
    )
    return _posthog_server_capture_configured() and os.environ.get(env_name, "").strip().lower() == "true"


def _posthog_revenue_uuid(event_name: str, object_type: str, object_id: str) -> str:
    semantic_key = f"{event_name}:{object_type}:{object_id}"
    return str(uuid.uuid5(_POSTHOG_REVENUE_UUID_NAMESPACE, semantic_key))


def _posthog_major_amount(amount_minor: int, currency: str) -> Decimal:
    normalized_currency = str(currency or "").upper()
    if normalized_currency in _POSTHOG_ZERO_DECIMAL_CURRENCIES:
        exponent = 0
    elif normalized_currency in _POSTHOG_THREE_DECIMAL_CURRENCIES:
        exponent = 3
    elif normalized_currency in _POSTHOG_TWO_DECIMAL_CURRENCIES:
        exponent = 2
    else:
        raise ValueError("unsupported currency for PostHog revenue conversion")
    return Decimal(int(amount_minor)).scaleb(-exponent)


def _posthog_capture_url() -> Optional[str]:
    host = os.environ.get("POSTHOG_HOST", "").strip()
    if not host:
        return None
    try:
        parsed = urlparse(host)
        port = parsed.port
    except ValueError:
        return None
    if (
        parsed.scheme != "https"
        or not parsed.hostname
        or port is not None
        or parsed.username
        or parsed.password
        or parsed.query
        or parsed.fragment
        or parsed.path.rstrip("/")
    ):
        return None
    hostname = parsed.hostname.lower()
    configured_hosts = {
        item.strip().lower()
        for item in os.environ.get("POSTHOG_ALLOWED_HOSTS", "").split(",")
        if item.strip()
    }
    canonical_host = hostname in {"us.i.posthog.com", "eu.i.posthog.com"}
    if not canonical_host and hostname not in configured_hosts:
        return None
    return f"https://{hostname}/capture/"


def _posthog_stripe_timestamp(event: Dict[str, Any]) -> Optional[str]:
    try:
        created = int(event.get("created"))
    except (TypeError, ValueError):
        return None
    return datetime.fromtimestamp(created, tz=timezone.utc).isoformat().replace("+00:00", "Z")


def _posthog_price_properties(source: Dict[str, Any]) -> Dict[str, Any]:
    lines = ((source.get("lines") or {}).get("data") or [])
    if not lines:
        lines = ((source.get("items") or {}).get("data") or [])
    line = lines[0] if lines else {}
    price = line.get("price") or {}
    if isinstance(price, str):
        price = {"id": price}
    product = price.get("product")
    properties = {
        "price_id": _stripe_object_id(price),
        "product_id": _stripe_object_id(product),
    }
    metadata = source.get("metadata") or {}
    plan = metadata.get("plan") if isinstance(metadata, dict) else None
    if plan:
        properties["plan"] = str(plan)
    return {key: value for key, value in properties.items() if value}


async def _posthog_existing_user_id(*candidate_ids: Any) -> Optional[str]:
    for candidate in candidate_ids:
        user_id = str(candidate or "").strip()
        if not user_id:
            continue
        found = await db.users.find_one({"user_id": user_id}, {"_id": 0, "user_id": 1})
        if found and found.get("user_id") == user_id:
            return user_id
    return None


async def _posthog_resolve_billing_user_id(
    *,
    metadata: Optional[Dict[str, Any]] = None,
    subscription_id: Optional[str] = None,
    customer_id: Optional[str] = None,
) -> Optional[str]:
    metadata_user_id = (metadata or {}).get("user_id")
    resolved = await _posthog_existing_user_id(metadata_user_id)
    if resolved:
        return resolved
    if subscription_id:
        found = await db.users.find_one(
            {"billing.stripe_subscription_id": subscription_id},
            {"_id": 0, "user_id": 1},
        )
        if found and found.get("user_id"):
            return str(found["user_id"])
    if customer_id:
        found = await db.users.find_one(
            {"billing.stripe_customer_id": customer_id},
            {"_id": 0, "user_id": 1},
        )
        if found and found.get("user_id"):
            return str(found["user_id"])
    return None


async def _capture_posthog_server_event(
    *,
    event_name: str,
    distinct_id: str,
    timestamp: str,
    semantic_uuid: str,
    properties: Dict[str, Any],
) -> bool:
    api_key = os.environ.get("POSTHOG_SERVER_API_KEY", "").strip()
    capture_url = _posthog_capture_url()
    if not api_key or not capture_url:
        logger.warning("posthog_server_capture_skipped_invalid_config event=%s", event_name)
        return False

    body = {
        "api_key": api_key,
        "event": event_name,
        "distinct_id": distinct_id,
        "timestamp": timestamp,
        "uuid": semantic_uuid,
        "properties": properties,
    }
    timeout = httpx.Timeout(connect=0.25, read=0.50, write=0.50, pool=0.25)
    started = time.monotonic()
    try:
        async def _send() -> int:
            async with httpx.AsyncClient(timeout=timeout) as client:
                response = await client.post(capture_url, json=body)
                response.raise_for_status()
                return response.status_code

        status_code = await asyncio.wait_for(_send(), timeout=0.75)
        logger.info(
            "posthog_server_capture_succeeded event=%s status_class=%s latency_ms=%s",
            event_name,
            f"{status_code // 100}xx",
            int((time.monotonic() - started) * 1000),
        )
        return True
    except Exception as exc:
        logger.warning(
            "posthog_server_capture_failed event=%s error_type=%s latency_ms=%s",
            event_name,
            type(exc).__name__,
            int((time.monotonic() - started) * 1000),
        )
        return False


async def _capture_posthog_invoice_payment(
    event: Dict[str, Any],
    invoice: Dict[str, Any],
    subscription: Optional[Dict[str, Any]] = None,
) -> bool:
    if not _posthog_revenue_enabled("payment_succeeded"):
        return False
    invoice_id = str(invoice.get("id") or "").strip()
    stripe_event_id = str(event.get("id") or "").strip()
    amount_minor = invoice.get("amount_paid")
    currency = str(invoice.get("currency") or "").upper()
    timestamp = _posthog_stripe_timestamp(event)
    subscription_id = _stripe_object_id(invoice.get("subscription"))
    customer_id = _stripe_object_id(invoice.get("customer"))
    metadata = {
        **((subscription or {}).get("metadata") or {}),
        **(invoice.get("metadata") or {}),
    }
    user_id = await _posthog_resolve_billing_user_id(
        metadata=metadata,
        subscription_id=subscription_id,
        customer_id=customer_id,
    )
    if not all((stripe_event_id, invoice_id, user_id, timestamp, currency)) or not isinstance(amount_minor, int) or amount_minor <= 0:
        logger.warning(
            "posthog_revenue_skipped_invalid event=payment_succeeded stripe_event_id=%s invoice_id=%s",
            event.get("id"),
            invoice_id or None,
        )
        return False

    properties = {
        "revenue": float(_posthog_major_amount(amount_minor, currency)),
        "currency": currency,
        "amount_minor": amount_minor,
        "stripe_event_id": stripe_event_id,
        "invoice_id": invoice_id,
        "subscription_id": subscription_id,
        **_posthog_price_properties(invoice),
        **(_posthog_price_properties(subscription or {}) if subscription else {}),
        "source": "stripe_webhook",
        "$process_person_profile": False,
    }
    properties = {key: value for key, value in properties.items() if value is not None}
    return await _capture_posthog_server_event(
        event_name="payment_succeeded",
        distinct_id=user_id,
        timestamp=timestamp,
        semantic_uuid=_posthog_revenue_uuid("payment_succeeded", "invoice", invoice_id),
        properties=properties,
    )


def _posthog_refund_success_confirmed(event: Dict[str, Any], refund: Dict[str, Any]) -> bool:
    if str(refund.get("status") or "").lower() != "succeeded":
        return False
    event_type = event.get("type")
    if event_type == "refund.created":
        return True
    if event_type != "refund.updated":
        return False
    previous_status = ((event.get("data") or {}).get("previous_attributes") or {}).get("status")
    return bool(previous_status and str(previous_status).lower() != "succeeded")


async def _posthog_stripe_retrieve(resource: Any, object_id: Any) -> Dict[str, Any]:
    normalized_id = _stripe_object_id(object_id)
    if not normalized_id:
        return {}
    response = await asyncio.wait_for(
        asyncio.to_thread(resource.retrieve, normalized_id),
        timeout=_POSTHOG_STRIPE_ENRICHMENT_TIMEOUT_SECONDS,
    )
    return _stripe_to_dict(response)


async def _posthog_refund_context(refund: Dict[str, Any]) -> Dict[str, Any]:
    charge = refund.get("charge")
    payment_intent = refund.get("payment_intent")
    charge_obj: Dict[str, Any] = {}
    payment_intent_obj: Dict[str, Any] = {}
    invoice_obj: Dict[str, Any] = {}
    subscription_obj: Dict[str, Any] = {}

    if isinstance(charge, dict):
        charge_obj = charge
    elif charge:
        charge_obj = await _posthog_stripe_retrieve(stripe.Charge, charge)
    payment_intent = payment_intent or charge_obj.get("payment_intent")
    if isinstance(payment_intent, dict):
        payment_intent_obj = payment_intent
    elif payment_intent:
        payment_intent_obj = await _posthog_stripe_retrieve(stripe.PaymentIntent, payment_intent)

    invoice = charge_obj.get("invoice") or payment_intent_obj.get("invoice")
    if isinstance(invoice, dict):
        invoice_obj = invoice
    elif invoice:
        invoice_obj = await _posthog_stripe_retrieve(stripe.Invoice, invoice)
    subscription_id = _stripe_object_id(invoice_obj.get("subscription"))
    if subscription_id:
        subscription_obj = await _posthog_stripe_retrieve(stripe.Subscription, subscription_id)
    customer_id = (
        _stripe_object_id(invoice_obj.get("customer"))
        or _stripe_object_id(payment_intent_obj.get("customer"))
        or _stripe_object_id(charge_obj.get("customer"))
        or _stripe_object_id(subscription_obj.get("customer"))
    )
    metadata = {
        **(subscription_obj.get("metadata") or {}),
        **(invoice_obj.get("metadata") or {}),
        **(payment_intent_obj.get("metadata") or {}),
        **(charge_obj.get("metadata") or {}),
        **(refund.get("metadata") or {}),
    }
    return {
        "invoice": invoice_obj,
        "subscription": subscription_obj,
        "invoice_id": _stripe_object_id(invoice_obj),
        "subscription_id": subscription_id,
        "customer_id": customer_id,
        "currency": str(
            refund.get("currency")
            or charge_obj.get("currency")
            or invoice_obj.get("currency")
            or ""
        ).upper(),
        "user_id": await _posthog_resolve_billing_user_id(
            metadata=metadata,
            subscription_id=subscription_id,
            customer_id=customer_id,
        ),
    }


async def _capture_posthog_refund(event: Dict[str, Any], refund: Dict[str, Any]) -> bool:
    if not _posthog_revenue_enabled("payment_refunded"):
        return False
    refund_id = str(refund.get("id") or "").strip()
    refund_status = str(refund.get("status") or "").lower()
    if refund_status not in {"succeeded", "pending", "requires_action", "failed", "canceled"}:
        logger.warning(
            "posthog_refund_revenue_status_unconfirmed stripe_event_id=%s refund_id=%s",
            event.get("id"),
            refund_id or None,
        )
        return False
    if not _posthog_refund_success_confirmed(event, refund):
        logger.info(
            "posthog_refund_revenue_nonterminal stripe_event_id=%s refund_id=%s status=%s",
            event.get("id"),
            refund.get("id"),
            str(refund.get("status") or "unknown").lower(),
        )
        return False
    stripe_event_id = str(event.get("id") or "").strip()
    amount_minor = refund.get("amount")
    timestamp = _posthog_stripe_timestamp(event)
    try:
        context = await _posthog_refund_context(refund)
    except Exception as exc:
        logger.warning(
            "posthog_revenue_skipped_invalid event=payment_refunded stripe_event_id=%s refund_id=%s error_type=%s",
            event.get("id"),
            refund_id or None,
            type(exc).__name__,
        )
        return False
    currency = context.get("currency")
    user_id = context.get("user_id")
    invoice_id = context.get("invoice_id")
    if not all((stripe_event_id, refund_id, user_id, invoice_id, timestamp, currency)) or not isinstance(amount_minor, int) or amount_minor <= 0:
        logger.warning(
            "posthog_revenue_skipped_invalid event=payment_refunded stripe_event_id=%s refund_id=%s",
            event.get("id"),
            refund_id or None,
        )
        return False

    properties = {
        "revenue": -float(_posthog_major_amount(amount_minor, currency)),
        "currency": currency,
        "amount_minor": amount_minor,
        "stripe_event_id": stripe_event_id,
        "refund_id": refund_id,
        "invoice_id": invoice_id,
        "subscription_id": context.get("subscription_id"),
        **_posthog_price_properties(context.get("invoice") or {}),
        **_posthog_price_properties(context.get("subscription") or {}),
        "source": "stripe_webhook",
        "$process_person_profile": False,
    }
    properties = {key: value for key, value in properties.items() if value is not None}
    return await _capture_posthog_server_event(
        event_name="payment_refunded",
        distinct_id=user_id,
        timestamp=timestamp,
        semantic_uuid=_posthog_revenue_uuid("payment_refunded", "refund", refund_id),
        properties=properties,
    )


async def _capture_posthog_revenue_fail_open(
    capture: Any,
    *args: Any,
    event_id: Optional[str],
    event_name: str,
) -> bool:
    try:
        return bool(
            await asyncio.wait_for(
                capture(*args),
                timeout=_POSTHOG_ANALYTICS_TIMEOUT_SECONDS,
            )
        )
    except Exception as exc:
        logger.warning(
            "posthog_revenue_adapter_failed event=%s stripe_event_id=%s error_type=%s",
            event_name,
            event_id,
            type(exc).__name__,
        )
        return False


@api_router.post("/stripe/webhook")
async def stripe_webhook(request: Request):
    _stripe_secret_key()
    webhook_secret = os.environ.get("STRIPE_WEBHOOK_SECRET", "").strip()
    if not webhook_secret:
        raise HTTPException(status_code=503, detail="Stripe webhook is not configured")
    payload = await request.body()
    signature = request.headers.get("Stripe-Signature", "")
    try:
        raw_event = stripe.Webhook.construct_event(payload, signature, webhook_secret)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid payload")
    except stripe.SignatureVerificationError:
        raise HTTPException(status_code=400, detail="Invalid signature")

    # Recent stripe-python releases no longer make StripeObject a dict subclass
    # (no .get() method at all), while every helper below is written against
    # plain dict semantics -- convert once, at the boundary, so nothing further
    # downstream needs to know or care that this came from the Stripe SDK.
    event = _stripe_to_dict(raw_event)
    event_type = event["type"]
    event_id = event.get("id")
    if await _stripe_event_already_processed(event_id):
        logger.info("stripe_webhook_duplicate_ignored event_id=%s type=%s", event_id, event_type)
        return {"received": True, "duplicate": True}

    obj = event["data"]["object"]
    try:
        subscription = None
        if event_type in {"checkout.session.completed", "checkout.session.async_payment_succeeded"}:
            await _apply_checkout_session_billing(obj, event_id=event_id)
        elif event_type in {"customer.subscription.created", "customer.subscription.updated", "customer.subscription.deleted"}:
            await _handle_subscription_event(obj)
        elif event_type in {"invoice.payment_succeeded", "invoice.payment_failed"}:
            subscription_id = obj.get("subscription")
            customer_id = obj.get("customer")
            payment_status = "failed" if event_type == "invoice.payment_failed" else "succeeded"
            if subscription_id:
                try:
                    subscription = _stripe_to_dict(stripe.Subscription.retrieve(subscription_id))
                    await _handle_subscription_event(subscription, last_payment_status=payment_status)
                except Exception:
                    if customer_id:
                        await _update_user_billing_by_customer_id(customer_id, {"last_payment_status": payment_status})
                    else:
                        logger.warning("stripe_invoice_event_missing_customer event_id=%s type=%s", event_id, event_type)
            elif customer_id:
                await _update_user_billing_by_customer_id(customer_id, {"last_payment_status": payment_status})
            else:
                logger.warning("stripe_invoice_event_missing_subscription_customer event_id=%s type=%s", event_id, event_type)
            if event_type == "invoice.payment_succeeded":
                await _capture_posthog_revenue_fail_open(
                    _capture_posthog_invoice_payment,
                    event,
                    obj,
                    subscription,
                    event_id=event_id,
                    event_name="payment_succeeded",
                )
        elif event_type in {"refund.created", "refund.updated", "refund.failed"}:
            await _capture_posthog_revenue_fail_open(
                _capture_posthog_refund,
                event,
                obj,
                event_id=event_id,
                event_name="payment_refunded",
            )
        else:
            logger.info("stripe_webhook ignored event_type=%s", event_type)
    except Exception:
        # Anything unexpected here previously crashed the whole request with an
        # opaque 500 and no logged traceback -- log it fully now so a real bug
        # is visible in Railway logs, and let Stripe's own retry schedule
        # reattempt delivery rather than silently losing the event.
        logger.exception("stripe_webhook_processing_failed event_id=%s type=%s", event_id, event_type)
        raise HTTPException(status_code=500, detail="Webhook processing failed") from None

    await _record_processed_stripe_event(event)
    return {"received": True}


# ===================== CV parsing =====================

CV_UPLOAD_EXTENSIONS = {".pdf", ".png", ".docx", ".txt", ".rtf", ".jpg", ".jpeg", ".heic", ".heif", ".webp"}
CV_IMAGE_FORMATS = {"png", "jpeg", "webp", "heic"}
# Legacy formats we can detect but cannot reliably parse — surfaced as a clear,
# actionable error instead of silently falling through to a garbled "txt" read.
CV_UNSUPPORTED_FORMATS = {"doc"}


def _detect_cv_format(filename: str, content: bytes, content_type: Optional[str] = None) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf") or content[:4] == b"%PDF":
        return "pdf"
    if name.endswith(".png") or content[:8] == b"\x89PNG\r\n\x1a\n":
        return "png"
    if content[:3] == b"\xff\xd8\xff" or name.endswith((".jpg", ".jpeg")):
        return "jpeg"
    if len(content) >= 12 and content[:4] == b"RIFF" and content[8:12] == b"WEBP":
        return "webp"
    if len(content) >= 12 and content[4:8] == b"ftyp":
        brand = content[8:12]
        if brand in {b"heic", b"heix", b"hevc", b"hevx", b"mif1", b"msf1"}:
            return "heic"
    if name.endswith((".heic", ".heif")):
        return "heic"
    if name.endswith(".webp"):
        return "webp"
    if name.endswith(".docx") or content[:2] == b"PK":
        return "docx"
    # Legacy Word 97-2003 binary format (OLE2 Compound File magic bytes).
    if name.endswith(".doc") or content[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
        return "doc"
    if name.endswith(".rtf") or content[:5] == b"{\\rtf":
        return "rtf"
    if name.endswith(".txt"):
        return "txt"
    ct = (content_type or "").lower().split(";")[0].strip()
    if ct == "application/pdf":
        return "pdf"
    if ct == "image/png":
        return "png"
    if ct in {"image/jpeg", "image/jpg"}:
        return "jpeg"
    if ct in {"image/heic", "image/heif"}:
        return "heic"
    if ct == "image/webp":
        return "webp"
    if ct == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return "docx"
    if ct in {"application/msword"}:
        return "doc"
    if ct in {"application/rtf", "text/rtf"}:
        return "rtf"
    if ct == "text/plain":
        return "txt"
    return "txt"


def _extract_pdf_text_pypdf(content: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(content))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                pass
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    except Exception as exc:
        logger.warning("cv_pdf_pypdf_extract_failed error=%s", exc)
        return ""


def _extract_pdf_text_pymupdf(content: bytes) -> str:
    try:
        import fitz
    except ImportError:
        return ""
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        parts = [(page.get_text() or "").strip() for page in doc]
        return "\n".join(part for part in parts if part).strip()
    except Exception as exc:
        logger.warning("cv_pdf_pymupdf_extract_failed error=%s", exc)
        return ""


def _extract_pdf_text(content: bytes) -> str:
    candidates = [_extract_pdf_text_pypdf(content), _extract_pdf_text_pymupdf(content)]
    return max(candidates, key=len, default="")


def _cv_text_looks_usable(text: str, min_chars: int = 40) -> bool:
    cleaned = " ".join((text or "").split())
    if len(cleaned) >= min_chars:
        return True
    return len(cleaned.split()) >= 8


def _profile_has_usable_phone(profile: Optional[Dict[str, Any]]) -> bool:
    if not profile:
        return False
    contact = profile.get("contact") or {}
    phone = str(contact.get("phone") or "").strip()
    digits = re.sub(r"\D", "", phone)
    return len(digits) >= 8


def _guess_image_mime(image_bytes: bytes) -> str:
    if image_bytes[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"
    if image_bytes[:3] == b"\xff\xd8\xff":
        return "image/jpeg"
    if len(image_bytes) >= 12 and image_bytes[:4] == b"RIFF" and image_bytes[8:12] == b"WEBP":
        return "image/webp"
    return "image/png"


def _dedupe_image_candidates(images: list[bytes]) -> list[bytes]:
    seen: set[tuple[int, bytes]] = set()
    unique: list[bytes] = []
    for image_bytes in images:
        if not image_bytes:
            continue
        fingerprint = (len(image_bytes), image_bytes[:32])
        if fingerprint in seen:
            continue
        seen.add(fingerprint)
        unique.append(image_bytes)
    return unique


def _pdf_ocr_image_candidates(content: bytes, max_pages: int = 4, dpi: int = 200) -> list[bytes]:
    """Prefer rendered pages for OCR — more reliable than raw embedded XObject bytes."""
    rendered = _rasterize_pdf_pages(content, max_pages=max_pages, dpi=dpi)
    embedded = _extract_pdf_image_bytes(content)
    return _dedupe_image_candidates([*rendered, *embedded])


def _extract_pdf_image_bytes(content: bytes) -> list[bytes]:
    """Collect embedded page images (e.g. scanned PDF pages)."""
    images: list[bytes] = []
    try:
        reader = PdfReader(io.BytesIO(content))
        for page in reader.pages:
            for image in getattr(page, "images", []) or []:
                data = getattr(image, "data", None)
                if data:
                    images.append(data)
            if images:
                continue
            resources = page.get("/Resources")
            if not resources:
                continue
            xobjects = resources.get("/XObject")
            if not xobjects:
                continue
            xobjects = xobjects.get_object()
            for ref in xobjects:
                xobj = xobjects[ref].get_object()
                if xobj.get("/Subtype") != "/Image":
                    continue
                try:
                    data = xobj.get_data()
                except Exception:
                    data = None
                if data:
                    images.append(data)
    except Exception:
        pass
    return images


def _looks_like_headshot(image_bytes: bytes) -> bool:
    """Best-effort filter for "is this embedded image plausibly a candidate
    photo" vs a logo/banner/decorative graphic -- no face detection, just
    size + aspect ratio. Deliberately conservative: false negatives (missing
    a real photo) are far less costly than false positives (grabbing a logo).
    """
    try:
        from PIL import Image

        with Image.open(io.BytesIO(image_bytes)) as image:
            width, height = image.size
    except Exception:
        return False
    if min(width, height) < 120:
        return False
    aspect = width / height if height else 0
    return 0.7 <= aspect <= 1.3


def _extract_docx_inline_images(content: bytes) -> list[bytes]:
    """Pull embedded picture bytes out of a DOCX via its inline shapes --
    python-docx has no direct `.image` accessor, so resolve each shape's
    relationship id to the underlying media part manually."""
    images: list[bytes] = []
    try:
        document = docx_lib.Document(io.BytesIO(content))
        for shape in document.inline_shapes:
            try:
                blip = shape._inline.graphic.graphicData.pic.blipFill.blip
                rid = blip.embed
                part = document.part.related_parts[rid]
                images.append(part.blob)
            except Exception:
                continue
    except Exception:
        pass
    return images


def _extract_cv_photo(fmt: str, content: bytes) -> tuple[Optional[bytes], Optional[str]]:
    """Best-effort candidate-photo extraction from an uploaded CV, used to
    put the user's own photo (not their OAuth login avatar) on the header of
    their AI-tailored CV. Never raises -- a missed photo just means the
    generated CV renders with no photo at all."""
    try:
        if fmt == "pdf":
            candidates = _extract_pdf_image_bytes(content)
        elif fmt == "docx":
            candidates = _extract_docx_inline_images(content)
        else:
            return None, None
        from PIL import Image

        for candidate in candidates:
            if not _looks_like_headshot(candidate):
                continue
            with Image.open(io.BytesIO(candidate)) as image:
                buffer = io.BytesIO()
                image.convert("RGB").save(buffer, format="JPEG", quality=90)
                return buffer.getvalue(), "image/jpeg"
    except Exception as exc:
        logger.warning("cv_photo_extraction_failed error=%s", exc)
    return None, None


def _rasterize_pdf_pages(content: bytes, max_pages: int = 4, dpi: int = 200) -> list[bytes]:
    """Render PDF pages to PNG when no text layer exists (common for mobile scans)."""
    try:
        import fitz
    except ImportError:
        return []
    images: list[bytes] = []
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        for page_index in range(min(len(doc), max_pages)):
            pix = doc[page_index].get_pixmap(dpi=dpi, alpha=False)
            images.append(pix.tobytes("png"))
    except Exception as exc:
        logger.warning("cv_pdf_rasterize_failed error=%s", exc)
    return images


def _extract_docx_text(content: bytes) -> str:
    document = docx_lib.Document(io.BytesIO(content))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                if text:
                    parts.append(text)
    return "\n".join(parts)


def _extract_rtf_text(content: bytes) -> str:
    """Lightweight RTF-to-text conversion (no extra dependency) — strips control
    words/groups. Not pixel-perfect but sufficient for downstream AI parsing."""
    try:
        text = content.decode("latin-1", errors="ignore")
    except Exception:
        return ""
    text = re.sub(r"\\par[d]?\b", "\n", text)
    text = re.sub(r"\\tab\b", "\t", text)
    text = re.sub(r"\\'[0-9a-fA-F]{2}", "", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d*\s?", " ", text)
    text = re.sub(r"[{}]", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s+", "\n", text)
    return text.strip()


def _cv_image_mime(fmt: str, filename: str, content_type: Optional[str]) -> str:
    ct = (content_type or "").lower().split(";")[0].strip()
    if ct.startswith("image/"):
        return ct
    name = (filename or "").lower()
    if fmt == "jpeg" or name.endswith((".jpg", ".jpeg")):
        return "image/jpeg"
    if fmt == "webp" or name.endswith(".webp"):
        return "image/webp"
    if fmt == "heic" or name.endswith((".heic", ".heif")):
        return "image/heic"
    return "image/png"


def _heic_to_jpeg_bytes(content: bytes) -> tuple[bytes, str]:
    try:
        import pillow_heif
        from PIL import Image

        pillow_heif.register_heif_opener()
        image = Image.open(io.BytesIO(content))
        buffer = io.BytesIO()
        image.convert("RGB").save(buffer, format="JPEG", quality=92)
        return buffer.getvalue(), "image/jpeg"
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail="Could not read HEIC photo. Export as JPG or PDF and try again.",
        ) from exc


def extract_text_from_upload(filename: str, content: bytes, content_type: Optional[str] = None) -> str:
    fmt = _detect_cv_format(filename, content, content_type)
    if fmt == "pdf":
        return _extract_pdf_text(content)
    if fmt == "docx":
        return _extract_docx_text(content)
    if fmt == "rtf":
        return _extract_rtf_text(content)
    if fmt in CV_IMAGE_FORMATS or fmt in CV_UNSUPPORTED_FORMATS:
        return ""
    try:
        return content.decode("utf-8", errors="ignore")
    except Exception:
        return ""


async def _ocr_cv_image_bytes(content: bytes, mime: str) -> str:
    return await extract_text_from_image_bytes(content, mime)


async def extract_cv_text_from_upload(
    filename: str,
    content: bytes,
    content_type: Optional[str] = None,
) -> str:
    fmt = _detect_cv_format(filename, content, content_type)
    if fmt == "pdf":
        text = _extract_pdf_text(content)
        if _cv_text_looks_usable(text):
            return text
        image_parts: list[str] = []
        for image_bytes in _pdf_ocr_image_candidates(content):
            mime = _guess_image_mime(image_bytes)
            try:
                part = await _ocr_cv_image_bytes(image_bytes, mime)
            except LLMProviderNotConfigured:
                raise
            except Exception as exc:
                logger.warning("cv_pdf_image_ocr_failed error=%s", exc)
                continue
            if part.strip():
                image_parts.append(part.strip())
        ocr_text = "\n\n".join(image_parts)
        if _cv_text_looks_usable(ocr_text):
            return ocr_text
        if text.strip():
            return text
        return ocr_text
    if fmt in CV_IMAGE_FORMATS:
        image_bytes = content
        mime = _cv_image_mime(fmt, filename, content_type)
        if fmt == "heic":
            image_bytes, mime = _heic_to_jpeg_bytes(content)
        return await _ocr_cv_image_bytes(image_bytes, mime)
    return extract_text_from_upload(filename, content, content_type)


def _parse_json_from_llm(text: str) -> Dict[str, Any]:
    """Extract JSON object from Claude response (handles ```json fences)."""
    text = text.strip()
    # remove code fences
    m = re.search(r"```(?:json)?\s*([\s\S]*?)```", text)
    if m:
        text = m.group(1).strip()
    # find first { ... last }
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1:
        text = text[start:end + 1]
    return json.loads(text)


async def claude_extract_profile(cv_text: str) -> Dict[str, Any]:
    system_message = (
        "You are an expert resume parser. Return ONLY valid JSON. "
        "No prose, no markdown fences, no commentary."
    )

    prompt = f"""Extract the candidate profile from this CV. Return JSON with this exact schema:
{{
  "contact": {{
    "name": "Full name or empty string",
    "email": "email or empty",
    "phone": "phone or empty",
    "location": "city, country or empty",
    "city": "city explicitly listed or empty",
    "country": "country explicitly listed or empty",
    "linkedin": "linkedin url or empty",
    "website": "personal site url or empty",
    "portfolio": "portfolio url or empty"
  }},
  "summary": "1-2 sentence professional summary",
  "skills": ["skill1", "skill2", ...max 15],
  "languages": ["Language - proficiency, only if explicitly present"],
  "experience": [{{"role": "...", "company": "...", "duration": "...", "location": "...", "highlights": ["...", "..."]}}],
  "education": [{{"degree": "...", "school": "...", "discipline": "field of study/major or empty", "field_of_study": "same if explicit or empty", "graduation_year": "YYYY or empty", "year": "YYYY or empty"}}],
  "target_roles": ["job title 1", "job title 2", "job title 3"],
  "seniority": "junior" | "mid" | "senior" | "lead" | "principal",
  "template_style": "modern" | "classic" | "minimal" | "two_column"
}}

For template_style, infer the layout aesthetic of the original CV: "two_column" if sidebar+main, "classic" if centered headers/serif feel, "minimal" if heavy whitespace and thin dividers, otherwise "modern".
Extract only facts explicitly present in the CV. Do not infer work authorization, sponsorship, demographic data, pronouns, or legal eligibility.
Keep languages separate from professional skills. Do not invent proficiency levels; use only what the CV states.

CV:
---
{cv_text[:8000]}
---
Return ONLY the JSON object."""
    response = await complete_json_text(system_message, prompt)
    return _parse_json_from_llm(response)


def _clean_string(value: Any, max_len: int = 300) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    return text[:max_len]


def _valid_email(value: Any) -> str:
    text = _clean_string(value, 200)
    return text if re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", text) else ""


def _normalize_profile_url(value: Any, *, linkedin: bool = False) -> str:
    text = _clean_string(value, 400)
    if not text:
        return ""
    if re.search(r"\s", text):
        return ""
    if linkedin and "linkedin." not in text.lower():
        return ""
    if not re.match(r"^https?://", text, flags=re.I):
        text = "https://" + text
    parsed = urlparse(text)
    if parsed.scheme not in ("http", "https") or not parsed.netloc or "." not in parsed.netloc:
        return ""
    return text


def _graduation_year(value: Any) -> str:
    match = re.search(r"\b(19|20)\d{2}\b", str(value or ""))
    return match.group(0) if match else ""


def _split_location(value: Any) -> tuple[str, str]:
    text = _clean_string(value, 160)
    if not text:
        return "", ""
    parts = [part.strip() for part in text.split(",") if part.strip()]
    if len(parts) >= 2:
        return parts[0], parts[-1]
    return text, ""


def _phone_country_code_from_explicit(phone: str) -> str:
    text = _clean_string(phone, 80)
    match = re.match(r"^\s*(\+\d{1,4})\b", text)
    return match.group(1) if match else ""


def _first_education_item(education: Any) -> Dict[str, Any]:
    items = education if isinstance(education, list) else []
    for item in items:
        if isinstance(item, dict):
            return item
    return {}


def _cv_lines(cv_text: str) -> List[str]:
    return [_clean_string(line, 500) for line in (cv_text or "").splitlines() if _clean_string(line, 500)]


def _fallback_extract_profile_from_cv(cv_text: str) -> Dict[str, Any]:
    lines = _cv_lines(cv_text)
    first_lines = lines[:12]
    text = cv_text or ""
    email_match = re.search(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", text, flags=re.I)
    phone_match = re.search(r"(?:\+\d{1,4}[\s.-]?)?(?:\(?\d{2,4}\)?[\s.-]?){2,5}\d{2,4}", text)
    linkedin_match = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/[^\s,;]+", text, flags=re.I)
    website_match = re.search(r"https?://(?![^/\s]*linkedin\.com)[^\s,;]+", text, flags=re.I)

    name = ""
    for line in first_lines:
        lower = line.lower()
        if "@" in line or "linkedin" in lower or re.search(r"\d{3,}", line):
            continue
        words = [part for part in re.split(r"\s+", line) if part]
        if 2 <= len(words) <= 5:
            name = line
            break

    location = ""
    for line in first_lines:
        if any(marker in line.lower() for marker in ("remote", "linkedin", "@")):
            continue
        if "," in line and not re.search(r"\d{3,}", line):
            location = line
            break

    known_skills = [
        "Python", "JavaScript", "TypeScript", "React", "Node.js", "FastAPI", "Django",
        "SQL", "PostgreSQL", "MongoDB", "Supabase", "AWS", "GCP", "Azure", "Docker",
        "Kubernetes", "Excel", "Power BI", "Tableau", "Salesforce", "Figma",
        "Marketing", "Sales", "Finance", "Accounting", "Operations", "Project Management",
        "Product Management", "Customer Success", "Data Analysis", "Machine Learning",
    ]
    lower_text = text.lower()
    skills = [skill for skill in known_skills if skill.lower() in lower_text][:15]
    languages = []
    language_patterns = (
        ("French", r"\bfrench\b|\bfran[cç]ais\b"),
        ("English", r"\benglish\b|\banglais\b"),
        ("Spanish", r"\bspanish\b|\bespa[nñ]ol\b"),
        ("Arabic", r"\barabic\b|\barabe\b"),
    )
    for language, pattern in language_patterns:
        if re.search(pattern, lower_text, flags=re.I):
            languages.append(language)
    contact = {
        "name": name,
        "email": email_match.group(0) if email_match else "",
        "phone": phone_match.group(0).strip() if phone_match else "",
        "location": location,
        "linkedin": linkedin_match.group(0) if linkedin_match else "",
        "website": website_match.group(0) if website_match else "",
        "portfolio": "",
        "city": "",
        "country": "",
    }
    return {
        "contact": contact,
        "summary": "",
        "skills": skills,
        "languages": languages,
        "experience": [],
        "education": [],
        "target_roles": [],
        "seniority": None,
        "template_style": "modern",
        "extraction_fallback_reason": "no_ai_provider",
    }


def _line_evidence(lines: List[str], value: Any, *, window: int = 1) -> str:
    needle = _clean_string(value, 200).lower()
    if not needle:
        return ""
    for idx, line in enumerate(lines):
        if needle in line.lower():
            start = max(0, idx - window)
            end = min(len(lines), idx + window + 1)
            return " | ".join(lines[start:end])
    return ""


def _normalize_country_name(value: Any) -> str:
    text = _clean_string(value, 120)
    key = text.lower()
    aliases = {
        "royaume-uni": "United Kingdom",
        "royaume uni": "United Kingdom",
        "uk": "United Kingdom",
        "u.k.": "United Kingdom",
        "united kingdom": "United Kingdom",
        "usa": "United States",
        "us": "United States",
        "u.s.": "United States",
        "united states": "United States",
    }
    return aliases.get(key, text)


def _education_degree_parts(value: Any) -> tuple[str, str]:
    text = _clean_string(value, 200)
    if not text:
        return "", ""
    patterns = [
        (r"\b(bsc|b\.sc\.?|bachelor(?:'s)?|ba|b\.a\.?)\b\s*(?:in|of)?\s*(.*)", "Bachelor"),
        (r"\b(msc|m\.sc\.?|master(?:'s)?|ma|m\.a\.?)\b\s*(?:in|of)?\s*(.*)", "Master"),
        (r"\b(phd|ph\.d\.?|doctorate)\b\s*(?:in|of)?\s*(.*)", "PhD"),
    ]
    for pattern, degree in patterns:
        match = re.search(pattern, text, flags=re.I)
        if match:
            discipline = _clean_string(match.group(2), 160)
            return degree, discipline
    return text, ""


def _education_candidate_from_cv(cv_text: str) -> Optional[Dict[str, Any]]:
    lines = _cv_lines(cv_text)
    lower_lines = [line.lower() for line in lines]
    try:
        start = lower_lines.index("education") + 1
    except ValueError:
        return None
    end = len(lines)
    for marker in ("work experience", "experience", "skills", "certificates"):
        if marker in lower_lines[start:]:
            end = min(end, start + lower_lines[start:].index(marker))
    edu_lines = lines[start:end]
    candidates: List[Dict[str, Any]] = []
    school_terms = ("university", "college", "lycée", "lycee", "school", "institute", "academy")
    degree_terms = ("bsc", "b.sc", "bachelor", "ba ", "msc", "m.sc", "master", "phd", "doctorate", "degree")
    low_level_terms = ("primary school", "middle school", "high school", "hight school")
    for idx, line in enumerate(edu_lines):
        lower = line.lower()
        if re.search(r"\d{2}/\d{4}|present|^\W*$", lower):
            continue
        next_line = edu_lines[idx + 1] if idx + 1 < len(edu_lines) else ""
        following = " | ".join(edu_lines[idx:idx + 5])
        has_degree = any(term in lower for term in degree_terms)
        low_level = any(term in lower for term in low_level_terms)
        next_is_institution = any(term in next_line.lower() for term in school_terms)
        if not (has_degree or next_is_institution):
            continue
        score = 0.35
        if has_degree:
            score += 0.35
        if "present" in following.lower():
            score += 0.2
        if low_level:
            score -= 0.45
        degree, discipline = _education_degree_parts(line)
        if low_level and not has_degree:
            degree = ""
        candidates.append({
            "school": next_line if next_is_institution else "",
            "degree": degree,
            "discipline": discipline,
            "field_of_study": discipline,
            "graduation_year": "" if "present" in following.lower() else _graduation_year(following),
            "year": "" if "present" in following.lower() else _graduation_year(following),
            "confidence": min(max(score, 0.0), 0.98),
            "evidence": following,
        })
    if not candidates:
        return None
    return max(candidates, key=lambda item: item.get("confidence", 0.0))


def _build_profile_intelligence(extracted: Dict[str, Any], cv_text: str) -> Dict[str, Any]:
    contact_raw = extracted.get("contact") or {}
    if not isinstance(contact_raw, dict):
        contact_raw = {}
    fields_extracted: List[str] = []
    fields_skipped: List[str] = []
    field_confidence: Dict[str, float] = {}
    field_evidence: Dict[str, str] = {}
    review_suggestions: Dict[str, Dict[str, Any]] = {}
    auto_saved_fields: List[str] = []
    lines = _cv_lines(cv_text)

    def save_field(field: str, value: Any, confidence: float, evidence: str = "") -> bool:
        if value in (None, ""):
            return False
        field_confidence[field] = round(float(confidence), 2)
        field_evidence[field] = evidence or _line_evidence(lines, value)
        if confidence >= 0.8:
            fields_extracted.append(field)
            auto_saved_fields.append(field)
            return True
        review_suggestions[field] = {
            "value": value,
            "confidence": round(float(confidence), 2),
            "evidence": field_evidence[field],
            "reason": "below_auto_save_threshold",
        }
        fields_skipped.append(f"{field}_low_confidence")
        return False

    contact: Dict[str, Any] = {}
    for key in ("name", "phone", "location"):
        value = _clean_string(contact_raw.get(key))
        confidence = 0.9 if key in ("name", "phone") else 0.82
        if save_field(f"contact.{key}", value, confidence):
            contact[key] = value
    email = _valid_email(contact_raw.get("email"))
    if email and save_field("contact.email", email, 0.95):
        contact["email"] = email
    elif contact_raw.get("email"):
        fields_skipped.append("contact.email_invalid")
    linkedin = _normalize_profile_url(contact_raw.get("linkedin"), linkedin=True)
    if linkedin and save_field("contact.linkedin", linkedin, 0.9):
        contact["linkedin"] = linkedin
    elif contact_raw.get("linkedin"):
        review_suggestions["contact.linkedin"] = {
            "value": _clean_string(contact_raw.get("linkedin")),
            "confidence": 0.3,
            "evidence": _line_evidence(lines, contact_raw.get("linkedin")),
            "reason": "invalid_url_or_contains_whitespace",
        }
        fields_skipped.append("contact.linkedin_invalid")
    website = _normalize_profile_url(contact_raw.get("website") or contact_raw.get("portfolio"))
    if website and save_field("contact.website", website, 0.85):
        contact["website"] = website
    elif contact_raw.get("website") or contact_raw.get("portfolio"):
        fields_skipped.append("contact.website_invalid")

    city = _clean_string(contact_raw.get("city"))
    country = _normalize_country_name(contact_raw.get("country"))
    if not city or not country:
        loc_city, loc_country = _split_location(contact.get("location"))
        city = city or loc_city
        country = country or _normalize_country_name(loc_country)
    location_data = {}
    if save_field("location.city", city, 0.82):
        location_data["city"] = city
    if save_field("location.country", country, 0.82):
        location_data["country"] = country

    education_item = _education_candidate_from_cv(cv_text) or _first_education_item(extracted.get("education"))
    education_confidence = float(education_item.get("confidence") or 0.6) if education_item else 0.0
    education: List[Dict[str, Any]] = []
    if education_item:
        degree, discipline_from_degree = _education_degree_parts(education_item.get("degree") or education_item.get("qualification"))
        edu = {
            "school": _clean_string(education_item.get("school") or education_item.get("institution") or education_item.get("university")),
            "degree": _clean_string(degree),
            "discipline": _clean_string(education_item.get("discipline") or education_item.get("field_of_study") or education_item.get("major") or discipline_from_degree),
            "field_of_study": _clean_string(education_item.get("field_of_study") or education_item.get("discipline") or education_item.get("major") or discipline_from_degree),
            "graduation_year": _graduation_year(education_item.get("graduation_year") or education_item.get("year") or education_item.get("end_date")),
            "year": _graduation_year(education_item.get("year") or education_item.get("graduation_year") or education_item.get("end_date")),
        }
        evidence = education_item.get("evidence") or ""
        clean_edu = {}
        for key, value in edu.items():
            confidence = education_confidence
            if key in ("graduation_year", "year") and not value:
                continue
            if key in ("graduation_year", "year") and "present" in evidence.lower():
                confidence = 0.2
            if save_field(f"education.{key}", value, confidence, evidence):
                clean_edu[key] = value
        if clean_edu:
            education.append(clean_edu)

    experience_items = extracted.get("experience") if isinstance(extracted.get("experience"), list) else []
    first_exp = next((item for item in experience_items if isinstance(item, dict)), {})
    experience_summary = {
        "current_title": _clean_string(first_exp.get("role") or first_exp.get("title")),
        "current_company": _clean_string(first_exp.get("company")),
    }
    experience_summary = {key: value for key, value in experience_summary.items() if value}
    if experience_summary:
        fields_extracted.extend([f"experience_summary.{key}" for key in experience_summary])
        auto_saved_fields.extend([f"experience_summary.{key}" for key in experience_summary])

    skills = [
        _clean_string(skill, 80)
        for skill in (extracted.get("skills") or [])
        if _clean_string(skill, 80)
    ][:30]
    if skills:
        fields_extracted.append("skills")

    defaults: Dict[str, Any] = {}
    if education:
        edu = education[0]
        mapping = {
            "education_school": edu.get("school"),
            "education_degree": edu.get("degree"),
            "education_discipline": edu.get("discipline") or edu.get("field_of_study"),
            "education_graduation_year": edu.get("graduation_year") or edu.get("year"),
        }
        defaults.update({key: value for key, value in mapping.items() if value})
    if website:
        defaults["website_url"] = website
    if linkedin:
        defaults["linkedin_url"] = linkedin
    if city:
        defaults["city"] = city
        defaults["current_location_city"] = city
    if country:
        defaults["country"] = country
        defaults["current_location_country"] = country
    phone_code = _phone_country_code_from_explicit(contact.get("phone", ""))
    if phone_code:
        defaults["phone_country_code"] = phone_code

    cv_extraction = {
        "extracted_at": datetime.now(timezone.utc).isoformat(),
        "fields_extracted": sorted(set(fields_extracted)),
        "fields_skipped": sorted(set(fields_skipped)),
        "field_confidence": field_confidence,
        "field_evidence": field_evidence,
        "auto_saved_fields": sorted(set(auto_saved_fields)),
        "review_suggestions": review_suggestions,
        "confidence_summary": {
            "mode": "explicit_cv_facts_with_confidence_thresholds",
            "saved_defaults": sorted(defaults.keys()),
            "suggested_defaults": sorted(review_suggestions.keys()),
            "cv_text_length": len(cv_text or ""),
        },
    }
    return {
        "contact": contact,
        "location_data": location_data,
        "education": education,
        "experience_summary": experience_summary,
        "skills": skills,
        "application_defaults": defaults,
        "cv_extraction": cv_extraction,
        "extraction_suggestions": review_suggestions,
    }


async def claude_score_jobs(profile: Dict[str, Any], jobs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Score a batch of jobs in a single LLM call. Returns list of {job_id, score, reasons}."""
    system_message = (
        "You are a job matching expert. For each job, score fit 0-100 and give "
        "2-3 short bullet reasons why this candidate is a great fit. "
        "Return ONLY valid JSON, no prose."
    )

    candidate = {
        "summary": profile.get("summary", ""),
        "skills": profile.get("skills", []),
        "experience": [
            {"role": e.get("role"), "company": e.get("company")} for e in profile.get("experience", [])
        ][:5],
        "seniority": profile.get("seniority"),
        "target_role": profile.get("target_role"),
        "target_location": profile.get("target_location"),
        "remote_preference": profile.get("remote_preference"),
    }
    job_summaries = [
        {
            "job_id": j["job_id"],
            "title": j["title"],
            "company": j["company"],
            "location": j["location"],
            "remote": j["remote"],
            "tech_stack": j.get("tech_stack", []),
            "requirements": j.get("requirements", [])[:6],
            "seniority": j.get("seniority"),
        }
        for j in jobs
    ]

    prompt = f"""Candidate:
{json.dumps(candidate, indent=2)}

Jobs:
{json.dumps(job_summaries, indent=2)}

Return JSON: {{"matches": [{{"job_id": "...", "score": 0-100, "reasons": ["...", "..."]}}]}}"""
    response = await complete_json_text(system_message, prompt)
    parsed = _parse_json_from_llm(response)
    return parsed.get("matches", [])


# ATS-specific formatting constraints passed to Claude for each known provider.
_ATS_HINTS: Dict[str, str] = {
    "greenhouse": (
        "This job is posted on Greenhouse ATS. "
        "Greenhouse parses structured fields (role title, dates, company name) and plain text bullets. "
        "Avoid tables, columns, text boxes, graphics, headers/footers, and icons. "
        "Use standard section titles: Summary, Experience, Education, Skills. "
        "Dates must follow MM/YYYY or YYYY format. "
        "Bullet points should start with strong action verbs and include measurable outcomes where possible."
    ),
    "lever": (
        "This job is posted on Lever ATS. "
        "Lever uses keyword matching against the job description. "
        "Avoid multi-column layouts, tables, and embedded images. "
        "Section titles must be standard (Experience, Education, Skills). "
        "Mirror the exact terminology used in the job description for skills and tools."
    ),
    "ashby": (
        "This job is posted on Ashby ATS. "
        "Ashby parses single-column PDF/DOCX resumes cleanly. "
        "Use standard section order: Summary, Experience, Education, Skills. "
        "Dates in MM/YYYY format. Bullet action verbs + quantified results preferred."
    ),
    "workday": (
        "This job is posted on Workday ATS. "
        "Workday is strict on parsing: use a single-column layout, plain text bullets, no tables or graphics. "
        "Job titles must be standard industry titles, not creative/custom ones. "
        "Dates: Month YYYY — Month YYYY format. "
        "Workday scores keywords heavily — mirror exact phrases from the job description."
    ),
    "icims": (
        "This job is posted on iCIMS ATS. "
        "iCIMS performs aggressive keyword matching. Use the exact skill/tool names from the job description. "
        "Single-column layout, no tables. Standard section titles only."
    ),
    "smartrecruiters": (
        "This job is posted on SmartRecruiters ATS. "
        "SmartRecruiters parses standard single-column resumes. "
        "Mirror job description keywords precisely. Standard section titles. No graphics."
    ),
    "teamtailor": (
        "This job is posted on Teamtailor ATS. "
        "Teamtailor performs basic keyword extraction. Plain text, single-column format recommended. "
        "Mirror exact skill names from the job description."
    ),
    "workable": (
        "This job is posted on Workable ATS. "
        "Workable scores resumes on keyword density vs. the job description. "
        "Single-column layout, plain text bullets. Mirror exact terminology from the job requirements."
    ),
    "successfactors": (
        "This job is posted on SAP SuccessFactors ATS. "
        "SuccessFactors is strict on format: single-column, no tables/graphics, standard section titles. "
        "Dates: MM/YYYY. Job titles must be standard. Heavy keyword scoring against description."
    ),
    "personio": (
        "This job is posted on Personio ATS. "
        "Personio uses keyword-based screening. Plain text, single-column. "
        "Mirror exact skill and role terminology from the job posting."
    ),
    "bamboohr": (
        "This job is posted on BambooHR ATS. "
        "BambooHR parses single-column resumes. Standard sections, plain bullets, no graphics."
    ),
    "flatchr": (
        "This job is posted on Flatchr ATS (French market). "
        "Flatchr uses keyword matching. Mirror exact French and English terms from the job description. "
        "Single-column layout, standard section titles."
    ),
    "taleez": (
        "This job is posted on Taleez ATS (French market). "
        "Taleez scores on keyword coverage. Use exact French terminology from the job posting."
    ),
    "recruitee": (
        "This job is posted on Recruitee ATS. "
        "Single-column format, plain text bullets, standard section titles. Mirror job description keywords."
    ),
}

_ATS_DEFAULT_HINT = (
    "The ATS used by this company is unknown. "
    "Apply conservative ATS-safe defaults: single-column layout, no tables or graphics, "
    "standard section titles (Summary, Experience, Education, Skills), plain text bullets, "
    "and mirror the exact keywords and skill names from the job description."
)


def _ats_hint_for_job(job: Dict[str, Any]) -> str:
    provider = (job.get("ats_provider") or "").strip().lower()
    return _ATS_HINTS.get(provider, _ATS_DEFAULT_HINT)


async def claude_generate_application(profile: Dict[str, Any], job: Dict[str, Any]) -> Dict[str, Any]:
    system_message = (
        "Tu es un moteur d'optimisation candidature. Pour le CV, applique des regles ATS strictes : "
        "parsing, mots-cles, structure, correspondance factuelle avec l'offre et scoring. "
        "Pour la lettre de motivation, redige un texte convaincant pour un recruteur humain en restant "
        "100% factuel sur le parcours du candidat. "
        "Retourne UNIQUEMENT du JSON valide. "
        "N'invente jamais d'entreprise, poste, diplome, date, certification, outil, chiffre ou resultat."
    )

    job_description = job.get("clean_description") or job.get("description") or ""
    job_requirements = job.get("requirements") or []
    ats_hint = _ats_hint_for_job(job)
    ats_provider_label = (job.get("ats_provider") or "unknown").upper()
    company_name = job.get("company") or "cette entreprise"
    job_title = job.get("title") or "ce poste"
    cover_letter_rules = build_cover_letter_prompt_section(company_name, job_title, job)
    cv_tailoring_rules = build_cv_tailoring_prompt_section(job_title)
    candidate_payload = {
        "contact": profile.get("contact", {}),
        "cv_text": profile.get("cv_text", "")[:12000],
        "summary": profile.get("summary"),
        "skills": profile.get("skills", []),
        "languages": profile.get("languages", []),
        "experience": profile.get("experience", []),
        "education": profile.get("education", []),
        "seniority": profile.get("seniority"),
        "target_role": profile.get("target_role"),
        "target_roles": profile.get("target_roles", []),
        "target_location": profile.get("target_location"),
        "remote_preference": profile.get("remote_preference"),
        "template_style": profile.get("template_style", "modern"),
        "cover_letter_reference": (profile.get("cover_letter_text") or "")[:6000],
    }
    job_payload = {
        "title": job.get("title"),
        "company": company_name,
        "location": job.get("location"),
        "remote": job.get("remote"),
        "ats_provider": ats_provider_label,
        "description": job_description,
        "requirements": job_requirements,
        "tech_stack": job.get("tech_stack", []),
    }

    prompt = f"""Agis comme l'ATS utilise par {company_name}, c'est-a-dire {ats_provider_label}.

REGLES SPECIFIQUES A CET ATS
{ats_hint}

Tu analyses les CV uniquement via des regles automatiques : parsing, mots-cles, structure, correspondance avec l'offre et scoring.
Tu ne dois faire aucune interpretation humaine, aucun conseil subjectif de recruteur, et aucune recommandation hors objectif ATS.

CONTEXTE

Voici l'offre d'emploi a analyser :
{json.dumps(job_payload, ensure_ascii=False, indent=2)}

Voici mon CV :
{json.dumps(candidate_payload, ensure_ascii=False, indent=2)}

OBJECTIF
1. Calculer un score ATS global sur 100 pour le CV actuel.
2. Identifier les mots-cles critiques presents dans l'offre.
3. Verifier lesquels sont presents, absents ou partiels dans le CV.
4. Detecter les competences, outils, experiences, intitules, verbes d'action et mots-cles metiers manquants.
5. Detecter les problemes de structure ou de parsing qui pourraient bloquer ou reduire le score ATS.
6. Identifier les sections faibles ou mal optimisees.
7. Proposer les optimisations necessaires pour atteindre un score ATS superieur a 80 %, si cela est realiste.
8. Si un score superieur a 80 % n'est pas atteignable sans inventer d'experience, explique clairement pourquoi.
9. Generer une lettre de motivation personnalisee pour {company_name} et le poste "{job_title}".
10. Produire resume_tailoring (modifications MINIMALES du CV — voir regles ci-dessous).

{cv_tailoring_rules}

{cover_letter_rules}

SORTIE ATTENDUE

Presente ta reponse dans cet ordre :

1. Score ATS global estime sur 100

Indique :
- le score actuel
- le niveau de compatibilite avec l'offre
- les principales raisons du score

2. Tableau des mots-cles critiques

Cree un tableau avec les colonnes suivantes :
- Mot-cle / competence issue de l'offre
- Importance : elevee / moyenne / faible
- Present dans mon CV : oui / non / partiellement
- Section ou il apparait
- Recommandation d'integration

3. Problemes ATS detectes

Liste les problemes eventuels lies a :
- structure du CV
- titres de sections
- lisibilite ATS
- mots-cles manquants
- intitules de poste
- format des dates
- competences techniques
- experience insuffisamment alignee
- contenu trop generique
- elements difficiles a parser

4. Optimisations prioritaires

Donne une liste classee par priorite :
- priorite 1 : changements indispensables
- priorite 2 : changements fortement recommandes
- priorite 3 : ameliorations secondaires

5. Analyse de la section Experience (SANS reecriture)

Analyse l'alignement des experiences du profil candidat avec l'offre.
Ne reecris PAS les experiences dans la sortie JSON : le code conserve le texte original.
Tu peux uniquement proposer experience_order dans resume_tailoring pour remonter l'experience la plus pertinente.

6. Resume professionnel (summary uniquement)

Propose un resume professionnel de 4-5 lignes MAX dans resume_tailoring.summary.
Conserve les faits marquants (Big Four, duree, contexte, filiales, equipes).
Ton humain et factuel — jamais de consigne interne dans le texte final.

7. Liste finale des modifications a appliquer

Donne-moi une checklist claire des changements a faire dans mon CV avant de postuler.

8. Lettre de motivation optimisee pour {company_name}

Redige une lettre de motivation en francais (sauf si l'offre est clairement en anglais), personnalisee pour {company_name} et le poste "{job_title}".
Applique strictement toutes les regles LETTRE DE MOTIVATION ci-dessus, en particulier :
- expliquer pourquoi le parcours du candidat est coherent avec le poste (regle #1) ;
- citer des competences/outils de l'offre avec preuves transferables du CV (regle #2) ;
- inclure une phrase specifique a {company_name} (regle #3).

CONTRAINTES GENERALES
- Reste factuel.
- Ne reecris pas le CV complet : resume_tailoring uniquement (headline, summary, skills_order, experience_order).
- Ne supprime jamais d'informations importantes du parcours candidat.
- Ne raccourcis pas le CV.
- Ne donne pas de conseils RH generiques.
- Ne modifie pas le sens reel du parcours.
- Ne mens pas.
- Ne gonfle pas artificiellement l'experience.
- Les analyses ATS vont dans ats_analysis ; les consignes internes restent dans ats_analysis et ne doivent JAMAIS apparaitre dans resume_tailoring.summary ou headline.

FORMAT TECHNIQUE OBLIGATOIRE

Retourne uniquement du JSON valide. Mappe la sortie attendue ci-dessus dans ce schema exact :
{{
  "resume_tailoring": {{
    "headline": "Intitule adapte au poste, segments separes par | (max 4), ex: Auditeur Financier | Data Analysis | Controle Interne | Excel Avance",
    "summary": "Resume professionnel 4-5 lignes MAX, factuel, humain, conserve Big Four/Deloitte/durees/contexte — AUCUNE consigne interne",
    "skills_order": [0, 3, 1, 2],
    "experience_order": [0, 1, 2],
    "template_recommendation": "ats_classic, modern_pro, executive_compact, luxe_minimal, studio_slate, or blue_split",
    "role_keywords": ["mot-cle deja present ou prouve dans le CV, max 10"]
  }},
  "match_score": 0-100,
  "match_reasons": ["raison ATS courte 1", "raison ATS courte 2", "raison ATS courte 3"],
  "interview_prep": [],
  "ats_score_before": 0-100,
  "ats_score_after": 0-100,
  "ats_provider": "{ats_provider_label}",
  "ats_analysis": {{
    "score_current": 0-100,
    "score_after_optimization": 0-100,
    "compatibility_level": "faible, moyen, eleve",
    "score_reasons": ["principale raison du score"],
    "critical_keywords": [
      {{"keyword": "mot-cle / competence issue de l'offre", "importance": "elevee, moyenne, faible", "present_in_cv": "oui, non, partiellement", "section": "section ou il apparait", "integration_recommendation": "recommandation d'integration"}}
    ],
    "ats_issues": ["probleme ATS detecte: structure, titres, dates, mots-cles, parsing ou contenu trop generique"],
    "priority_optimizations": {{
      "priority_1": ["changements indispensables"],
      "priority_2": ["changements fortement recommandes"],
      "priority_3": ["ameliorations secondaires"]
    }},
    "optimized_experience_notes": ["analyse d'alignement uniquement — le texte des experiences n'est pas reecrit par l'IA"],
    "optimized_summary_notes": ["resume de ce qui change dans resume_tailoring.summary"],
    "score_over_80_realistic": true,
    "score_over_80_explanation": "Explique clairement si 80+ est realiste sans inventer d'experience.",
    "final_checklist": ["modification claire a appliquer avant de postuler"]
  }},
  "keywords_gap": [
    {{"keyword": "mot-cle exact de l'offre", "present_in_original": true/false, "added_in_optimized": true/false}}
  ],
  "resume_quality_checks": {{
    "ats_parse_safe": true/false,
    "uses_standard_headings": true/false,
    "no_ai_phrasing": true/false,
    "no_unsupported_claims": true/false,
    "no_keyword_stuffing": true/false,
    "ats_readable": true/false,
    "notes": ["specific remaining issue, if any"]
  }},
  "tailored_cover_letter": {{
    "template": "french_formal",
    "sender_name": "Nom complet du candidat",
    "sender_address": "Adresse du candidat",
    "sender_phone": "Telephone",
    "sender_email": "Email",
    "recipient_attention": "A l'attention du Service des Ressources Humaines",
    "recipient_company": "{company_name}",
    "recipient_address": "Adresse ou ville de l'entreprise si connue, sinon ville du poste",
    "date_line": "A [ville candidat], le [date du jour en francais]",
    "subject": "Candidature pour le poste de [titre du poste] - {company_name}",
    "greeting": "Madame, Monsieur,",
    "paragraphs": [
      "Paragraphe 1 : accroche + motivation pour le poste + phrase specifique a {company_name}",
      "Paragraphe 2 : pont parcours candidat → poste (experiences CV reliees aux exigences, ex. volumes de donnees, qualite, rigueur)",
      "Paragraphe 3 : competences/outils de l'offre + preuves transferables du CV (Python, SQL, PySpark, ETL… selon l'offre)",
      "Paragraphe 4 (optionnel) : disponibilite + contribution concrete — sans phrase de faiblesse"
    ],
    "sign_off": "Je vous prie de recevoir, Madame, Monsieur, l'expression de mes sinceres salutations.",
    "signature_name": "Nom complet du candidat"
  }}
}}"""
    response = await complete_json_text(system_message, prompt)
    parsed = _parse_json_from_llm(response)
    parsed = _merge_generated_resume_with_profile(parsed, profile, job)
    if "tailored_cover_letter" in parsed and "cover_letter" not in parsed:
        parsed["cover_letter"] = parsed["tailored_cover_letter"]
    return normalize_application_generation(parsed)


def _merge_generated_resume_with_profile(
    parsed: Dict[str, Any],
    profile: Dict[str, Any],
    job: Dict[str, Any],
) -> Dict[str, Any]:
    """Apply minimal tailoring deltas on top of the stored profile resume."""
    tailored = apply_minimal_resume_tailoring(profile, parsed, job)
    parsed["tailored_resume_structured"] = tailored
    parsed["tailored_resume"] = tailored
    parsed["resume_preservation_report"] = validate_minimal_tailoring_preserved(profile, tailored)
    return parsed


async def _generate_application_doc(user: User, profile: Dict[str, Any], job: Dict[str, Any]) -> Dict[str, Any]:
    set_llm_user_context(user.user_id)
    profile = prepare_profile_for_application_generation(profile, user)
    try:
        gen = await claude_generate_application(profile, job)
    except Exception as exc:
        logger.exception("Application generation failed")
        return _pending_application_doc(user, job, exc)

    return _build_generated_application_doc(user, profile, job, gen)


def _build_generated_application_doc(
    user: User,
    profile: Dict[str, Any],
    job: Dict[str, Any],
    gen: Dict[str, Any],
    package_builder: Any = build_application_package,
) -> Dict[str, Any]:
    profile = prepare_profile_for_application_generation(sanitize_docx_text(profile), user)
    gen = sanitize_docx_text(normalize_application_generation(gen))
    gen = _merge_generated_resume_with_profile(gen, profile, job)
    gen = attach_cover_letter_quality_report(gen, job)
    cv_text = str(profile.get("cv_text") or "")
    job_description = str(job.get("clean_description") or job.get("description") or "")
    tailored_resume = enrich_tailored_resume_contact(
        gen.get("tailored_resume_structured") or gen.get("tailored_resume") or {},
        profile,
    )
    gen["tailored_resume_structured"] = tailored_resume
    gen["tailored_resume"] = tailored_resume
    cover_letter = enrich_cover_letter_from_profile(
        gen.get("tailored_cover_letter") or gen.get("cover_letter") or {},
        profile,
        user,
    )
    gen["tailored_cover_letter"] = cover_letter
    gen["cover_letter"] = cover_letter
    resume_quality_report = gen.get("resume_quality_report") or validate_resume_quality(tailored_resume)
    preservation_report = gen.get("resume_preservation_report") or {}
    tailored_resume_length = len(json.dumps(tailored_resume, default=str))
    cover_letter_length = len(cover_letter_to_text(cover_letter))
    cover_letter_quality = gen.get("cover_letter_quality_report") or {}
    generation_mode = "ai" if gen else "fallback"
    logger.info(
        "application_generation_quality user_id=%s job_id=%s has_cv_text=%s cv_text_length=%s job_description_length=%s tailored_resume_length=%s cover_letter_length=%s cover_letter_quality=%s resume_preservation=%s match_score=%s generation_mode=%s",
        user.user_id,
        job.get("job_id"),
        bool(cv_text),
        len(cv_text),
        len(job_description),
        tailored_resume_length,
        cover_letter_length,
        cover_letter_quality.get("status"),
        preservation_report.get("status"),
        gen.get("match_score"),
        generation_mode,
    )
    package_status = "generated"
    generation_status = "generated"
    generation_error = None
    application_package: Dict[str, Any]
    data_quality_status = None
    if len(cv_text) < 300:
        data_quality_status = "needs_profile_data"
    elif len(job_description) < 300:
        data_quality_status = "needs_job_data"

    try:
        if data_quality_status:
            application_package = {
                "tailored_resume_structured": tailored_resume,
                "tailored_cover_letter": cover_letter,
                "application_answers": gen.get("application_answers") or [],
                "tailored_cv_file_b64": None,
                "tailored_cv_filename": None,
                "tailored_cv_mime": None,
                "template_preservation_status": "not_supported",
                "template_preservation_notes": (
                    "Application text was generated, but inputs were too thin to claim a tailored CV package."
                ),
            }
            package_status = data_quality_status
            generation_error = data_quality_status
        else:
            gen["job_title"] = job.get("title")
            application_package = package_builder(profile, gen)
    except Exception as exc:
        logger.exception("DOCX_BUILD_FAILED")
        generation_error = "docx_build_failed"
        application_package = {
            "tailored_resume_structured": tailored_resume,
            "tailored_cover_letter": cover_letter,
            "application_answers": gen.get("application_answers") or [],
            "tailored_cv_file_b64": None,
            "tailored_cv_filename": None,
            "tailored_cv_mime": None,
            "template_preservation_status": "not_supported",
            "template_preservation_notes": (
                "DOCX generation failed after sanitizing AI output. "
                "Generated text content was saved without a tailored file."
            ),
        }
        package_status = "generated_text_only"

    # Preserved so the "use original CV instead" toggle (cv-source endpoint)
    # can restore the AI-tailored file after the active tailored_cv_file_b64
    # fields get temporarily overwritten with the user's original upload.
    ai_tailored_cv_file_b64 = application_package.get("tailored_cv_file_b64")
    ai_tailored_cv_filename = application_package.get("tailored_cv_filename")
    ai_tailored_cv_mime = application_package.get("tailored_cv_mime")

    now = datetime.now(timezone.utc).isoformat()
    has_reviewable_documents = bool(
        tailored_resume
        and cover_letter
        and generation_status == "generated"
    )
    return {
        "application_id": f"app_{uuid.uuid4().hex[:12]}",
        "user_id": user.user_id,
        "job_id": job["job_id"],
        "status": "applied",
        "admin_status": "manual_review_needed",
        "manual_status": "manual_review_needed",
        "manual_status_updated_at": now,
        "admin_timeline": [{
            "event_id": f"timeline_{uuid.uuid4().hex[:12]}",
            "type": "manual_status",
            "manual_status": "manual_review_needed",
            "note": "Queued for manual fulfillment after user swipe.",
            "created_at": now,
        }],
        "package_status": package_status,
        "generation_status": generation_status,
        "generation_error": generation_error,
        "submission_status": "not_submitted",
        "submitted_at": None,
        "submission_provider": None,
        "submission_response_id": None,
        "submission_error": None,
        "tailored_resume": tailored_resume,
        "cover_letter": cover_letter,
        **application_package,
        "ai_tailored_cv_file_b64": ai_tailored_cv_file_b64,
        "ai_tailored_cv_filename": ai_tailored_cv_filename,
        "ai_tailored_cv_mime": ai_tailored_cv_mime,
        "cv_source": "tailored",
        "match_score": gen.get("match_score", 75),
        "match_reasons": gen.get("match_reasons", []),
        "interview_prep": gen.get("interview_prep", []),
        "ats_provider": gen.get("ats_provider") or job.get("ats_provider") or None,
        "ats_analysis": gen.get("ats_analysis") or {},
        "ats_score_before": gen.get("ats_score_before"),
        "ats_score_after": gen.get("ats_score_after"),
        "keywords_gap": gen.get("keywords_gap") or [],
        "resume_quality_report": resume_quality_report,
        "resume_preservation_report": preservation_report,
        "cover_letter_quality_report": cover_letter_quality,
        "generation_pipeline": APPLICATION_GENERATION_PIPELINE,
        "awaiting_review_at": now if has_reviewable_documents else None,
        "document_review_status": "awaiting_user" if has_reviewable_documents else None,
        "document_review_approved_at": None,
        "created_at": now,
        "updated_at": now,
    }


def _pending_application_doc(user: User, job: Dict[str, Any], error: Optional[Exception] = None) -> Dict[str, Any]:
    now = datetime.now(timezone.utc).isoformat()
    error_message = None
    if error is not None:
        message = str(error).strip() or error.__class__.__name__
        error_message = message[:500]
    return {
        "application_id": f"app_{uuid.uuid4().hex[:12]}",
        "user_id": user.user_id,
        "job_id": job["job_id"],
        "status": "applied",
        "admin_status": "manual_review_needed",
        "manual_status": "manual_review_needed",
        "manual_status_updated_at": now,
        "admin_timeline": [{
            "event_id": f"timeline_{uuid.uuid4().hex[:12]}",
            "type": "manual_status",
            "manual_status": "manual_review_needed",
            "note": "Queued for manual fulfillment after user swipe; package generation is pending.",
            "created_at": now,
        }],
        "package_status": "pending_generation",
        "submission_status": "not_submitted",
        "submitted_at": None,
        "submission_provider": None,
        "submission_response_id": None,
        "submission_error": error_message,
        "tailored_resume": {},
        "cover_letter": {},
        "tailored_resume_structured": {},
        "tailored_cover_letter": {},
        "application_answers": [],
        "match_score": None,
        "match_reasons": [],
        "interview_prep": [],
        "generation_status": "pending_generation",
        "generation_error": error_message,
        "created_at": now,
        "updated_at": now,
    }


def _normalize_application_status_fields(app_doc: Dict[str, Any]) -> Dict[str, Any]:
    app = dict(app_doc)
    if not app.get("package_status"):
        has_package = any([
            app.get("tailored_resume_structured"),
            app.get("tailored_cover_letter"),
            app.get("tailored_cv_file_b64"),
        ])
        app["package_status"] = "generated" if has_package else "not_generated"
    if not app.get("submission_status"):
        app["submission_status"] = "not_submitted"
    app.setdefault("submitted_at", None)
    app.setdefault("submission_provider", None)
    app.setdefault("submission_response_id", None)
    app.setdefault("submission_error", None)
    if app.get("prepared_missing_information"):
        payload = app.get("prepared_application_payload") or {}
        app["prepared_missing_information"] = _normalize_missing_information(
            app.get("prepared_missing_information") or [],
            _all_payload_fields(payload),
        )
    return app


def _application_text_lengths(app_doc: Dict[str, Any]) -> Dict[str, int]:
    tailored_resume = app_doc.get("tailored_resume_structured") or app_doc.get("tailored_resume") or {}
    cover_letter = app_doc.get("tailored_cover_letter") or app_doc.get("cover_letter") or {}
    return {
        "tailored_resume_length": len(json.dumps(tailored_resume, default=str)) if tailored_resume else 0,
        "cover_letter_length": len(cover_letter_to_text(cover_letter)) if cover_letter else 0,
    }


def _application_generation_update(generated_doc: Dict[str, Any], existing_app: Dict[str, Any]) -> Dict[str, Any]:
    now = datetime.now(timezone.utc).isoformat()
    timeline = _append_admin_timeline(existing_app, {
        "event_id": f"timeline_{uuid.uuid4().hex[:12]}",
        "type": "generation_completed",
        "package_status": generated_doc.get("package_status"),
        "generation_status": generated_doc.get("generation_status"),
        "generation_error": generated_doc.get("generation_error"),
        "created_at": now,
    })
    preserved = {
        "application_id": existing_app.get("application_id"),
        "user_id": existing_app.get("user_id"),
        "job_id": existing_app.get("job_id"),
        "created_at": existing_app.get("created_at"),
        "admin_status": existing_app.get("admin_status") or "manual_review_needed",
        "manual_status": existing_app.get("manual_status") or "manual_review_needed",
        "manual_status_updated_at": existing_app.get("manual_status_updated_at") or now,
        "admin_timeline": timeline,
        "updated_at": now,
    }
    update_doc = {
        key: value
        for key, value in generated_doc.items()
        if key not in {"application_id", "user_id", "job_id", "created_at", "admin_status", "manual_status", "manual_status_updated_at", "admin_timeline"}
    }
    if update_doc.get("generation_status") == "pending_generation" and update_doc.get("generation_error"):
        update_doc["generation_status"] = "failed"
        update_doc["package_status"] = "failed"
    return {**update_doc, **preserved}


async def _process_application_generation_queue(user_id: str) -> None:
    lock = _application_generation_locks.setdefault(user_id, asyncio.Lock())
    async with lock:
        while True:
            pending = await db.applications.find(
                {"user_id": user_id, "generation_status": {"$in": ["pending_generation", "generating"]}},
                {"_id": 0},
            ).sort("created_at", 1).limit(1).to_list(1)
            if not pending:
                return

            app_doc = pending[0]
            application_id = app_doc.get("application_id")
            job_id = app_doc.get("job_id")
            now = datetime.now(timezone.utc).isoformat()
            await db.applications.update_one(
                {"application_id": application_id, "user_id": user_id},
                {"$set": {"generation_status": "generating", "generation_started_at": now, "updated_at": now}},
            )
            try:
                user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0})
                profile = await db.profiles.find_one({"user_id": user_id}, {"_id": 0})
                job = await db.jobs.find_one({"job_id": job_id}, {"_id": 0})
                if not user_doc or not profile or not job:
                    raise RuntimeError("Missing user, profile, or job for queued application generation")
                queued_user = User(
                    user_id=user_doc.get("user_id"),
                    email=user_doc.get("email") or "",
                    name=user_doc.get("name") or "",
                    picture=user_doc.get("picture"),
                    demo_account=bool(user_doc.get("demo_account")),
                    training_access=bool(user_doc.get("training_access")),
                )
                profile = prepare_profile_for_application_generation(profile, queued_user)
                generated_doc = await _generate_application_doc(queued_user, profile, job)
                latest_app = await db.applications.find_one({"application_id": application_id, "user_id": user_id}, {"_id": 0}) or app_doc
                update_doc = _application_generation_update(generated_doc, latest_app)
                await db.applications.update_one(
                    {"application_id": application_id, "user_id": user_id},
                    {"$set": update_doc},
                )
                logger.info(
                    "queued_application_generation_complete user_id=%s application_id=%s job_id=%s package_status=%s",
                    user_id,
                    application_id,
                    job_id,
                    update_doc.get("package_status"),
                )
                latest_for_queue = await db.applications.find_one(
                    {"application_id": application_id, "user_id": user_id},
                    {"_id": 0},
                ) or {**latest_app, **update_doc}
                if update_doc.get("package_status") in {"generated", "generated_text_only"}:
                    if auto_apply_queue.provider_for_job(job):
                        try:
                            await auto_apply_queue.enqueue_application(
                                db,
                                latest_for_queue,
                                job,
                                user_doc=user_doc,
                            )
                        except Exception as exc:
                            logger.warning(
                                "auto_apply_enqueue_failed user_id=%s application_id=%s job_id=%s error=%s",
                                user_id, application_id, job_id, str(exc)[:300],
                            )
                    elif _agent_auto_prepare_enabled(job):
                        try:
                            await _run_agent_apply(job_id, queued_user, click_submit=False)
                            logger.info("agent_auto_prepare_complete user_id=%s application_id=%s job_id=%s", user_id, application_id, job_id)
                        except Exception as exc:
                            logger.warning(
                                "agent_auto_prepare_failed user_id=%s application_id=%s job_id=%s error=%s",
                                user_id, application_id, job_id, str(exc)[:300],
                            )
            except Exception as exc:
                logger.exception("queued_application_generation_failed user_id=%s application_id=%s job_id=%s", user_id, application_id, job_id)
                failed_at = datetime.now(timezone.utc).isoformat()
                latest_app = await db.applications.find_one({"application_id": application_id, "user_id": user_id}, {"_id": 0}) or app_doc
                timeline = _append_admin_timeline(latest_app, {
                    "event_id": f"timeline_{uuid.uuid4().hex[:12]}",
                    "type": "generation_failed",
                    "generation_error": (str(exc).strip() or exc.__class__.__name__)[:500],
                    "created_at": failed_at,
                })
                await db.applications.update_one(
                    {"application_id": application_id, "user_id": user_id},
                    {"$set": {
                        "generation_status": "failed",
                        "generation_error": (str(exc).strip() or exc.__class__.__name__)[:500],
                        "package_status": "failed",
                        "submission_error": (str(exc).strip() or exc.__class__.__name__)[:500],
                        "admin_timeline": timeline,
                        "updated_at": failed_at,
                    }},
                )


def _schedule_application_generation(user_id: str) -> None:
    task = _application_generation_tasks.get(user_id)
    if task and not task.done():
        return
    _application_generation_tasks[user_id] = asyncio.create_task(_process_application_generation_queue(user_id))


async def _resume_pending_application_generation() -> None:
    try:
        rows = await db.applications.find(
            {"generation_status": {"$in": ["pending_generation", "generating"]}},
            {"_id": 0, "user_id": 1},
        ).limit(500).to_list(500)
        for user_id in sorted({row.get("user_id") for row in rows if row.get("user_id")}):
            _schedule_application_generation(user_id)
        if rows:
            logger.info("resumed_pending_application_generation users=%s applications=%s", len({row.get("user_id") for row in rows if row.get("user_id")}), len(rows))
    except Exception as exc:
        logger.warning("resume_pending_application_generation_failed: %s", exc)


# ===================== Career Coach (Interviews + Improve) =====================
from coach import (  # noqa: E402
    claude_interview_prep, claude_interview_score, claude_improve_analysis,
    is_fresh, stamp,
)


async def _require_profile(user: User) -> Dict[str, Any]:
    profile = await db.profiles.find_one({"user_id": user.user_id}, {"_id": 0, "cv_original_b64": 0})
    if not profile or not profile.get("cv_text"):
        raise HTTPException(status_code=400, detail="Upload your CV first to unlock coaching.")
    return profile


async def _save_coach(user_id: str, key: str, stamped: Dict[str, Any]) -> None:
    """Persist an already-stamped payload to profile.coach.<key>."""
    await db.profiles.update_one(
        {"user_id": user_id},
        {"$set": {f"coach.{key}": stamped, "updated_at": stamped.get("_cached_at") or datetime.now(timezone.utc).isoformat()}},
        upsert=True,
    )


@api_router.get("/coach/interview")
async def coach_interview_prep(refresh: bool = False, user: User = Depends(get_current_user)):
    profile = await _require_profile(user)
    cached = (profile.get("coach") or {}).get("interview")
    if not refresh and is_fresh(cached):
        return cached
    try:
        data = await claude_interview_prep(profile)
    except LLMProviderNotConfigured as e:
        raise HTTPException(status_code=502, detail=str(e))
    except Exception as e:
        logger.exception("interview prep failed")
        raise HTTPException(status_code=502, detail=f"AI coach is unavailable: {e}")
    stamped = stamp(data)
    await _save_coach(user.user_id, "interview", stamped)
    return stamped


class InterviewScoreBody(BaseModel):
    questions: List[str]
    answers: List[str]


@api_router.post("/coach/interview/score")
async def coach_interview_score(body: InterviewScoreBody, user: User = Depends(get_current_user)):
    if len(body.questions) != len(body.answers) or not body.questions:
        raise HTTPException(status_code=400, detail="questions and answers must be same non-empty length")
    profile = await _require_profile(user)
    try:
        result = await claude_interview_score(profile, body.questions, body.answers)
    except LLMProviderNotConfigured as e:
        raise HTTPException(status_code=502, detail=str(e))
    except Exception as e:
        logger.exception("interview scoring failed")
        raise HTTPException(status_code=502, detail=f"AI coach is unavailable: {e}")

    # streak — count distinct days a mock was completed
    now = datetime.now(timezone.utc).isoformat()
    history = (profile.get("coach") or {}).get("interview_history") or []
    history.append({"finished_at": now, "overall": result.get("overall", 0)})
    await db.profiles.update_one(
        {"user_id": user.user_id},
        {"$set": {"coach.interview_history": history[-30:]}},
        upsert=True,
    )
    return result


@api_router.get("/coach/improve")
async def coach_improve(refresh: bool = False, user: User = Depends(get_current_user)):
    profile = await _require_profile(user)
    cached = (profile.get("coach") or {}).get("improve")
    if not refresh and is_fresh(cached):
        return cached
    try:
        data = await claude_improve_analysis(profile)
    except LLMProviderNotConfigured as e:
        raise HTTPException(status_code=502, detail=str(e))
    except Exception as e:
        logger.exception("improve analysis failed")
        raise HTTPException(status_code=502, detail=f"AI coach is unavailable: {e}")
    stamped = stamp(data)
    await _save_coach(user.user_id, "improve", stamped)
    return stamped


@api_router.get("/coach/streak")
async def coach_streak(user: User = Depends(get_current_user)):
    profile = await db.profiles.find_one({"user_id": user.user_id}, {"_id": 0, "coach.interview_history": 1})
    history = ((profile or {}).get("coach") or {}).get("interview_history") or []
    if not history:
        return {"streak": 0, "sessions_total": 0, "sessions_week": 0, "best": 0}
    today = datetime.now(timezone.utc).date()
    days = set()
    for h in history:
        try:
            d = datetime.fromisoformat(h["finished_at"]).date()
            days.add(d)
        except (KeyError, ValueError):
            pass
    streak = 0
    cursor = today
    while cursor in days:
        streak += 1
        cursor = cursor - timedelta(days=1)
    sessions_week = sum(1 for h in history if (today - datetime.fromisoformat(h["finished_at"]).date()).days < 7)
    best = max((h.get("overall", 0) for h in history), default=0)
    return {"streak": streak, "sessions_total": len(history), "sessions_week": sessions_week, "best": best}




CV_UPLOAD_ACCEPTED_FORMATS = {"pdf", "png", "docx", "txt", "rtf", "jpeg", "webp", "heic"}


@api_router.post("/profile/cv")
async def upload_cv(file: UploadFile = File(...), user: User = Depends(get_current_user)):
    import base64
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="The file appears to be empty. Please choose a different file.")
    if len(content) > MAX_PROFILE_DOCUMENT_BYTES:
        limit_mb = MAX_PROFILE_DOCUMENT_BYTES // (1024 * 1024)
        raise HTTPException(status_code=400, detail=f"File must be {limit_mb}MB or smaller")
    filename = file.filename or "cv"
    fmt = _detect_cv_format(filename, content, file.content_type)
    if fmt == "doc":
        raise HTTPException(
            status_code=400,
            detail="Legacy .doc files aren't supported. Please re-save your resume as PDF or DOCX and try again.",
        )
    if fmt not in CV_UPLOAD_ACCEPTED_FORMATS:
        raise HTTPException(
            status_code=400,
            detail="Please upload a PDF, DOCX, RTF, TXT, or image (PNG/JPG/HEIC/WEBP) resume.",
        )
    photo_bytes, photo_mime = _extract_cv_photo(fmt, content)
    set_llm_user_context(user.user_id)
    try:
        cv_text = await extract_cv_text_from_upload(filename, content, file.content_type)
    except LLMProviderNotConfigured:
        raise HTTPException(
            status_code=503,
            detail="AI provider is not configured. Scanned or photo resumes require vision OCR.",
        )
    except LLMRateLimitError as e:
        # Scanned/photo CVs have no non-AI extraction path, so an exhausted quota
        # blocks them outright — surface a clean, actionable message instead of
        # letting the raw provider error (or an unhandled 500) reach the user.
        logger.error("cv_upload_ai_quota_exceeded_ocr user_id=%s error=%s", user.user_id, str(e))
        raise HTTPException(
            status_code=503,
            detail="Our AI provider is temporarily at capacity. Please try again in a few minutes, or upload a PDF/DOCX version of your CV instead of a photo scan.",
        )
    if not cv_text.strip():
        raise HTTPException(
            status_code=400,
            detail="Could not extract text from CV. Try a clearer PDF or a DOCX export from your editor.",
        )

    try:
        extracted = await claude_extract_profile(cv_text)
    except LLMProviderNotConfigured as e:
        logger.warning("cv_upload_ai_provider_not_configured user_id=%s error=%s", user.user_id, str(e))
        extracted = _fallback_extract_profile_from_cv(cv_text)
    except LLMRateLimitError as e:
        # Text-based CVs (the common case) have a heuristic fallback, so an
        # exhausted AI quota degrades the profile quality instead of blocking
        # the upload entirely.
        logger.error("cv_upload_ai_quota_exceeded user_id=%s error=%s", user.user_id, str(e))
        extracted = _fallback_extract_profile_from_cv(cv_text)
    except Exception as e:
        logger.exception("CV extraction failed")
        raise HTTPException(status_code=500, detail="We couldn't finish analyzing your CV. Please try again in a moment.")

    mime = _mime_from_profile_document_filename(filename)

    intelligence = _build_profile_intelligence(extracted, cv_text)
    if extracted.get("extraction_fallback_reason"):
        intelligence.setdefault("cv_extraction", {}).setdefault("confidence_summary", {})["mode"] = "fallback_no_ai_provider"
        intelligence.setdefault("cv_extraction", {}).setdefault("fields_skipped", []).append("ai_provider_not_configured")
    existing_profile = await db.profiles.find_one({"user_id": user.user_id}, {"_id": 0, "application_defaults": 1, "contact": 1}) or {}
    existing_defaults = existing_profile.get("application_defaults") or {}
    cv_managed_default_keys = {
        "city",
        "country",
        "current_location_city",
        "current_location_country",
        "education_school",
        "education_degree",
        "education_discipline",
        "education_graduation_year",
        "linkedin_url",
        "website_url",
        "portfolio_url",
    }
    retained_defaults = {
        key: value
        for key, value in existing_defaults.items()
        if key not in cv_managed_default_keys
    }
    merged_defaults = {**retained_defaults, **(intelligence.get("application_defaults") or {})}
    contact = intelligence.get("contact") or {}
    if user.email:
        contact["email"] = user.email

    profile_doc = {
        "user_id": user.user_id,
        "cv_text": cv_text,
        "cv_filename": filename,
        "cv_original_b64": base64.b64encode(content).decode("ascii"),
        "cv_mime": mime,
        "cv_photo_b64": base64.b64encode(photo_bytes).decode("ascii") if photo_bytes else None,
        "cv_photo_mime": photo_mime,
        "contact": contact,
        "summary": extracted.get("summary", ""),
        "skills": intelligence.get("skills") or extracted.get("skills", []),
        "languages": extracted.get("languages", []),
        "experience": extracted.get("experience", []),
        "education": intelligence.get("education") or extracted.get("education", []),
        "experience_summary": intelligence.get("experience_summary") or {},
        "location_data": intelligence.get("location_data") or {},
        "application_defaults": merged_defaults,
        "cv_extraction": intelligence.get("cv_extraction") or {},
        "extraction_suggestions": intelligence.get("extraction_suggestions") or {},
        "target_roles": extracted.get("target_roles", []),
        "seniority": extracted.get("seniority"),
        "template_style": extracted.get("template_style", "modern"),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.profiles.update_one(
        {"user_id": user.user_id},
        {"$set": profile_doc},
        upsert=True,
    )
    if _posthog_client is not None:
        with new_context():
            identify_context(user.user_id)
            posthog_capture(
                "cv_uploaded",
                properties={
                    "file_format": fmt,
                    "has_photo": bool(photo_bytes),
                    "skills_count": len(extracted.get("skills") or []),
                    "experience_count": len(extracted.get("experience") or []),
                    "used_ai_fallback": bool(extracted.get("extraction_fallback_reason")),
                },
            )
    # don't ship the heavy fields back
    profile_doc.pop("cv_text", None)
    profile_doc.pop("cv_original_b64", None)
    return profile_doc


@api_router.get("/profile/cv/original")
async def download_original_cv(user: User = Depends(get_current_user)):
    """Stream back the user's original CV file."""
    import base64
    from fastapi.responses import Response as FastAPIResponse
    profile = await db.profiles.find_one(
        {"user_id": user.user_id},
        {"_id": 0, "cv_original_b64": 1, "cv_mime": 1, "cv_filename": 1},
    )
    if not profile or not profile.get("cv_original_b64"):
        raise HTTPException(status_code=404, detail="No original CV stored")
    content = base64.b64decode(profile["cv_original_b64"])
    filename = profile.get("cv_filename") or "cv"
    return FastAPIResponse(
        content=content,
        media_type=profile.get("cv_mime", "application/octet-stream"),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


MAX_PROFILE_DOCUMENT_BYTES = 20 * 1024 * 1024
PROFILE_DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".txt", ".rtf", ".png", ".jpg", ".jpeg", ".heic", ".heif", ".webp"}
COVER_LETTER_EXTENSIONS = {".pdf", ".docx", ".txt", ".rtf"}


def _mime_from_profile_document_filename(filename: str) -> str:
    name_lower = (filename or "").lower()
    if name_lower.endswith(".pdf"):
        return "application/pdf"
    if name_lower.endswith(".docx"):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if name_lower.endswith(".png"):
        return "image/png"
    if name_lower.endswith(".jpg") or name_lower.endswith(".jpeg"):
        return "image/jpeg"
    if name_lower.endswith(".webp"):
        return "image/webp"
    if name_lower.endswith((".heic", ".heif")):
        return "image/heic"
    return "text/plain"


def _public_profile_document(doc: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "id": doc.get("id"),
        "name": doc.get("name"),
        "mime": doc.get("mime"),
        "uploaded_at": doc.get("uploaded_at"),
        "size": doc.get("size"),
    }


def _sanitize_profile_for_client(profile: Optional[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    if not profile:
        return profile
    docs = profile.get("additional_documents") or []
    if docs:
        profile = {**profile, "additional_documents": [_public_profile_document(doc) for doc in docs if isinstance(doc, dict)]}
    return profile


@api_router.post("/profile/documents")
async def upload_profile_document(file: UploadFile = File(...), user: User = Depends(get_current_user)):
    import base64

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="The file appears to be empty. Please choose a different file.")
    if len(content) > MAX_PROFILE_DOCUMENT_BYTES:
        limit_mb = MAX_PROFILE_DOCUMENT_BYTES // (1024 * 1024)
        raise HTTPException(status_code=400, detail=f"File must be {limit_mb}MB or smaller")
    filename = file.filename or "document"
    ext = filename.lower()[filename.rfind("."):] if "." in filename else ""
    if ext not in PROFILE_DOCUMENT_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    entry = {
        "id": str(uuid.uuid4()),
        "name": filename,
        "mime": _mime_from_profile_document_filename(filename),
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
        "size": len(content),
        "file_b64": base64.b64encode(content).decode("ascii"),
    }
    await db.profiles.update_one(
        {"user_id": user.user_id},
        {
            "$push": {"additional_documents": entry},
            "$set": {"updated_at": datetime.now(timezone.utc).isoformat()},
        },
        upsert=True,
    )
    return {"ok": True, "document": _public_profile_document(entry)}


@api_router.get("/profile/documents/{document_id}")
async def download_profile_document(document_id: str, user: User = Depends(get_current_user)):
    import base64
    from fastapi.responses import Response as FastAPIResponse

    profile = await db.profiles.find_one(
        {"user_id": user.user_id},
        {"_id": 0, "additional_documents": 1},
    )
    for doc in (profile or {}).get("additional_documents") or []:
        if doc.get("id") != document_id:
            continue
        payload = doc.get("file_b64")
        if not payload:
            raise HTTPException(status_code=404, detail="Document not found")
        content = base64.b64decode(payload)
        filename = doc.get("name") or "document"
        mime = doc.get("mime") or "application/octet-stream"
        disposition = "inline" if mime.startswith("image/") or mime == "application/pdf" else "attachment"
        return FastAPIResponse(
            content=content,
            media_type=mime,
            headers={"Content-Disposition": f'{disposition}; filename="{filename}"'},
        )
    raise HTTPException(status_code=404, detail="Document not found")


@api_router.delete("/profile/documents/{document_id}")
async def delete_profile_document(document_id: str, user: User = Depends(get_current_user)):
    profile = await db.profiles.find_one(
        {"user_id": user.user_id},
        {"_id": 0, "additional_documents.id": 1},
    )
    docs = (profile or {}).get("additional_documents") or []
    if not any(doc.get("id") == document_id for doc in docs):
        raise HTTPException(status_code=404, detail="Document not found")
    await db.profiles.update_one(
        {"user_id": user.user_id},
        {
            "$pull": {"additional_documents": {"id": document_id}},
            "$set": {"updated_at": datetime.now(timezone.utc).isoformat()},
        },
    )
    return {"ok": True}


@api_router.post("/profile/cover-letter")
async def upload_cover_letter(file: UploadFile = File(...), user: User = Depends(get_current_user)):
    import base64

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="The file appears to be empty. Please choose a different file.")
    if len(content) > MAX_PROFILE_DOCUMENT_BYTES:
        limit_mb = MAX_PROFILE_DOCUMENT_BYTES // (1024 * 1024)
        raise HTTPException(status_code=400, detail=f"File must be {limit_mb}MB or smaller")
    filename = file.filename or "cover_letter"
    ext = filename.lower()[filename.rfind(".") :] if "." in filename else ""
    if ext not in COVER_LETTER_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Please upload a PDF, DOCX, RTF, or TXT cover letter")

    cover_letter_text = extract_text_from_upload(filename, content, file.content_type).strip()
    mime = _mime_from_profile_document_filename(filename)
    uploaded_at = datetime.now(timezone.utc).isoformat()

    await db.profiles.update_one(
        {"user_id": user.user_id},
        {
            "$set": {
                "cover_letter_filename": filename,
                "cover_letter_mime": mime,
                "cover_letter_original_b64": base64.b64encode(content).decode("ascii"),
                "cover_letter_text": cover_letter_text,
                "cover_letter_uploaded_at": uploaded_at,
                "updated_at": uploaded_at,
            }
        },
        upsert=True,
    )
    return {
        "ok": True,
        "cover_letter_filename": filename,
        "cover_letter_uploaded_at": uploaded_at,
        "has_cover_letter_text": bool(cover_letter_text),
    }


@api_router.get("/profile/cover-letter/original")
async def download_cover_letter_original(user: User = Depends(get_current_user)):
    import base64
    from fastapi.responses import Response as FastAPIResponse

    profile = await db.profiles.find_one(
        {"user_id": user.user_id},
        {"_id": 0, "cover_letter_original_b64": 1, "cover_letter_mime": 1, "cover_letter_filename": 1},
    )
    if not profile or not profile.get("cover_letter_original_b64"):
        raise HTTPException(status_code=404, detail="No cover letter stored")
    content = base64.b64decode(profile["cover_letter_original_b64"])
    filename = profile.get("cover_letter_filename") or "cover_letter"
    return FastAPIResponse(
        content=content,
        media_type=profile.get("cover_letter_mime", "application/octet-stream"),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@api_router.delete("/profile/cover-letter")
async def delete_cover_letter(user: User = Depends(get_current_user)):
    profile = await db.profiles.find_one(
        {"user_id": user.user_id},
        {"_id": 0, "cover_letter_filename": 1},
    )
    if not profile or not profile.get("cover_letter_filename"):
        raise HTTPException(status_code=404, detail="No cover letter stored")
    await db.profiles.update_one(
        {"user_id": user.user_id},
        {
            "$unset": {
                "cover_letter_filename": "",
                "cover_letter_mime": "",
                "cover_letter_original_b64": "",
                "cover_letter_text": "",
                "cover_letter_uploaded_at": "",
            },
            "$set": {"updated_at": datetime.now(timezone.utc).isoformat()},
        },
    )
    return {"ok": True}


@api_router.get("/profile")
async def get_profile(user: User = Depends(get_current_user)):
    profile = await db.profiles.find_one(
        {"user_id": user.user_id},
        {"_id": 0, "cv_original_b64": 0, "cover_letter_original_b64": 0},
    )
    if not profile:
        return None
    profile = _sanitize_profile_for_client(profile)
    profile["profile_completion"] = _profile_completion(profile)
    return profile


def _profile_completion(profile: Dict[str, Any]) -> Dict[str, Any]:
    defaults = profile.get("application_defaults") or {}
    contact = profile.get("contact") or {}
    checks = [
        ("cv_uploaded", bool(profile.get("cv_text") or profile.get("cv_filename"))),
        ("target_roles", bool(profile.get("target_role") or profile.get("target_roles"))),
        ("location", bool(profile.get("target_location") or profile.get("target_location_data"))),
        ("linkedin_or_portfolio", bool(contact.get("linkedin") or contact.get("website") or contact.get("github"))),
        ("application_defaults", all(
            defaults.get(key) not in (None, "")
            for key in (
                "work_authorized_countries",
                "requires_sponsorship",
                "current_location_city",
                "former_employer_restriction_or_noncompete",
            )
        ) or bool(defaults.get("prefer_not_to_say_demographics"))),
    ]
    completed = sum(1 for _, ok in checks if ok)
    return {
        "percentage": round((completed / len(checks)) * 100),
        "completed": completed,
        "total": len(checks),
        "items": {key: ok for key, ok in checks},
    }


async def _cancel_stripe_subscription_for_user(user_doc: Dict[str, Any]) -> None:
    billing = user_doc.get("billing") or {}
    sub_id = billing.get("stripe_subscription_id")
    if not sub_id or str(sub_id).startswith("master_code_"):
        return
    try:
        _stripe_secret_key()
        stripe.Subscription.cancel(str(sub_id))
    except HTTPException:
        pass
    except Exception as exc:
        logger.warning(
            "account_delete_stripe_cancel_failed user_id=%s subscription_id=%s error=%s",
            user_doc.get("user_id"),
            sub_id,
            str(exc)[:200],
        )


async def _delete_supabase_auth_user(user_doc: Dict[str, Any]) -> None:
    supabase_user_id = user_doc.get("supabase_user_id")
    if not supabase_user_id:
        return
    try:
        response = await _supabase_admin_request("DELETE", f"/auth/v1/admin/users/{supabase_user_id}")
        if response.status_code not in (200, 204):
            logger.warning(
                "account_delete_supabase_auth_failed user_id=%s supabase_user_id=%s status=%s",
                user_doc.get("user_id"),
                supabase_user_id,
                response.status_code,
            )
    except HTTPException:
        pass
    except Exception as exc:
        logger.warning(
            "account_delete_supabase_auth_failed user_id=%s supabase_user_id=%s error=%s",
            user_doc.get("user_id"),
            supabase_user_id,
            str(exc)[:200],
        )


async def _delete_all_user_data(user_id: str) -> None:
    """Remove every app row tied to this user."""
    await db.application_emails.delete_many({"user_id": user_id})
    await db.browser_submission_runs.delete_many({"user_id": user_id})
    await db.applications.delete_many({"user_id": user_id})
    await db.swipes.delete_many({"user_id": user_id})
    await db.profiles.delete_many({"user_id": user_id})
    await db.gmail_connections.delete_many({"user_id": user_id})
    await db.analytics_events.delete_many({"user_id": user_id})
    await db.training_enrollments.delete_many({"user_id": user_id})
    await db.training_creators.delete_many({"user_id": user_id})
    await db.user_feedback.delete_many({"user_id": user_id})
    await db.user_sessions.delete_many({"user_id": user_id})

    templates_col = getattr(db, "interview_simulator_templates", None)
    if templates_col is not None:
        await templates_col.delete_many({"created_by_user_id": user_id})

    await db.creator_invites.update_many(
        {"redeemed_by_user_id": user_id},
        {"$set": {"redeemed_by_user_id": None}},
    )
    await db.users.delete_one({"user_id": user_id})


@api_router.delete("/profile")
async def delete_account(user: User = Depends(get_current_user)):
    """Wipe everything the user created. Sessions are revoked too."""
    user_doc = await db.users.find_one({"user_id": user.user_id}, {"_id": 0}) or {}
    logger.info("account_delete_start user_id=%s", user.user_id)
    if _posthog_client is not None:
        with new_context():
            identify_context(user.user_id)
            posthog_capture("account_deleted")
    await _cancel_stripe_subscription_for_user(user_doc)
    await _delete_all_user_data(user.user_id)
    await _delete_supabase_auth_user(user_doc)
    logger.info("account_delete_complete user_id=%s", user.user_id)
    return {"ok": True}


@api_router.put("/profile/application-defaults")
async def update_application_defaults(body: ApplicationDefaultsUpdate, user: User = Depends(get_current_user)):
    defaults = body.application_defaults or {}
    if not isinstance(defaults, dict):
        raise HTTPException(status_code=400, detail="application_defaults must be an object")
    allowed_keys = {
        "country",
        "city",
        "phone_country_code",
        "linkedin_url",
        "website_url",
        "portfolio_url",
        "education_school",
        "education_degree",
        "education_discipline",
        "education_graduation_year",
        "work_authorized_countries",
        "requires_sponsorship",
        "willing_to_relocate",
        "current_location_country",
        "current_location_city",
        "referral_source",
        "privacy_consent",
        "eeo_gender",
        "eeo_race",
        "eeo_veteran",
        "eeo_disability",
        "eeo_lgbtq",
        "prefer_not_to_say_demographics",
        "former_company_history",
        "former_employer_restriction_or_noncompete",
    }
    clean = {key: value for key, value in defaults.items() if key in allowed_keys}
    now = datetime.now(timezone.utc).isoformat()
    update_fields = {f"application_defaults.{key}": value for key, value in clean.items()}
    update_fields["updated_at"] = now
    await db.profiles.update_one(
        {"user_id": user.user_id},
        {"$set": update_fields},
        upsert=True,
    )
    profile = await db.profiles.find_one({"user_id": user.user_id}, {"_id": 0, "cv_original_b64": 0})
    if profile:
        profile["profile_completion"] = _profile_completion(profile)
    return {"ok": True, "application_defaults": (profile or {}).get("application_defaults") or {}, "profile": profile}


@api_router.put("/profile/structured-data")
async def update_structured_profile_data(body: StructuredProfileDataUpdate, user: User = Depends(get_current_user)):
    existing = await db.profiles.find_one({"user_id": user.user_id}, {"_id": 0}) or {}
    contact_payload = body.contact or {}
    contact = {**(existing.get("contact") or {})}
    for key in ("name", "phone", "location", "linkedin", "website"):
        value = contact_payload.get(key)
        if value is not None:
            contact[key] = _clean_string(value)
    if contact.get("linkedin"):
        normalized = _normalize_profile_url(contact.get("linkedin"), linkedin=True)
        if normalized:
            contact["linkedin"] = normalized
    if contact.get("website"):
        normalized = _normalize_profile_url(contact.get("website"))
        if normalized:
            contact["website"] = normalized
    contact["email"] = user.email

    education = []
    for item in body.education or []:
        if not isinstance(item, dict):
            continue
        clean_item = {
            "school": _clean_string(item.get("school")),
            "degree": _clean_string(item.get("degree")),
            "discipline": _clean_string(item.get("discipline") or item.get("field_of_study")),
            "field_of_study": _clean_string(item.get("field_of_study") or item.get("discipline")),
            "graduation_year": _graduation_year(item.get("graduation_year") or item.get("year")),
            "year": _graduation_year(item.get("year") or item.get("graduation_year")),
        }
        clean_item = {key: value for key, value in clean_item.items() if value}
        if clean_item:
            education.append(clean_item)
    if not education:
        education = existing.get("education") or []

    defaults = {**(existing.get("application_defaults") or {})}
    incoming_defaults = body.application_defaults or {}
    allowed_default_keys = {
        "linkedin_url",
        "website_url",
        "portfolio_url",
        "education_school",
        "education_degree",
        "education_discipline",
        "education_graduation_year",
        "phone_country_code",
        "city",
        "country",
        "current_location_city",
        "current_location_country",
    }
    for key, value in incoming_defaults.items():
        if key in allowed_default_keys:
            defaults[key] = value
    if education:
        edu = education[0]
        for key, value in {
            "education_school": edu.get("school"),
            "education_degree": edu.get("degree"),
            "education_discipline": edu.get("discipline") or edu.get("field_of_study"),
            "education_graduation_year": edu.get("graduation_year") or edu.get("year"),
        }.items():
            if value:
                defaults[key] = value
    if contact.get("linkedin"):
        defaults["linkedin_url"] = contact["linkedin"]
    if contact.get("website"):
        defaults["website_url"] = contact["website"]
    if contact.get("phone"):
        code = _phone_country_code_from_explicit(contact["phone"])
        if code:
            defaults["phone_country_code"] = code
    city, country = _split_location(contact.get("location"))
    if city:
        defaults["city"] = city
        defaults["current_location_city"] = city
    if country:
        defaults["country"] = country
        defaults["current_location_country"] = country

    update = {
        "contact": contact,
        "education": education,
        "application_defaults": defaults,
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    if body.skills:
        update["skills"] = [_clean_string(skill, 80) for skill in body.skills if _clean_string(skill, 80)][:30]
    if body.experience_summary:
        update["experience_summary"] = {
            key: _clean_string(value)
            for key, value in body.experience_summary.items()
            if key in {"current_title", "current_company", "years_experience"} and value not in (None, "")
        }
    await db.profiles.update_one({"user_id": user.user_id}, {"$set": update}, upsert=True)
    profile = await db.profiles.find_one({"user_id": user.user_id}, {"_id": 0, "cv_original_b64": 0})
    if profile:
        profile["profile_completion"] = _profile_completion(profile)
    return {"ok": True, "profile": profile}


@api_router.patch("/profile/extras")
async def patch_profile_extras(payload: Dict[str, Any], user: User = Depends(get_current_user)):
    """Merge-update the `profile.extras` dict — holds user-managed CV sections that AI
    didn't extract automatically (volunteer, projects, references, languages, …)."""
    if not isinstance(payload, dict):
        raise HTTPException(status_code=400, detail="Body must be a JSON object")
    existing = await db.profiles.find_one({"user_id": user.user_id}, {"_id": 0, "extras": 1}) or {}
    merged = {**(existing.get("extras") or {}), **payload}
    await db.profiles.update_one(
        {"user_id": user.user_id},
        {"$set": {"extras": merged, "updated_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True,
    )
    return {"ok": True, "extras": merged}


@api_router.get("/locations/search")
async def locations_search(
    q: str = Query("", min_length=0),
    limit: int = Query(10, ge=1, le=15),
    country_codes: Optional[str] = Query(None, max_length=8),
):
    """Worldwide city/region search via OpenStreetMap + optional Google Places."""
    query = (q or "").strip()
    if len(query) < 1:
        raise HTTPException(status_code=400, detail="Query parameter q is required")
    codes = (country_codes or "").strip().lower() or None
    return await search_locations(query, limit=limit, country_codes=codes)


@api_router.post("/onboarding/suggest-categories")
async def onboarding_suggest_categories(body: OnboardingSuggestCategoriesRequest):
    """Region + contract-aware job categories for onboarding."""
    if not body.location.strip():
        raise HTTPException(status_code=400, detail="Location is required")
    if not body.contract_type.strip():
        raise HTTPException(status_code=400, detail="Contract type is required")
    return await suggest_categories(
        body.location.strip(),
        body.contract_type.strip(),
        body.location_data,
    )


@api_router.post("/creators/apply")
async def submit_creator_application(body: CreatorApplicationRequest):
    """Public, unauthenticated submission from the /creators program landing page."""
    email = body.email.strip().lower()
    first_name = body.first_name.strip()
    last_name = body.last_name.strip()
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="A valid email is required")
    if not first_name or not last_name:
        raise HTTPException(status_code=400, detail="First and last name are required")
    if not (body.tiktok_handle or "").strip() and not (body.instagram_handle or "").strip():
        raise HTTPException(status_code=400, detail="At least one social profile is required")

    doc = await create_creator_application(
        db,
        email=email,
        first_name=first_name,
        last_name=last_name,
        tiktok_handle=(body.tiktok_handle or "").strip() or None,
        instagram_handle=(body.instagram_handle or "").strip() or None,
        has_company=body.has_company,
        whatsapp_country=body.whatsapp_country,
        whatsapp_number=(body.whatsapp_number or "").strip() or None,
        country=body.country,
        referred_by=(body.referred_by or "").strip() or None,
        message=(body.message or "").strip() or None,
    )
    return {"ok": True, "creator_application_id": doc["creator_application_id"]}


@api_router.post("/onboarding/suggest-roles")
async def onboarding_suggest_roles(body: OnboardingSuggestRolesRequest):
    """Specific role titles for selected onboarding categories."""
    if not body.categories:
        raise HTTPException(status_code=400, detail="At least one category is required")
    categories = [category.model_dump() for category in body.categories]
    return await suggest_roles(
        body.location.strip(),
        body.contract_type.strip(),
        categories,
        body.location_data,
    )


@api_router.put("/account/settings")
async def update_account_settings(body: AccountSettingsUpdate, user: User = Depends(get_current_user)):
    if body.require_review_before_send is not None:
        await _set_user_require_review_before_send(user.user_id, body.require_review_before_send)
    language = body.language if body.language in {"en", "fr"} else None
    if language is not None:
        await _set_user_language(user.user_id, language)
    return {
        "require_review_before_send": (
            body.require_review_before_send
            if body.require_review_before_send is not None
            else user.require_review_before_send
        ),
        "language": language if language is not None else user.language,
    }


@api_router.put("/profile/preferences")
async def update_preferences(prefs: PreferencesUpdate, user: User = Depends(get_current_user)):
    payload = prefs.model_dump(exclude_unset=True)
    update = {
        k: v
        for k, v in payload.items()
        if v is not None or k == "target_location_data"
    }
    target_role = update.get("target_role")
    target_roles = update.get("target_roles")
    if isinstance(target_roles, list) and target_roles:
        cleaned = [role.strip() for role in target_roles if isinstance(role, str) and role.strip()]
        if cleaned:
            update["target_roles"] = cleaned
            update["target_role"] = cleaned[0]
    elif isinstance(target_role, str) and target_role.strip():
        update["target_role"] = target_role.strip()
        update["target_roles"] = [target_role.strip()]
    update["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.profiles.update_one(
        {"user_id": user.user_id},
        {"$set": update},
        upsert=True,
    )
    profile = await db.profiles.find_one({"user_id": user.user_id}, {"_id": 0})
    return profile


@api_router.put("/profile/contact")
async def update_contact(contact: ContactUpdate, user: User = Depends(get_current_user)):
    """Update the contact block. Email is ALWAYS forced to the authenticated user's
    registered email — the UI never allows changing it, per product spec."""
    data = {k: v for k, v in contact.model_dump().items() if v is not None}
    data["email"] = user.email  # force-overwrite
    existing = await db.profiles.find_one({"user_id": user.user_id}, {"_id": 0, "contact": 1}) or {}
    merged = {**(existing.get("contact") or {}), **data}
    first_name = _clean_string(merged.get("first_name"))
    last_name = _clean_string(merged.get("last_name"))
    if first_name or last_name:
        merged["name"] = " ".join(part for part in (first_name, last_name) if part).strip()
    await db.profiles.update_one(
        {"user_id": user.user_id},
        {"$set": {"contact": merged, "updated_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True,
    )
    return {"ok": True, "contact": merged}


# ===================== Jobs / Swipe =====================

@api_router.get("/jobs/{job_id}/rome-profile")
async def get_job_rome_profile(job_id: str, user: User = Depends(get_current_user)):
    """Official France Travail ROME 4.0 occupation profile for a job card."""
    job = await db.jobs.find_one(
        {"job_id": job_id},
        {"_id": 0, "rome_code": 1, "rome_label": 1, "title": 1},
    )
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    rome_code = normalize_rome_code(job.get("rome_code"))
    if not rome_code:
        return {"available": False, "reason": "no_rome_code"}
    if not rome_profile_enabled():
        return {"available": False, "reason": "rome_profile_disabled"}
    return await get_rome_profile(
        db,
        rome_code,
        rome_label=job.get("rome_label") or job.get("title"),
    )


@api_router.get("/jobs/feed")
async def get_feed(
    user: User = Depends(get_current_user),
    limit: int = 5,
    min_salary: int = 0,
    posted_within: Optional[str] = None,            # any | 1d | 7d | 30d
    work_location: Optional[List[str]] = Query(None),   # remote | hybrid | onsite
    job_type: Optional[List[str]] = Query(None),        # full_time | part_time | internship | fixed_term | apprenticeship | summer_job | seasonal | freelance
    experience: Optional[List[str]] = Query(None),      # entry | mid | senior | executive
    location: Optional[List[str]] = Query(None),        # free-text city/country tokens, OR-matched on `location` field
    only_company: Optional[List[str]] = Query(None),
    hide_company: Optional[List[str]] = Query(None),
    only_industry: Optional[List[str]] = Query(None),   # placeholder (jobs lack industry field today)
    hide_industry: Optional[List[str]] = Query(None),   # placeholder
    include_unknown_location: bool = True,
    include_unknown_salary: bool = True,
    include_non_auto_apply: bool = False,
    search_radius: str = "50km",
    locations_json: Optional[str] = None,
    only_my_country: bool = False,
    location_label: Optional[str] = None,
    place_id: Optional[str] = None,
    country: Optional[str] = None,
    country_code: Optional[str] = None,
    lat: Optional[float] = None,
    lng: Optional[float] = None,
    force_provider_refresh: bool = False,
    prefetch: bool = False,
    score: bool = False,                                  # opt-in AI scoring (slow); default off for snappy UX
    search_role: Optional[str] = None,                    # override profile target_role for this feed request
    audit_mode: bool = False,
):
    started_at = time.perf_counter()
    logger.info(
        "jobs/feed start: user_id=%s limit=%s search_radius=%s include_non_auto_apply=%s locations_json=%s location=%s",
        user.user_id,
        limit,
        search_radius,
        include_non_auto_apply,
        bool(locations_json),
        location,
    )
    profile = await db.profiles.find_one({"user_id": user.user_id}, {"_id": 0})
    if user.demo_account:
        profile = _demo_profile_for_feed(user.user_id, profile)
    elif not profile or not profile.get("cv_text"):
        logger.info(
            "jobs/feed cv_readiness_failed user_id=%s profile_exists=%s has_cv_text=%s has_cv_filename=%s profile_keys=%s",
            user.user_id,
            profile is not None,
            bool((profile or {}).get("cv_text")),
            bool((profile or {}).get("cv_filename")),
            sorted(list((profile or {}).keys()))[:30],
        )
        raise HTTPException(status_code=400, detail="Upload CV first")

    if not _profile_has_usable_phone(profile):
        logger.info(
            "jobs/feed phone_readiness_failed user_id=%s has_phone=%s",
            user.user_id,
            bool((profile.get("contact") or {}).get("phone")),
        )
        raise HTTPException(status_code=400, detail="Add your phone number to apply")

    if search_role is not None and not str(search_role).strip():
        search_role = None
    if search_role is not None:
        feed_target_role = search_role.strip()
    else:
        feed_target_role = resolve_profile_target_role(profile)

    async def _legacy_jsearch_only_feed() -> Dict[str, Any]:
        legacy_started = time.perf_counter()
        provider_name = primary_job_provider_name()
        if not is_job_provider_configured(provider_name):
            logger.warning("jobs/feed legacy_provider missing_api_key provider=%s", provider_name)
            return {
                "jobs": [],
                "total": 0,
                "feed_mode": "legacy_jsearch_only",
                "fallback_reason": "missing_job_provider_credentials",
                "provider_rate_limited": False,
                "refresh_results": [{"attempted": False, "reason": "missing_api_key"}],
            }
        api_key = os.environ.get("JSEARCH_API_KEY") or ""
        selected_location_data = None
        selected_location_label = None
        if locations_json:
            try:
                parsed_locations = json.loads(locations_json)
                if isinstance(parsed_locations, list) and parsed_locations:
                    first = parsed_locations[0]
                    if isinstance(first, dict):
                        selected_location_data = first
                        selected_location_label = first.get("location_label")
            except json.JSONDecodeError:
                pass
        if location_label:
            selected_location_label = location_label
            selected_location_data = {
                "location_label": location_label,
                "place_id": place_id,
                "country": country,
                "country_code": country_code,
                "lat": lat,
                "lng": lng,
            }
        elif location and not selected_location_label:
            selected_location_label = location[0]

        query = build_profile_job_query(
            profile,
            location_override=selected_location_label,
            location_data_override=selected_location_data,
            search_radius=search_radius,
            role_override=feed_target_role,
        )
        query = JobSearchQuery(
            role=query.role,
            location=query.location,
            remote_preference=query.remote_preference,
            country=query.country,
            language=query.language,
            limit=max(5, min(int(limit or 5) * 4, _env_int("JOBS_FEED_LEGACY_JSEARCH_LIMIT", 25))),
            raw_query=query.raw_query,
            max_pages=max(1, min(_env_int("JOBS_FEED_LEGACY_JSEARCH_MAX_PAGES", 1), 3)),
            page_size=max(5, min(_env_int("JOBS_FEED_LEGACY_JSEARCH_PAGE_SIZE", 20), 50)),
        )
        provider = get_job_provider(provider_name, api_key)
        jsearch_attempted = True
        jsearch_error = None
        import_stats = None
        try:
            if is_france_travail_provider(provider_name):
                import_stats = await jobs_service_module._import_provider_jobs(
                    db, provider, query
                )
                raw_jobs = import_stats.get("jobs", [])
            else:
                result = await provider.search(query)
                raw_jobs = result.jobs
        except Exception as exc:
            raw_jobs = []
            jsearch_error = f"{exc.__class__.__name__}: {str(exc)[:160]}"
            logger.warning("jobs/feed legacy_jsearch provider_error=%s", jsearch_error)
        if raw_jobs:
            try:
                if import_stats is None:
                    import_stats = await jobs_service_module.upsert_imported_jobs(
                        db, raw_jobs
                    )
                logger.info(
                    "jobs/feed legacy_jsearch_upserted fetched=%s imported=%s auto_apply_supported=%s",
                    len(raw_jobs),
                    import_stats.get("total_imported", 0),
                    import_stats.get("auto_apply_supported_imported", 0),
                )
            except Exception as exc:
                logger.warning(
                    "jobs/feed legacy_jsearch_upsert_failed fetched=%s error=%s",
                    len(raw_jobs),
                    f"{exc.__class__.__name__}: {str(exc)[:160]}",
                )

        swiped_rows = await db.swipes.find({"user_id": user.user_id}, {"_id": 0, "job_id": 1}).limit(500).to_list(500)
        swiped_ids = {row.get("job_id") for row in swiped_rows if row.get("job_id")}
        jobs = []
        for job in raw_jobs:
            if job.get("job_id") in swiped_ids:
                continue
            enriched = _with_apply_fulfillment_fields(job)
            clean = {key: value for key, value in enriched.items() if not key.startswith("_")}
            jobs.append({
                **clean,
                "match_score": clean.get("match_score") or random.randint(72, 94),
                "match_reasons": clean.get("match_reasons") or ["Fresh JSearch result from emergency legacy feed mode."],
            })
            if len(jobs) >= max(1, min(int(limit or 5), 25)):
                break
        elapsed = int((time.perf_counter() - legacy_started) * 1000)
        debug = {
            "mode": "legacy_jsearch_only",
            "request": {
                "search_role": search_role,
                "search_radius": search_radius,
                "work_location": work_location,
                "job_type": job_type,
                "experience": experience,
                "posted_within": posted_within,
                "min_salary": min_salary,
                "only_my_country": only_my_country,
                "include_unknown_location": include_unknown_location,
                "include_unknown_salary": include_unknown_salary,
            },
            "jsearch_attempted": jsearch_attempted,
            "jsearch_count": len(raw_jobs),
            "jsearch_error": jsearch_error,
            "swiped_count": len(swiped_ids),
            "final_returned_count": len(jobs),
            "duration_ms": elapsed,
        }
        logger.info(
            "jobs/feed legacy_jsearch_complete user_id_hash=%s jsearch_count=%s returned=%s elapsed_ms=%s error=%s",
            hashlib.sha1(user.user_id.encode("utf-8")).hexdigest()[:12],
            len(raw_jobs),
            len(jobs),
            elapsed,
            bool(jsearch_error),
        )
        response = {
            "jobs": jobs,
            "total": len(jobs),
            "feed_mode": "legacy_jsearch_only",
            "auto_apply_count": sum(1 for job in jobs if job.get("auto_apply_supported") is True),
            "total_count": len(raw_jobs),
            "fallback_reason": "legacy_jsearch_only",
            "searched_location": query.location,
            "searched_locations": [query.location] if query.location else [],
            "search_radius": search_radius,
            "suggested_next_radius": None if jobs else "worldwide",
            "only_my_country": only_my_country,
            "widened_search": search_radius in ("worldwide", "remote/worldwide"),
            "original_location": selected_location_label or profile.get("target_location"),
            "final_location_used": query.location or "worldwide",
            "provider_rate_limited": False,
            "provider_cooldown_until": None,
            "refresh_results": [{"attempted": jsearch_attempted, "reason": "legacy_jsearch_only", "count": len(raw_jobs), "error": jsearch_error}],
            "matched_role": feed_target_role or None,
            "matched_location": [query.location] if query.location else [],
            "companies_returned": sorted({job.get("company") for job in jobs if job.get("company")}),
            "filters_applied": {"target_role": feed_target_role or None, "search_radius": search_radius, "legacy_jsearch_only": True},
            "feed_elapsed_ms": elapsed,
            "fallback_used": "legacy_jsearch_only",
            "explicit_filters": True,
        }
        if _env_bool("JOBS_FEED_DEBUG_DIAGNOSTICS", False) and _is_admin_email(user.email):
            response["debug"] = debug
        return response

    if _env_bool("JOBS_FEED_LEGACY_JSEARCH_ONLY", False) and not audit_mode:
        return await _legacy_jsearch_only_feed()

    # Raised alongside sync_refresh_max_seconds/total_seconds below (was 16.0
    # paired with a 12s sync budget, itself raised from an 8.0/4s pairing).
    # Kept a few seconds above the new 30s sync/refresh budget so the later
    # widen/backfill stages (which check _timed_out()) still get some
    # post-refresh runway instead of being starved the instant the
    # synchronous JSearch+France Travail refresh finishes.
    max_elapsed_seconds = 34.0
    ats_supported = ["greenhouse", "lever", "ashby"]

    def _elapsed_ms() -> int:
        return int((time.perf_counter() - started_at) * 1000)

    def _timed_out() -> bool:
        return (time.perf_counter() - started_at) >= max_elapsed_seconds

    def _tokens(value: str) -> List[str]:
        value = unicodedata.normalize("NFKD", value or "").encode("ascii", "ignore").decode("ascii")
        stop = {
            "and", "or", "the", "a", "an", "of", "for", "to", "in", "with",
            "remote", "jobs", "job", "cdi", "cdd", "full", "time",
        } | ACADEMIC_LEVEL_STOPWORDS
        return [
            token
            for token in re.findall(r"[a-z0-9]+", (value or "").lower())
            if len(token) > 2 and token not in stop
        ]

    def _job_text(job: Dict[str, Any]) -> str:
        text = " ".join([
            str(job.get("title") or ""),
            str(job.get("company") or ""),
            str(job.get("description") or ""),
            str(job.get("clean_description") or ""),
            " ".join(str(item) for item in (job.get("requirements") or [])),
        ]).lower()
        return unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")

    # Broad, Indeed/LinkedIn-scale occupational taxonomy (EN + FR keywords) so
    # the feed recognizes essentially any common job field, not just
    # office/tech roles. This single dict drives both category compatibility
    # (_category_compatible) and role-family token expansion for scoring
    # (_role_family_tokens derives its expansion straight from it, so the two
    # can't drift out of sync the way two hand-maintained lists would).
    # "chef" alone is excluded from french-catering keywords: in French it
    # means "lead" (chef de projet, chef de chantier), not a kitchen chef.
    # Use "cuisine"/"cuisinier" instead.
    role_category_keywords = {
        "technology": {"software", "developer", "developpeur", "engineer", "ingenieur", "frontend", "backend", "fullstack", "full-stack", "javascript", "python", "java", "devops", "cloud", "qa", "informatique", "programmeur", "programmer", "sysadmin", "reseau", "network", "cybersecurity", "cybersecurite", "it"},
        "data_analytics": {"analyst", "analyste", "analytics", "analysis", "data", "donnees", "scientist", "scientifique", "statistician", "statisticien", "bi", "insights", "reporting"},
        "product": {"product", "produit", "product owner", "product manager"},
        "design_creative": {"design", "designer", "ux", "ui", "graphiste", "graphic", "creative", "createur", "illustrator", "illustrateur", "artiste", "artist", "motion"},
        "marketing_communications": {"marketing", "communication", "community", "seo", "sem", "brand", "marque", "contenu", "content", "digital", "growth", "social", "attache de presse", "pr", "publicite", "advertising"},
        "sales": {"sales", "commercial", "commerciale", "vente", "vendeur", "vendeuse", "account executive", "account manager", "business developer", "prospection", "closer"},
        "customer_service": {"customer", "support", "success", "client", "clientele", "service client", "help desk", "assistance", "relation client"},
        "hr_recruiting": {"hr", "rh", "human resources", "ressources humaines", "recruiter", "recruteur", "recrutement", "recruiting", "talent", "paie", "payroll", "formation", "headhunter"},
        "administrative": {"administrative", "administratif", "assistant", "assistante", "direction", "reception", "receptionniste", "receptionist", "office", "secretaire", "secretary", "back office"},
        "executive_management": {"executive", "director", "directeur", "directrice", "ceo", "coo", "cfo", "cto", "vp", "vice president", "president"},
        "project_management": {"project", "projet", "program manager", "programme", "scrum", "agile", "chef de projet"},
        "general_management": {"manager", "management", "responsable", "gerant", "gerante", "supervisor", "superviseur", "lead", "team lead"},
        "operations": {"operations", "operationnel", "ops"},
        "finance_accounting": {"finance", "financier", "accountant", "comptable", "comptabilite", "accounting", "auditor", "audit", "controleur de gestion", "controller", "treasury", "tresorerie", "credit"},
        "banking_insurance": {"bank", "banque", "banking", "bancaire", "insurance", "assurance", "assureur", "actuary", "actuaire", "courtier"},
        "legal": {"legal", "juridique", "lawyer", "avocat", "avocate", "attorney", "notaire", "notary", "paralegal", "juriste", "compliance", "conformite", "droit"},
        "consulting": {"consultant", "consultante", "consulting", "conseil", "advisory"},
        "supply_chain_logistics": {"warehouse", "logistics", "logistique", "magasinier", "preparateur", "supply chain", "approvisionnement", "inventory", "stock", "cariste", "forklift"},
        "transportation_driving": {"driver", "delivery", "chauffeur", "livreur", "livreuse", "trucking", "routier", "taxi", "vtc", "coursier"},
        "aviation_maritime": {"pilot", "pilote", "aviation", "airline", "flight attendant", "steward", "hotesse", "maritime", "marin", "seafarer"},
        "manufacturing_production": {"manufacturing", "production", "factory", "usine", "assembly", "assemblage", "machiniste", "machinist", "technicien de production"},
        "construction_trades": {"construction", "batiment", "chantier", "electricien", "electrician", "plumber", "plombier", "carpenter", "charpentier", "mason", "macon", "welder", "soudeur", "roofer", "couvreur", "peintre", "hvac", "chauffagiste", "cvc", "climatisation", "chauffage", "climaticien", "genie climatique"},
        "engineering_technical": {"mechanical engineer", "electrical engineer", "civil engineer", "industrial engineer", "ingenieur mecanique", "ingenieur electrique", "ingenieur civil", "technicien", "technician", "maintenance"},
        "real_estate": {"real estate", "immobilier", "realtor", "property manager", "syndic"},
        "hospitality_food_service": {"hospitality", "hotellerie", "hotel", "restaurant", "waiter", "serveur", "serveuse", "barista", "barman", "bartender", "kitchen", "cuisine", "cuisinier", "cuisiniere", "commis de cuisine", "plongeur", "concierge"},
        "retail": {"retail", "store", "magasin", "boutique", "cashier", "caissier", "caissiere", "employe polyvalent", "sales associate", "merchandiser"},
        "beauty_personal_care": {"barber", "barbier", "hairdresser", "coiffeur", "coiffeuse", "coiffure", "esthetician", "esthetique", "beautician", "manucure", "nail", "spa", "cosmetologist", "maquilleur", "makeup artist"},
        "cleaning_facilities": {"cleaner", "cleaning", "nettoyage", "menage", "housekeeper", "housekeeping", "femme de menage", "agent d'entretien", "janitor", "facilities"},
        "security_law_enforcement": {"security", "securite", "guard", "gardien", "police", "policier", "gendarme", "firefighter", "pompier", "surveillance"},
        "healthcare_medical": {"nurse", "infirmier", "infirmiere", "medical", "medecin", "doctor", "physician", "sante", "care", "soignant", "soignante", "aide-soignant", "pharmacy", "pharmacien", "dentist", "dentiste", "surgeon", "chirurgien", "radiologist", "radiologue", "physiotherapist", "kinesitherapeute", "caregiver"},
        "mental_health_social_work": {"psychologist", "psychologue", "psychiatrist", "psychiatre", "therapist", "therapeute", "counselor", "social worker", "travailleur social", "assistant social", "assistante sociale", "psychotherapy", "psychotherapie"},
        "veterinary_animal_care": {"veterinarian", "veterinaire", "vet tech", "toiletteur", "groomer", "kennel", "pet sitter"},
        "education_training": {"teacher", "enseignant", "enseignante", "professeur", "professor", "trainer", "formateur", "formatrice", "teaching", "tutor", "tuteur", "instructor", "moniteur", "educateur", "school", "ecole", "university", "universite"},
        "science_research": {"research", "researcher", "recherche", "chercheur", "scientist", "scientifique", "laboratory", "laboratoire", "biologist", "biologiste", "chemist", "chimiste", "physicist", "physicien", "r&d"},
        "arts_entertainment_media": {"artist", "artiste", "musician", "musicien", "actor", "acteur", "actrice", "journalist", "journaliste", "writer", "ecrivain", "editor", "redacteur", "redactrice", "photographer", "photographe", "videographer", "video", "producer", "producteur", "media", "entertainment", "divertissement"},
        "sports_fitness": {"coach", "fitness", "personal trainer", "gym", "athlete", "sport instructor", "moniteur sportif"},
        "childcare_family_services": {"childcare", "nanny", "nounou", "babysitter", "assistante maternelle", "daycare", "creche", "au pair"},
        "agriculture_farming": {"agriculture", "farming", "agricole", "farmer", "agriculteur", "agricultrice", "viticulture", "vigneron", "elevage", "ouvrier agricole"},
        "energy_environment": {"energy", "energie", "renewable", "renouvelable", "solar", "solaire", "wind", "eolien", "environment", "environnement", "sustainability", "durabilite", "utilities", "electricite"},
        "nonprofit_government": {"nonprofit", "non-profit", "association", "ngo", "ong", "humanitarian", "humanitaire", "government", "gouvernement", "public sector", "fonction publique", "fonctionnaire"},
    }

    def _normalize_ascii(value: str) -> str:
        return unicodedata.normalize("NFKD", (value or "").lower()).encode("ascii", "ignore").decode("ascii")

    def _keyword_matches(keyword: str, normalized_text: str, word_set: set) -> bool:
        # Whole-word match, not plain substring: several category keywords
        # are short ("ui", "it", "hr", "pr", "bi", "vp") and a naive `in`
        # check matches them inside unrelated words (e.g. "ui" inside
        # "cuisinier", "it" inside "recruiter"). Multi-word phrases
        # ("account executive") aren't single tokens, so those still need a
        # padded substring check.
        if " " in keyword:
            return f" {keyword} " in f" {normalized_text} "
        return keyword in word_set

    def _text_category(normalized_text: str) -> Optional[str]:
        word_set = set(normalized_text.split())
        # Multi-word phrases first: a generic single-word keyword shared
        # across categories (e.g. "developer" in technology, "manager" in
        # general_management, "trainer" in education_training) would
        # otherwise win over a much more specific phrase match elsewhere
        # ("business developer" -> sales, "property manager" -> real_estate,
        # "personal trainer" -> sports_fitness).
        for category, keywords in role_category_keywords.items():
            if any(" " in keyword and _keyword_matches(keyword, normalized_text, word_set) for keyword in keywords):
                return category
        for category, keywords in role_category_keywords.items():
            if any(" " not in keyword and _keyword_matches(keyword, normalized_text, word_set) for keyword in keywords):
                return category
        return None

    def _role_category_from_text(role: str) -> Optional[str]:
        # A few role labels are ambiguous/homonymous across languages (e.g.
        # "Chef"/"Cuisinier"); resolve to their safe token list first so
        # category detection isn't neutralized by the raw ambiguous word
        # (bare "chef" is deliberately absent from every category's keyword
        # set -- see role_category_keywords comment above).
        override_tokens = resolve_role_match_tokens(role)
        text_for_category = " ".join(override_tokens) if override_tokens else role
        return _text_category(_normalize_ascii(text_for_category))

    def _role_family_tokens(role: str) -> List[str]:
        override_tokens = resolve_role_match_tokens(role)
        tokens = override_tokens if override_tokens is not None else _tokens(role)
        family = list(tokens)
        normalized = _normalize_ascii(" ".join(tokens) if override_tokens else role)
        word_set = set(normalized.split())
        for keywords in role_category_keywords.values():
            if any(_keyword_matches(keyword, normalized, word_set) for keyword in keywords):
                family.extend(keywords)
        return list(dict.fromkeys(token for token in family if token))

    target_role_category = _role_category_from_text(feed_target_role)
    feed_strict_tokens = resolve_role_match_tokens(feed_target_role) or _tokens(feed_target_role)

    def _job_role_category(job: Dict[str, Any]) -> Optional[str]:
        title_category = _text_category(_normalize_ascii(str(job.get("title") or "")))
        if title_category:
            return title_category
        return _text_category(_job_text(job))

    # Sibling categories close enough that either should satisfy the other
    # (e.g. a "Sales" search shouldn't hide Marketing-adjacent postings that
    # often do the same job under a different title).
    _COMPATIBLE_CATEGORY_PAIRS = {
        frozenset({"sales", "marketing_communications"}),
        frozenset({"sales", "administrative"}),
        frozenset({"sales", "customer_service"}),
        frozenset({"hr_recruiting", "administrative"}),
        frozenset({"general_management", "executive_management"}),
        frozenset({"general_management", "project_management"}),
        frozenset({"general_management", "operations"}),
        frozenset({"operations", "supply_chain_logistics"}),
        frozenset({"finance_accounting", "banking_insurance"}),
        frozenset({"legal", "consulting"}),
        frozenset({"engineering_technical", "manufacturing_production"}),
        frozenset({"engineering_technical", "construction_trades"}),
        frozenset({"healthcare_medical", "mental_health_social_work"}),
        frozenset({"beauty_personal_care", "hospitality_food_service"}),
        frozenset({"cleaning_facilities", "hospitality_food_service"}),
        frozenset({"education_training", "childcare_family_services"}),
        frozenset({"science_research", "technology"}),
        frozenset({"data_analytics", "technology"}),
        frozenset({"product", "technology"}),
    }

    def _category_compatible(job: Dict[str, Any]) -> bool:
        if not target_role_category:
            return True
        job_category = _job_role_category(job)
        if not job_category:
            return True
        if target_role_category == job_category:
            return True
        if frozenset({target_role_category, job_category}) in _COMPATIBLE_CATEGORY_PAIRS:
            return True
        # Free-text roles (not in the suggestion list): keep jobs whose title or
        # description contains the user's search tokens even across categories.
        if feed_strict_tokens:
            title = _normalize_ascii(str(job.get("title") or ""))
            if any(token in title for token in feed_strict_tokens):
                return True
            if any(token in _job_text(job) for token in feed_strict_tokens):
                return True
        return False

    def _role_score(job: Dict[str, Any], strict_tokens: List[str], family_tokens: List[str]) -> int:
        title = unicodedata.normalize("NFKD", (job.get("title") or "").lower()).encode("ascii", "ignore").decode("ascii")
        text = _job_text(job)
        if not _category_compatible(job):
            return 0
        if not strict_tokens and not family_tokens:
            return 10
        strict_title_hits = sum(1 for token in strict_tokens if token in title)
        strict_text_hits = sum(1 for token in strict_tokens if token in text)
        family_title_hits = sum(1 for token in family_tokens if token in title)
        family_text_hits = sum(1 for token in family_tokens if token in text)
        return strict_title_hits * 30 + min(strict_text_hits, 4) * 8 + family_title_hits * 14 + min(family_text_hits, 6) * 3

    def _profile_feed_location_data() -> Dict[str, Any]:
        return resolve_profile_target_location_data(profile)

    def _profile_feed_location_label() -> str:
        return resolve_profile_target_location_label(profile)

    def _location_terms() -> Dict[str, List[str]]:
        raw_locations: List[str] = []
        explicit_request_location = False
        if locations_json:
            try:
                parsed = json.loads(locations_json)
                if isinstance(parsed, list):
                    parsed_labels = [str(item.get("location_label") or "") for item in parsed if isinstance(item, dict)]
                    raw_locations.extend(parsed_labels)
                    explicit_request_location = any(label for label in parsed_labels)
            except json.JSONDecodeError:
                pass
        if location_label:
            raw_locations.append(location_label)
            explicit_request_location = True
        if location:
            raw_locations.extend(location)
            explicit_request_location = True
        target_location_data = _profile_feed_location_data()
        if not explicit_request_location:
            if target_location_data.get("location_label"):
                raw_locations.append(str(target_location_data.get("location_label")))
            elif _profile_feed_location_label():
                raw_locations.append(_profile_feed_location_label())
        country_code_value = (
            (country_code or "")
            or ("" if explicit_request_location else str(target_location_data.get("country_code") or ""))
        ).lower().strip()
        country_value = (
            (country or "")
            or ("" if explicit_request_location else str(target_location_data.get("country") or ""))
        ).lower().strip()
        aliases = {
            "fr": ["france", "paris", "ile de france", "ile-de-france", "île-de-france"],
            "gb": ["united kingdom", "uk", "england", "london"],
            "us": ["united states", "usa", "new york", "san francisco"],
            "ma": ["morocco", "maroc", "casablanca"],
        }
        # Tokenize only the city part (before the comma) so the country token
        # ("france") never counts as a city match for jobs in other cities.
        known_country_tokens = {"france", "morocco", "maroc", "england", "uk", "usa"}
        city_terms = list(dict.fromkeys(
            token
            for label in raw_locations
            for token in _tokens(re.split(r"[,/|]", label, maxsplit=1)[0])
            if token not in known_country_tokens
        ))
        country_terms = list(dict.fromkeys([country_value, *aliases.get(country_code_value, [])]))
        return {"labels": [label for label in raw_locations if label], "city": city_terms, "country": [term for term in country_terms if term]}

    def _location_score(job: Dict[str, Any], terms: Dict[str, List[str]], worldwide: bool = False) -> int:
        if worldwide:
            return 10
        job_location = normalize_place_name(" ".join([
            str(job.get("city") or ""),
            str(job.get("region") or ""),
            str(job.get("location") or ""),
        ]))
        remote_value = str(job.get("remote") or "").lower()
        score_value = 0
        if any(term and term in job_location for term in terms["city"]):
            score_value += 45
        if any(term and term in job_location for term in terms["country"]):
            score_value += 25
        if remote_value == "remote":
            score_value += 8
        return score_value

    def _recency_score(job: Dict[str, Any]) -> int:
        raw = job.get("posted_at") or job.get("imported_at") or job.get("last_seen_at")
        if not raw:
            return 0
        try:
            parsed = datetime.fromisoformat(str(raw).replace("Z", "+00:00"))
            if parsed.tzinfo is None:
                parsed = parsed.replace(tzinfo=timezone.utc)
        except ValueError:
            return 0
        age_days = max(0, (datetime.now(timezone.utc) - parsed).days)
        if age_days <= 7:
            return 20
        if age_days <= 30:
            return 10
        return 0

    async def _fast_cached_feed() -> Dict[str, Any]:
        global _feed_sync_refresh_cooldown_until, _feed_sync_refresh_cooldowns
        requested_limit = max(1, min(int(limit or 5), 25))
        db_first_enabled = _env_bool("JOBS_DB_FIRST_ENABLED", True)
        db_min_good_results = max(1, _env_int("JOBS_DB_MIN_GOOD_RESULTS_BEFORE_JSEARCH", 30))
        db_weak_results_threshold = max(0, _env_int("JOBS_DB_WEAK_RESULTS_THRESHOLD", 10))
        allow_unknown_tier = _env_bool("JOBS_ALLOW_UNKNOWN_TIER_IN_FEED", False)
        sync_refresh_enabled = not audit_mode and _env_bool("JOBS_FEED_SYNC_REFRESH_ENABLED", True)
        # Was 4s. Confirmed live that JSearch's own response time regularly
        # exceeds that for non-French markets (Marseille/Lyon: ~4-5s; London/
        # New York: 10-20s, reliably blowing a 4s budget before any answer
        # comes back). France barely ever hits this path at all (its DB cache
        # is deep enough that JOBS_DB_FIRST_ENABLED usually short-circuits
        # before reaching here), so raising this doesn't change France's
        # outcome -- it just gives thinner markets enough time to actually
        # get a response instead of failing before JSearch replies. Set a
        # bit above JSearchProvider's own httpx client timeout (10s default,
        # JSEARCH_HTTP_TIMEOUT_SECONDS) so a genuinely slow call fails on its
        # own terms instead of always hitting this wrapper's cutoff first.
        # Raised 12 -> 30: 12s was cutting off calls that would otherwise have
        # confirmed a real error or a clean zero-results response given more
        # room -- a call should only be treated as a timeout once it's
        # genuinely still pending at 30s, not chopped off pre-emptively.
        sync_refresh_max_seconds = max(1, min(_env_int("JOBS_FEED_SYNC_REFRESH_MAX_SECONDS", 30), 30))
        # page_size default raised 10 -> 50: JSearch's per-call latency is
        # dominated by the network round-trip, not by how many results come
        # back in that one response (its own cap is 100/call), so this is a
        # free 5x increase in jobs-per-request within the same single HTTP
        # call and the same sync_refresh_max_seconds budget. max_results
        # raised to match, otherwise the extra results fetched would just
        # get discarded unused.
        sync_refresh_max_results = max(1, min(_env_int("JOBS_FEED_SYNC_REFRESH_MAX_RESULTS", 50), 50))
        sync_refresh_max_pages = max(1, min(_env_int("JSEARCH_FEED_FALLBACK_MAX_PAGES", 1), 3))
        sync_refresh_page_size = max(1, min(_env_int("JSEARCH_FEED_FALLBACK_PAGE_SIZE", 50), 50))
        sync_refresh_cooldown_seconds = max(30, min(_env_int("JOBS_FEED_SYNC_REFRESH_COOLDOWN_SECONDS", 300), 1800))
        sync_refresh_total_seconds = max(2, min(_env_int("JOBS_FEED_SYNC_REFRESH_TOTAL_SECONDS", 30), 30))
        sync_refresh_attempts_per_city = max(1, min(_env_int("JOBS_FEED_PROVIDER_ATTEMPTS_PER_CITY", 2), 4))
        if primary_job_provider_name() == "france_travail":
            sync_refresh_max_pages = max(sync_refresh_max_pages, min(_env_int("FRANCE_TRAVAIL_MAX_PAGES", 2), 4))
            sync_refresh_page_size = max(sync_refresh_page_size, min(_env_int("FRANCE_TRAVAIL_PAGE_SIZE", 50), 150))
            sync_refresh_max_results = max(sync_refresh_max_results, min(_env_int("FRANCE_TRAVAIL_PAGE_SIZE", 50), 100))
            sync_refresh_attempts_per_city = 1
        target_role = feed_target_role
        profile_contract_type = resolve_profile_contract_type(profile)
        effective_job_types = list(job_type or [])
        if not effective_job_types and profile_contract_type:
            effective_job_types = contract_type_to_job_types(profile_contract_type)
        profile_for_refresh = dict(profile or {})
        if profile_contract_type and not profile_for_refresh.get("contract_type"):
            profile_for_refresh["contract_type"] = profile_contract_type
        strict_tokens = feed_strict_tokens
        family_tokens = _role_family_tokens(target_role)
        terms = _location_terms()
        radius_scope = (search_radius or "50km").lower().strip()
        is_worldwide_radius = radius_scope in ("worldwide", "remote/worldwide")
        radius_km_match = re.match(r"^(\d+)\s*km$", radius_scope)
        radius_km = int(radius_km_match.group(1)) if radius_km_match else None

        def _parse_selected_locations() -> List[Dict[str, Any]]:
            selected: List[Dict[str, Any]] = []
            if locations_json:
                try:
                    parsed = json.loads(locations_json)
                    if isinstance(parsed, list):
                        selected.extend(
                            loc
                            for loc in parsed
                            if isinstance(loc, dict) and (loc.get("location_label") or loc.get("country") or loc.get("country_code"))
                        )
                except json.JSONDecodeError:
                    pass
            if location_label:
                selected.append({
                    "location_label": location_label,
                    "place_id": place_id,
                    "country": country,
                    "country_code": country_code,
                    "lat": lat,
                    "lng": lng,
                })
            if location:
                selected.extend({"location_label": item} for item in location if item)
            return [_enrich_feed_location(loc) for loc in selected]

        def _enrich_feed_location(loc: Dict[str, Any]) -> Dict[str, Any]:
            enriched = dict(loc)
            label = str(enriched.get("location_label") or enriched.get("label") or "").strip()
            if not label:
                return enriched
            country_code = str(enriched.get("country_code") or "").lower().strip()
            if not country_code and jobs_service_module._looks_like_france_location(label):
                country_code = "fr"
                enriched["country_code"] = country_code
                enriched["country"] = enriched.get("country") or "France"
                if "france" not in label.lower():
                    city = label.split(",")[0].strip()
                    enriched["location_label"] = f"{city}, France"
            return enriched

        request_selected_locations = _parse_selected_locations()
        selected_locations = list(request_selected_locations)
        if not selected_locations and _profile_feed_location_data():
            selected_locations = [_profile_feed_location_data()]
        elif not selected_locations and _profile_feed_location_label():
            selected_locations = [{"location_label": _profile_feed_location_label()}]
        explicit_local_intent = bool(request_selected_locations) and radius_km is not None and not is_worldwide_radius
        explicit_location_filter = bool(
            selected_locations
            or location
            or location_label
            or place_id
            or country
            or country_code
            or lat is not None
            or lng is not None
        )
        explicit_filters = bool(
            min_salary > 0
            or (posted_within and posted_within != "any")
            or work_location
            or effective_job_types
            or experience
            or explicit_location_filter
            or only_company
            or hide_company
            or only_industry
            or hide_industry
            or include_unknown_location is False
            or include_unknown_salary is False
            or only_my_country
            or radius_scope not in ("50km", "50 km")
        )
        remote_explicitly_selected = bool(
            work_location
            and any(str(item).lower() == "remote" for item in work_location)
        )

        base_query: Dict[str, Any] = {}
        # Hard cap avoids pulling 1000+ full JSONB job docs into memory per request.
        feed_candidate_hard_cap = max(120, min(_env_int("JOBS_FEED_CANDIDATE_HARD_CAP", 350), 600))
        if prefetch:
            candidate_limit = min(max(40, requested_limit * 6), 120)
        else:
            candidate_limit = max(80, requested_limit * 16) if explicit_filters else max(120, requested_limit * 24)
        if explicit_location_filter and not prefetch:
            candidate_limit = max(candidate_limit, min(400, feed_candidate_hard_cap))
        if is_worldwide_radius and not prefetch:
            candidate_limit = max(candidate_limit, min(requested_limit * 40, feed_candidate_hard_cap))
        candidate_limit = min(candidate_limit, feed_candidate_hard_cap)

        refresh_results = []
        provider_search_keys: List[str] = []
        direct_refresh_jobs: List[Dict[str, Any]] = []

        def _is_current_provider_job(job: Dict[str, Any]) -> bool:
            return (
                bool(provider_search_keys)
                and job.get("provider") == "jsearch"
                and str(job.get("provider_search_key") or "") in provider_search_keys
            )

        def _compact_refresh_result(item: Dict[str, Any]) -> Dict[str, Any]:
            compact: Dict[str, Any] = {}
            for key, value in (item or {}).items():
                if key == "jobs" and isinstance(value, list):
                    compact["jobs_count"] = len(value)
                elif key == "sample_jobs" and isinstance(value, list):
                    compact["sample_jobs_count"] = len(value)
                elif key in ("greenhouse", "lever") and isinstance(value, dict):
                    compact[key] = _compact_refresh_result(value)
                else:
                    compact[key] = value
            return compact

        def _country_terms_from_locations() -> List[str]:
            aliases = {
                "fr": ["france"],
                "gb": ["united kingdom", "uk", "england"],
                "us": ["united states", "usa"],
                "ma": ["morocco", "maroc"],
            }
            values: List[str] = []
            for loc in selected_locations:
                code = str(loc.get("country_code") or "").lower().strip()
                country_name = str(loc.get("country") or "").lower().strip()
                if country_name:
                    values.append(country_name)
                values.extend(aliases.get(code, []))
            if only_my_country:
                profile_location_data = _profile_feed_location_data()
                profile_code = str(profile_location_data.get("country_code") or "").lower().strip()
                profile_country = str(profile_location_data.get("country") or "").lower().strip()
                if profile_country:
                    values.append(profile_country)
                values.extend(aliases.get(profile_code, []))
            return list(dict.fromkeys(value for value in values if value))

        def _city_terms_from_locations() -> List[str]:
            values: List[str] = []
            for loc in selected_locations:
                label = str(loc.get("location_label") or "").strip()
                if not label:
                    continue
                city_part = re.split(r"[,/|-]", label, maxsplit=1)[0]
                values.extend(_tokens(city_part))
            if not selected_locations and location:
                for label in location:
                    city_part = re.split(r"[,/|-]", str(label), maxsplit=1)[0]
                    values.extend(_tokens(city_part))
            return list(dict.fromkeys(values))

        selected_city_terms = _city_terms_from_locations()
        selected_country_terms = _country_terms_from_locations()

        def _country_codes_from_locations() -> List[str]:
            values: List[str] = []
            for loc in selected_locations:
                code = str(loc.get("country_code") or "").lower().strip()
                if code:
                    values.append(code)
                    continue
                label = str(loc.get("location_label") or loc.get("label") or "")
                if label and jobs_service_module._looks_like_france_location(label):
                    values.append("fr")
            if only_my_country:
                profile_location_data = _profile_feed_location_data()
                code = str(profile_location_data.get("country_code") or "").lower().strip()
                if code:
                    values.append(code)
            return list(dict.fromkeys(values))

        selected_country_codes = _country_codes_from_locations()
        # When searching France with France Travail configured, boost timeouts so
        # the OAuth + commune-code lookup + API pagination have enough headroom.
        _ft_configured = is_job_provider_configured("france_travail")
        _search_is_france = bool(selected_country_codes) and all(cc == "fr" for cc in selected_country_codes)
        if _search_is_france and _ft_configured:
            sync_refresh_max_pages = max(sync_refresh_max_pages, min(_env_int("FRANCE_TRAVAIL_MAX_PAGES", 2), 4))
            sync_refresh_page_size = max(sync_refresh_page_size, min(_env_int("FRANCE_TRAVAIL_PAGE_SIZE", 50), 150))
            sync_refresh_max_results = max(sync_refresh_max_results, min(_env_int("FRANCE_TRAVAIL_PAGE_SIZE", 50), 100))
            sync_refresh_attempts_per_city = 1
            sync_refresh_max_seconds = max(sync_refresh_max_seconds, _env_int("JOBS_FEED_SYNC_REFRESH_MAX_SECONDS", 4))
            sync_refresh_total_seconds = max(sync_refresh_total_seconds, _env_int("JOBS_FEED_SYNC_REFRESH_TOTAL_SECONDS", 4))
        location_intelligence_enabled = _env_bool("JOBS_LOCATION_INTELLIGENCE_ENABLED", True)
        location_max_expanded_cities = max(1, min(_env_int("JOBS_LOCATION_MAX_EXPANDED_CITIES", 10), 50))
        location_min_radius_km = max(1, _env_int("JOBS_LOCATION_MIN_RADIUS_KM", 10))
        location_include_cross_border = _env_bool("JOBS_LOCATION_INCLUDE_CROSS_BORDER", True) and not only_my_country
        location_min_population = max(0, _env_int("JOBS_LOCATION_MIN_POPULATION", 1000))

        async def _build_location_context() -> Dict[str, Any]:
            context: Dict[str, Any] = {
                "enabled": False,
                "used": False,
                "radius_km": radius_km,
                "origin_locations": [],
                "expanded_places": [],
                "city_terms": [],
                "country_codes": selected_country_codes,
                "include_cross_border": location_include_cross_border,
                "reason": None,
            }
            if not location_intelligence_enabled:
                context["reason"] = "disabled"
                return context
            if radius_km is None or radius_scope in ("worldwide", "remote", "remote/worldwide"):
                context["reason"] = "non_numeric_or_global_radius"
                return context
            if radius_km < location_min_radius_km:
                context["reason"] = "radius_below_minimum"
                return context
            if not selected_locations:
                context["reason"] = "no_selected_locations"
                return context

            context["enabled"] = True
            expanded: List[Dict[str, Any]] = []
            seen = set()
            per_location_cap = max(location_max_expanded_cities, 1)
            for loc in selected_locations[:5]:
                label = str(loc.get("location_label") or loc.get("label") or "").strip()
                country_hint = str(loc.get("country_code") or "").lower().strip() or None
                loc_lat = loc.get("lat")
                loc_lng = loc.get("lng")
                context["origin_locations"].append({
                    "location_label": label,
                    "country_code": country_hint,
                    "has_coordinates": loc_lat not in (None, "") and loc_lng not in (None, ""),
                })
                try:
                    places = await expand_location_radius(
                        location_label=label,
                        lat=loc_lat,
                        lng=loc_lng,
                        country_hint=country_hint,
                        radius_km=radius_km,
                        max_cities=per_location_cap,
                        include_cross_border=location_include_cross_border,
                        min_population=location_min_population,
                        db=db,
                    )
                except Exception as exc:
                    logger.warning(
                        "jobs/feed location_intelligence_failed label=%s radius_km=%s error=%s",
                        label,
                        radius_km,
                        f"{exc.__class__.__name__}: {str(exc)[:160]}",
                    )
                    continue
                for place in places:
                    key = f"{normalize_place_name(place.get('name') or place.get('ascii_name'))}:{str(place.get('country_code') or '').lower()}"
                    if not key or key in seen:
                        continue
                    seen.add(key)
                    expanded.append(place)
                    if len(expanded) >= location_max_expanded_cities:
                        break
                if len(expanded) >= location_max_expanded_cities:
                    break

            city_terms: List[str] = []
            country_codes: List[str] = []
            for place in expanded:
                names = [
                    place.get("name"),
                    place.get("ascii_name"),
                    place.get("normalized_name"),
                    *(place.get("alternate_names") or []),
                ]
                for name in names:
                    normalized = normalize_place_name(str(name or ""))
                    if normalized:
                        city_terms.append(normalized)
                code = str(place.get("country_code") or "").lower().strip()
                if code:
                    country_codes.append(code)
            context["expanded_places"] = expanded
            context["city_terms"] = list(dict.fromkeys(city_terms))
            context["country_codes"] = list(dict.fromkeys(country_codes or selected_country_codes))
            context["used"] = bool(expanded)
            context["reason"] = "expanded" if expanded else "no_places_found"
            return context

        location_context = await _build_location_context()
        expanded_city_terms = location_context.get("city_terms") or []
        expanded_country_codes = location_context.get("country_codes") or selected_country_codes
        debug_diagnostics = _env_bool("JOBS_FEED_DEBUG_DIAGNOSTICS", False)
        request_trace: Dict[str, Any] = {
            "received_query_params": {
                "limit": limit,
                "search_role": search_role,
                "search_radius": search_radius,
                "locations_json_present": bool(locations_json),
                "location": location,
                "work_location": work_location,
                "only_my_country": only_my_country,
                "include_unknown_location": include_unknown_location,
                "include_unknown_salary": include_unknown_salary,
                "posted_within": posted_within,
                "min_salary": min_salary,
                "force_provider_refresh": force_provider_refresh,
            },
            "parsed_locations_count": len(request_selected_locations),
            "parsed_locations": [
                {
                    "location_label": loc.get("location_label"),
                    "country_code": loc.get("country_code"),
                    "has_coordinates": loc.get("lat") not in (None, "") and loc.get("lng") not in (None, ""),
                }
                for loc in request_selected_locations[:10]
            ],
            "search_radius_raw": search_radius,
            "search_radius_km": radius_km,
            "explicit_local_intent": explicit_local_intent,
            "work_location_raw": work_location,
            "remote_explicitly_selected": remote_explicitly_selected,
            "include_unknown_location": include_unknown_location,
            "location_intelligence_enabled": location_context.get("enabled"),
            "location_intelligence_used": location_context.get("used"),
            "expanded_places": [
                {
                    "name": place.get("name"),
                    "country_code": place.get("country_code"),
                    "distance_km": place.get("distance_km"),
                }
                for place in (location_context.get("expanded_places") or [])[:20]
            ],
            "expanded_country_codes": expanded_country_codes,
            "db_local_candidate_count_before_filters": 0,
            "db_local_candidate_count_after_filters": 0,
            "manual_candidate_count": 0,
            "auto_apply_candidate_count": 0,
            "blocked_hidden_count": 0,
            "local_inventory_weak": False,
            "local_jsearch_discovery_should_run": False,
            "local_jsearch_discovery_attempted": False,
            "local_jsearch_skip_reason": None,
            "jsearch_queries_planned": [],
            "jsearch_queries_executed": [],
            "jsearch_attempts_per_city": sync_refresh_attempts_per_city,
            "jsearch_refresh_results": [],
            "jsearch_results_count": 0,
            "final_jobs_count": 0,
            "frontend_relevant_application_modes": [],
        }
        logger.info(
            "jobs/feed location_intelligence user_id=%s enabled=%s used=%s radius_km=%s origins=%s expanded=%s countries=%s cross_border=%s reason=%s",
            user.user_id,
            location_context.get("enabled"),
            location_context.get("used"),
            location_context.get("radius_km"),
            len(location_context.get("origin_locations") or []),
            len(location_context.get("expanded_places") or []),
            expanded_country_codes,
            location_context.get("include_cross_border"),
            location_context.get("reason"),
        )

        def _validated_cache_query(*, include_unknown: bool = False) -> Dict[str, Any]:
            query: Dict[str, Any] = {
                "applyability_tier": {"$in": ["A", "B", "C"] if include_unknown else ["A", "B"]},
            }
            if not include_unknown:
                query["validation_status"] = "valid"
            country_filter_codes = expanded_country_codes if location_context.get("used") else selected_country_codes
            if country_filter_codes and radius_scope not in ("worldwide", "remote", "remote/worldwide"):
                query["country_code"] = {"$in": country_filter_codes}
            if posted_within and posted_within != "any":
                days_map = {"1d": 1, "7d": 7, "30d": 30}
                days = days_map.get(posted_within)
                if days:
                    query["posted_at"] = {"$gte": (datetime.now(timezone.utc) - timedelta(days=days)).isoformat()}
            if min_salary and min_salary > 0 and not include_unknown_salary:
                query["salary_max"] = {"$gte": min_salary}
            if work_location:
                wanted = {
                    ("onsite" if str(item).lower() in ("in-person", "in_person", "on-site", "onsite") else str(item).lower())
                    for item in work_location
                }
                if wanted == {"remote"}:
                    query["remote"] = True
            return query

        def _tier_rank(job: Dict[str, Any]) -> int:
            return {"A": 0, "B": 1, "C": 2}.get(str(job.get("applyability_tier") or "").upper(), 9)

        def _quality_sort_key(job: Dict[str, Any]) -> tuple:
            score_value = job.get("applyability_score")
            try:
                applyability_score = float(score_value if score_value is not None else 0)
            except (TypeError, ValueError):
                applyability_score = 0.0
            return (
                _tier_rank(job),
                -applyability_score,
                0 if job.get("auto_apply_supported") is True else 1,
                0 if job.get("manual_fulfillment_ready") is True else 1,
                -_recency_score(job),
            )

        def _job_work_location(job: Dict[str, Any]) -> str:
            raw = job.get("remote")
            if isinstance(raw, bool):
                return "remote" if raw else "onsite"
            text = " ".join([
                str(raw or ""),
                str(job.get("location") or ""),
                str(job.get("workplace_type") or ""),
                str(job.get("work_location") or ""),
            ]).lower()
            if "hybrid" in text:
                return "hybrid"
            if "remote" in text or "work from home" in text:
                return "remote"
            if "onsite" in text or "on-site" in text or "in person" in text or "office" in text:
                return "onsite"
            return "unknown"

        extra_country_name_to_code = {
            "ireland": "ie",
            "india": "in",
            "brazil": "br",
            "poland": "pl",
            "canada": "ca",
            "australia": "au",
            "singapore": "sg",
            "switzerland": "ch",
            "belgium": "be",
            "austria": "at",
            "denmark": "dk",
            "sweden": "se",
            "norway": "no",
            "finland": "fi",
            "czech republic": "cz",
            "czechia": "cz",
            "romania": "ro",
            "hungary": "hu",
            "mexico": "mx",
            "argentina": "ar",
            "chile": "cl",
            "colombia": "co",
            "japan": "jp",
            "south korea": "kr",
            "korea": "kr",
            "china": "cn",
            "hong kong": "hk",
            "united arab emirates": "ae",
            "uae": "ae",
        }
        country_name_to_code = {**COUNTRY_NAME_TO_CODE, **extra_country_name_to_code}

        def _visible_job_location_text(job: Dict[str, Any]) -> str:
            return normalize_place_name(" ".join([
                str(job.get("city") or ""),
                str(job.get("region") or ""),
                str(job.get("location") or ""),
            ]))

        def _raw_job_location_text(job: Dict[str, Any]) -> str:
            data = job.get("data") if isinstance(job.get("data"), dict) else {}
            raw_parts = [
                data.get("location"),
                data.get("job_city"),
                data.get("job_state"),
                data.get("job_country"),
                data.get("job_location"),
            ]
            return normalize_place_name(" ".join(str(part or "") for part in raw_parts))

        def _normalized_job_location_text(job: Dict[str, Any]) -> str:
            visible_text = _visible_job_location_text(job)
            return visible_text or _raw_job_location_text(job)

        def _country_code_from_location_text(text: str) -> Optional[str]:
            normalized = normalize_place_name(text)
            if not normalized:
                return None
            padded = f" {normalized} "
            for country_name, code in sorted(country_name_to_code.items(), key=lambda item: len(item[0]), reverse=True):
                if f" {normalize_place_name(country_name)} " in padded:
                    return code
            return None

        def _known_outside_expanded_country(job: Dict[str, Any]) -> bool:
            if not explicit_local_intent:
                return False
            allowed = {str(code).lower().strip() for code in expanded_country_codes if code}
            if not allowed:
                return False
            job_country_code = str(job.get("country_code") or "").lower().strip()
            if job_country_code:
                return job_country_code not in allowed
            visible_country_code = _country_code_from_location_text(_visible_job_location_text(job))
            return bool(visible_country_code and visible_country_code not in allowed)

        def _expanded_location_match(job: Dict[str, Any]) -> Optional[Dict[str, Any]]:
            if not location_context.get("used"):
                return None
            if _known_outside_expanded_country(job):
                return None
            visible_text = _visible_job_location_text(job)
            job_text = visible_text or _raw_job_location_text(job)
            job_country_code = str(job.get("country_code") or "").lower().strip()
            if not job_text and not job_country_code:
                return None
            for place in location_context.get("expanded_places") or []:
                place_country = str(place.get("country_code") or "").lower().strip()
                if place_country and job_country_code and place_country != job_country_code:
                    continue
                names = [
                    place.get("name"),
                    place.get("ascii_name"),
                    place.get("normalized_name"),
                    *(place.get("alternate_names") or []),
                ]
                normalized_names = [normalize_place_name(str(name or "")) for name in names]
                if any(name and name in job_text for name in normalized_names):
                    return place
            selected_terms = [term for term in selected_city_terms if len(term) >= 4]
            if selected_terms and any(term in job_text for term in selected_terms):
                places = location_context.get("expanded_places") or []
                return places[0] if places else {"name": "selected_location", "distance_km": 0, "is_origin": True}
            return None

        def _matches_work_location(job: Dict[str, Any]) -> bool:
            if not work_location:
                return True
            wanted = {
                ("onsite" if str(item).lower() in ("in-person", "in_person", "on-site", "onsite") else str(item).lower())
                for item in work_location
            }
            actual = _job_work_location(job)
            if actual == "unknown":
                return include_unknown_location
            if "remote" in wanted and actual == "remote":
                return True
            if "hybrid" in wanted and actual == "hybrid":
                return True
            if "onsite" in wanted and actual == "onsite":
                return True
            return False

        def _explicit_local_remote_allowed(job: Dict[str, Any]) -> bool:
            if _job_work_location(job) != "remote":
                return False
            if not work_location:
                return False
            wanted = {
                ("onsite" if str(item).lower() in ("in-person", "in_person", "on-site", "onsite") else str(item).lower())
                for item in work_location
            }
            return "remote" in wanted

        def _matches_location(job: Dict[str, Any]) -> bool:
            if not explicit_location_filter and not only_my_country:
                return True
            normalized_job_location = _normalized_job_location_text(job)
            job_country_code = str(job.get("country_code") or "").lower().strip()
            if not normalized_job_location and not job_country_code:
                return include_unknown_location
            if explicit_local_intent and _explicit_local_remote_allowed(job):
                return True
            if explicit_local_intent and _known_outside_expanded_country(job):
                return False
            city_match = bool(
                selected_city_terms
                and any(term in normalized_job_location for term in selected_city_terms)
            )
            if city_match:
                return True
            expanded_match = _expanded_location_match(job)
            if expanded_match:
                return True
            job_location = str(job.get("location") or "").lower()
            country_match = bool(
                (selected_country_terms and any(term in job_location for term in selected_country_terms))
                or any(str(loc.get("country_code") or "").lower().strip() == job_country_code for loc in selected_locations if loc.get("country_code"))
            )
            if only_my_country:
                profile_location_data = _profile_feed_location_data()
                profile_code = str(profile_location_data.get("country_code") or "").lower().strip()
                if profile_code and job_country_code == profile_code:
                    return True
                return country_match
            if radius_scope in ("worldwide", "remote", "remote/worldwide"):
                return True
            if radius_km is not None:
                if location_context.get("used"):
                    return False
                # A specific city with a radius must never widen to the whole
                # country: "Paris + 50km" should not surface Lyon jobs.
                if explicit_local_intent and selected_city_terms:
                    return False
                return country_match if selected_country_terms else False
            return country_match

        def _passes_explicit_local_hard_constraint(job: Dict[str, Any]) -> bool:
            if not explicit_local_intent:
                return True
            has_known_location = bool(job.get("location") or job.get("city") or job.get("region") or job.get("country_code"))
            if not has_known_location:
                return False
            if _explicit_local_remote_allowed(job):
                return True
            if _known_outside_expanded_country(job):
                return False
            return _matches_location(job)

        def _matches_salary(job: Dict[str, Any]) -> bool:
            if not min_salary or min_salary <= 0:
                return True
            salary_max = job.get("salary_max")
            if salary_max in (None, ""):
                return include_unknown_salary
            try:
                return int(salary_max) >= int(min_salary)
            except (TypeError, ValueError):
                return include_unknown_salary

        def _matches_posted_date(job: Dict[str, Any]) -> bool:
            if not posted_within or posted_within == "any":
                return True
            days_map = {"1d": 1, "7d": 7, "30d": 30}
            days = days_map.get(posted_within)
            if not days:
                return True
            raw = job.get("posted_at") or job.get("imported_at") or job.get("last_seen_at")
            if not raw:
                return False
            try:
                parsed = datetime.fromisoformat(str(raw).replace("Z", "+00:00"))
                if parsed.tzinfo is None:
                    parsed = parsed.replace(tzinfo=timezone.utc)
            except ValueError:
                return False
            return parsed >= datetime.now(timezone.utc) - timedelta(days=days)

        def _matches_job_type(job: Dict[str, Any]) -> bool:
            return job_matches_job_types(job, effective_job_types)

        def _matches_experience(job: Dict[str, Any]) -> bool:
            if not experience:
                return True
            exp_map = {
                "entry": ["junior", "entry"],
                "mid": ["mid", "intermediate"],
                "senior": ["senior"],
                "executive": ["lead", "principal", "executive", "director"],
            }
            wanted: List[str] = []
            for item in experience:
                wanted.extend(exp_map.get(item, [item]))
            seniority = str(job.get("seniority") or "").lower()
            title = str(job.get("title") or "").lower()
            return any(item in seniority or item in title for item in wanted)

        def _matches_company(job: Dict[str, Any]) -> bool:
            company = str(job.get("company") or "").lower()
            if only_company and not any(company == str(item).lower() for item in only_company):
                return False
            if hide_company and any(str(item).lower() in company for item in hide_company):
                return False
            return True

        def _matches_industry(job: Dict[str, Any]) -> bool:
            if not only_industry and not hide_industry:
                return True
            text = " ".join([
                str(job.get("industry") or ""),
                " ".join(str(item) for item in (job.get("industries") or [])),
                str(job.get("company") or ""),
                str(job.get("description") or ""),
                str(job.get("clean_description") or ""),
            ]).lower()
            if only_industry and not any(str(item).lower() in text for item in only_industry):
                return False
            if hide_industry and any(str(item).lower() in text for item in hide_industry):
                return False
            return True

        def _matches_explicit_filters(job: Dict[str, Any]) -> bool:
            return (
                _matches_work_location(job)
                and _matches_location(job)
                and _matches_salary(job)
                and _matches_posted_date(job)
                and _matches_job_type(job)
                and _matches_experience(job)
                and _matches_company(job)
                and _matches_industry(job)
            )

        def _matches_non_location_filters(job: Dict[str, Any]) -> bool:
            return (
                _matches_work_location(job)
                and _matches_salary(job)
                and _matches_posted_date(job)
                and _matches_job_type(job)
                and _matches_experience(job)
                and _matches_company(job)
                and _matches_industry(job)
            )

        def _role_compatible_for_threshold(job: Dict[str, Any]) -> bool:
            if not _category_compatible(job):
                return False
            if not strict_tokens and not family_tokens:
                return True
            return _role_score(job, strict_tokens, family_tokens) > 0

        swiped_rows = await db.swipes.find({"user_id": user.user_id}, {"_id": 0, "job_id": 1}).limit(500).to_list(500)
        swiped_ids = {row.get("job_id") for row in swiped_rows if row.get("job_id")}

        async def _load_db_candidates(*, include_unknown: bool = False) -> tuple[List[Dict[str, Any]], Dict[str, Any], int, int, int]:
            query = _validated_cache_query(include_unknown=include_unknown)
            db_started = time.perf_counter()
            rows = await _get_feed_job_candidates(query, candidate_limit)
            if include_unknown_location and "country_code" in query:
                unknown_query = {key: value for key, value in query.items() if key != "country_code"}
                unknown_rows = await _get_feed_job_candidates(unknown_query, min(candidate_limit, 500))
                seen_ids = {row.get("job_id") for row in rows}
                for row in unknown_rows:
                    has_location = bool(row.get("location") or row.get("city") or row.get("region") or row.get("country_code"))
                    if not has_location and row.get("job_id") not in seen_ids:
                        rows.append(row)
                        seen_ids.add(row.get("job_id"))
            db_elapsed_ms = int((time.perf_counter() - db_started) * 1000)
            rows = [job for job in rows if job.get("job_id") not in swiped_ids]
            unfiltered = len(rows)
            rejected = sum(
                1
                for job in rows
                if str(job.get("validation_status") or "").lower() == "invalid"
                or str(job.get("applyability_tier") or "").upper() in {"D", "E"}
            )
            rows = [_with_apply_fulfillment_fields(job) for job in rows if _job_is_applyable(job)]
            applyable = len(rows)
            rows = [job for job in rows if _matches_explicit_filters(job)]
            rows.sort(key=_quality_sort_key)
            logger.info(
                "jobs/feed db_query_complete: user_id=%s elapsed_ms=%s include_unknown=%s query=%s raw_rows=%s applyable=%s filtered=%s",
                user.user_id,
                db_elapsed_ms,
                include_unknown,
                query,
                unfiltered,
                applyable,
                len(rows),
            )
            return rows, query, unfiltered, applyable, rejected

        async def _load_legacy_direct_ats_candidates() -> tuple[List[Dict[str, Any]], Dict[str, Any], int, int, int]:
            query: Dict[str, Any] = {
                "provider": {"$in": ats_supported},
                "auto_apply_supported": True,
            }
            db_started = time.perf_counter()
            rows = await _get_feed_job_candidates(query, candidate_limit)
            db_elapsed_ms = int((time.perf_counter() - db_started) * 1000)
            rows = [job for job in rows if job.get("job_id") not in swiped_ids]
            unfiltered = len(rows)
            rejected = sum(
                1
                for job in rows
                if str(job.get("validation_status") or "").lower() == "invalid"
                or str(job.get("applyability_tier") or "").upper() in {"D", "E"}
            )
            rows = [
                _with_apply_fulfillment_fields(job)
                for job in rows
                if not job.get("validation_status")
                and str(job.get("applyability_tier") or "").upper() not in {"D", "E"}
                and str(job.get("provider") or "").lower() in set(ats_supported)
                and _job_is_applyable(job)
            ]
            applyable = len(rows)
            rows = [job for job in rows if _matches_explicit_filters(job)]
            rows.sort(key=_quality_sort_key)
            logger.info(
                "jobs/feed legacy_direct_ats_query_complete: user_id=%s elapsed_ms=%s query=%s raw_rows=%s applyable=%s filtered=%s",
                user.user_id,
                db_elapsed_ms,
                query,
                unfiltered,
                applyable,
                len(rows),
            )
            return rows, query, unfiltered, applyable, rejected

        async def _load_local_visible_candidates() -> tuple[List[Dict[str, Any]], Dict[str, Any], int, int, int]:
            country_filter_codes = expanded_country_codes if location_context.get("used") else selected_country_codes
            query: Dict[str, Any] = {}
            if country_filter_codes and radius_scope not in ("worldwide", "remote", "remote/worldwide"):
                query["country_code"] = {"$in": country_filter_codes}
            db_started = time.perf_counter()
            rows = await _get_feed_job_candidates(query, candidate_limit)
            # Avoid unbounded find({}) when country-filtered inventory is already usable.
            if explicit_local_intent and len(rows) < max(20, requested_limit * 2):
                broad_limit = min(candidate_limit, 150)
                broad_rows = await _get_feed_job_candidates({}, broad_limit)
                seen_ids = {row.get("job_id") for row in rows}
                for row in broad_rows:
                    job_id = row.get("job_id")
                    if job_id and job_id not in seen_ids:
                        rows.append(row)
                        seen_ids.add(job_id)
            db_elapsed_ms = int((time.perf_counter() - db_started) * 1000)
            rows = [job for job in rows if job.get("job_id") not in swiped_ids]
            unfiltered = len(rows)
            rejected = sum(1 for job in rows if _job_is_blocked_for_feed(job))
            rows = [
                _with_application_capability_fields(job)
                for job in rows
                if not _job_is_blocked_for_feed(job) and _job_has_usable_apply_url(job)
            ]
            visible = len(rows)
            rows = [job for job in rows if _matches_explicit_filters(job)]
            rows.sort(key=_quality_sort_key)
            logger.info(
                "jobs/feed local_visible_query_complete: user_id=%s elapsed_ms=%s query=%s raw_rows=%s visible=%s filtered=%s rejected=%s",
                user.user_id,
                db_elapsed_ms,
                query,
                unfiltered,
                visible,
                len(rows),
                rejected,
            )
            return rows, query, unfiltered, visible, rejected

        base_query = _validated_cache_query(include_unknown=False)
        candidates: List[Dict[str, Any]] = []
        unfiltered_count = 0
        applyable_count = 0
        local_visible_count = 0
        local_manual_count = 0
        local_blocked_hidden_count = 0
        pre_filter_candidates: List[Dict[str, Any]] = []
        db_good_count = 0
        db_rejected_count = 0
        feed_source = "db_first_disabled"
        jsearch_fallback_triggered = False

        if db_first_enabled:
            candidates, base_query, unfiltered_count, applyable_count, db_rejected_count = await _load_db_candidates(include_unknown=False)
            db_good_count = sum(1 for job in candidates if _role_compatible_for_threshold(job))
            pre_filter_candidates = list(candidates)
            feed_source = "db_cache"
            logger.info(
                "jobs/feed db_first: user_id=%s elapsed_ms=%s good_count=%s candidate_count=%s unfiltered_count=%s applyable_count=%s rejected_de_count=%s query=%s min_good=%s weak_threshold=%s allow_unknown=%s",
                user.user_id,
                _elapsed_ms(),
                db_good_count,
                len(candidates),
                unfiltered_count,
                applyable_count,
                db_rejected_count,
                base_query,
                db_min_good_results,
                db_weak_results_threshold,
                allow_unknown_tier,
            )
            if allow_unknown_tier and db_good_count < requested_limit:
                c_candidates, c_query, c_unfiltered, c_applyable, c_rejected = await _load_db_candidates(include_unknown=True)
                seen_ids = {job.get("job_id") for job in candidates}
                candidates.extend(job for job in c_candidates if job.get("job_id") not in seen_ids)
                pre_filter_candidates = list(candidates)
                unfiltered_count += c_unfiltered
                applyable_count += c_applyable
                db_rejected_count += c_rejected
                base_query = c_query
                logger.info(
                    "jobs/feed db_first_unknown_tier: user_id=%s elapsed_ms=%s combined_count=%s c_unfiltered=%s",
                    user.user_id,
                    _elapsed_ms(),
                    len(candidates),
                    c_unfiltered,
                )
            if db_good_count == 0 and not allow_unknown_tier:
                legacy_candidates, legacy_query, legacy_unfiltered, legacy_applyable, legacy_rejected = await _load_legacy_direct_ats_candidates()
                if legacy_candidates:
                    candidates = legacy_candidates
                    base_query = legacy_query
                    unfiltered_count = legacy_unfiltered
                    applyable_count = legacy_applyable
                    db_rejected_count += legacy_rejected
                    db_good_count = sum(1 for job in candidates if _role_compatible_for_threshold(job))
                    pre_filter_candidates = list(candidates)
                    feed_source = "legacy_direct_ats_cache"
                    logger.info(
                        "jobs/feed legacy_direct_ats_used: user_id=%s good_count=%s candidate_count=%s",
                        user.user_id,
                        db_good_count,
                        len(candidates),
                    )
            if explicit_local_intent and len(candidates) < requested_limit:
                visible_candidates, visible_query, visible_unfiltered, visible_applyable, visible_rejected = await _load_local_visible_candidates()
                seen_ids = {job.get("job_id") for job in candidates}
                additions = [job for job in visible_candidates if job.get("job_id") not in seen_ids]
                if additions:
                    candidates.extend(additions)
                    pre_filter_candidates = list(candidates)
                    base_query = visible_query if not base_query else base_query
                    feed_source = f"{feed_source}+local_visible"
                unfiltered_count += visible_unfiltered
                applyable_count += visible_applyable
                db_rejected_count += visible_rejected
                local_visible_count += len(visible_candidates)
                local_manual_count += sum(1 for job in visible_candidates if job.get("application_mode") != "auto_apply")
                local_blocked_hidden_count += visible_rejected
                logger.info(
                    "jobs/feed local_visible_used: user_id=%s visible_count=%s added=%s manual_count=%s",
                    user.user_id,
                    len(visible_candidates),
                    len(additions),
                    local_manual_count,
                )

        pre_refresh_candidate_count = len(candidates)
        if explicit_local_intent and candidates:
            candidates = [job for job in candidates if _passes_explicit_local_hard_constraint(job)]
            if len(candidates) != pre_refresh_candidate_count:
                db_good_count = sum(1 for job in candidates if _role_compatible_for_threshold(job))
                pre_filter_candidates = list(candidates)
                logger.info(
                    "jobs/feed explicit_local_pre_refresh_constraint: user_id=%s before=%s after=%s db_good_count=%s",
                    user.user_id,
                    pre_refresh_candidate_count,
                    len(candidates),
                    db_good_count,
                )
            if candidates:
                hydrated_candidates = await _hydrate_feed_jobs(candidates)
                hydrated_candidates = [_with_application_capability_fields(job) for job in hydrated_candidates]
                hydrated_candidates = [job for job in hydrated_candidates if _passes_explicit_local_hard_constraint(job)]
                if len(hydrated_candidates) != len(candidates):
                    logger.info(
                        "jobs/feed explicit_local_pre_refresh_hydrated_constraint: user_id=%s before=%s after=%s",
                        user.user_id,
                        len(candidates),
                        len(hydrated_candidates),
                    )
                candidates = hydrated_candidates
                db_good_count = sum(1 for job in candidates if _role_compatible_for_threshold(job))
                pre_filter_candidates = list(candidates)
        local_inventory_count = (
            sum(1 for job in candidates if _role_compatible_for_threshold(job))
            if explicit_local_intent
            else len(candidates)
        )

        request_trace["db_local_candidate_count_before_filters"] = pre_refresh_candidate_count
        request_trace["db_local_candidate_count_after_filters"] = len(candidates)
        request_trace["local_role_compatible_count"] = local_inventory_count
        request_trace["manual_candidate_count"] = sum(1 for job in candidates if job.get("application_mode") in {"manual", "assisted"})
        request_trace["auto_apply_candidate_count"] = sum(1 for job in candidates if job.get("application_mode") == "auto_apply")
        request_trace["blocked_hidden_count"] = local_blocked_hidden_count or db_rejected_count
        request_trace["local_inventory_weak"] = bool(explicit_local_intent and local_inventory_count < requested_limit)

        def _explicit_local_cooldown_key() -> str:
            location_parts: List[str] = []
            if location_context.get("expanded_places"):
                for place in (location_context.get("expanded_places") or [])[:6]:
                    name = normalize_place_name(str(place.get("name") or place.get("ascii_name") or ""))
                    country_value = str(place.get("country_code") or "").lower().strip()
                    if name:
                        location_parts.append(f"{name}:{country_value}")
            if not location_parts:
                for loc in request_selected_locations[:5]:
                    label = normalize_place_name(str(loc.get("location_label") or loc.get("label") or ""))
                    country_value = str(loc.get("country_code") or loc.get("country") or "").lower().strip()
                    if label:
                        location_parts.append(f"{label}:{country_value}")
            role_value = normalize_place_name(target_role or "")
            return "|".join([
                "explicit_local",
                role_value,
                radius_scope,
                ",".join(location_parts) or "unknown_location",
            ])

        def _prune_feed_cooldowns(now_value: float) -> None:
            expired = [key for key, until in _feed_sync_refresh_cooldowns.items() if until <= now_value]
            for key in expired:
                _feed_sync_refresh_cooldowns.pop(key, None)

        cooldown_enabled = _env_bool("JOBS_FEED_FALLBACK_COOLDOWN_ENABLED", True)
        cooldown_now = time.monotonic()
        _prune_feed_cooldowns(cooldown_now)
        cooldown_key = _explicit_local_cooldown_key() if explicit_local_intent else "global"
        cooldown_until = (
            _feed_sync_refresh_cooldowns.get(cooldown_key, 0.0)
            if explicit_local_intent
            else _feed_sync_refresh_cooldown_until
        )
        # Cards we can serve right now without waiting on providers
        # (role-compatible only, so we never serve irrelevant cards instead of refreshing).
        has_showable_cards = (
            local_inventory_count > 0 if explicit_local_intent else db_good_count > 0
        )
        # A prior timeout/error cooldown only suppresses *further* refresh attempts
        # while we already have something reasonable cached to show. It must never
        # suppress the one attempt that stands between the user and an incorrect
        # "no jobs found" response -- that check has to run every time.
        cooldown_active = (
            cooldown_enabled
            and cooldown_now < cooldown_until
            and not (force_provider_refresh and explicit_local_intent)
            and has_showable_cards
        )
        request_trace["feed_provider_cooldown_key"] = cooldown_key
        request_trace["feed_provider_cooldown_active"] = bool(cooldown_active)
        request_trace["feed_provider_cooldown_remaining_seconds"] = max(0, int(cooldown_until - cooldown_now))
        explicit_local_after_budget_refresh = (
            explicit_local_intent
            and local_inventory_count < requested_limit
            and _env_bool("JOBS_FEED_EXPLICIT_LOCAL_DISCOVERY_AFTER_DB_TIMEOUT", True)
        )
        refresh_budget_available = not _timed_out() or explicit_local_after_budget_refresh
        inventory_is_weak = (
            not db_first_enabled
            or db_good_count == 0
            or (explicit_local_intent and local_inventory_count < requested_limit)
        )
        explicit_local_empty_inventory = bool(explicit_local_intent and local_inventory_count == 0)
        if explicit_local_empty_inventory and not force_provider_refresh:
            # A brand-new location with zero cached jobs relies entirely on
            # this one live call succeeding -- same JSearch latency reality
            # as the general sync_refresh budget above, so it gets the same
            # bump (was 5s, too tight for the same reason).
            explicit_local_sync_seconds = max(1, min(_env_int("JOBS_FEED_EXPLICIT_LOCAL_SYNC_SECONDS", 30), 30))
            sync_refresh_total_seconds = explicit_local_sync_seconds
            sync_refresh_max_seconds = explicit_local_sync_seconds
            sync_refresh_attempts_per_city = 1
            request_trace["explicit_local_quick_sync_seconds"] = explicit_local_sync_seconds
        needs_provider_refresh = (
            sync_refresh_enabled
            and not cooldown_active
            and ((force_provider_refresh and explicit_local_intent) or inventory_is_weak)
        )
        # Blocking (sync) refresh when worldwide DB is empty, client forced refresh,
        # explicit local search has zero local inventory (quick API lookup, ~5s cap),
        # or (per product decision) local A/B-tier inventory is merely thin -- the
        # user should see a real live-refresh attempt in this same response instead
        # of a "check back later" background task, since that background task's
        # own results were getting silently discarded by the same tier filter
        # anyway (confirmed live: it was importing relevant jobs, just all tier C).
        should_refresh = (
            not prefetch
            and needs_provider_refresh
            and (refresh_budget_available or (force_provider_refresh and explicit_local_intent))
            and (
                (force_provider_refresh and explicit_local_intent)
                or (not has_showable_cards and not explicit_local_intent)
                or explicit_local_empty_inventory
                or (explicit_local_intent and inventory_is_weak)
            )
        )
        background_refresh_wanted = (
            needs_provider_refresh
            and not should_refresh
            and not force_provider_refresh
            and not explicit_local_empty_inventory
            and _env_bool("JOBS_FEED_BACKGROUND_REFRESH_ENABLED", True)
        )
        if prefetch and explicit_local_intent and local_inventory_count < requested_limit and not has_showable_cards:
            should_refresh = True
            background_refresh_wanted = False
            request_trace["prefetch"] = True
            request_trace["prefetch_forced_local_discovery"] = True
        elif prefetch:
            request_trace["prefetch"] = True
            request_trace["local_jsearch_skip_reason"] = "prefetch_db_only"
        if audit_mode:
            should_refresh = False
            background_refresh_wanted = False
            request_trace["local_jsearch_skip_reason"] = "coverage_audit_mode"
        request_trace["local_jsearch_discovery_should_run"] = bool(should_refresh and explicit_local_intent)
        if explicit_local_intent and not should_refresh:
            if not sync_refresh_enabled:
                request_trace["local_jsearch_skip_reason"] = "sync_refresh_disabled"
            elif _timed_out() and not explicit_local_after_budget_refresh:
                request_trace["local_jsearch_skip_reason"] = "feed_timeout_budget_exhausted"
            elif cooldown_active:
                request_trace["local_jsearch_skip_reason"] = "provider_cooldown_active"
            elif local_inventory_count >= requested_limit:
                request_trace["local_jsearch_skip_reason"] = "local_candidate_count_sufficient"
            elif db_first_enabled and db_good_count > 0:
                request_trace["local_jsearch_skip_reason"] = "db_good_count_nonzero"
            else:
                request_trace["local_jsearch_skip_reason"] = "unknown_skip_condition"
        if db_good_count > 0 and db_good_count < db_weak_results_threshold:
            logger.info(
                "jobs/feed sync_refresh_skipped_db_has_results: user_id=%s db_good_count=%s weak_threshold=%s",
                user.user_id,
                db_good_count,
                db_weak_results_threshold,
            )
        elif cooldown_active:
            logger.info(
                "jobs/feed sync_refresh_skipped_cooldown: user_id=%s cooldown_key=%s cooldown_remaining_seconds=%s",
                user.user_id,
                cooldown_key,
                max(0, int(cooldown_until - time.monotonic())),
            )
        def _compute_refresh_locations() -> List[Any]:
            max_refresh_locations = max(1, min(int(os.environ.get("FEED_MAX_REFRESH_LOCATIONS", "1")), 3))
            if explicit_local_intent:
                max_refresh_locations = max(1, min(_env_int("JOBS_FEED_LOCAL_DISCOVERY_MAX_CITIES", 3), 4))
            if is_worldwide_radius:
                return [None]
            # Empty local inventory: query the user's city first (fast path for Avignon, etc.).
            if explicit_local_empty_inventory and selected_locations and not force_provider_refresh:
                origin_cap = max(1, min(_env_int("JOBS_FEED_EXPLICIT_LOCAL_SYNC_MAX_CITIES", 1), 2))
                return selected_locations[:origin_cap]
            if explicit_local_intent and location_context.get("expanded_places"):
                def _provider_place_sort_key(place: Dict[str, Any]) -> tuple:
                    population = int(place.get("population") or 0)
                    distance = float(place.get("distance_km") or 0)
                    is_origin = bool(place.get("is_origin"))
                    origin_penalty = 1 if is_origin and population < 10000 else 0
                    return (origin_penalty, -population, distance)

                expanded_for_provider = sorted(
                    location_context.get("expanded_places") or [],
                    key=_provider_place_sort_key,
                )
                return [
                    {
                        "location_label": place.get("name") or place.get("ascii_name") or place.get("query_label"),
                        "country_code": str(place.get("country_code") or "").lower().strip(),
                        "lat": place.get("latitude"),
                        "lng": place.get("longitude"),
                    }
                    for place in expanded_for_provider
                    if place.get("name") or place.get("ascii_name")
                ][:max_refresh_locations]
            return (selected_locations or [None])[:max_refresh_locations]

        if background_refresh_wanted:
            bg_signature = (
                _explicit_local_cooldown_key()
                if explicit_local_intent
                else "|".join(["global", normalize_place_name(target_role or ""), radius_scope])
            )
            bg_scheduled = schedule_feed_background_refresh(
                bg_signature,
                profile_for_refresh,
                _compute_refresh_locations(),
                search_radius=search_radius,
                role_override=target_role,
                requested_limit=requested_limit,
                max_results=sync_refresh_max_results,
                max_pages=sync_refresh_max_pages,
                page_size=sync_refresh_page_size,
                attempts_per_city=sync_refresh_attempts_per_city,
                user_id=user.user_id,
            )
            request_trace["background_refresh_scheduled"] = bool(bg_scheduled)
            if bg_scheduled:
                request_trace["local_jsearch_skip_reason"] = "background_refresh_scheduled"
                logger.info(
                    "jobs/feed background_refresh_scheduled: user_id=%s signature=%s inventory=%s requested=%s",
                    user.user_id,
                    bg_signature,
                    local_inventory_count,
                    requested_limit,
                )

        if should_refresh:
            jsearch_fallback_triggered = True
            refresh_locations = _compute_refresh_locations()
            request_trace["jsearch_queries_planned"] = [
                {
                    "location_label": loc.get("location_label") if isinstance(loc, dict) else None,
                    "country_code": loc.get("country_code") if isinstance(loc, dict) else None,
                }
                for loc in refresh_locations
            ]
            provider_force_refresh = force_provider_refresh or os.environ.get("JOB_FEED_ON_DEMAND_JSEARCH", "true").lower() in ("1", "true", "yes", "on")
            logger.info(
                "jobs/feed jsearch_fallback_start: user_id=%s db_good_count=%s weak_threshold=%s refresh_locations=%s force_provider_refresh=%s max_seconds=%s total_seconds=%s max_results=%s max_pages=%s page_size=%s attempts_per_city=%s",
                user.user_id,
                db_good_count,
                db_weak_results_threshold,
                len(refresh_locations),
                provider_force_refresh,
                sync_refresh_max_seconds,
                sync_refresh_total_seconds,
                sync_refresh_max_results,
                sync_refresh_max_pages,
                sync_refresh_page_size,
                sync_refresh_attempts_per_city,
            )
            provider_refresh_deadline = time.perf_counter() + sync_refresh_total_seconds

            # Fire France Travail concurrently with the JSearch loop below instead
            # of only as JSearch's own internal last-resort (which only tries FT
            # once a given location's JSearch attempt found zero relevant jobs).
            # Product decision: both providers should be tried at once whenever
            # local A/B inventory is thin, not FT gated behind JSearch failing
            # first. Skipped entirely when France Travail is already the primary
            # provider (its main loop iteration below already hits it directly).
            france_travail_parallel_task = None
            if primary_job_provider_name() != "france_travail":
                primary_refresh_loc = refresh_locations[0] if refresh_locations and isinstance(refresh_locations[0], dict) else None
                ft_query = build_profile_job_query(
                    profile_for_refresh,
                    location_override=primary_refresh_loc.get("location_label") if primary_refresh_loc else None,
                    location_data_override=primary_refresh_loc,
                    search_radius=search_radius,
                    role_override=target_role,
                )
                france_travail_parallel_task = asyncio.create_task(
                    jobs_service_module._attempt_france_travail_fallback(
                        db, ft_query, search_radius, primary_job_provider_name(),
                    )
                )

            for loc_data in refresh_locations:
                remaining_seconds = provider_refresh_deadline - time.perf_counter()
                if remaining_seconds <= 0.5:
                    request_trace["local_jsearch_budget_exhausted"] = True
                    logger.info(
                        "jobs/feed jsearch_fallback_budget_exhausted: user_id=%s attempted_locations=%s total_seconds=%s",
                        user.user_id,
                        len(refresh_results),
                        sync_refresh_total_seconds,
                    )
                    break
                loc_label = loc_data.get("location_label") if isinstance(loc_data, dict) else None
                request_trace["jsearch_queries_executed"].append({
                    "location_label": loc_label,
                    "country_code": loc_data.get("country_code") if isinstance(loc_data, dict) else None,
                })
                refresh_started = time.perf_counter()
                try:
                    per_location_timeout = max(0.5, min(float(sync_refresh_max_seconds), remaining_seconds))
                    refresh_result = await asyncio.wait_for(
                        refresh_jobs_for_profile_if_needed(
                            db,
                            profile_for_refresh,
                            require_auto_apply=False,
                            target_auto_apply_count=min(sync_refresh_max_results, max(requested_limit, 1)),
                            location_override=loc_label,
                            location_data_override=loc_data if isinstance(loc_data, dict) else None,
                            search_radius=search_radius,
                            role_override=target_role,
                            force_provider_refresh=provider_force_refresh,
                            query_limit_override=sync_refresh_max_results,
                            provider_max_pages=sync_refresh_max_pages,
                            provider_page_size=sync_refresh_page_size,
                            max_provider_requests_override=sync_refresh_attempts_per_city,
                            max_direct_apply_requests_override=0,
                        ),
                        timeout=per_location_timeout,
                    )
                except asyncio.TimeoutError:
                    if cooldown_enabled and not (force_provider_refresh and explicit_local_intent):
                        next_until = time.monotonic() + sync_refresh_cooldown_seconds
                        if explicit_local_intent:
                            _feed_sync_refresh_cooldowns[cooldown_key] = next_until
                        else:
                            _feed_sync_refresh_cooldown_until = next_until
                    refresh_result = {
                        "attempted": True,
                        "ok": False,
                        "reason": "feed_sync_refresh_timeout",
                        "elapsed_ms": int((time.perf_counter() - refresh_started) * 1000),
                    }
                    logger.warning(
                        "jobs/feed jsearch_fallback_timeout: user_id=%s elapsed_ms=%s cooldown_seconds=%s location=%s",
                        user.user_id,
                        refresh_result["elapsed_ms"],
                        sync_refresh_cooldown_seconds,
                        loc_label,
                    )
                except Exception as exc:
                    if cooldown_enabled and not (force_provider_refresh and explicit_local_intent):
                        next_until = time.monotonic() + sync_refresh_cooldown_seconds
                        if explicit_local_intent:
                            _feed_sync_refresh_cooldowns[cooldown_key] = next_until
                        else:
                            _feed_sync_refresh_cooldown_until = next_until
                    refresh_result = {
                        "attempted": True,
                        "ok": False,
                        "reason": "feed_sync_refresh_error",
                        "error_type": exc.__class__.__name__,
                        "elapsed_ms": int((time.perf_counter() - refresh_started) * 1000),
                    }
                    logger.warning(
                        "jobs/feed jsearch_fallback_error: user_id=%s elapsed_ms=%s error=%s",
                        user.user_id,
                        refresh_result["elapsed_ms"],
                        exc,
                    )
                else:
                    logger.info(
                        "jobs/feed jsearch_fallback_location_complete: user_id=%s elapsed_ms=%s reason=%s imported=%s",
                        user.user_id,
                        int((time.perf_counter() - refresh_started) * 1000),
                        refresh_result.get("reason"),
                        refresh_result.get("jobs_imported", refresh_result.get("count")),
                    )
                refresh_results.append(refresh_result)
                if refresh_result.get("provider_rate_limited"):
                    break
                if refresh_result.get("reason") in {"feed_sync_refresh_timeout", "feed_sync_refresh_error"}:
                    if not explicit_local_intent:
                        break
                    continue

            # If the short blocking budget did not finish all planned locations
            # (or imported nothing), continue discovery in the background so the
            # response returns quickly and the client can silently re-poll.
            sync_imported_total = sum(
                int(item.get("jobs_imported", item.get("imported", item.get("count") or 0)) or 0)
                for item in refresh_results
                if isinstance(item, dict)
            )
            remaining_locations = refresh_locations[len(refresh_results):]
            if (
                _env_bool("JOBS_FEED_BACKGROUND_REFRESH_ENABLED", True)
                and not explicit_local_empty_inventory
                and (remaining_locations or sync_imported_total == 0)
            ):
                continuation_signature = (
                    _explicit_local_cooldown_key()
                    if explicit_local_intent
                    else "|".join(["global", normalize_place_name(target_role or ""), radius_scope])
                ) + "|continue"
                continuation_scheduled = schedule_feed_background_refresh(
                    continuation_signature,
                    profile_for_refresh,
                    remaining_locations or refresh_locations,
                    search_radius=search_radius,
                    role_override=target_role,
                    requested_limit=requested_limit,
                    max_results=sync_refresh_max_results,
                    max_pages=sync_refresh_max_pages,
                    page_size=sync_refresh_page_size,
                    attempts_per_city=sync_refresh_attempts_per_city,
                    user_id=user.user_id,
                )
                if continuation_scheduled:
                    request_trace["background_refresh_scheduled"] = True
                    logger.info(
                        "jobs/feed sync_refresh_continued_in_background: user_id=%s remaining_locations=%s sync_imported=%s",
                        user.user_id,
                        len(remaining_locations),
                        sync_imported_total,
                    )

            if france_travail_parallel_task is not None:
                ft_remaining = max(0.1, provider_refresh_deadline - time.perf_counter())
                try:
                    ft_result = await asyncio.wait_for(france_travail_parallel_task, timeout=ft_remaining)
                    refresh_results.append({
                        "attempted": True,
                        "ok": bool(ft_result.get("used")),
                        "reason": "france_travail_parallel_fallback",
                        "jobs_imported": ft_result.get("total_imported", 0),
                        "relevant_imported": ft_result.get("relevant_imported", 0),
                        "manual_ready_imported": ft_result.get("manual_ready_imported", 0),
                        "jobs": ft_result.get("jobs") or [],
                        "search_key": None,
                        "search_keys": [],
                    })
                    logger.info(
                        "jobs/feed france_travail_parallel_complete: user_id=%s used=%s total_imported=%s relevant_imported=%s",
                        user.user_id,
                        ft_result.get("used"),
                        ft_result.get("total_imported"),
                        ft_result.get("relevant_imported"),
                    )
                except asyncio.TimeoutError:
                    france_travail_parallel_task.cancel()
                    logger.warning("jobs/feed france_travail_parallel_timeout: user_id=%s", user.user_id)
                except Exception as exc:
                    logger.warning("jobs/feed france_travail_parallel_error: user_id=%s error=%s", user.user_id, exc)

            for refresh_result in refresh_results:
                direct_refresh_jobs.extend(job for job in (refresh_result.get("jobs") or []) if isinstance(job, dict))
                for key in refresh_result.get("search_keys") or []:
                    if key:
                        provider_search_keys.append(str(key))
                if refresh_result.get("search_key"):
                    provider_search_keys.append(str(refresh_result.get("search_key")))
            provider_search_keys = list(dict.fromkeys(provider_search_keys))
            request_trace["local_jsearch_discovery_attempted"] = bool(explicit_local_intent)
            request_trace["jsearch_refresh_results"] = [
                {
                    "reason": item.get("reason"),
                    "ok": item.get("ok"),
                    "imported": item.get("jobs_imported", item.get("imported", item.get("count", 0))),
                    "relevant_imported": item.get("relevant_imported"),
                    "manual_ready_imported": item.get("manual_ready_imported"),
                    "provider_requests": item.get("provider_requests"),
                    "provider_errors": item.get("provider_errors"),
                    "provider_rate_limited": item.get("provider_rate_limited"),
                    "elapsed_ms": item.get("elapsed_ms"),
                    "final_location_used": item.get("final_location_used"),
                }
                for item in refresh_results
                if isinstance(item, dict)
            ]
            total_provider_imported = sum(
                int(item.get("jobs_imported", item.get("imported", item.get("count") or 0)) or 0)
                for item in refresh_results
                if isinstance(item, dict)
            )
            request_trace["jsearch_results_count"] = total_provider_imported

            skip_candidate_reload = (
                total_provider_imported == 0
                and not direct_refresh_jobs
                and bool(candidates)
            )
            if skip_candidate_reload:
                logger.info(
                    "jobs/feed jsearch_fallback_no_new_jobs_skip_reload: user_id=%s elapsed_ms=%s existing_candidates=%s",
                    user.user_id,
                    _elapsed_ms(),
                    len(candidates),
                )
            if not skip_candidate_reload:
                _clear_feed_job_pool_cache()
                candidates, base_query, unfiltered_count, applyable_count, db_rejected_count = await _load_db_candidates(include_unknown=False)
                # Post-refresh tier-C fallback: only after the live JSearch +
                # France Travail attempt above has had its chance to find more
                # A/B-tier jobs. Tier A/B is still always preferred (both by
                # append order here and by _quality_sort_key's own tier_rank
                # sort downstream) -- this only fills remaining slots once A/B
                # is confirmed too thin, it never displaces an A/B result.
                if len(candidates) < requested_limit:
                    c_candidates, c_query, c_unfiltered, c_applyable, c_rejected = await _load_db_candidates(include_unknown=True)
                    seen_ids = {job.get("job_id") for job in candidates}
                    candidates.extend(job for job in c_candidates if job.get("job_id") not in seen_ids)
                    unfiltered_count += c_unfiltered
                    applyable_count += c_applyable
                    db_rejected_count += c_rejected
                    base_query = c_query
                if explicit_local_intent and len(candidates) < requested_limit:
                    visible_candidates, visible_query, visible_unfiltered, visible_applyable, visible_rejected = await _load_local_visible_candidates()
                    seen_ids = {job.get("job_id") for job in candidates}
                    additions = [job for job in visible_candidates if job.get("job_id") not in seen_ids]
                    candidates.extend(additions)
                    unfiltered_count += visible_unfiltered
                    applyable_count += visible_applyable
                    db_rejected_count += visible_rejected
                    local_visible_count += len(visible_candidates)
                    local_manual_count += sum(1 for job in visible_candidates if job.get("application_mode") != "auto_apply")
                    local_blocked_hidden_count += visible_rejected
                    if additions:
                        base_query = visible_query if not base_query else base_query
                        feed_source = "db_after_jsearch_fallback+local_visible"
                if direct_refresh_jobs:
                    seen_ids = {job.get("job_id") for job in candidates}
                    fresh_additions: List[Dict[str, Any]] = []
                    for job in direct_refresh_jobs:
                        job_id = job.get("job_id")
                        if job_id and job_id in seen_ids:
                            continue
                        if job_id and job_id in swiped_ids:
                            continue
                        if _job_is_blocked_for_feed(job) or not _job_has_usable_apply_url(job):
                            continue
                        candidate = _with_application_capability_fields(job)
                        if explicit_local_intent and not _passes_explicit_local_hard_constraint(candidate):
                            continue
                        if not _matches_explicit_filters(candidate):
                            continue
                        fresh_additions.append(candidate)
                        if job_id:
                            seen_ids.add(job_id)
                    if fresh_additions:
                        candidates.extend(fresh_additions)
                        local_visible_count += len(fresh_additions)
                        local_manual_count += sum(1 for job in fresh_additions if job.get("application_mode") != "auto_apply")
                        feed_source = f"{feed_source}+fresh_provider_jobs"
                pre_filter_candidates = list(candidates)
                if not str(feed_source).startswith("db_after_jsearch_fallback"):
                    feed_source = "db_after_jsearch_fallback"
            logger.info(
                "jobs/feed jsearch_fallback_complete: user_id=%s imported_results=%s provider_search_keys=%s db_after_count=%s db_after_role_good_count=%s unfiltered_count=%s applyable_count=%s rejected_de_count=%s",
                user.user_id,
                [item.get("jobs_imported", item.get("count")) for item in refresh_results if isinstance(item, dict)],
                len(provider_search_keys),
                len(candidates),
                sum(1 for job in candidates if _role_compatible_for_threshold(job)),
                unfiltered_count,
                applyable_count,
                db_rejected_count,
            )
        elif not db_first_enabled:
            candidates, base_query, unfiltered_count, applyable_count, db_rejected_count = await _load_db_candidates(include_unknown=allow_unknown_tier)
            pre_filter_candidates = list(candidates)

        deduped_candidates: List[Dict[str, Any]] = []
        seen_candidate_ids = set()
        for job in candidates:
            key = job.get("job_id") or job.get("fingerprint") or f"{job.get('provider')}:{job.get('external_id')}"
            if key in seen_candidate_ids:
                continue
            seen_candidate_ids.add(key)
            deduped_candidates.append(job)
        candidates = deduped_candidates
        if explicit_local_intent and candidates:
            ranked_pre_hydrate_count = len(candidates)
            candidates = await _hydrate_feed_jobs(candidates)
            candidates = [_with_application_capability_fields(job) for job in candidates]
            candidates = [job for job in candidates if _passes_explicit_local_hard_constraint(job)]
            if len(candidates) != ranked_pre_hydrate_count:
                logger.info(
                    "feed_filter_stage user_id=%s stage=explicit_local_candidate_hydrated_constraint elapsed_ms=%s before=%s after=%s rejected=%s",
                    user.user_id,
                    _elapsed_ms(),
                    ranked_pre_hydrate_count,
                    len(candidates),
                    ranked_pre_hydrate_count - len(candidates),
                )
        pre_filter_candidates = list(candidates)

        logger.info(
            "feed_filter_stage user_id=%s stage=candidate_pool elapsed_ms=%s source=%s jsearch_fallback=%s query=%s count=%s db_good_count=%s unfiltered_count=%s applyable_count=%s rejected_de_count=%s swiped_count=%s explicit_filters=%s",
            user.user_id,
            _elapsed_ms(),
            feed_source,
            jsearch_fallback_triggered,
            base_query,
            len(candidates),
            db_good_count,
            unfiltered_count,
            applyable_count,
            db_rejected_count,
            len(swiped_ids),
            explicit_filters,
        )

        def _feed_location_score(job: Dict[str, Any], *, worldwide: bool) -> int:
            if worldwide:
                return 10
            if explicit_local_intent and _explicit_local_remote_allowed(job):
                return 25
            expanded_match = _expanded_location_match(job)
            if expanded_match:
                distance = float(expanded_match.get("distance_km") or 0)
                if expanded_match.get("is_origin"):
                    return 75
                distance_score = max(0, 45 - int(distance))
                population_bonus = min(15, int((int(expanded_match.get("population") or 0) ** 0.5) / 40))
                return 30 + distance_score + population_bonus
            if explicit_local_intent and selected_city_terms:
                job_text = _normalized_job_location_text(job)
                if any(term in job_text for term in selected_city_terms if len(term) >= 3):
                    return 60
            if include_unknown_location and not (job.get("location") or job.get("city") or job.get("region") or job.get("country_code")):
                return 5
            return _location_score(job, terms, worldwide=False)

        def rank(pool: List[Dict[str, Any]], *, worldwide: bool, broad: bool, any_role: bool = False) -> List[Dict[str, Any]]:
            ranked = []
            for job in pool:
                if not any_role and not _category_compatible(job):
                    continue
                role_score = 10 if any_role else _role_score(job, strict_tokens if not broad else [], family_tokens)
                location_score = _feed_location_score(job, worldwide=worldwide)
                if not worldwide and _is_current_provider_job(job):
                    location_score = max(location_score, 35)
                if not worldwide and terms["labels"] and location_score <= 0:
                    continue
                if not any_role and role_score <= 0:
                    continue
                ranked.append({
                    **enrich_job_employment_kind(job),
                    "_feed_rank_score": role_score * 3 + location_score * 2 + _recency_score(job) + max(0, 30 - _tier_rank(job) * 10) + employment_kind_rank_bonus(job, effective_job_types),
                    "_role_match_score": role_score,
                    "_location_match_score": location_score,
                })
            ranked.sort(key=lambda row: row.get("_feed_rank_score", 0), reverse=True)
            diverse: List[Dict[str, Any]] = []
            deferred: List[Dict[str, Any]] = []
            seen_companies = set()
            for job in ranked:
                company_key = (job.get("company") or "").strip().lower()
                if len(diverse) < min(5, requested_limit) and company_key in seen_companies:
                    deferred.append(job)
                    continue
                diverse.append(job)
                if company_key:
                    seen_companies.add(company_key)
                if len(diverse) >= requested_limit:
                    break
            for job in deferred:
                if len(diverse) >= requested_limit:
                    break
                diverse.append(job)
            return diverse[:requested_limit]

        fallback_used = "none"
        if is_worldwide_radius:
            fallback_used = "worldwide_radius"
            jobs = rank(candidates, worldwide=True, broad=False)
            logger.info("feed_filter_stage user_id=%s stage=worldwide_radius_strict_role elapsed_ms=%s count=%s", user.user_id, _elapsed_ms(), len(jobs))
            if len(jobs) < requested_limit and not _timed_out():
                fallback_used = "worldwide_radius_role_family"
                jobs = rank(candidates, worldwide=True, broad=True)
                logger.info("feed_filter_stage user_id=%s stage=worldwide_radius_role_family elapsed_ms=%s count=%s", user.user_id, _elapsed_ms(), len(jobs))
            if len(jobs) < requested_limit and not _timed_out():
                fallback_used = "worldwide_radius_auto_apply"
                jobs = rank(candidates, worldwide=True, broad=True, any_role=True)
                logger.info(
                    "feed_filter_stage user_id=%s stage=worldwide_radius_auto_apply elapsed_ms=%s count=%s candidate_pool=%s source=%s",
                    user.user_id,
                    _elapsed_ms(),
                    len(jobs),
                    len(candidates),
                    feed_source,
                )
        else:
            jobs = rank(candidates, worldwide=False, broad=False)
            logger.info("feed_filter_stage user_id=%s stage=strict_role_location elapsed_ms=%s count=%s", user.user_id, _elapsed_ms(), len(jobs))
            # Allow relaxed fallback even when timed out: user already waited, show something
            if len(jobs) < requested_limit and (not _timed_out() or len(jobs) == 0):
                relaxed_local = rank(candidates, worldwide=False, broad=True)
                if relaxed_local:
                    fallback_used = "explicit_local_role_family"
                    jobs = relaxed_local
                logger.info(
                    "feed_filter_stage user_id=%s stage=explicit_local_role_family elapsed_ms=%s count=%s",
                    user.user_id,
                    _elapsed_ms(),
                    len(jobs),
                )
            if len(jobs) < requested_limit and direct_refresh_jobs and not _timed_out():
                fallback_used = "direct_provider_relaxed_role"
                jobs = rank(candidates, worldwide=False, broad=True)
                logger.info(
                    "feed_filter_stage user_id=%s stage=direct_provider_relaxed_role elapsed_ms=%s count=%s direct_refresh_jobs=%s",
                    user.user_id,
                    _elapsed_ms(),
                    len(jobs),
                    len(direct_refresh_jobs),
                )
        allow_location_widening = (not explicit_location_filter or radius_km is None) and not explicit_local_intent
        if not is_worldwide_radius and explicit_filters and allow_location_widening and len(jobs) < requested_limit and explicit_location_filter and selected_country_terms:
            relaxed = [
                job for job in pre_filter_candidates
                if _matches_non_location_filters(job)
                and _role_score(job, strict_tokens, family_tokens) > 0
                and (
                    _matches_location(job)
                    or bool(selected_country_terms and any(term in str(job.get("location") or "").lower() for term in selected_country_terms))
                    or str(job.get("remote") or "").lower() == "remote"
                )
            ]
            if relaxed:
                relaxed_jobs = rank(relaxed, worldwide=True, broad=False)
                if len(relaxed_jobs) < requested_limit:
                    relaxed_jobs = rank(relaxed, worldwide=True, broad=True)
                if relaxed_jobs:
                    fallback_used = "relaxed_country"
                    jobs = relaxed_jobs
                logger.info(
                    "feed_filter_stage user_id=%s stage=relaxed_country elapsed_ms=%s count=%s",
                    user.user_id,
                    _elapsed_ms(),
                    len(jobs),
                )
        if not is_worldwide_radius and explicit_filters and allow_location_widening and len(jobs) < requested_limit and not _timed_out():
            widened = [
                job for job in pre_filter_candidates
                if _matches_non_location_filters(job)
                and _role_score(job, strict_tokens, family_tokens) > 0
            ]
            widened_jobs = rank(widened, worldwide=True, broad=True)
            if widened_jobs:
                fallback_used = "explicit_filters_worldwide_role_family"
                jobs = widened_jobs
            logger.info("feed_filter_stage user_id=%s stage=explicit_filters_worldwide_role_family elapsed_ms=%s count=%s", user.user_id, _elapsed_ms(), len(jobs))
        if explicit_filters and len(jobs) < requested_limit:
            fallback_used = "none_due_to_explicit_filters" if fallback_used == "none" else fallback_used
        if not is_worldwide_radius and not explicit_filters and len(jobs) < requested_limit and not _timed_out():
            fallback_used = "worldwide_role_family"
            jobs = rank(candidates, worldwide=True, broad=True)
            logger.info("feed_filter_stage user_id=%s stage=worldwide_role_family elapsed_ms=%s count=%s", user.user_id, _elapsed_ms(), len(jobs))
        if not is_worldwide_radius and not explicit_filters and len(jobs) < requested_limit and not _timed_out():
            fallback_used = "worldwide_auto_apply"
            jobs = rank(candidates, worldwide=True, broad=True, any_role=not bool(strict_tokens or family_tokens))
            logger.info("feed_filter_stage user_id=%s stage=worldwide_auto_apply elapsed_ms=%s count=%s", user.user_id, _elapsed_ms(), len(jobs))

        hard_location_pre_count = len(jobs)
        hard_location_rejected_count = 0
        global_fallback_suppressed = False
        if explicit_local_intent:
            constrained_jobs = [job for job in jobs if _passes_explicit_local_hard_constraint(job)]
            hard_location_rejected_count = hard_location_pre_count - len(constrained_jobs)
            if hard_location_rejected_count:
                global_fallback_suppressed = (
                    fallback_used in {
                        "relaxed_country",
                        "explicit_filters_worldwide_role_family",
                        "worldwide_role_family",
                        "worldwide_auto_apply",
                        "worldwide_radius_auto_apply",
                        "direct_provider_relaxed_role",
                    }
                    or feed_source == "legacy_direct_ats_cache"
                )
                jobs = constrained_jobs
                logger.info(
                    "feed_filter_stage user_id=%s stage=explicit_local_hard_constraint elapsed_ms=%s before=%s after=%s rejected=%s fallback_used=%s source=%s",
                    user.user_id,
                    _elapsed_ms(),
                    hard_location_pre_count,
                    len(jobs),
                    hard_location_rejected_count,
                    fallback_used,
                    feed_source,
                )

        jobs = await _hydrate_feed_jobs(jobs)
        jobs = [_with_application_capability_fields(job) for job in jobs]
        if explicit_local_intent:
            hydrated_pre_count = len(jobs)
            jobs = [job for job in jobs if _passes_explicit_local_hard_constraint(job)]
            hydrated_rejected = hydrated_pre_count - len(jobs)
            if hydrated_rejected:
                hard_location_rejected_count += hydrated_rejected
                logger.info(
                    "feed_filter_stage user_id=%s stage=explicit_local_hydrated_hard_constraint elapsed_ms=%s before=%s after=%s rejected=%s",
                    user.user_id,
                    _elapsed_ms(),
                    hydrated_pre_count,
                    len(jobs),
                    hydrated_rejected,
                )

        clean_jobs = []
        for job in jobs[:requested_limit]:
            clean = {key: value for key, value in job.items() if not key.startswith("_")}
            clean_jobs.append({
                **clean,
                "match_score": clean.get("match_score") or random.randint(78, 96),
                "match_reasons": clean.get("match_reasons") or ["Auto-apply compatible role from a supported ATS."],
            })
        elapsed = _elapsed_ms()
        logger.info(
            "jobs/feed fast complete: user_id=%s feed_elapsed_ms=%s returned=%s fallback_used=%s timed_out=%s direct_refresh_jobs=%s provider_search_keys=%s unfiltered_count=%s candidates=%s",
            user.user_id,
            elapsed,
            len(clean_jobs),
            fallback_used,
            _timed_out(),
            len(direct_refresh_jobs),
            len(provider_search_keys),
            unfiltered_count,
            len(candidates),
        )
        safe_refresh_results = [_compact_refresh_result(item) for item in refresh_results if isinstance(item, dict)]
        local_candidate_count = (
            len([job for job in pre_filter_candidates if _passes_explicit_local_hard_constraint(job)])
            if explicit_local_intent
            else len(pre_filter_candidates)
        )
        empty_reason = None
        if explicit_local_intent and not clean_jobs:
            empty_reason = {
                "code": "NO_LOCAL_AUTO_APPLY_JOBS",
                "message": "No verified auto-apply jobs found in this location yet.",
                "local_candidates_count": local_candidate_count,
                "local_c_tier_count": sum(1 for job in pre_filter_candidates if str(job.get("applyability_tier") or "").upper() == "C"),
                "local_blocked_count": local_blocked_hidden_count or db_rejected_count,
                "suggested_action": "broaden_location_or_enable_assisted_jobs",
            }
        feed_summary = {
            "auto_apply_count": sum(1 for job in clean_jobs if job.get("application_mode") == "auto_apply"),
            "assisted_count": sum(1 for job in clean_jobs if job.get("application_mode") == "assisted"),
            "manual_count": sum(1 for job in clean_jobs if job.get("application_mode") == "manual"),
            "blocked_hidden_count": local_blocked_hidden_count or db_rejected_count,
            "location_expansion_used": bool(location_context.get("used")),
            "local_visible_count": local_visible_count,
            "local_manual_count": local_manual_count,
        }
        request_trace["manual_candidate_count"] = feed_summary["manual_count"]
        request_trace["auto_apply_candidate_count"] = feed_summary["auto_apply_count"]
        request_trace["blocked_hidden_count"] = feed_summary["blocked_hidden_count"]
        request_trace["final_jobs_count"] = len(clean_jobs)
        request_trace["frontend_relevant_application_modes"] = sorted({
            str(job.get("application_mode") or "")
            for job in clean_jobs
            if job.get("application_mode")
        })
        response = {
            "jobs": clean_jobs,
            "total": len(clean_jobs),
            "empty_reason": empty_reason,
            "feed_summary": feed_summary,
            "feed_mode": "mixed",
            "background_refresh_scheduled": bool(request_trace.get("background_refresh_scheduled")),
            "auto_apply_count": feed_summary["auto_apply_count"],
            "total_count": len(candidates),
            "fallback_reason": None if fallback_used == "none" else fallback_used,
            "searched_location": terms["labels"][0] if terms["labels"] else profile.get("target_location"),
            "searched_locations": terms["labels"],
            "search_radius": search_radius,
            "suggested_next_radius": None if clean_jobs else "worldwide",
            "only_my_country": only_my_country,
            "widened_search": fallback_used.startswith("worldwide"),
            "original_location": terms["labels"][0] if terms["labels"] else profile.get("target_location"),
            "final_location_used": "worldwide" if fallback_used.startswith("worldwide") else (terms["labels"][0] if terms["labels"] else profile.get("target_location")),
            "provider_rate_limited": any(item.get("provider_rate_limited") for item in refresh_results),
            "provider_cooldown_until": next((item.get("provider_cooldown_until") for item in refresh_results if item.get("provider_cooldown_until")), None),
            "refresh_results": safe_refresh_results,
            "matched_role": target_role or None,
            "matched_location": terms["labels"],
            "companies_returned": sorted({job.get("company") for job in clean_jobs if job.get("company")}),
            "filters_applied": {
                "target_role": target_role or None,
                "role_tokens": strict_tokens,
                "broader_role_tokens": family_tokens,
                "auto_apply_supported": None,
                "ats_provider": None,
                "auto_apply_fallback": False,
                "provider_search_keys": provider_search_keys,
                "search_radius": search_radius,
                "explicit_filters": explicit_filters,
                "explicit_local_intent": explicit_local_intent,
                "hard_location_pre_count": hard_location_pre_count,
                "hard_location_post_count": len(jobs),
                "hard_location_rejected_count": hard_location_rejected_count,
                "global_fallback_suppressed": global_fallback_suppressed,
                "local_visible_count": local_visible_count,
                "local_manual_count": local_manual_count,
                "local_blocked_hidden_count": local_blocked_hidden_count,
                "manual_fulfillment_ready": True,
                "work_location": work_location,
                "job_type": effective_job_types or None,
                "profile_contract_type": profile_contract_type or None,
                "locations": [loc.get("location_label") for loc in selected_locations if loc.get("location_label")],
                "selected_city_terms": selected_city_terms,
                "selected_country_terms": selected_country_terms,
                "selected_country_codes": selected_country_codes,
                "location_intelligence": {
                    "enabled": location_context.get("enabled"),
                    "used": location_context.get("used"),
                    "reason": location_context.get("reason"),
                    "radius_km": location_context.get("radius_km"),
                    "include_cross_border": location_context.get("include_cross_border"),
                    "expanded_country_codes": expanded_country_codes,
                    "expanded_places": [
                        {
                            "name": place.get("name"),
                            "country_code": place.get("country_code"),
                            "distance_km": place.get("distance_km"),
                            "population": place.get("population"),
                            "is_origin": bool(place.get("is_origin")),
                        }
                        for place in (location_context.get("expanded_places") or [])[:20]
                    ],
                },
                "only_my_country": only_my_country,
                "min_salary": min_salary,
                "posted_within": posted_within,
                "job_type": job_type,
                "experience": experience,
                "only_company": only_company,
                "hide_company": hide_company,
                "only_industry": only_industry,
                "hide_industry": hide_industry,
                "include_unknown_location": include_unknown_location,
                "include_unknown_salary": include_unknown_salary,
            },
            "feed_elapsed_ms": elapsed,
            "fallback_used": fallback_used,
            "explicit_filters": explicit_filters,
        }
        if debug_diagnostics:
            response["request_trace"] = request_trace
        return response

    return await _fast_cached_feed()

    provider_name = primary_job_provider_name()
    provider_enabled = is_job_provider_enabled(provider_name)
    provider_configured = is_job_provider_configured(provider_name)
    fallback_mock = os.environ.get("JOB_PROVIDER_FALLBACK_MOCK", "false").lower() in ("1", "true", "yes", "on")
    profile_location_data = profile.get("target_location_data") or {}
    profile_country_code = (profile_location_data.get("country_code") or "").strip().lower()
    profile_country = profile_location_data.get("country")

    selected_locations: List[Dict[str, Any]] = []
    if locations_json:
        try:
            parsed_locations = json.loads(locations_json)
            if isinstance(parsed_locations, list):
                selected_locations = [loc for loc in parsed_locations if isinstance(loc, dict) and loc.get("location_label")]
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="locations_json must be a JSON array")

    request_location_data = None
    if location_label:
        request_location_data = {
            "location_label": location_label,
            "place_id": place_id,
            "country": country,
            "country_code": country_code,
            "lat": lat,
            "lng": lng,
        }
        selected_locations = [request_location_data]

    if only_my_country:
        if profile_country_code:
            selected_locations = [
                {**loc, "country_code": profile_country_code, "country": profile_country or loc.get("country")}
                for loc in selected_locations
            ]
            if not selected_locations and profile_location_data:
                selected_locations = [profile_location_data]
        elif not selected_locations and profile_location_data:
            selected_locations = [profile_location_data]

    if not selected_locations and profile_location_data:
        selected_locations = [profile_location_data]

    request_location = (
        selected_locations[0].get("location_label")
        if selected_locations
        else (location[0] if location else None)
    )

    async def _count_auto_for_labels(labels: List[str]) -> int:
        count_query: Dict[str, Any] = {"auto_apply_supported": True}
        if labels and search_radius not in ("worldwide", "remote", "remote/worldwide"):
            count_query["$or"] = [{"location": {"$regex": re.escape(label), "$options": "i"}} for label in labels]
        return await db.jobs.count_documents(count_query)

    refresh_results = []
    max_feed_refresh_locations = max(1, min(int(os.environ.get("FEED_MAX_REFRESH_LOCATIONS", "3")), 10))
    refresh_locations = (selected_locations or [None])[:max_feed_refresh_locations]
    for loc_data in refresh_locations:
        loc_label = loc_data.get("location_label") if loc_data else None
        refresh_result = await refresh_jobs_for_profile_if_needed(
            db,
            profile,
            require_auto_apply=not include_non_auto_apply,
            target_auto_apply_count=limit,
            location_override=loc_label,
            location_data_override=loc_data,
            search_radius=search_radius,
        )
        refresh_results.append(refresh_result)
        if refresh_result.get("provider_rate_limited"):
            break
        refreshed_labels = [
            item.get("location_label")
            for item in refresh_locations[:len(refresh_results)]
            if isinstance(item, dict) and item.get("location_label")
        ]
        if refreshed_labels and await _count_auto_for_labels(refreshed_labels) >= limit:
            break
    refresh_result = refresh_results[-1] if refresh_results else {"attempted": False, "reason": "no_refresh"}
    logger.info(
        "jobs/feed refresh complete: user_id=%s elapsed_ms=%s refresh_results=%s",
        user.user_id,
        int((time.perf_counter() - started_at) * 1000),
        refresh_results,
    )

    # exclude jobs already swiped
    swiped = await db.swipes.find({"user_id": user.user_id}, {"_id": 0, "job_id": 1}).to_list(2000)
    swiped_ids = {s["job_id"] for s in swiped}

    query: Dict[str, Any] = {}

    if swiped_ids:
        query["job_id"] = {"$nin": list(swiped_ids)}

    # Work location filter
    if work_location:
        if include_unknown_location:
            query["$or"] = [
                {"remote": {"$in": work_location}},
                {"remote": {"$in": [None, ""]}},
            ]
        else:
            query["remote"] = {"$in": work_location}

    # Experience / seniority filter (entry→junior, executive→principal/lead)
    if experience:
        exp_map = {
            "entry": ["junior"],
            "mid": ["mid"],
            "senior": ["senior"],
            "executive": ["lead", "principal"],
        }
        wanted: List[str] = []
        for e in experience:
            wanted.extend(exp_map.get(e, [e]))
        query["seniority"] = {"$in": wanted}

    # Location free-text filter - case-insensitive substring OR across tokens
    location_filter_clause = None
    selected_location_labels = [
        loc.get("location_label")
        for loc in selected_locations
        if isinstance(loc, dict) and loc.get("location_label")
    ]
    if search_radius in ("worldwide", "remote", "remote/worldwide"):
        filter_locations = None
    elif selected_location_labels:
        filter_locations = selected_location_labels
    elif location:
        filter_locations = location
    elif only_my_country and profile_country:
        filter_locations = [profile_country]
    else:
        filter_locations = None
    if filter_locations:
        expanded_filter_locations = set()
        for loc in filter_locations:
            expanded_filter_locations.add(loc)
            for part in re.split(r"[,/|-]", loc):
                part = part.strip()
                if len(part) >= 3:
                    expanded_filter_locations.add(part)
        loc_clauses = [{"location": {"$regex": re.escape(loc), "$options": "i"}} for loc in expanded_filter_locations]
        location_filter_clause = {"$or": loc_clauses}
        query.setdefault("$and", []).append(location_filter_clause)

    # Posted within
    if posted_within and posted_within != "any":
        days_map = {"1d": 1, "7d": 7, "30d": 30}
        days = days_map.get(posted_within)
        if days:
            cutoff = (datetime.now(timezone.utc) - timedelta(days=days)).isoformat()
            query["posted_at"] = {"$gte": cutoff}

    # Min salary filter — match any job whose top-of-range >= min_salary
    if min_salary and min_salary > 0:
        salary_clause: Dict[str, Any] = {"salary_max": {"$gte": min_salary}}
        if include_unknown_salary:
            query.setdefault("$and", []).append({
                "$or": [salary_clause, {"salary_max": {"$in": [None, 0]}}, {"salary_max": {"$exists": False}}],
            })
        else:
            query["salary_max"] = {"$gte": min_salary}

    # Only / hide companies
    if only_company:
        regexes = [{"company": {"$regex": f"^{re.escape(c)}$", "$options": "i"}} for c in only_company]
        query.setdefault("$and", []).append({"$or": regexes})
    if hide_company:
        query["company"] = {"$nin": []}  # init
        query["company"] = {"$not": {"$regex": "|".join(re.escape(c) for c in hide_company), "$options": "i"}}

    remote_pref = profile.get("remote_preference") or "any"
    target_role = feed_target_role
    radius_scope = (search_radius or "50km").lower().strip()

    def _tokens(value: str) -> List[str]:
        stop = {"and", "or", "the", "a", "an", "of", "for", "to", "in", "with", "remote", "jobs", "job"} | ACADEMIC_LEVEL_STOPWORDS
        return [token for token in re.findall(r"[a-z0-9]+", (value or "").lower()) if len(token) > 2 and token not in stop]

    # Some role labels are ambiguous/homonymous across languages (e.g.
    # "Chef"/"Cuisinier": bare "chef" in French means "lead", colliding
    # with unrelated titles like "chef de mission"). When target_role is a
    # known case, score against its safe token list instead of naively
    # tokenizing the raw (possibly ambiguous) text -- see role_query_terms.py.
    _role_match_override = resolve_role_match_tokens(target_role)
    role_tokens = _role_match_override if _role_match_override is not None else _tokens(target_role)
    broader_role_tokens = []
    if role_tokens:
        broader_role_tokens = role_tokens[-1:]
        if "analyst" in role_tokens:
            broader_role_tokens = list(dict.fromkeys([*role_tokens, "analytics", "analysis", "insights", "scientist"]))
        if "software" in role_tokens and "engineer" in role_tokens:
            broader_role_tokens = ["software", "engineer"]

    selected_country_codes = [
        (loc.get("country_code") or "").strip().lower()
        for loc in selected_locations
        if isinstance(loc, dict) and loc.get("country_code")
    ]
    selected_countries = [
        (loc.get("country") or "").strip().lower()
        for loc in selected_locations
        if isinstance(loc, dict) and loc.get("country")
    ]
    country_aliases = {
        "gb": ["united kingdom", "uk", "england", "london"],
        "us": ["united states", "usa", "new york", "san francisco"],
        "ma": ["morocco", "maroc", "casablanca"],
    }
    selected_country_terms = set(selected_countries)
    for code in selected_country_codes or ([profile_country_code] if profile_country_code else []):
        selected_country_terms.update(country_aliases.get(code, []))
    location_terms = set()
    for label in selected_location_labels or (location or []):
        location_terms.update(_tokens(label))
    if profile_country:
        selected_country_terms.add(profile_country.lower())

    def _role_score(job: Dict[str, Any], tokens: List[str]) -> int:
        if not tokens:
            return 0
        title = (job.get("title") or "").lower()
        body = " ".join([
            job.get("description") or "",
            job.get("clean_description") or "",
            " ".join(job.get("requirements") or []),
        ]).lower()
        title_hits = sum(1 for token in tokens if token in title)
        body_hits = sum(1 for token in tokens if token in body)
        # Skip the raw-phrase bonus for roles with a safe-token override --
        # the override's own tokens above already provide accurate matching
        # without re-introducing the ambiguous full phrase (e.g. "chef").
        exact_bonus = 25 if not _role_match_override and target_role and target_role.lower() in title else 0
        return exact_bonus + title_hits * 20 + min(body_hits, len(tokens)) * 5

    def _location_score(job: Dict[str, Any]) -> int:
        if radius_scope in ("worldwide", "remote/worldwide"):
            return 10
        job_location = (job.get("location") or "").lower()
        job_remote = (job.get("remote") or "").lower()
        if radius_scope == "remote" or remote_pref == "remote":
            return 35 if job_remote == "remote" else 5
        score = 0
        if location_terms and any(term in job_location for term in location_terms):
            score += 40
        if selected_country_terms and any(term and term in job_location for term in selected_country_terms):
            score += 25
        if not location_terms and selected_country_terms and score:
            score += 10
        if job_remote == "remote":
            score += 8
        return score

    def _recency_score(job: Dict[str, Any]) -> int:
        raw = job.get("posted_at") or job.get("imported_at")
        if not raw:
            return 0
        try:
            parsed = datetime.fromisoformat(str(raw).replace("Z", "+00:00"))
            if parsed.tzinfo is None:
                parsed = parsed.replace(tzinfo=timezone.utc)
        except ValueError:
            return 0
        age_days = max(0, (datetime.now(timezone.utc) - parsed).days)
        if age_days <= 7:
            return 20
        if age_days <= 30:
            return 10
        return 0

    def _rank_jobs(candidates: List[Dict[str, Any]], wanted: int) -> List[Dict[str, Any]]:
        strict_matches = []
        broader_matches = []
        for job in candidates:
            strict_role_score = _role_score(job, role_tokens)
            broad_role_score = _role_score(job, broader_role_tokens)
            location_match_score = _location_score(job)
            recency = _recency_score(job)
            base_score = strict_role_score * 3 + location_match_score * 2 + recency
            ranked_job = {
                **job,
                "_feed_rank_score": base_score,
                "_role_match_score": strict_role_score,
                "_location_match_score": location_match_score,
            }
            if not role_tokens or strict_role_score >= max(20, len(role_tokens) * 12):
                strict_matches.append(ranked_job)
            elif broad_role_score >= 20:
                ranked_job["_feed_rank_score"] = broad_role_score * 2 + location_match_score * 2 + recency
                broader_matches.append(ranked_job)

        strict_matches.sort(key=lambda j: j["_feed_rank_score"], reverse=True)
        broader_matches.sort(key=lambda j: j["_feed_rank_score"], reverse=True)
        ranked = strict_matches if len(strict_matches) >= wanted else [*strict_matches, *broader_matches]

        diverse = []
        deferred = []
        seen_companies = set()
        for job in ranked:
            company_key = (job.get("company") or "").strip().lower()
            if len(diverse) < min(5, wanted) and company_key in seen_companies:
                deferred.append(job)
                continue
            diverse.append(job)
            if company_key:
                seen_companies.add(company_key)
            if len(diverse) >= wanted:
                break
        if len(diverse) < wanted:
            for job in deferred:
                diverse.append(job)
                if len(diverse) >= wanted:
                    break
        return diverse[:wanted]

    def _without_location_filter(src: Dict[str, Any]) -> Dict[str, Any]:
        widened = {**src}
        if location_filter_clause and "$and" in widened:
            remaining = [clause for clause in widened["$and"] if clause != location_filter_clause]
            if remaining:
                widened["$and"] = remaining
            else:
                widened.pop("$and", None)
        return widened

    async def _fetch(q: Dict[str, Any], wanted: int) -> List[Dict[str, Any]]:
        candidate_limit = max(wanted * 100, 500)
        rows = await db.jobs.find(q, {"_id": 0}).limit(candidate_limit).to_list(candidate_limit)
        rows = [_with_apply_fulfillment_fields(row) for row in rows if _job_is_applyable(row)]
        return _rank_jobs(rows, wanted)

    base_query = {**query}
    total_all = await db.jobs.count_documents(base_query)
    auto_query = {
        **base_query,
        "auto_apply_supported": True,
        "ats_provider": {"$in": ["greenhouse", "lever", "ashby"]},
    }

    feed_mode = "auto_apply_only"
    fallback_reason = None
    widened_search = False
    final_location_used = request_location or profile.get("target_location")
    provider_rate_limited = bool(refresh_result.get("provider_rate_limited"))
    provider_cooldown_until = refresh_result.get("provider_cooldown_until")

    if include_non_auto_apply:
        feed_mode = "mixed"
        jobs = await _fetch(base_query, limit)
    else:
        jobs = await _fetch(auto_query, limit)

    jobs = jobs[:limit]
    auto_apply_count = sum(1 for j in jobs if j.get("auto_apply_supported") is True)
    total = await db.jobs.count_documents(auto_query if not include_non_auto_apply else base_query)
    companies_returned = sorted({j.get("company") for j in jobs if j.get("company")})
    filters_applied = {
        "target_role": target_role or None,
        "role_tokens": role_tokens,
        "broader_role_tokens": broader_role_tokens,
        "locations": selected_location_labels,
        "country_code": selected_country_codes or ([profile_country_code] if profile_country_code else []),
        "search_radius": search_radius,
        "remote_preference": remote_pref,
        "auto_apply_supported": not include_non_auto_apply,
        "ats_provider": ["greenhouse", "lever", "ashby"] if not include_non_auto_apply else None,
        "manual_fulfillment_ready": True,
    }
    if provider_rate_limited:
        fallback_reason = "provider_rate_limited"

    if not jobs:
        if include_non_auto_apply and provider_enabled and provider_configured and fallback_mock:
            mock_query = {**base_query}
            mock_query.pop("provider", None)
            jobs = await _fetch(mock_query, limit)
            total = await db.jobs.count_documents(mock_query)
            auto_apply_count = sum(1 for j in jobs if j.get("auto_apply_supported") is True)
        if not jobs:
            logger.info("No jobs found for feed; provider refresh result=%s", refresh_result)
            empty_fallback_reason = (
                "provider_rate_limited"
                if provider_rate_limited
                else "no_auto_apply_jobs_found" if not include_non_auto_apply else "No jobs found with these filters. Try widening your search distance or changing your location."
            )
            return {
                "jobs": [],
                "total": 0,
                "feed_mode": "auto_apply_only" if not include_non_auto_apply else "mixed",
                "auto_apply_count": 0,
                "total_count": 0,
                "fallback_reason": empty_fallback_reason,
                "searched_location": refresh_result.get("searched_location") or request_location or profile.get("target_location"),
                "searched_locations": selected_location_labels,
                "search_radius": search_radius,
                "suggested_next_radius": refresh_result.get("suggested_next_radius"),
                "only_my_country": only_my_country,
                "widened_search": widened_search or bool(refresh_result.get("widened_search")),
                "original_location": request_location or profile.get("target_location"),
                "final_location_used": final_location_used,
                "provider_rate_limited": provider_rate_limited,
                "provider_cooldown_until": provider_cooldown_until,
                "matched_role": target_role or None,
                "matched_location": selected_location_labels or ([request_location] if request_location else []),
                "companies_returned": [],
                "filters_applied": filters_applied,
            }

    # AI scoring batch — slow (5-12s for Claude). Off by default for snappy UX.
    score_map: Dict[str, Dict[str, Any]] = {}
    if score:
        try:
            matches = await claude_score_jobs(profile, jobs)
            score_map = {m["job_id"]: m for m in matches}
        except LLMProviderNotConfigured as e:
            raise HTTPException(status_code=502, detail=str(e))
        except Exception as e:
            logger.exception("Match scoring failed")
            raise HTTPException(status_code=502, detail="AI job scoring failed")

    enriched = []
    for j in jobs:
        m = score_map.get(j["job_id"], {})
        clean_job = {k: v for k, v in j.items() if not k.startswith("_feed_") and not k.startswith("_role_") and not k.startswith("_location_")}
        enriched.append({
            **clean_job,
            "match_score": m.get("score") or random.randint(78, 96),
            "match_reasons": m.get("reasons") or ["Strong alignment with your skills."],
        })
    if fallback_reason is None and feed_mode == "auto_apply_only":
        fallback_reason = None
    logger.info(
        "jobs/feed complete: user_id=%s elapsed_ms=%s returned=%s total=%s feed_mode=%s fallback_reason=%s",
        user.user_id,
        int((time.perf_counter() - started_at) * 1000),
        len(enriched),
        total,
        feed_mode,
        fallback_reason,
    )
    return {
        "jobs": enriched,
        "total": total,
        "feed_mode": feed_mode,
        "auto_apply_count": auto_apply_count,
        "total_count": total_all,
        "fallback_reason": fallback_reason,
        "searched_location": refresh_result.get("searched_location") or request_location or profile.get("target_location"),
        "searched_locations": selected_location_labels,
        "search_radius": search_radius,
        "suggested_next_radius": refresh_result.get("suggested_next_radius"),
        "only_my_country": only_my_country,
        "widened_search": widened_search or bool(refresh_result.get("widened_search")),
        "original_location": request_location or profile.get("target_location"),
        "final_location_used": final_location_used,
        "provider_rate_limited": provider_rate_limited,
        "provider_cooldown_until": provider_cooldown_until,
        "matched_role": target_role or None,
        "matched_location": selected_location_labels or ([request_location] if request_location else []),
        "companies_returned": companies_returned,
        "filters_applied": filters_applied,
    }


@api_router.post("/swipe")
async def swipe(req: SwipeRequest, user: User = Depends(get_current_user)):
    phase = "swipe_start"

    def log_phase(name: str, **extra: Any) -> None:
        safe_extra = {
            key: value
            for key, value in extra.items()
            if key not in {"cv_text", "cover_letter", "tailored_cover_letter", "prompt", "tokens"}
        }
        logger.info(
            "swipe_phase=%s user_id=%s job_id=%s direction=%s provider=%s extra=%s",
            name,
            user.user_id,
            req.job_id,
            req.direction,
            DATABASE_PROVIDER,
            safe_extra,
        )

    def safe_error(exc: Exception) -> HTTPException:
        message = str(exc).strip() or "Swipe failed"
        if len(message) > 500:
            message = message[:500]
        logger.exception(
            "swipe_failed phase=%s user_id=%s job_id=%s direction=%s exception=%s",
            phase,
            user.user_id,
            req.job_id,
            req.direction,
            exc.__class__.__name__,
        )
        return HTTPException(
            status_code=getattr(exc, "status_code", 500) if isinstance(exc, HTTPException) else 500,
            detail={
                "phase": phase,
                "exception_class": exc.__class__.__name__,
                "message": message,
            },
        )

    try:
        log_phase("swipe_start")
        phase = "auth_ok"
        log_phase("auth_ok")

        if req.direction == "right" and user.demo_account:
            existing = await db.swipes.find_one(
                {"user_id": user.user_id, "job_id": req.job_id},
                {"_id": 0, "swipe_id": 1},
            )
            if not existing:
                await db.swipes.insert_one(_swipe_insert_doc(user.user_id, req.job_id, req.direction))
            log_phase("demo_account_apply_blocked")
            return {
                "ok": True,
                "applied": True,
                "submitted": False,
                "demo_account": True,
                "duplicate": bool(existing),
            }

        phase = "job_loaded"
        job = await db.jobs.find_one({"job_id": req.job_id}, {"_id": 0})
        if not job:
            if req.direction != "right":
                existing = await db.swipes.find_one({"user_id": user.user_id, "job_id": req.job_id}, {"_id": 0})
                if not existing:
                    await db.swipes.insert_one({
                        "user_id": user.user_id,
                        "job_id": req.job_id,
                        "direction": req.direction,
                        "created_at": datetime.now(timezone.utc).isoformat(),
                    })
                log_phase("left_swipe_missing_job_recorded", duplicate=bool(existing))
                return {"ok": True, "applied": False, "duplicate": bool(existing), "missing_job": True}
            raise HTTPException(status_code=404, detail="Job not found")
        job = _with_apply_fulfillment_fields(job)
        log_phase(
            "job_loaded",
            job_provider=job.get("provider"),
            ats_provider=job.get("ats_provider"),
            apply_fulfillment_status=job.get("apply_fulfillment_status"),
            has_description=bool(job.get("description") or job.get("clean_description")),
            job_keys=sorted(list(job.keys()))[:60],
        )

        billing_credit_status = None
        if req.direction == "right":
            phase = "pre_apply_validation"
            log_phase("pre_apply_validation_start")
            pre_apply = await validate_job_before_application(job)
            job = _with_apply_fulfillment_fields(pre_apply.get("job") or job)
            if not pre_apply.get("allowed"):
                log_phase(
                    "pre_apply_validation_blocked",
                    reason=pre_apply.get("reason"),
                    validation_status=pre_apply.get("validation_status"),
                    applyability_tier=pre_apply.get("applyability_tier"),
                    selected_apply_url=pre_apply.get("selected_apply_url"),
                    requires_login=pre_apply.get("requires_login"),
                    requires_account_creation=pre_apply.get("requires_account_creation"),
                    captcha_detected=pre_apply.get("captcha_detected"),
                    credit_skipped=True,
                    application_skipped=True,
                )
                raise HTTPException(
                    status_code=422,
                    detail={
                        "success": False,
                        "ok": False,
                        "blocked": True,
                        "reason": pre_apply.get("reason"),
                        "message": pre_apply.get("reason"),
                        "job_id": req.job_id,
                        "validation_status": pre_apply.get("validation_status"),
                        "applyability_tier": pre_apply.get("applyability_tier"),
                        "selected_apply_url": pre_apply.get("selected_apply_url"),
                        "requires_login": pre_apply.get("requires_login"),
                        "requires_account_creation": pre_apply.get("requires_account_creation"),
                        "captcha_detected": pre_apply.get("captcha_detected"),
                    },
                )
            log_phase(
                "pre_apply_validation_allowed",
                validation_status=pre_apply.get("validation_status"),
                applyability_tier=pre_apply.get("applyability_tier"),
                selected_apply_url=pre_apply.get("selected_apply_url"),
            )
            if not _job_is_applyable(job):
                log_phase(
                    "apply_fulfillment_blocked",
                    apply_fulfillment_status=job.get("apply_fulfillment_status"),
                    apply_fulfillment_reason=job.get("apply_fulfillment_reason"),
                    source=job.get("source"),
                    external_url=job.get("external_url"),
                )
                raise HTTPException(
                    status_code=422,
                    detail={
                        "message": "This job cannot be fulfilled by Hirly.",
                        "apply_fulfillment_status": job.get("apply_fulfillment_status"),
                        "apply_fulfillment_reason": job.get("apply_fulfillment_reason"),
                    },
                )
            phase = "billing_credit_check"
            billing_credit_status = await _billing_apply_credit_status(user)
            if not billing_credit_status.get("is_premium") or int(billing_credit_status.get("credits_remaining") or 0) <= 0:
                log_phase("billing_credit_blocked", billing=billing_credit_status)
                raise HTTPException(
                    status_code=402,
                    detail={
                        "message": "No application credits remaining",
                        "billing": billing_credit_status,
                    },
                )
            log_phase("billing_credit_ok", billing=billing_credit_status)

        phase = "swipe_record_insert_start"
        existing = await db.swipes.find_one({"user_id": user.user_id, "job_id": req.job_id}, {"_id": 0})
        log_phase("swipe_record_insert_start", duplicate=bool(existing))
        if not existing:
            await db.swipes.insert_one(_swipe_insert_doc(user.user_id, req.job_id, req.direction, job))
        phase = "swipe_record_insert_done"
        log_phase("swipe_record_insert_done", duplicate=bool(existing))

        if req.direction != "right":
            phase = "response_build_start"
            log_phase("response_build_start")
            response = {"ok": True, "applied": False, "duplicate": bool(existing)}
            phase = "response_build_done"
            log_phase("response_build_done")
            phase = "swipe_complete"
            log_phase("swipe_complete")
            if _posthog_client is not None and not existing:
                with new_context():
                    identify_context(user.user_id)
                    posthog_capture(
                        "job_dismissed",
                        properties={
                            "job_provider": job.get("provider"),
                            "ats_provider": job.get("ats_provider"),
                        },
                    )
            return response

        profile = await db.profiles.find_one({"user_id": user.user_id}, {"_id": 0})
        if not profile:
            raise HTTPException(status_code=400, detail="Profile required")
        if not profile.get("cv_text"):
            raise HTTPException(status_code=400, detail="Upload CV first")
        if not _profile_has_usable_phone(profile):
            raise HTTPException(status_code=400, detail="Add your phone number before applying")

        existing_app = await db.applications.find_one({"user_id": user.user_id, "job_id": req.job_id}, {"_id": 0})
        if existing_app:
            normalized = _normalize_application_status_fields(existing_app)
            if normalized.get("generation_status") in {"pending_generation", "generating"}:
                _schedule_application_generation(user.user_id)
            duplicate_lengths = _application_text_lengths(normalized)
            phase = "response_build_start"
            log_phase(
                "response_build_start",
                duplicate_application=True,
                application_id=normalized.get("application_id"),
                ats_provider=job.get("ats_provider"),
                package_status=normalized.get("package_status"),
                submission_status=normalized.get("submission_status"),
                **duplicate_lengths,
            )
            response = {
                "ok": True,
                "applied": True,
                "submitted": normalized.get("submission_status") == "submitted",
                "duplicate": bool(existing),
                "application_id": normalized["application_id"],
                "package_status": normalized["package_status"],
                "application_status": normalized["package_status"],
                "submission_status": normalized["submission_status"],
                "manual_status": _effective_manual_status(normalized),
                "manual_fulfillment": _manual_application_fulfillment_enabled(),
                "queued_for_generation": normalized.get("generation_status") in {"pending_generation", "generating"},
                "generation_status": normalized.get("generation_status"),
                "billing": billing_credit_status,
            }
            phase = "response_build_done"
            log_phase(
                "response_build_done",
                duplicate_application=True,
                application_id=normalized.get("application_id"),
                status_returned_to_frontend=normalized.get("submission_status"),
            )
            phase = "swipe_complete"
            log_phase(
                "swipe_complete",
                duplicate_application=True,
                application_id=normalized.get("application_id"),
                status_returned_to_frontend=normalized.get("submission_status"),
            )
            return response

        phase = "application_queue_start"
        log_phase("application_queue_start")
        doc = _pending_application_doc(user, job)

        phase = "application_upsert_start"
        log_phase("application_upsert_start", application_id=doc.get("application_id"))
        await db.applications.update_one(
            {"user_id": user.user_id, "job_id": req.job_id},
            {"$setOnInsert": doc},
            upsert=True,
        )
        saved_doc = None
        application_id_doc = None
        user_job_doc = None
        if doc.get("application_id"):
            application_id_doc = await db.applications.find_one({"application_id": doc["application_id"]}, {"_id": 0})
        log_phase(
            "application_verify_by_application_id_found",
            application_id=doc.get("application_id"),
            found=bool(application_id_doc),
        )
        user_job_doc = await db.applications.find_one({"user_id": user.user_id, "job_id": req.job_id}, {"_id": 0})
        log_phase(
            "application_verify_by_user_job_found",
            application_id=(user_job_doc or {}).get("application_id"),
            found=bool(user_job_doc),
        )
        if application_id_doc:
            saved_doc = application_id_doc
            log_phase("application_verify_fallback_used", fallback="application_id_primary")
        elif user_job_doc:
            saved_doc = user_job_doc
            log_phase("application_verify_fallback_used", fallback="user_job_secondary")
        if not saved_doc:
            logger.warning(
                "application_verification_failed user_id=%s job_id=%s application_id=%s insert_returned_without_readback=true",
                user.user_id,
                req.job_id,
                doc.get("application_id"),
            )
            log_phase("verification_failed", application_id=doc.get("application_id"))
            saved_doc = doc
        saved_doc = _normalize_application_status_fields(saved_doc)
        browser_prepare_started = False

        phase = "application_upsert_done"
        credit_result = await _consume_application_credit(user.user_id)
        _schedule_application_generation(user.user_id)
        log_phase(
            "application_upsert_done",
            application_id=saved_doc.get("application_id"),
            ats_provider=job.get("ats_provider"),
            package_status=saved_doc.get("package_status"),
            **_application_text_lengths(saved_doc),
            browser_prepare_started=browser_prepare_started,
            final_submission_status_saved=saved_doc.get("submission_status"),
            submission_status=saved_doc.get("submission_status"),
            credits_remaining=credit_result.get("credits_remaining"),
            manual_fulfillment=_manual_application_fulfillment_enabled(),
            queued_for_generation=True,
        )

        phase = "response_build_start"
        log_phase(
            "response_build_start",
            application_id=saved_doc.get("application_id"),
            package_status=saved_doc.get("package_status"),
            submission_status=saved_doc.get("submission_status"),
        )
        response = {
            "ok": True,
            "applied": True,
            "submitted": False,
            "duplicate": bool(existing),
            "application_id": saved_doc["application_id"],
            "package_status": saved_doc["package_status"],
            "application_status": saved_doc["package_status"],
            "submission_status": saved_doc["submission_status"],
            "generation_status": saved_doc.get("generation_status"),
            "queued_for_generation": True,
            "manual_status": _effective_manual_status(saved_doc),
            "manual_fulfillment": _manual_application_fulfillment_enabled(),
            "billing": credit_result,
        }
        phase = "response_build_done"
        log_phase(
            "response_build_done",
            application_id=saved_doc.get("application_id"),
            status_returned_to_frontend=response.get("submission_status"),
        )
        phase = "swipe_complete"
        log_phase(
            "swipe_complete",
            application_id=saved_doc.get("application_id"),
            status_returned_to_frontend=response.get("submission_status"),
        )
        if _posthog_client is not None and not existing_app:
            with new_context():
                identify_context(user.user_id)
                posthog_capture(
                    "job_application_created",
                    properties={
                        "application_id": saved_doc.get("application_id"),
                        "job_provider": job.get("provider"),
                        "ats_provider": job.get("ats_provider"),
                        "package_status": saved_doc.get("package_status"),
                    },
                )
        return response
    except LLMProviderNotConfigured as exc:
        phase = "application_generation_start"
        raise HTTPException(
            status_code=502,
            detail={
                "phase": phase,
                "exception_class": exc.__class__.__name__,
                "message": str(exc),
            },
        ) from exc
    except HTTPException as exc:
        if exc.status_code == 402 or (exc.status_code == 422 and isinstance(exc.detail, dict) and exc.detail.get("blocked") is True):
            raise exc
        raise safe_error(exc) from exc
    except Exception as exc:
        raise safe_error(exc) from exc


@api_router.get("/swipes/history")
async def swipes_history(
    user: User = Depends(get_current_user),
    direction: Optional[str] = None,  # "left" | "right" | None=all
    limit: int = 100,
):
    """Return the user's swipe history with the joined job doc.
    direction='left' → SKIP (we'll surface as 'Skipped Jobs')
    direction='right' → APPLY (alias for /applications minimal)."""
    q: Dict[str, Any] = {"user_id": user.user_id}
    if direction in ("left", "right"):
        q["direction"] = direction
    rows = await db.swipes.find(q, {"_id": 0}).sort("created_at", -1).limit(limit).to_list(limit)
    if not rows:
        return {"swipes": []}
    job_ids = list({r["job_id"] for r in rows})
    jobs = await db.jobs.find({"job_id": {"$in": job_ids}}, {"_id": 0}).to_list(len(job_ids))
    job_map = {j["job_id"]: j for j in jobs}
    return {
        "swipes": [
            {
                **r,
                "job": job_map.get(r["job_id"]) or r.get("job_snapshot"),
            }
            for r in rows
        ],
    }


@api_router.post("/swipes/{job_id}/apply-from-passed")
async def apply_from_passed(job_id: str, user: User = Depends(get_current_user)):
    """Convert a passed (left) swipe into an application with a tailored package."""
    swipe_row = await db.swipes.find_one(
        {"user_id": user.user_id, "job_id": job_id},
        {"_id": 0},
    )
    if not swipe_row or swipe_row.get("direction") != "left":
        raise HTTPException(status_code=404, detail={"message": "Passed job not found"})

    restored = await _restore_job_from_swipe_snapshot(job_id, swipe_row)
    if not restored:
        raise HTTPException(
            status_code=404,
            detail={"message": "This job is no longer available. It may have expired."},
        )

    await db.swipes.delete_one({"user_id": user.user_id, "job_id": job_id})
    return await swipe(SwipeRequest(job_id=job_id, direction="right"), user)


@api_router.delete("/swipes/{job_id}")
async def delete_swipe(job_id: str, user: User = Depends(get_current_user)):
    """Remove a swipe so the job can re-enter the feed (used by 'Apply Now' from history)."""
    res = await db.swipes.delete_one({"user_id": user.user_id, "job_id": job_id})
    if res.deleted_count == 0:
        raise HTTPException(status_code=404, detail="No such swipe")
    return {"ok": True}


@api_router.post("/swipe/undo")
async def undo_swipe(user: User = Depends(get_current_user)):
    last = await db.swipes.find({"user_id": user.user_id}, {"_id": 0}).sort("created_at", -1).limit(1).to_list(1)
    if not last:
        return {"ok": False}
    last = last[0]
    await db.swipes.delete_one({"user_id": user.user_id, "job_id": last["job_id"]})
    if last["direction"] == "right":
        await db.applications.delete_one({"user_id": user.user_id, "job_id": last["job_id"]})
    return {"ok": True, "job_id": last["job_id"]}


# ===================== Applications =====================

def _application_effective_timestamp(app_doc: Dict[str, Any]) -> datetime:
    for key in ("submitted_at", "manual_status_updated_at", "updated_at", "created_at"):
        value = app_doc.get(key)
        if not value:
            continue
        if isinstance(value, datetime):
            parsed = value
        else:
            try:
                parsed = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
            except ValueError:
                continue
        if parsed.tzinfo is None:
            parsed = parsed.replace(tzinfo=timezone.utc)
        return parsed
    return datetime.min.replace(tzinfo=timezone.utc)


def _sort_applications_newest_first(apps: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    return sorted(apps, key=_application_effective_timestamp, reverse=True)


@api_router.get("/applications")
async def list_applications(user: User = Depends(get_current_user)):
    apps = await db.applications.find({"user_id": user.user_id}, {"_id": 0}).sort("created_at", -1).to_list(500)
    apps = _sort_applications_newest_first(apps)
    # join job data
    job_ids = list({a["job_id"] for a in apps if a.get("job_id")})
    jobs = await db.jobs.find({"job_id": {"$in": job_ids}}, {"_id": 0}).to_list(500) if job_ids else []
    job_map = {j["job_id"]: j for j in jobs}
    result = []
    for a in apps:
        a = _normalize_application_status_fields(a)
        result.append(_public_application_doc(a, job_map.get(a["job_id"])))
    return {"applications": result}


@api_router.get("/applications/auto-apply-queue")
async def list_auto_apply_queue(user: User = Depends(get_current_user)):
    """Waiting list + recent outcomes for the production auto-apply worker."""
    return await auto_apply_queue.list_queue_for_user(db, user.user_id)


@api_router.post("/applications/{application_id}/approve-documents")
async def approve_application_documents(application_id: str, user: User = Depends(get_current_user)):
    """Mark CV and cover letter as reviewed by the candidate."""
    app_doc = await db.applications.find_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"_id": 0},
    )
    if not app_doc:
        raise HTTPException(status_code=404, detail="Application not found")

    now = datetime.now(timezone.utc).isoformat()
    await db.applications.update_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"$set": {
            "document_review_status": "approved",
            "document_review_approved_at": now,
            "updated_at": now,
        }},
    )
    updated = await db.applications.find_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"_id": 0},
    ) or app_doc
    try:
        await auto_apply_queue.release_after_document_approval(db, updated)
        updated = await db.applications.find_one(
            {"application_id": application_id, "user_id": user.user_id},
            {"_id": 0},
        ) or updated
    except Exception as exc:
        logger.warning(
            "auto_apply_release_after_review_failed application_id=%s error=%s",
            application_id,
            str(exc)[:200],
        )
    job = await db.jobs.find_one({"job_id": updated.get("job_id")}, {"_id": 0})
    return _public_application_doc(_normalize_application_status_fields(updated), job)


def _reset_review_status_if_approved(update_fields: Dict[str, Any], app_doc: Dict[str, Any]) -> None:
    """Editing the CV or cover letter after approval must require the user
    to re-approve before it can be submitted again -- otherwise a stale
    approval could wave through content the user never actually reviewed.
    """
    if app_doc.get("document_review_status") == "approved":
        update_fields["document_review_status"] = "awaiting_user"
        update_fields["document_review_approved_at"] = None


@api_router.post("/applications/{application_id}/cv-source")
async def set_application_cv_source(application_id: str, body: CvSourceUpdate, user: User = Depends(get_current_user)):
    """Switch which CV file gets submitted for this application: the
    AI-tailored one (default) or the user's originally-uploaded CV."""
    app_doc = await db.applications.find_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"_id": 0},
    )
    if not app_doc:
        raise HTTPException(status_code=404, detail="Application not found")

    now = datetime.now(timezone.utc).isoformat()
    update_fields: Dict[str, Any] = {"cv_source": body.source, "updated_at": now}

    if body.source == "original":
        profile = await db.profiles.find_one(
            {"user_id": user.user_id},
            {"_id": 0, "cv_original_b64": 1, "cv_filename": 1, "cv_mime": 1},
        )
        if not profile or not profile.get("cv_original_b64"):
            raise HTTPException(status_code=400, detail="Upload a CV to your profile before using it here.")
        update_fields["tailored_cv_file_b64"] = profile["cv_original_b64"]
        update_fields["tailored_cv_filename"] = profile.get("cv_filename")
        update_fields["tailored_cv_mime"] = profile.get("cv_mime")
    else:
        ai_file_b64 = app_doc.get("ai_tailored_cv_file_b64")
        ai_filename = app_doc.get("ai_tailored_cv_filename")
        ai_mime = app_doc.get("ai_tailored_cv_mime")
        if not ai_file_b64:
            # Applications generated before this backup field existed never had
            # one saved. The structured resume data was never touched by the
            # source switch, though, so rebuild the DOCX from it instead of
            # failing -- and save the result as the backup for next time.
            profile = await db.profiles.find_one({"user_id": user.user_id}, {"_id": 0})
            rebuilt = build_application_package(profile or {}, app_doc) if profile else {}
            ai_file_b64 = rebuilt.get("tailored_cv_file_b64")
            ai_filename = rebuilt.get("tailored_cv_filename")
            ai_mime = rebuilt.get("tailored_cv_mime")
            if not ai_file_b64:
                raise HTTPException(status_code=400, detail="No AI-tailored CV is available for this application.")
            update_fields["ai_tailored_cv_file_b64"] = ai_file_b64
            update_fields["ai_tailored_cv_filename"] = ai_filename
            update_fields["ai_tailored_cv_mime"] = ai_mime
        update_fields["tailored_cv_file_b64"] = ai_file_b64
        update_fields["tailored_cv_filename"] = ai_filename
        update_fields["tailored_cv_mime"] = ai_mime

    _reset_review_status_if_approved(update_fields, app_doc)
    await db.applications.update_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"$set": update_fields},
    )
    updated = await db.applications.find_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"_id": 0},
    ) or app_doc
    job = await db.jobs.find_one({"job_id": updated.get("job_id")}, {"_id": 0})
    return _public_application_doc(_normalize_application_status_fields(updated), job)


@api_router.patch("/applications/{application_id}/cover-letter")
async def edit_application_cover_letter(application_id: str, body: CoverLetterEditRequest, user: User = Depends(get_current_user)):
    """Let the user rewrite the AI-generated cover letter body (greeting +
    paragraphs + sign-off + signature) as free text. Letterhead fields
    (subject/sender/recipient/date) are left untouched."""
    app_doc = await db.applications.find_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"_id": 0},
    )
    if not app_doc:
        raise HTTPException(status_code=404, detail="Application not found")

    text = (body.body_text or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="Cover letter text cannot be empty.")
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]

    now = datetime.now(timezone.utc).isoformat()
    update_fields: Dict[str, Any] = {"updated_at": now}
    for key in ("cover_letter", "tailored_cover_letter"):
        letter = dict(app_doc.get(key) or {})
        letter["paragraphs"] = paragraphs
        letter["greeting"] = ""
        letter["sign_off"] = ""
        letter["signature_name"] = ""
        letter["cover_letter_edited"] = True
        update_fields[key] = letter

    _reset_review_status_if_approved(update_fields, app_doc)
    await db.applications.update_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"$set": update_fields},
    )
    updated = await db.applications.find_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"_id": 0},
    ) or app_doc
    job = await db.jobs.find_one({"job_id": updated.get("job_id")}, {"_id": 0})
    return _public_application_doc(_normalize_application_status_fields(updated), job)


@api_router.patch("/applications/{application_id}/resume")
async def edit_application_resume(application_id: str, body: ResumeEditRequest, user: User = Depends(get_current_user)):
    """Let the user rewrite the experience/education/languages content of
    their AI-tailored CV, then regenerate the actual DOCX so what they see
    matches what would be submitted."""
    app_doc = await db.applications.find_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"_id": 0},
    )
    if not app_doc:
        raise HTTPException(status_code=404, detail="Application not found")
    profile = await db.profiles.find_one({"user_id": user.user_id}, {"_id": 0})
    if not profile:
        raise HTTPException(status_code=400, detail="Profile not found")
    job = await db.jobs.find_one({"job_id": app_doc.get("job_id")}, {"_id": 0}) or {}

    existing_resume = dict(app_doc.get("tailored_resume_structured") or app_doc.get("tailored_resume") or {})
    existing_resume["experience"] = [item.model_dump() for item in body.experience]
    existing_resume["education"] = [item.model_dump() for item in body.education]
    existing_resume["languages"] = [lang.strip() for lang in body.languages if lang and lang.strip()]

    generated = {
        "tailored_resume_structured": existing_resume,
        "tailored_cover_letter": app_doc.get("tailored_cover_letter") or app_doc.get("cover_letter") or {},
        "job_title": job.get("title") or "",
    }
    result = build_application_package(profile, generated)

    now = datetime.now(timezone.utc).isoformat()
    update_fields: Dict[str, Any] = {
        "tailored_resume_structured": result["tailored_resume_structured"],
        "tailored_resume": result["tailored_resume_structured"],
        "ai_tailored_cv_file_b64": result["tailored_cv_file_b64"],
        "ai_tailored_cv_filename": result["tailored_cv_filename"],
        "ai_tailored_cv_mime": result["tailored_cv_mime"],
        "updated_at": now,
    }
    # Only overwrite the active submission file if the user isn't currently
    # using their own original CV -- editing the AI text shouldn't silently
    # switch them back to it, but the backup above keeps the edit available
    # whenever they do switch back.
    if (app_doc.get("cv_source") or "tailored") != "original":
        update_fields["tailored_cv_file_b64"] = result["tailored_cv_file_b64"]
        update_fields["tailored_cv_filename"] = result["tailored_cv_filename"]
        update_fields["tailored_cv_mime"] = result["tailored_cv_mime"]

    _reset_review_status_if_approved(update_fields, app_doc)
    await db.applications.update_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"$set": update_fields},
    )
    updated = await db.applications.find_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"_id": 0},
    ) or app_doc
    job_doc = await db.jobs.find_one({"job_id": updated.get("job_id")}, {"_id": 0})
    return _public_application_doc(_normalize_application_status_fields(updated), job_doc)


async def require_admin_user(user: User = Depends(get_current_user)) -> User:
    if _is_admin_email(user.email) or bool(getattr(user, "is_admin", False)):
        return user

    raise HTTPException(status_code=403, detail="Admin access denied")


def _env_email_set(name: str) -> set[str]:
    return {
        item.strip().lower()
        for item in os.environ.get(name, "").split(",")
        if item.strip()
    }


def _env_enabled(name: str, default: str = "false") -> bool:
    return os.environ.get(name, default).strip().lower() in ("1", "true", "yes", "on")


def _agent_real_submit_allowed_emails() -> set[str]:
    explicit = _env_email_set("REAL_SUBMIT_ALLOWED_EMAILS")
    return explicit or _env_email_set("ADMIN_EMAILS")


def _manual_application_fulfillment_enabled() -> bool:
    return _env_enabled("MANUAL_APPLICATION_FULFILLMENT", "true")


def _agent_auto_prepare_enabled(job: Dict[str, Any]) -> bool:
    """Gate for running the apply agent automatically right after package
    generation (Phase 1: supervised prepare, never an automatic submit).

    Kill-switch defaults to on; flip AGENT_AUTO_PREPARE_ENABLED=false on
    Railway if the Nix-chromium launch path turns out broken -- this runs
    inside the same per-user background queue as CV generation, so a hung
    Playwright launch would delay every subsequent queued application for
    that user, not just fail silently.

    Skipped for jobs already known to require login or be otherwise blocked
    (tier D/E, or job_board_account_required) -- launching a browser for
    those would just burn time/cost to rediscover something we already know.
    """
    if not _env_enabled("AGENT_AUTO_PREPARE_ENABLED", "true"):
        return False
    if job.get("job_board_account_required"):
        return False
    tier = str(job.get("applyability_tier") or "").upper()
    if tier in {"D", "E"}:
        return False
    return True


def _require_job_maintenance_enabled() -> None:
    if not job_cache_env_bool("JOBS_MAINTENANCE_ENABLED", True):
        raise HTTPException(status_code=403, detail="Job maintenance endpoints are disabled.")


def _require_ats_direct_enabled() -> None:
    _require_job_maintenance_enabled()
    if not job_cache_env_bool("JOBS_ATS_DIRECT_ENABLED", True):
        raise HTTPException(status_code=403, detail="Direct ATS refresh endpoints are disabled.")


@api_router.post("/admin/jobs/clear-feed-provider-cooldown")
async def admin_clear_feed_provider_cooldown(admin: User = Depends(require_admin_user)):
    global _feed_sync_refresh_cooldown_until, _feed_sync_refresh_cooldowns
    now = time.monotonic()
    previous_feed_until = _feed_sync_refresh_cooldown_until
    previous_scoped_cooldowns = dict(_feed_sync_refresh_cooldowns)
    provider_cooldowns = {
        key: value.isoformat()
        for key, value in getattr(jobs_service_module, "_PROVIDER_COOLDOWN_UNTIL", {}).items()
    }
    _feed_sync_refresh_cooldown_until = 0.0
    _feed_sync_refresh_cooldowns = {}
    if hasattr(jobs_service_module, "_PROVIDER_COOLDOWN_UNTIL"):
        jobs_service_module._PROVIDER_COOLDOWN_UNTIL.clear()
    logger.info(
        "admin_clear_feed_provider_cooldown admin=%s previous_feed_active=%s scoped_cooldowns=%s provider_cooldowns=%s",
        admin.email,
        previous_feed_until > now,
        len(previous_scoped_cooldowns),
        list(provider_cooldowns.keys()),
    )
    return {
        "cleared": True,
        "previous_feed_cooldown_active": previous_feed_until > now or any(until > now for until in previous_scoped_cooldowns.values()),
        "previous_feed_cooldown_remaining_seconds": max(0, int(previous_feed_until - now)),
        "previous_scoped_feed_cooldowns": {
            key: max(0, int(until - now))
            for key, until in previous_scoped_cooldowns.items()
            if until > now
        },
        "previous_provider_cooldowns": provider_cooldowns,
    }


@api_router.post("/admin/jobs/refresh")
async def admin_jobs_refresh(body: AdminJobsRefreshRequest, admin: User = Depends(require_admin_user)):
    _require_job_maintenance_enabled()
    logger.info(
        "admin_jobs_refresh_requested admin=%s role=%s location=%s country_code=%s dry_run=%s",
        admin.email,
        body.search_role,
        body.location,
        body.country_code,
        body.dry_run,
    )
    return await refresh_jobs_for_query_or_filters(
        db,
        search_role=body.search_role,
        location=body.location,
        country_code=body.country_code,
        lat=body.lat,
        lng=body.lng,
        search_radius=body.search_radius,
        include_cross_border=body.include_cross_border,
        discover_ats_sources=body.discover_ats_sources,
        refresh_discovered_ats_sources=body.refresh_discovered_ats_sources,
        ats_refresh_limit=body.ats_refresh_limit,
        remote=body.remote,
        limit=body.limit,
        dry_run=body.dry_run,
    )


@api_router.post("/admin/jobs/revalidate")
async def admin_jobs_revalidate(body: AdminJobsRevalidateRequest, admin: User = Depends(require_admin_user)):
    _require_job_maintenance_enabled()
    logger.info(
        "admin_jobs_revalidate_requested admin=%s validation_status=%s tier=%s country_code=%s dry_run=%s",
        admin.email,
        body.validation_status,
        body.applyability_tier,
        body.country_code,
        body.dry_run,
    )
    return await revalidate_cached_jobs(
        db,
        validation_status=body.validation_status,
        applyability_tier=body.applyability_tier,
        older_than_hours=body.older_than_hours,
        country_code=body.country_code,
        limit=body.limit,
        dry_run=body.dry_run,
    )


@api_router.post("/admin/jobs/expire-stale")
async def admin_jobs_expire_stale(body: AdminJobsExpireStaleRequest, admin: User = Depends(require_admin_user)):
    _require_job_maintenance_enabled()
    logger.info(
        "admin_jobs_expire_stale_requested admin=%s older_than_days=%s provider=%s country_code=%s dry_run=%s",
        admin.email,
        body.older_than_days,
        body.provider,
        body.country_code,
        body.dry_run,
    )
    return await expire_stale_jobs(
        db,
        older_than_days=body.older_than_days,
        provider=body.provider,
        country_code=body.country_code,
        limit=body.limit,
        dry_run=body.dry_run,
        completeness_run_id=body.completeness_run_id,
    )


@api_router.post("/admin/jobs/purge-invalid")
async def admin_jobs_purge_invalid(body: AdminJobsPurgeInvalidRequest, admin: User = Depends(require_admin_user)):
    """Expire stale jobs then hard-delete invalid/D/E inventory (egress + storage)."""
    _require_job_maintenance_enabled()
    logger.info(
        "admin_jobs_purge_invalid_requested admin=%s older_than_days=%s tiers=%s expire_first=%s dry_run=%s",
        admin.email,
        body.older_than_days,
        body.applyability_tiers,
        body.expire_first,
        body.dry_run,
    )
    return await purge_invalid_jobs(
        db,
        older_than_days=body.older_than_days,
        applyability_tiers=body.applyability_tiers,
        expire_first=body.expire_first,
        country_code=body.country_code,
        limit=body.limit,
        dry_run=body.dry_run,
        completeness_run_id=body.completeness_run_id,
    )


@api_router.post("/admin/jobs/maintenance")
async def admin_jobs_maintenance(body: AdminJobsMaintenanceRequest, admin: User = Depends(require_admin_user)):
    _require_job_maintenance_enabled()
    logger.info(
        "admin_jobs_maintenance_requested admin=%s dry_run=%s refresh_popular=%s",
        admin.email,
        body.dry_run,
        body.refresh_popular,
    )
    return await run_job_cache_maintenance(
        db,
        dry_run=body.dry_run,
        refresh_popular=body.refresh_popular,
    )


@api_router.get("/admin/jobs/cache-status")
async def admin_jobs_cache_status(admin: User = Depends(require_admin_user)):
    _require_job_maintenance_enabled()
    logger.info("admin_jobs_cache_status_requested admin=%s", admin.email)
    return await job_cache_status(db)


@api_router.get("/admin/jobs/inventory")
async def admin_jobs_inventory(
    admin: User = Depends(require_admin_user),
    days: int = Query(30, ge=7, le=90),
):
    _require_job_maintenance_enabled()
    logger.info("admin_jobs_inventory_requested admin=%s days=%s", admin.email, days)
    return await job_inventory_analytics(db, days=days)


@api_router.post("/admin/jobs/france-travail/harvest")
async def admin_jobs_france_travail_harvest(
    admin: User = Depends(require_admin_user),
    max_queries: Optional[int] = None,
    dry_run: bool = False,
):
    _require_job_maintenance_enabled()
    logger.info(
        "admin_jobs_ft_harvest_requested admin=%s max_queries=%s dry_run=%s",
        admin.email,
        max_queries,
        dry_run,
    )
    return await harvest_france_travail(db, max_queries=max_queries, dry_run=dry_run)


@api_router.get("/admin/jobs/france-travail/harvest-status")
async def admin_jobs_france_travail_harvest_status(admin: User = Depends(require_admin_user)):
    _require_job_maintenance_enabled()
    return {
        "harvest_enabled": ft_harvest_enabled(),
        "last_run": ft_last_harvest_summary(),
    }


@api_router.post("/admin/jobs/jsearch/harvest")
async def admin_jobs_jsearch_harvest(
    admin: User = Depends(require_admin_user),
    max_queries: Optional[int] = None,
    dry_run: bool = False,
    start_offset: Optional[int] = None,
    mode: Optional[str] = None,
    cities: Optional[str] = None,
    roles: Optional[str] = None,
    date_posted: Optional[str] = None,
    page_size: Optional[int] = None,
    max_pages: Optional[int] = None,
):
    _require_job_maintenance_enabled()
    cities_list = [c.strip() for c in cities.split(",") if c.strip()] if cities else None
    roles_list = [r.strip() for r in roles.split(",") if r.strip()] if roles else None
    if mode == "aggressive":
        cities_list = cities_list or list(AGGRESSIVE_HARVEST_CITIES)
        roles_list = roles_list or list(AGGRESSIVE_HARVEST_ROLES)
        date_posted = date_posted or "month"
        page_size = page_size if page_size is not None else 100
        max_pages = max_pages if max_pages is not None else 5
    logger.info(
        "admin_jobs_jsearch_harvest_requested admin=%s max_queries=%s dry_run=%s start_offset=%s mode=%s cities=%s roles=%s date_posted=%s page_size=%s max_pages=%s",
        admin.email,
        max_queries,
        dry_run,
        start_offset,
        mode,
        len(cities_list) if cities_list else None,
        len(roles_list) if roles_list else None,
        date_posted,
        page_size,
        max_pages,
    )
    return await harvest_jsearch(
        db,
        max_queries=max_queries,
        dry_run=dry_run,
        start_offset=start_offset,
        cities=cities_list,
        roles=roles_list,
        date_posted=date_posted,
        page_size=page_size,
        max_pages=max_pages,
    )


@api_router.get("/admin/jobs/jsearch/harvest-status")
async def admin_jobs_jsearch_harvest_status(admin: User = Depends(require_admin_user)):
    _require_job_maintenance_enabled()
    return {
        "harvest_enabled": jsearch_harvest_enabled(),
        "last_run": jsearch_last_harvest_summary(),
    }


@api_router.post("/admin/jobs/company-discovery")
async def admin_jobs_company_discovery(
    admin: User = Depends(require_admin_user),
    dry_run: bool = False,
    start_offset: Optional[int] = None,
    max_companies: Optional[int] = None,
):
    _require_job_maintenance_enabled()
    logger.info(
        "admin_company_discovery_requested admin=%s dry_run=%s start_offset=%s max_companies=%s",
        admin.email,
        dry_run,
        start_offset,
        max_companies,
    )
    return await run_company_discovery(
        db,
        dry_run=dry_run,
        company_list_start_offset=start_offset,
        company_list_max_companies=max_companies,
    )


@api_router.get("/admin/jobs/company-discovery/status")
async def admin_jobs_company_discovery_status(admin: User = Depends(require_admin_user)):
    _require_job_maintenance_enabled()
    return {"last_run": company_discovery_last_summary()}


class AdminAutoApplyExecuteRequest(BaseModel):
    job_id: str
    user_id: Optional[str] = None
    dry_run: bool = False
    headless: Optional[bool] = None


async def _auto_apply_target_user(user_id: Optional[str], admin: User) -> User:
    if not user_id or user_id == admin.user_id:
        return admin
    loaded = await db.users.find_one({"user_id": user_id})
    if not loaded:
        raise HTTPException(status_code=404, detail="user_not_found")
    return User(user_id=loaded["user_id"], email=loaded.get("email", ""), name=loaded.get("name", ""))


_AUTO_APPLY_BACKGROUND_TASKS: set = set()

# Real (non-dry-run) auto-apply outcomes that should be mirrored onto the
# application record so the admin Applications tab reflects what happened —
# success means no more manual work is needed; the rest routes into the
# existing manual-review queue (_effective_manual_status already promotes
# submission_status="failed" to "manual_review_needed").
_AUTO_APPLY_RESULT_SUBMITTED_STATUSES = {"submitted_success", "already_submitted"}
_AUTO_APPLY_RESULT_NEEDS_MANUAL_STATUSES = {
    "submit_failed", "verification_failed", "error", "unsupported", "needs_user_input",
}


async def _sync_application_status_from_auto_apply_result(
    application_id: str, user_id: str, result: Dict[str, Any],
) -> None:
    status = (result or {}).get("status")
    if status not in _AUTO_APPLY_RESULT_SUBMITTED_STATUSES and status not in _AUTO_APPLY_RESULT_NEEDS_MANUAL_STATUSES:
        return
    now = datetime.now(timezone.utc).isoformat()
    if status in _AUTO_APPLY_RESULT_SUBMITTED_STATUSES:
        update = {"submission_status": "submitted", "submitted_at": now, "updated_at": now}
    else:
        update = {"submission_status": "failed", "updated_at": now}
    await db.applications.update_one(
        {"application_id": application_id, "user_id": user_id},
        {"$set": update},
    )


async def _admin_auto_apply_background_run(
    *,
    job_id: str,
    target_user: User,
    dry_run: bool,
    headless: Optional[bool],
    run_id: Optional[str] = None,
) -> None:
    """Run prepare + execute off the HTTP request so Railway's gateway cannot 502."""
    from auto_apply.debug_report import format_run_error, transport_error_report

    application_id: Optional[str] = None
    result: Dict[str, Any]
    try:
        job, profile, app_doc = await _load_or_create_agent_application(
            job_id,
            target_user,
            require_tailored_package=False,
        )
        application_id = app_doc.get("application_id")
        result = await auto_apply_execute_application(
            db, job, profile, app_doc, target_user.model_dump(mode="json"),
            dry_run=dry_run,
            headless=_resolve_auto_apply_headless(headless),
        )
        if run_id and isinstance(result, dict):
            result.setdefault("debug", {})
            if isinstance(result["debug"], dict):
                result["debug"]["run_id"] = run_id
            result["run_id"] = run_id
    except HTTPException as http_exc:
        detail = http_exc.detail
        message = detail if isinstance(detail, str) else (
            (detail or {}).get("message") if isinstance(detail, dict) else str(detail)
        )
        result = transport_error_report(
            message=str(message or "Request failed"),
            phase="prepare",
            stage="driver",
            http_status=http_exc.status_code,
            exception_class="HTTPException",
        )
    except Exception as exc:
        error_detail = format_run_error(exc, checkpoint="execute")
        logger.warning(
            "admin_auto_apply_background_failed job=%s error=%s",
            job_id,
            error_detail.get("message", str(exc))[:300],
        )
        result = transport_error_report(
            message=error_detail.get("message") or str(exc),
            phase=str(error_detail.get("phase") or "execute"),
            stage="driver",
            exception_class=error_detail.get("exception_class") or exc.__class__.__name__,
            extra={"traceback": error_detail.get("traceback"), "hint": error_detail.get("hint")},
        )

    try:
        await auto_apply_persist_execution_report(db, target_user.user_id, job_id, result)
    except Exception as exc:
        logger.warning(
            "admin_auto_apply_persist_report_failed job=%s error=%s",
            job_id,
            str(exc)[:200],
        )

    if not dry_run and application_id:
        try:
            await _sync_application_status_from_auto_apply_result(application_id, target_user.user_id, result)
        except Exception as exc:
            logger.warning(
                "admin_auto_apply_status_sync_failed job=%s application_id=%s error=%s",
                job_id,
                application_id,
                str(exc)[:200],
            )


@api_router.post("/admin/auto-apply/execute")
async def admin_auto_apply_execute(body: AdminAutoApplyExecuteRequest, admin: User = Depends(require_admin_user)):
    """Start an auto-apply run in the background and return immediately.

    Browser + residential-proxy retries regularly exceed Railway's HTTP gateway
    limit (~100s). The admin console polls `/admin/auto-apply/status` until the
    attempt leaves `in_flight` and reads `execution_report`.
    """
    from auto_apply.debug_report import format_run_error, transport_error_report

    try:
        target_user = await _auto_apply_target_user(body.user_id, admin)
        job = await db.jobs.find_one({"job_id": body.job_id}, {"_id": 0, "job_id": 1})
        if not job:
            return {
                "accepted": False,
                "polling": False,
                "result": transport_error_report(
                    message="Job not found",
                    phase="prepare",
                    stage="driver",
                    http_status=404,
                    exception_class="HTTPException",
                ),
                "attempt": None,
            }

        started_at = datetime.now(timezone.utc).isoformat()
        run_id = uuid.uuid4().hex
        # Orphan in_flight rows (killed deploy/worker) block the unique claim index.
        try:
            from auto_apply.metrics import release_stale_in_flight as _release_stale

            released = await _release_stale_in_flight(
                db, target_user.user_id, body.job_id, max_age_s=120.0, force=False,
            )
            if released:
                logger.info(
                    "admin_auto_apply_released_stale user=%s job=%s n=%s",
                    target_user.user_id,
                    body.job_id,
                    released,
                )
        except Exception as release_exc:
            logger.warning("admin_auto_apply_release_stale_failed error=%s", str(release_exc)[:200])

        task = asyncio.create_task(
            _admin_auto_apply_background_run(
                job_id=body.job_id,
                target_user=target_user,
                dry_run=body.dry_run,
                headless=body.headless,
                run_id=run_id,
            )
        )
        _AUTO_APPLY_BACKGROUND_TASKS.add(task)
        task.add_done_callback(_AUTO_APPLY_BACKGROUND_TASKS.discard)

        return {
            "accepted": True,
            "polling": True,
            "started_at": started_at,
            "run_id": run_id,
            "job_id": body.job_id,
            "user_id": target_user.user_id,
            "result": {
                "status": "in_flight",
                "stage_reached": "driver",
                "reason": "run_started",
                "debug": {
                    "timeline": [{
                        "stage": "driver",
                        "status": "ok",
                        "detail": "Run accepted — browser pipeline running in background",
                    }],
                },
            },
            "attempt": {
                "status": "in_flight",
                "stage_reached": "driver",
                "job_id": body.job_id,
                "user_id": target_user.user_id,
                "claimed_at": started_at,
                "run_id": run_id,
            },
        }
    except HTTPException as http_exc:
        detail = http_exc.detail
        message = detail if isinstance(detail, str) else str(detail)
        return {
            "accepted": False,
            "polling": False,
            "result": transport_error_report(
                message=str(message or "Request failed"),
                phase="prepare",
                stage="driver",
                http_status=http_exc.status_code,
                exception_class="HTTPException",
            ),
            "attempt": None,
        }
    except Exception as exc:
        error_detail = format_run_error(exc, checkpoint="execute_accept")
        logger.exception("admin_auto_apply_execute_failed job=%s", body.job_id)
        return {
            "accepted": False,
            "polling": False,
            "result": transport_error_report(
                message=error_detail.get("message") or str(exc),
                phase=str(error_detail.get("phase") or "execute"),
                stage="driver",
                exception_class=error_detail.get("exception_class") or exc.__class__.__name__,
                extra={"traceback": error_detail.get("traceback"), "hint": error_detail.get("hint")},
            ),
            "attempt": None,
        }


@api_router.get("/admin/auto-apply/status")
async def admin_auto_apply_status(job_id: str, user_id: Optional[str] = None, admin: User = Depends(require_admin_user)):
    """Latest attempt record for a job -- exposes stage, status, reason,
    missing_fields, driver_version, blueprint_signature, lifecycle timestamps,
    and the persisted execution_report (when the async run finished)."""
    from auto_apply.debug_report import format_run_error, transport_error_report

    target = user_id or admin.user_id
    try:
        attempt = await auto_apply_latest_attempt(db, target, job_id)
        return {"attempt": auto_apply_status_safe_attempt(attempt)}
    except Exception as exc:
        error_detail = format_run_error(exc, checkpoint="status")
        logger.exception("admin_auto_apply_status_failed job=%s", job_id)
        return {
            "attempt": {
                "status": "error",
                "stage_reached": "driver",
                "reason": error_detail.get("message") or str(exc),
                "execution_report": transport_error_report(
                    message=error_detail.get("message") or str(exc),
                    phase="status",
                    stage="driver",
                    exception_class=error_detail.get("exception_class") or exc.__class__.__name__,
                ),
            },
        }


@api_router.get("/admin/auto-apply/metrics")
async def admin_auto_apply_metrics(provider: str = "greenhouse", admin: User = Depends(require_admin_user)):
    return await auto_apply_metrics_summary(db, provider)


@api_router.get("/admin/auto-apply/right-swipes")
async def admin_auto_apply_right_swipes(
    limit: int = Query(default=100, ge=1, le=300),
    admin: User = Depends(require_admin_user),
):
    """Recent right swipes across all users, joined with user + job data, driver
    support and the latest auto-apply attempt. Powers the admin auto-apply test
    bench: each row can be replayed through the production pipeline."""
    from auto_apply.driver import DRIVER_REGISTRY

    swipes = await _admin_safe_find(
        db.swipes, {"direction": "right"}, limit=limit, sort=[("created_at", -1)],
    )
    job_ids = list(dict.fromkeys(s.get("job_id") for s in swipes if s.get("job_id")))
    user_ids = list(dict.fromkeys(s.get("user_id") for s in swipes if s.get("user_id")))

    jobs = {j.get("job_id"): j for j in await _admin_jobs_for_ids(job_ids)}

    users: Dict[str, Dict[str, Any]] = {}
    chunk_size = 80
    for index in range(0, len(user_ids), chunk_size):
        chunk = user_ids[index:index + chunk_size]
        for row in await _admin_safe_find(db.users, {"user_id": {"$in": chunk}}, limit=len(chunk)):
            users[row.get("user_id")] = row

    latest_attempts: Dict[tuple, Dict[str, Any]] = {}
    for index in range(0, len(job_ids), chunk_size):
        chunk = job_ids[index:index + chunk_size]
        rows = await _admin_safe_find(db.auto_apply_attempts, {"job_id": {"$in": chunk}}, limit=2000)
        for row in rows:
            key = (row.get("user_id"), row.get("job_id"))
            previous = latest_attempts.get(key)
            if previous is None or (row.get("created_at") or "") > (previous.get("created_at") or ""):
                latest_attempts[key] = row

    applications_by_key: Dict[tuple, Dict[str, Any]] = {}
    for index in range(0, len(job_ids), chunk_size):
        chunk = job_ids[index:index + chunk_size]
        rows = await _admin_safe_find(db.applications, {"job_id": {"$in": chunk}}, limit=2000)
        for row in rows:
            key = (row.get("user_id"), row.get("job_id"))
            previous = applications_by_key.get(key)
            if previous is None or (row.get("created_at") or "") > (previous.get("created_at") or ""):
                applications_by_key[key] = row

    items: List[Dict[str, Any]] = []
    for swipe_row in swipes:
        job = jobs.get(swipe_row.get("job_id")) or {}
        user_row = users.get(swipe_row.get("user_id")) or {}
        attempt = latest_attempts.get((swipe_row.get("user_id"), swipe_row.get("job_id")))
        application = applications_by_key.get((swipe_row.get("user_id"), swipe_row.get("job_id")))
        normalized_application = _normalize_application_status_fields(application) if application else None
        items.append({
            "user_id": swipe_row.get("user_id"),
            "user_email": user_row.get("email") or "",
            "user_name": user_row.get("name") or "",
            "job_id": swipe_row.get("job_id"),
            "job_found": bool(job),
            "title": job.get("title") or "",
            "company": job.get("company") or "",
            "ats_provider": str(job.get("ats_provider") or job.get("provider") or "").lower() or "unknown",
            "application_id": (application or {}).get("application_id"),
            "has_application": bool(application),
            "submission_status": (normalized_application or {}).get("submission_status") or "not_submitted",
            # Prefer the public posting URL so admins can open the ATS and check expiry.
            "apply_url": (
                job.get("external_url")
                or job.get("selected_apply_url")
                or job.get("job_apply_link")
                or job.get("apply_url")
                or ""
            ),
            "external_url": job.get("external_url") or "",
            "driver_supported": bool(job) and DRIVER_REGISTRY.for_job(job) is not None,
            "swiped_at": swipe_row.get("created_at"),
            "latest_attempt": {
                "status": attempt.get("status"),
                "stage_reached": attempt.get("stage_reached"),
                "reason": attempt.get("reason"),
                "verdict": attempt.get("verdict"),
                "missing_fields": attempt.get("missing_fields") or [],
                "updated_at": attempt.get("updated_at"),
            } if attempt else None,
        })
    return {"swipes": items, "supported_providers": DRIVER_REGISTRY.providers()}


class AdminAutoApplyValidateRequest(BaseModel):
    greenhouse_url: str
    resume_b64: Optional[str] = None
    resume_filename: Optional[str] = None
    cover_letter_text: Optional[str] = None
    additional_answers: Dict[str, Any] = Field(default_factory=dict)
    contact: Dict[str, Any] = Field(default_factory=dict)
    dry_run: bool = True
    headless: Optional[bool] = None


def _greenhouse_job_from_url(url: str) -> Dict[str, Any]:
    """Build an ephemeral job doc from a pasted Greenhouse posting URL. Test-only
    plumbing for the validation flow -- the pipeline itself stays generic."""
    from urllib.parse import urlparse
    from job_providers.ats_adapters.greenhouse import GreenhouseAtsAdapter

    token = GreenhouseAtsAdapter().extract_source_key_from_url(url)
    parts = [p for p in urlparse(url).path.split("/") if p]
    job_id = None
    if "jobs" in parts:
        idx = parts.index("jobs")
        if idx + 1 < len(parts):
            job_id = parts[idx + 1]
    if not job_id:
        digits = [p for p in parts if p.isdigit()]
        job_id = digits[-1] if digits else None
    if not token or not job_id:
        raise HTTPException(status_code=400, detail="invalid_greenhouse_url")
    return {
        "job_id": f"gh:{token}:{job_id}",
        "ats_provider": "greenhouse",
        "provider": "greenhouse",
        "external_url": url,
        "board_token": token,
        "provider_job_id": job_id,
        "company": token,
        "title": "Greenhouse validation job",
    }


def _admin_contact(admin: User) -> Dict[str, Any]:
    parts = [p for p in (admin.name or "").split() if p]
    return {
        "first_name": parts[0] if parts else "",
        "last_name": " ".join(parts[1:]) if len(parts) > 1 else "",
        "email": admin.email or "",
    }


@api_router.post("/admin/auto-apply-lab/execute")
async def admin_auto_apply_lab_execute(body: AdminAutoApplyValidateRequest, admin: User = Depends(require_admin_user)):
    """Minimal end-to-end validation: build an ephemeral job + application context
    from a pasted Greenhouse URL, uploaded resume/cover letter, and typed answers,
    then run the production execute_application and return its ExecutionReport."""
    job = _greenhouse_job_from_url(body.greenhouse_url)
    profile = {
        "contact": {**_admin_contact(admin), **(body.contact or {})},
        "application_answers_profile": body.additional_answers or {},
    }
    app_doc: Dict[str, Any] = {"application_id": f"validation:{job['job_id']}"}
    if body.resume_b64:
        app_doc["tailored_cv_file_b64"] = body.resume_b64
        app_doc["tailored_cv_filename"] = body.resume_filename or "resume.pdf"
    if body.cover_letter_text:
        app_doc["cover_letter"] = {"paragraphs": [body.cover_letter_text]}
    return await auto_apply_execute_application(
        db, job, profile, app_doc, admin.model_dump(mode="json"),
        dry_run=body.dry_run,
        headless=_resolve_auto_apply_headless(body.headless),
    )


@api_router.post("/admin/jobs/feed-diagnostic")
async def admin_jobs_feed_diagnostic(body: AdminJobsFeedDiagnosticRequest, admin: User = Depends(require_admin_user)):
    _require_job_maintenance_enabled()
    started = time.perf_counter()
    strict_query = {"validation_status": "valid", "applyability_tier": {"$in": ["A", "B"]}}
    legacy_query = {"provider": {"$in": ["greenhouse", "lever", "ashby"]}, "auto_apply_supported": True}
    strict_rows = await db.jobs.find(strict_query, {"_id": 0}).limit(1000).to_list(1000)
    legacy_rows = await db.jobs.find(legacy_query, {"_id": 0}).limit(1000).to_list(1000)
    legacy_applyable = [_with_apply_fulfillment_fields(job) for job in legacy_rows if _job_is_applyable(job)]
    provider_attempted = False
    provider_count = 0
    provider_error = None
    if is_job_provider_configured(primary_job_provider_name()):
        try:
            provider = get_configured_job_provider()
            query = JobSearchQuery(
                role=body.search_role or "sales",
                location=body.location,
                country=(body.country_code or os.environ.get("JSEARCH_COUNTRY", "fr")).lower(),
                language=os.environ.get("JSEARCH_LANGUAGE", "fr"),
                limit=max(1, min(body.limit, 20)),
                max_pages=1,
                page_size=max(5, min(body.limit * 2, 20)),
            )
            provider_attempted = True
            result = await provider.search(query)
            provider_count = len(result.jobs)
        except Exception as exc:
            provider_error = f"{exc.__class__.__name__}: {str(exc)[:160]}"
    return {
        "db_strict_count": len(strict_rows),
        "db_legacy_direct_count": len(legacy_rows),
        "db_legacy_applyable_count": len(legacy_applyable),
        "jsearch_attempted": provider_attempted,
        "jsearch_count": provider_count,
        "jsearch_error": provider_error,
        "provider_attempted": provider_attempted,
        "provider_count": provider_count,
        "provider_error": provider_error,
        "provider_name": primary_job_provider_name(),
        "final_new_feed_count_estimate": min(len(strict_rows) or len(legacy_applyable), body.limit),
        "final_legacy_feed_count_estimate": min(provider_count, body.limit),
        "cooldown": {
            "feed_active": time.monotonic() < _feed_sync_refresh_cooldown_until,
            "feed_remaining_seconds": max(0, int(_feed_sync_refresh_cooldown_until - time.monotonic())),
            "feed_scoped_active_count": sum(1 for until in _feed_sync_refresh_cooldowns.values() if until > time.monotonic()),
            "feed_scoped_cooldowns": {
                key: max(0, int(until - time.monotonic()))
                for key, until in _feed_sync_refresh_cooldowns.items()
                if until > time.monotonic()
            },
            "provider_cooldowns": list(getattr(jobs_service_module, "_PROVIDER_COOLDOWN_UNTIL", {}).keys()),
        },
        "timing_ms": int((time.perf_counter() - started) * 1000),
    }


@api_router.post("/admin/jobs/feed-coverage-audit")
async def admin_jobs_feed_coverage_audit(
    body: AdminFeedCoverageAuditRequest,
    admin: User = Depends(require_admin_user),
):
    """Evaluate a fixed paid-user snapshot through the real feed with I/O disabled."""
    requested_user_ids = list(dict.fromkeys(user_id.strip() for user_id in body.user_ids if user_id.strip()))
    if not requested_user_ids:
        raise HTTPException(status_code=400, detail="At least one user_id is required")
    if len(requested_user_ids) > 100:
        raise HTTPException(status_code=400, detail="Coverage audits are limited to 100 users")

    evaluated_at = datetime.now(timezone.utc)
    user_rows = await db.users.find(
        {"user_id": {"$in": requested_user_ids}},
        {"_id": 0},
    ).limit(len(requested_user_ids)).to_list(len(requested_user_ids))
    users_by_id = {str(row.get("user_id")): row for row in user_rows if row.get("user_id")}
    snapshots: List[Dict[str, Any]] = []

    for user_id in requested_user_ids:
        user_doc = users_by_id.get(user_id)
        if not user_doc:
            snapshots.append({
                "user_id": hash_feed_coverage_user_id(user_id),
                "evaluated_at": evaluated_at.isoformat(),
                "terminal_reason": "USER_NOT_FOUND",
                "evaluator_version": FEED_COVERAGE_EVALUATOR_VERSION,
            })
            continue
        data = user_doc.get("data") if isinstance(user_doc.get("data"), dict) else {}
        billing = user_doc.get("billing") if isinstance(user_doc.get("billing"), dict) else data.get("billing") or {}
        if str(billing.get("subscription_status") or "").lower() not in {"active", "trialing"}:
            snapshots.append({
                "user_id": hash_feed_coverage_user_id(user_id),
                "evaluated_at": evaluated_at.isoformat(),
                "terminal_reason": "NOT_PAID_COHORT",
                "evaluator_version": FEED_COVERAGE_EVALUATOR_VERSION,
            })
            continue
        profile = await db.profiles.find_one({"user_id": user_id}, {"_id": 0}) or {}
        audit_user = User(
            user_id=user_id,
            email=str(user_doc.get("email") or f"{user_id}@coverage-audit.invalid"),
            name=str(user_doc.get("name") or "Coverage audit"),
            is_admin=False,
        )
        try:
            feed_response = await get_feed(
                user=audit_user,
                limit=max(1, min(body.limit, 25)),
                min_salary=0,
                posted_within=None,
                work_location=None,
                job_type=None,
                experience=None,
                location=None,
                only_company=None,
                hide_company=None,
                only_industry=None,
                hide_industry=None,
                include_unknown_location=True,
                include_unknown_salary=True,
                include_non_auto_apply=False,
                search_radius="50km",
                locations_json=None,
                only_my_country=False,
                location_label=None,
                place_id=None,
                country=None,
                country_code=None,
                lat=None,
                lng=None,
                force_provider_refresh=False,
                prefetch=False,
                score=False,
                search_role=None,
                audit_mode=True,
            )
        except HTTPException as exc:
            snapshots.append({
                "user_id": hash_feed_coverage_user_id(user_id),
                "evaluated_at": evaluated_at.isoformat(),
                "terminal_reason": f"PROFILE_NOT_READY_{exc.status_code}",
                "evaluator_version": FEED_COVERAGE_EVALUATOR_VERSION,
            })
            continue
        snapshots.append(build_feed_coverage_snapshot(
            user_id=user_id,
            profile=profile,
            feed_response=feed_response,
            evaluated_at=evaluated_at,
            freshness_window_days=body.freshness_window_days,
        ))

    coverage_run_id = hashlib.sha256(
        "|".join([
            FEED_COVERAGE_EVALUATOR_VERSION,
            evaluated_at.isoformat(),
            *(snapshot["user_id"] for snapshot in snapshots),
        ]).encode("utf-8")
    ).hexdigest()[:24]
    for snapshot in snapshots:
        snapshot["coverage_run_id"] = coverage_run_id
    return {
        "coverage_run_id": coverage_run_id,
        "evaluated_at": evaluated_at.isoformat(),
        "evaluator_version": FEED_COVERAGE_EVALUATOR_VERSION,
        "profile_count": len(snapshots),
        "snapshots": snapshots,
    }


@api_router.post("/admin/jobs/ats/discover-sources")
async def admin_jobs_ats_discover_sources(body: AdminAtsDiscoverSourcesRequest, admin: User = Depends(require_admin_user)):
    _require_ats_direct_enabled()
    logger.info(
        "admin_jobs_ats_discover_requested admin=%s provider=%s country_code=%s limit=%s dry_run=%s",
        admin.email,
        body.provider,
        body.country_code,
        body.limit,
        body.dry_run,
    )
    return await discover_ats_sources_from_cached_jobs(
        db,
        provider=body.provider,
        country_code=body.country_code,
        limit=body.limit,
        dry_run=body.dry_run,
    )


@api_router.post("/admin/jobs/ats/refresh-source")
async def admin_jobs_ats_refresh_source(body: AdminAtsRefreshSourceRequest, admin: User = Depends(require_admin_user)):
    _require_ats_direct_enabled()
    logger.info(
        "admin_jobs_ats_refresh_source_requested admin=%s provider=%s source_key=%s limit=%s dry_run=%s",
        admin.email,
        body.ats_provider,
        body.source_key,
        body.limit,
        body.dry_run,
    )
    return await refresh_ats_source(
        db,
        ats_provider=body.ats_provider,
        source_key=body.source_key,
        limit=body.limit,
        dry_run=body.dry_run,
    )


@api_router.post("/admin/jobs/ats/discover-friendly-company-pages")
async def admin_jobs_discover_friendly_company_pages(
    body: AdminDiscoverFriendlyCompanyPagesRequest, admin: User = Depends(require_admin_user)
):
    _require_ats_direct_enabled()
    logger.info(
        "admin_jobs_discover_friendly_company_pages_requested admin=%s limit=%s concurrency=%s dry_run=%s",
        admin.email,
        body.limit,
        body.concurrency,
        body.dry_run,
    )
    return await discover_friendly_company_career_pages(
        db,
        limit=body.limit,
        concurrency=body.concurrency,
        dry_run=body.dry_run,
    )


@api_router.post("/admin/jobs/ats/refresh-known-sources")
async def admin_jobs_ats_refresh_known_sources(body: AdminAtsRefreshKnownSourcesRequest, admin: User = Depends(require_admin_user)):
    _require_ats_direct_enabled()
    logger.info(
        "admin_jobs_ats_refresh_known_requested admin=%s provider=%s country_code=%s limit=%s older_than_hours=%s dry_run=%s",
        admin.email,
        body.provider,
        body.country_code,
        body.limit,
        body.older_than_hours,
        body.dry_run,
    )
    return await refresh_known_ats_sources(
        db,
        provider=body.provider,
        country_code=body.country_code,
        limit=body.limit,
        older_than_hours=body.older_than_hours,
        dry_run=body.dry_run,
    )


def _require_agent_real_submit_allowed(user: User) -> None:
    if not _env_enabled("AGENT_REAL_SUBMIT_ENABLED", "false"):
        raise HTTPException(status_code=403, detail="Real auto-apply submit is disabled.")
    allowed_emails = _agent_real_submit_allowed_emails()
    user_email = (user.email or "").strip().lower()
    if not user_email or user_email not in allowed_emails:
        logger.warning("agent_real_submit_denied user_id=%s email=%s", user.user_id, user.email)
        raise HTTPException(status_code=403, detail="Real auto-apply submit is not allowed for this account.")


def _admin_status_filter(filter_value: Optional[str]) -> Optional[set[str]]:
    if not filter_value or filter_value == "all":
        return None
    mapping = {
        "action_required": {"action_required"},
        "blocked": {"blocked"},
        "blocked_captcha": {"blocked_captcha"},
        "prepare_failed": {"prepare_failed"},
        "prepared": {"prepared", "ready"},
        "submitted": {"submitted"},
    }
    return mapping.get(filter_value)


MANUAL_STATUSES = {"manual_review_needed", "manual_in_progress", "manually_submitted", "manual_blocked", "needs_user_input", "offer_expired"}


def _has_remaining_user_questions(app_doc: Dict[str, Any]) -> bool:
    return bool(app_doc.get("prepared_missing_information") or app_doc.get("required_questions"))


def _effective_manual_status(app_doc: Dict[str, Any]) -> Optional[str]:
    manual_status = app_doc.get("manual_status")
    admin_status = app_doc.get("admin_status")
    if admin_status in MANUAL_STATUSES and admin_status != manual_status:
        admin_updated_at = _parse_dt(app_doc.get("admin_status_updated_at"))
        manual_updated_at = _parse_dt(app_doc.get("manual_status_updated_at"))
        if admin_updated_at and (not manual_updated_at or admin_updated_at >= manual_updated_at):
            return admin_status
    manual_status = manual_status or admin_status
    if manual_status in MANUAL_STATUSES:
        return manual_status
    # Captcha / bot / login walls are self-serve: user opens the employer site.
    # Do not swallow those into "manual_review_needed" (pending / team queue).
    submission = app_doc.get("submission_status")
    if submission in {"blocked_captcha", "blocked"} and not _has_remaining_user_questions(app_doc):
        return None
    if submission in {"prepare_failed", "failed"} and not _has_remaining_user_questions(app_doc):
        return "manual_review_needed"
    return None


def _user_facing_submission_status(app_doc: Dict[str, Any]) -> str:
    submission = app_doc.get("submission_status") or "not_submitted"
    # Keep security / employer-site blockers visible so the client can show
    # "Apply on company site" instead of a generic "pending" state.
    if submission == "blocked_captcha":
        return "blocked_captcha"
    if submission == "blocked" and not _has_remaining_user_questions(app_doc):
        return "blocked"
    manual_status = _effective_manual_status(app_doc)
    if manual_status == "offer_expired" or submission == "expired":
        return "expired"
    if submission == "submitted" or manual_status == "manually_submitted":
        return "submitted"
    if manual_status == "needs_user_input":
        return "action_required"
    queue_status = app_doc.get("auto_apply_queue_status")
    if queue_status in {"queued", "running", "awaiting_review"}:
        return "pending"
    if manual_status in {"manual_review_needed", "manual_in_progress", "manual_blocked"}:
        return "pending"
    return submission


def _public_application_doc(app_doc: Dict[str, Any], job_doc: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    public_doc = dict(_normalize_application_status_fields(app_doc))
    public_doc["user_facing_submission_status"] = _user_facing_submission_status(public_doc)
    failure = classify_application_failure(public_doc, job_doc=job_doc)
    if failure:
        public_doc["failure_code"] = failure.get("code")
        public_doc["failure_message_en"] = failure.get("user_message_en")
        public_doc["failure_message_fr"] = failure.get("user_message_fr")
    for key in (
        "admin_status",
        "manual_status",
        "assigned_to",
        "assigned_to_user_id",
        "assigned_at",
        "admin_notes",
        "admin_timeline",
        "admin_status_updated_by",
        "admin_status_updated_at",
        "manual_status_updated_by",
        "manual_status_updated_at",
    ):
        public_doc.pop(key, None)
    if job_doc is not None:
        public_doc["job"] = job_doc
    apply_url = _job_application_url(job_doc) or _job_application_url(
        public_doc.get("job") if isinstance(public_doc.get("job"), dict) else None
    )
    if apply_url:
        public_doc["apply_url"] = apply_url
    return public_doc


def _job_application_url(job_doc: Optional[Dict[str, Any]]) -> Optional[str]:
    if not job_doc:
        return None
    for key in ("application_url", "external_url", "source_url", "url", "absolute_url", "job_url"):
        value = job_doc.get(key)
        if value:
            return value
    data = job_doc.get("data") if isinstance(job_doc.get("data"), dict) else {}
    for key in ("application_url", "external_url", "source_url", "url", "absolute_url", "job_url"):
        value = data.get(key)
        if value:
            return value
    return None


def _append_admin_timeline(app_doc: Dict[str, Any], event: Dict[str, Any]) -> List[Dict[str, Any]]:
    timeline = list(app_doc.get("admin_timeline") or [])
    timeline.append(event)
    return timeline[-100:]


def _admin_doc_metadata(app_doc: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "tailored_cv_available": bool(app_doc.get("tailored_cv_file_b64")),
        "tailored_cv_text_available": bool(app_doc.get("tailored_resume_structured") or app_doc.get("tailored_resume")),
        "tailored_cv_filename": app_doc.get("tailored_cv_filename"),
        "tailored_cv_mime": app_doc.get("tailored_cv_mime"),
        "cover_letter_available": bool(app_doc.get("tailored_cover_letter") or app_doc.get("cover_letter")),
        "cover_letter_format": "text" if (app_doc.get("tailored_cover_letter") or app_doc.get("cover_letter")) else None,
        **_application_text_lengths(app_doc),
    }


def _admin_application_row(app_doc: Dict[str, Any], user_doc: Optional[Dict[str, Any]], job_doc: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    app_doc = _normalize_application_status_fields(app_doc)
    tailored_resume = app_doc.get("tailored_resume_structured") or app_doc.get("tailored_resume") or {}
    cover_letter = app_doc.get("tailored_cover_letter") or app_doc.get("cover_letter") or {}
    return {
        "application_id": app_doc.get("application_id"),
        "user_id": app_doc.get("user_id"),
        "user_email": (user_doc or {}).get("email"),
        "company": (job_doc or {}).get("company") or app_doc.get("company"),
        "title": (job_doc or {}).get("title") or app_doc.get("title"),
        "ats_provider": (job_doc or {}).get("ats_provider") or app_doc.get("submission_provider") or app_doc.get("ats_provider") or app_doc.get("auto_apply_provider"),
        "submission_status": app_doc.get("submission_status"),
        "user_facing_submission_status": _user_facing_submission_status(app_doc),
        "package_status": app_doc.get("package_status"),
        "admin_status": app_doc.get("admin_status"),
        "manual_status": _effective_manual_status(app_doc),
        "assigned_to": app_doc.get("assigned_to"),
        "assigned_at": app_doc.get("assigned_at"),
        "created_at": app_doc.get("created_at"),
        "updated_at": app_doc.get("updated_at") or app_doc.get("created_at"),
        "has_tailored_resume": bool(tailored_resume),
        "has_cover_letter": bool(cover_letter),
        "email_confirmed_outcome": app_doc.get("email_confirmed_outcome"),
        "auto_apply_queue_status": app_doc.get("auto_apply_queue_status"),
        "auto_apply_queue_reason": app_doc.get("auto_apply_queue_reason"),
        "auto_apply_provider": app_doc.get("auto_apply_provider"),
        "prepared_missing_information": app_doc.get("prepared_missing_information") or [],
        "submitted_at": app_doc.get("submitted_at"),
    }


def _parse_dt(value: Any) -> Optional[datetime]:
    if isinstance(value, datetime):
        dt = value
    elif isinstance(value, str) and value:
        try:
            dt = datetime.fromisoformat(value.replace("Z", "+00:00"))
        except ValueError:
            return None
    else:
        return None
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt


def _is_today(value: Any) -> bool:
    dt = _parse_dt(value)
    if not dt:
        return False
    return dt.astimezone(timezone.utc).date() == datetime.now(timezone.utc).date()


def _profile_completion(profile: Optional[Dict[str, Any]]) -> int:
    if not profile:
        return 0
    checks = [
        _admin_profile_has_cv(profile),
        bool(profile.get("target_role") or profile.get("target_roles")),
        bool(profile.get("target_location") or profile.get("target_location_data")),
        bool((profile.get("application_answers_profile") or {}) or (profile.get("application_defaults") or {})),
    ]
    return int(round((sum(1 for item in checks if item) / len(checks)) * 100))


def _admin_profile_has_cv(profile: Optional[Dict[str, Any]]) -> bool:
    if not profile:
        return False
    return any(profile.get(key) for key in (
        "cv_text",
        "cv_filename",
        "cv_original_filename",
        "cv_storage_path",
        "cv_file_id",
    ))


def _attention_status(status: Optional[str]) -> bool:
    return status in {"action_required", "blocked", "blocked_captcha", "prepare_failed", "failed"}


def _ats_bucket(value: Optional[str]) -> str:
    normalized = (value or "unknown").strip().lower()
    if normalized in {"greenhouse", "lever", "ashby"}:
        return normalized
    return "unknown"


def _admin_blocker_labels(app_doc: Dict[str, Any]) -> List[str]:
    raw = app_doc.get("prepared_missing_information") or app_doc.get("prepared_blockers")
    if raw is None:
        fallback = app_doc.get("submission_error") or app_doc.get("submission_status") or "Unknown blocker"
        return [str(fallback).strip() or "Unknown blocker"]
    if isinstance(raw, str):
        return [raw.strip() or "Unknown blocker"]
    if isinstance(raw, dict):
        label = raw.get("label") or raw.get("field_label") or raw.get("field_name") or raw.get("reason")
        return [str(label or raw).strip() or "Unknown blocker"]
    if not isinstance(raw, list):
        return [str(raw).strip() or "Unknown blocker"]
    labels: List[str] = []
    for blocker in raw:
        if isinstance(blocker, dict):
            label = blocker.get("label") or blocker.get("field_label") or blocker.get("field_name") or blocker.get("reason")
        else:
            label = str(blocker or "")
        label = (label or "Unknown blocker").strip()
        labels.append(label)
    return labels or ["Unknown blocker"]


async def _admin_safe_find(
    collection,
    filter: Optional[Dict[str, Any]] = None,
    limit: int = 10000,
    sort: Optional[Any] = None,
) -> List[Dict[str, Any]]:
    name = getattr(collection, "name", None) or getattr(collection, "table_name", "unknown")
    try:
        cursor = collection.find(filter or {}, {"_id": 0})
        if sort is not None and hasattr(cursor, "sort"):
            cursor = cursor.sort(sort)
        if hasattr(cursor, "limit"):
            cursor = cursor.limit(limit)
        return await cursor.to_list(limit)
    except Exception as exc:
        logger.warning("admin_safe_find_failed collection=%s error=%s", name, str(exc)[:200])
        return []


async def _admin_safe_read(
    collection,
    *,
    select: str,
    filter: Optional[Dict[str, Any]] = None,
    limit: int = 10000,
) -> List[Dict[str, Any]]:
    """Use compact PostgREST selects in production, with test/in-memory fallback."""
    name = getattr(collection, "name", None) or getattr(collection, "table_name", "unknown")
    read_with_select = getattr(collection, "read_with_select", None)
    if callable(read_with_select):
        try:
            return await read_with_select(filter or {}, limit, select=select)
        except Exception as exc:
            logger.warning("admin_safe_read_failed collection=%s error=%s", name, str(exc)[:200])
            return []
    return await _admin_safe_find(collection, filter, limit=limit)


async def _admin_safe_find_one(collection, filter: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    name = getattr(collection, "name", None) or getattr(collection, "table_name", "unknown")
    try:
        return await collection.find_one(filter, {"_id": 0})
    except Exception as exc:
        logger.warning("admin_safe_find_one_failed collection=%s error=%s", name, str(exc)[:200])
        return None


async def _admin_jobs_for_ids(job_ids: List[str]) -> List[Dict[str, Any]]:
    unique_ids = list(dict.fromkeys(job_id for job_id in job_ids if job_id))
    if not unique_ids:
        return []
    jobs: List[Dict[str, Any]] = []
    chunk_size = 80
    for index in range(0, len(unique_ids), chunk_size):
        chunk = unique_ids[index:index + chunk_size]
        jobs.extend(await _admin_safe_read(
            db.jobs,
            filter={"job_id": {"$in": chunk}},
            limit=len(chunk),
            select="job_id,title,company,location,ats_provider",
        ))
    return jobs


async def _admin_base_data(
    *,
    include_swipes: bool = True,
    include_jobs: bool = True,
) -> tuple[List[Dict[str, Any]], List[Dict[str, Any]], List[Dict[str, Any]], List[Dict[str, Any]], List[Dict[str, Any]]]:
    users_task = _admin_safe_read(
        db.users,
        select=(
            "user_id,email,name:data->>name,billing:data->billing,"
            "friend_referral:data->friend_referral,created_at:data->>created_at,"
            "updated_at:data->>updated_at,last_login_at:data->>last_login_at,"
            "demo_account:data->demo_account"
        ),
    )
    profiles_task = _admin_safe_read(
        db.profiles,
        select=(
            "user_id,target_role:data->>target_role,target_roles:data->target_roles,"
            "target_location:data->>target_location,target_location_data:data->target_location_data,"
            "application_defaults:data->application_defaults,"
            "application_answers_profile:data->application_answers_profile,extras:data->extras,"
            "contact:data->contact,contract_type:data->>contract_type,seniority:data->>seniority,"
            "cv_filename:data->>cv_filename,cv_original_filename:data->>cv_original_filename,"
            "cv_storage_path:data->>cv_storage_path,cv_file_id:data->>cv_file_id,"
            "updated_at:data->>updated_at"
        ),
    )
    applications_task = _admin_safe_read(
        db.applications,
        select=(
            "application_id,user_id,job_id,package_status:data->>package_status,"
            "submission_status:data->>submission_status,submission_provider:data->>submission_provider,"
            "admin_status:data->>admin_status,manual_status:data->>manual_status,"
            "created_at:data->>created_at,updated_at:data->>updated_at,submitted_at:data->>submitted_at,"
            "assigned_to:data->>assigned_to,assigned_at:data->>assigned_at,"
            "prepared_missing_information:data->prepared_missing_information,"
            "prepared_blockers:data->prepared_blockers,submission_error:data->>submission_error,"
            "email_confirmed_outcome:data->>email_confirmed_outcome"
        ),
    )
    swipes_task = (
        _admin_safe_read(
            db.swipes,
            select=(
                "user_id,job_id,direction:data->>direction,"
                "created_at:data->>created_at,updated_at:data->>updated_at"
            ),
        )
        if include_swipes
        else asyncio.sleep(0, result=[])
    )
    users, profiles, applications, swipes = await asyncio.gather(
        users_task,
        profiles_task,
        applications_task,
        swipes_task,
    )
    jobs = []
    if include_jobs:
        job_ids = list({app.get("job_id") for app in applications if app.get("job_id")})
        jobs = await _admin_jobs_for_ids(job_ids)
    return users, profiles, swipes, applications, jobs


async def _find_user_by_email(email: str) -> Optional[Dict[str, Any]]:
    normalized = (email or "").strip().lower()
    if not normalized:
        return None
    user_doc = await db.users.find_one({"email": normalized}, {"_id": 0})
    if user_doc:
        return user_doc
    users = await _admin_safe_find(db.users)
    for user_doc in users:
        if (user_doc.get("email") or "").strip().lower() == normalized:
            return user_doc
    return None


async def _set_user_demo_account(user_id: str, demo_account: bool) -> None:
    now = datetime.now(timezone.utc).isoformat()
    result = await db.users.update_one(
        {"user_id": user_id},
        {"$set": {"demo_account": demo_account, "updated_at": now}},
    )
    if getattr(result, "matched_count", 0) == 0:
        raise HTTPException(status_code=404, detail="User not found")


async def _set_user_training_access(user_id: str, training_access: bool) -> None:
    now = datetime.now(timezone.utc).isoformat()
    result = await db.users.update_one(
        {"user_id": user_id},
        {"$set": {"training_access": training_access, "updated_at": now}},
    )
    if getattr(result, "matched_count", 0) == 0:
        raise HTTPException(status_code=404, detail="User not found")


async def _set_user_require_review_before_send(user_id: str, value: bool) -> None:
    now = datetime.now(timezone.utc).isoformat()
    result = await db.users.update_one(
        {"user_id": user_id},
        {"$set": {"require_review_before_send": value, "updated_at": now}},
    )
    if getattr(result, "matched_count", 0) == 0:
        raise HTTPException(status_code=404, detail="User not found")


async def _set_user_language(user_id: str, value: str) -> None:
    now = datetime.now(timezone.utc).isoformat()
    result = await db.users.update_one(
        {"user_id": user_id},
        {"$set": {"language": value, "updated_at": now}},
    )
    if getattr(result, "matched_count", 0) == 0:
        raise HTTPException(status_code=404, detail="User not found")


def _admin_email_set() -> set[str]:
    return {"anto.delbos@gmail.com", "oudrhiriyouneslfim@gmail.com"} | _env_email_set("ADMIN_EMAILS")


def _is_admin_email(email: Optional[str]) -> bool:
    normalized = (email or "").strip().lower()
    return bool(normalized and normalized in _admin_email_set())


async def _resolve_training_access(user: User) -> bool:
    from training_access import user_has_training_access
    return await user_has_training_access(
        db,
        user,
        is_admin_email=_is_admin_email,
        is_training_creator=is_training_creator,
        tutorial_user_id=TUTORIAL_FILMING_USER_ID,
    )


async def _get_training_access_payload(user: User = Depends(get_current_user)) -> Dict[str, Any]:
    return await training_access_payload(
        db,
        user,
        is_admin_email=_is_admin_email,
        is_training_creator=is_training_creator,
        tutorial_user_id=TUTORIAL_FILMING_USER_ID,
    )


async def _require_training_user(user: User = Depends(get_current_user)) -> User:
    return await require_training_access(
        db,
        user,
        is_admin_email=_is_admin_email,
        is_training_creator=is_training_creator,
        tutorial_user_id=TUTORIAL_FILMING_USER_ID,
    )


async def _require_record_tools_user(user: User = Depends(get_current_user)) -> User:
    return await require_record_tools_user(
        db,
        user,
        is_admin_email=_is_admin_email,
        is_training_creator=is_training_creator,
        tutorial_user_id=TUTORIAL_FILMING_USER_ID,
    )


async def _analytics_events() -> List[Dict[str, Any]]:
    events = await _admin_safe_read(
        db.analytics_events,
        limit=10000,
        select=(
            "event_id,user_id,anonymous_id,event,page,source,created_at,"
            "properties:data->properties"
        ),
    )
    events.sort(key=lambda item: item.get("created_at") or "", reverse=True)
    return events


def _event_actor(event_doc: Dict[str, Any]) -> str:
    return str(event_doc.get("user_id") or event_doc.get("anonymous_id") or event_doc.get("event_id") or "")


def _unique_event_actors(events: List[Dict[str, Any]], event_name: str) -> set[str]:
    return {_event_actor(item) for item in events if item.get("event") == event_name and _event_actor(item)}


def _event_count(events: List[Dict[str, Any]], event_name: str) -> int:
    return sum(1 for item in events if item.get("event") == event_name)


def _series_counts(items: List[Dict[str, Any]], days: int, date_getter, predicate=lambda item: True) -> List[Dict[str, Any]]:
    today = datetime.now(timezone.utc).date()
    buckets = {(today - timedelta(days=offset)).isoformat(): 0 for offset in range(days - 1, -1, -1)}
    for item in items:
        if not predicate(item):
            continue
        dt = _parse_dt(date_getter(item))
        if not dt:
            continue
        key = dt.astimezone(timezone.utc).date().isoformat()
        if key in buckets:
            buckets[key] += 1
    return [{"date": key, "count": value} for key, value in buckets.items()]


def _rate(numerator: int, denominator: int) -> float:
    if not denominator:
        return 0.0
    return round((numerator / denominator) * 100, 1)


# Mirrors ONBOARDING_STEP_ORDER in frontend/src/components/onboarding/onboardingData.js.
# Used to figure out where a user dropped off in the onboarding flow.
ONBOARDING_STEP_ORDER: List[tuple[str, str]] = [
    ("intro", "Intro slides"),
    ("signup", "Sign up"),
    ("jobSearch", "Job search status"),
    ("jobGoal", "Job goal"),
    ("compare2x", "2× interviews comparison"),
    ("contractType", "Contract type"),
    ("otherApps", "Other apps used"),
    ("longTerm", "Long-term results"),
    ("categories", "Job categories"),
    ("experience", "Experience level"),
    ("location", "Target location"),
    ("contactPhone", "Phone number"),
    ("salary", "Salary expectations"),
    ("interviews", "Interviews per week"),
    ("jobTimeline", "Job search timeline"),
    ("interviewsConfirm", "Interviews confirmation"),
    ("jobBlocker", "Job search blocker"),
    ("jobAccomplish", "Job search goal"),
    ("potentialChart", "Interview potential"),
    ("attribution", "Acquisition source"),
    ("referralCode", "Referral code"),
    ("upload", "CV upload"),
    ("profileSetup", "Profile setup"),
    ("profileWelcome", "Profile welcome"),
    ("showcaseLanding", "Showcase — landing"),
    ("showcaseAllInOne", "Showcase — all-in-one"),
    ("showcasePricing", "Pricing / checkout"),
]
ONBOARDING_STEP_LABELS: Dict[str, str] = dict(ONBOARDING_STEP_ORDER)
ONBOARDING_STEP_KEYS: List[str] = [key for key, _label in ONBOARDING_STEP_ORDER]


def _group_events_by_user(events: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
    grouped: Dict[str, List[Dict[str, Any]]] = {}
    for event in events:
        uid = event.get("user_id")
        if not uid:
            continue
        grouped.setdefault(uid, []).append(event)
    return grouped


def _estimate_time_spent(timestamps: List[Any], gap_minutes: float = 20.0) -> Dict[str, Any]:
    """Rough 'time on app' estimate: cluster event timestamps into sessions (new session
    when the gap between two consecutive events exceeds `gap_minutes`), then sum each
    session's span. Isolated events count as a minimum 0.5 minute session."""
    parsed = sorted(dt for dt in (_parse_dt(ts) for ts in timestamps) if dt)
    if not parsed:
        return {"minutes": 0.0, "sessions": 0}
    total_minutes = 0.0
    sessions = 0
    session_start = parsed[0]
    last_ts = parsed[0]
    for ts in parsed[1:]:
        delta_minutes = (ts - last_ts).total_seconds() / 60
        if delta_minutes > gap_minutes:
            total_minutes += max(0.5, (last_ts - session_start).total_seconds() / 60)
            sessions += 1
            session_start = ts
        last_ts = ts
    total_minutes += max(0.5, (last_ts - session_start).total_seconds() / 60)
    sessions += 1
    return {"minutes": round(total_minutes, 1), "sessions": sessions}


def _onboarding_progress_for_events(user_events: List[Dict[str, Any]]) -> Dict[str, Any]:
    started_event = next((e for e in user_events if e.get("event") == "onboarding_started"), None)
    completed_event = next((e for e in user_events if e.get("event") == "onboarding_completed"), None)
    step_events = [e for e in user_events if e.get("event") == "onboarding_step_completed"]

    furthest_index = -1
    furthest_step = None
    for event in step_events:
        props = event.get("properties") or {}
        try:
            idx = int(props.get("step_index"))
        except (TypeError, ValueError):
            continue
        if idx > furthest_index:
            furthest_index = idx
            furthest_step = props.get("step")

    completed = bool(completed_event)
    drop_off_step = None
    if not completed:
        if furthest_index >= 0 and furthest_index + 1 < len(ONBOARDING_STEP_KEYS):
            drop_off_step = ONBOARDING_STEP_KEYS[furthest_index + 1]
        elif furthest_index < 0 and started_event:
            drop_off_step = ONBOARDING_STEP_KEYS[0]

    return {
        "started_at": started_event.get("created_at") if started_event else None,
        "completed_at": completed_event.get("created_at") if completed_event else None,
        "completed": completed,
        "furthest_step": furthest_step,
        "furthest_step_label": ONBOARDING_STEP_LABELS.get(furthest_step, furthest_step) if furthest_step else None,
        "furthest_step_index": furthest_index if furthest_index >= 0 else None,
        "drop_off_step": drop_off_step,
        "drop_off_step_label": ONBOARDING_STEP_LABELS.get(drop_off_step, drop_off_step) if drop_off_step else None,
    }


def _profile_onboarding_answers(profile: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    """Merge stored onboarding extras with profile preferences (fallback when checkout interrupted)."""
    if not profile:
        return {}
    extras = ((profile.get("extras") or {}).get("onboarding") or {})
    answers: Dict[str, Any] = dict(extras)
    if not answers.get("onboarding_location") and profile.get("target_location"):
        answers["onboarding_location"] = profile.get("target_location")
    if not answers.get("selected_roles") and profile.get("target_roles"):
        answers["selected_roles"] = profile.get("target_roles")
    elif not answers.get("selected_roles") and profile.get("target_role"):
        answers["selected_roles"] = [profile.get("target_role")]
    if not answers.get("contract_type") and profile.get("contract_type"):
        answers["contract_type"] = profile.get("contract_type")
    if not answers.get("seniority") and profile.get("seniority"):
        answers["seniority"] = profile.get("seniority")
    if not answers.get("experience") and answers.get("seniority"):
        answers["experience"] = answers.get("seniority")
    contact = profile.get("contact") or {}
    if not answers.get("phone") and contact.get("phone"):
        answers["phone"] = contact.get("phone")
    return answers


async def _admin_user_analytics_payload() -> Dict[str, Any]:
    users, profiles, swipes, applications, _jobs = await _admin_base_data(
        include_swipes=True,
        include_jobs=False,
    )
    events = await _analytics_events()
    events_by_user = _group_events_by_user(events)
    profile_map = {item.get("user_id"): item for item in profiles}

    app_counts: Dict[str, int] = {}
    last_app_at: Dict[str, Any] = {}
    for app_doc in applications:
        uid = app_doc.get("user_id")
        if not uid:
            continue
        app_counts[uid] = app_counts.get(uid, 0) + 1
        candidate = app_doc.get("updated_at") or app_doc.get("created_at")
        if not last_app_at.get(uid) or (_parse_dt(candidate) or datetime.min.replace(tzinfo=timezone.utc)) > (_parse_dt(last_app_at[uid]) or datetime.min.replace(tzinfo=timezone.utc)):
            last_app_at[uid] = candidate

    swipe_counts: Dict[str, int] = {}
    right_swipe_counts: Dict[str, int] = {}
    left_swipe_counts: Dict[str, int] = {}
    last_swipe_at: Dict[str, Any] = {}
    for swipe_doc in swipes:
        uid = swipe_doc.get("user_id")
        if not uid:
            continue
        swipe_counts[uid] = swipe_counts.get(uid, 0) + 1
        if swipe_doc.get("direction") == "right":
            right_swipe_counts[uid] = right_swipe_counts.get(uid, 0) + 1
        elif swipe_doc.get("direction") == "left":
            left_swipe_counts[uid] = left_swipe_counts.get(uid, 0) + 1
        candidate = swipe_doc.get("updated_at") or swipe_doc.get("created_at")
        if not last_swipe_at.get(uid) or (_parse_dt(candidate) or datetime.min.replace(tzinfo=timezone.utc)) > (_parse_dt(last_swipe_at[uid]) or datetime.min.replace(tzinfo=timezone.utc)):
            last_swipe_at[uid] = candidate

    onboarding_dropoff_counts: Dict[str, int] = {}
    onboarding_never_started = 0
    onboarding_in_progress = 0
    onboarding_completed_count = 0
    total_time_minutes = 0.0

    rows = []
    for user_doc in users:
        uid = user_doc.get("user_id")
        profile = profile_map.get(uid)
        billing = _billing_status_payload(user_doc)
        user_events = events_by_user.get(uid, [])
        activity_stats = _estimate_time_spent([event.get("created_at") for event in user_events])
        total_time_minutes += activity_stats["minutes"]
        last_event_at = max(
            (event.get("created_at") for event in user_events if event.get("created_at")),
            key=lambda value: _parse_dt(value) or datetime.min.replace(tzinfo=timezone.utc),
            default=None,
        )
        onboarding_progress = _onboarding_progress_for_events(user_events)
        onboarding_answers = _profile_onboarding_answers(profile)

        if onboarding_progress["completed"]:
            onboarding_completed_count += 1
        elif onboarding_progress["drop_off_step"]:
            onboarding_dropoff_counts[onboarding_progress["drop_off_step"]] = onboarding_dropoff_counts.get(onboarding_progress["drop_off_step"], 0) + 1
            onboarding_in_progress += 1
        else:
            onboarding_never_started += 1

        activity_candidates = [
            last_app_at.get(uid),
            last_swipe_at.get(uid),
            last_event_at,
            user_doc.get("last_login_at"),
            (profile or {}).get("updated_at"),
            user_doc.get("updated_at"),
            user_doc.get("created_at"),
        ]
        last_active_at = max(
            (value for value in activity_candidates if value),
            key=lambda value: _parse_dt(value) or datetime.min.replace(tzinfo=timezone.utc),
            default=None,
        )
        rows.append({
            "user_id": uid,
            "email": user_doc.get("email"),
            "name": user_doc.get("name"),
            "demo_account": bool(user_doc.get("demo_account")),
            "cv_uploaded": _admin_profile_has_cv(profile),
            "profile_completion": _profile_completion(profile),
            "total_applications": app_counts.get(uid, 0),
            "total_swipes": swipe_counts.get(uid, 0),
            "right_swipes": right_swipe_counts.get(uid, 0),
            "left_swipes": left_swipe_counts.get(uid, 0),
            "last_swipe_at": last_swipe_at.get(uid),
            "last_active_at": last_active_at,
            "last_login_at": user_doc.get("last_login_at"),
            "time_spent_minutes": activity_stats["minutes"],
            "sessions_count": activity_stats["sessions"],
            "onboarding_progress": onboarding_progress,
            "onboarding_answers": onboarding_answers,
            "created_at": user_doc.get("created_at"),
            "plan": billing.get("plan"),
            "is_premium": billing.get("is_premium"),
        })

    rows.sort(key=lambda item: _parse_dt(item.get("last_active_at")) or datetime.min.replace(tzinfo=timezone.utc), reverse=True)
    dropoff_by_step = sorted(
        (
            {"step": key, "label": ONBOARDING_STEP_LABELS.get(key, key), "count": count}
            for key, count in onboarding_dropoff_counts.items()
        ),
        key=lambda row: row["count"],
        reverse=True,
    )
    return {
        "summary": {
            "total_users": len(rows),
            "onboarding_completed": onboarding_completed_count,
            "onboarding_in_progress": onboarding_in_progress,
            "onboarding_never_started": onboarding_never_started,
            "avg_time_spent_minutes": round(total_time_minutes / len(rows), 1) if rows else 0,
            "total_swipes": sum(row["total_swipes"] for row in rows),
            "total_applications": sum(row["total_applications"] for row in rows),
        },
        "onboarding_dropoff": {
            "by_step": dropoff_by_step,
            "never_started": onboarding_never_started,
            "in_progress": onboarding_in_progress,
            "completed": onboarding_completed_count,
        },
        "users": rows,
    }


@api_router.get("/admin/overview")
async def admin_overview(admin: User = Depends(require_admin_user)):
    users, profiles, _swipes, applications, jobs = await _admin_base_data(include_swipes=False)
    normalized_apps = [_normalize_application_status_fields(app_doc) for app_doc in applications]
    user_map = {item.get("user_id"): item for item in users}
    job_map = {item.get("job_id"): item for item in jobs}
    prepared_statuses = {"ready", "prepared"}
    failed_blocked_statuses = {"failed", "blocked", "blocked_captcha", "prepare_failed"}
    generated_count = sum(1 for app_doc in normalized_apps if app_doc.get("package_status") in {"generated", "generated_text_only"})
    prepared_count = sum(1 for app_doc in normalized_apps if app_doc.get("submission_status") in prepared_statuses)
    submitted_count = sum(1 for app_doc in normalized_apps if app_doc.get("submission_status") == "submitted")

    blocker_counts: Dict[str, int] = {}
    for app_doc in normalized_apps:
        if not _attention_status(app_doc.get("submission_status")):
            continue
        for label in _admin_blocker_labels(app_doc):
            blocker_counts[label] = blocker_counts.get(label, 0) + 1

    attention_apps = [
        app_doc
        for app_doc in normalized_apps
        if _attention_status(app_doc.get("submission_status"))
    ]
    attention_apps.sort(key=lambda item: _parse_dt(item.get("updated_at") or item.get("created_at")) or datetime.min.replace(tzinfo=timezone.utc), reverse=True)

    return {
        "metrics": {
            "total_users": len(users),
            "new_users_today": sum(1 for user in users if _is_today(user.get("created_at"))),
            "applications_today": sum(1 for app_doc in normalized_apps if _is_today(app_doc.get("created_at"))),
            "prepared_applications": prepared_count,
            "action_required": sum(1 for app_doc in normalized_apps if app_doc.get("submission_status") == "action_required"),
            "failed_blocked": sum(1 for app_doc in normalized_apps if app_doc.get("submission_status") in failed_blocked_statuses),
            "submitted": submitted_count,
            "conversion": {
                "generated": generated_count,
                "prepared": prepared_count,
                "submitted": submitted_count,
            },
        },
        "top_blockers": [
            {"label": label, "count": count}
            for label, count in sorted(blocker_counts.items(), key=lambda item: item[1], reverse=True)[:8]
        ],
        "latest_attention": [
            _admin_application_row(app_doc, user_map.get(app_doc.get("user_id")), job_map.get(app_doc.get("job_id")))
            for app_doc in attention_apps[:10]
        ],
    }


@api_router.get("/admin/user-analytics")
async def admin_user_analytics(admin: User = Depends(require_admin_user)):
    return await _admin_user_analytics_payload()


@api_router.get("/admin/users")
async def admin_list_users(admin: User = Depends(require_admin_user)):
    users, profiles, swipes, applications, _jobs = await _admin_base_data(
        include_swipes=True,
        include_jobs=False,
    )
    profile_map = {item.get("user_id"): item for item in profiles}
    app_counts: Dict[str, int] = {}
    last_app_at: Dict[str, Any] = {}
    for app_doc in applications:
        uid = app_doc.get("user_id")
        if not uid:
            continue
        app_counts[uid] = app_counts.get(uid, 0) + 1
        candidate = app_doc.get("updated_at") or app_doc.get("created_at")
        if not last_app_at.get(uid) or (_parse_dt(candidate) or datetime.min.replace(tzinfo=timezone.utc)) > (_parse_dt(last_app_at[uid]) or datetime.min.replace(tzinfo=timezone.utc)):
            last_app_at[uid] = candidate

    swipe_counts: Dict[str, int] = {}
    right_swipe_counts: Dict[str, int] = {}
    left_swipe_counts: Dict[str, int] = {}
    last_swipe_at: Dict[str, Any] = {}
    for swipe_doc in swipes:
        uid = swipe_doc.get("user_id")
        if not uid:
            continue
        swipe_counts[uid] = swipe_counts.get(uid, 0) + 1
        if swipe_doc.get("direction") == "right":
            right_swipe_counts[uid] = right_swipe_counts.get(uid, 0) + 1
        elif swipe_doc.get("direction") == "left":
            left_swipe_counts[uid] = left_swipe_counts.get(uid, 0) + 1
        candidate = swipe_doc.get("updated_at") or swipe_doc.get("created_at")
        if not last_swipe_at.get(uid) or (_parse_dt(candidate) or datetime.min.replace(tzinfo=timezone.utc)) > (_parse_dt(last_swipe_at[uid]) or datetime.min.replace(tzinfo=timezone.utc)):
            last_swipe_at[uid] = candidate

    rows = []
    for user_doc in users:
        uid = user_doc.get("user_id")
        profile = profile_map.get(uid)
        billing = _billing_status_payload(user_doc)
        activity_candidates = [
            last_app_at.get(uid),
            last_swipe_at.get(uid),
            (profile or {}).get("updated_at"),
            user_doc.get("updated_at"),
            user_doc.get("created_at"),
        ]
        last_active_at = max(
            (value for value in activity_candidates if value),
            key=lambda value: _parse_dt(value) or datetime.min.replace(tzinfo=timezone.utc),
            default=None,
        )
        rows.append({
            "user_id": uid,
            "email": user_doc.get("email"),
            "name": user_doc.get("name"),
            "demo_account": bool(user_doc.get("demo_account")),
            "profile_completion": _profile_completion(profile),
            "cv_uploaded": _admin_profile_has_cv(profile),
            "total_applications": app_counts.get(uid, 0),
            "total_swipes": swipe_counts.get(uid, 0),
            "right_swipes": right_swipe_counts.get(uid, 0),
            "left_swipes": left_swipe_counts.get(uid, 0),
            "last_swipe_at": last_swipe_at.get(uid),
            "last_active_at": last_active_at,
            "created_at": user_doc.get("created_at"),
            "plan": billing.get("plan"),
            "is_premium": billing.get("is_premium"),
            "subscription_status": billing.get("subscription_status"),
            "credits_total": billing.get("credits_total"),
            "credits_remaining": billing.get("credits_remaining"),
        })
    rows.sort(key=lambda item: _parse_dt(item.get("last_active_at")) or datetime.min.replace(tzinfo=timezone.utc), reverse=True)
    return {"users": rows}


@api_router.get("/admin/users/{user_id}")
async def admin_get_user(user_id: str, admin: User = Depends(require_admin_user)):
    user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    if not user_doc:
        raise HTTPException(status_code=404, detail="User not found")
    profile = await db.profiles.find_one({"user_id": user_id}, {"_id": 0})
    apps = await db.applications.find({"user_id": user_id}, {"_id": 0}).sort("updated_at", -1).to_list(500)
    job_ids = list({app.get("job_id") for app in apps if app.get("job_id")})
    jobs = await db.jobs.find({"job_id": {"$in": job_ids}}, {"_id": 0}).to_list(500) if job_ids else []
    job_map = {item.get("job_id"): item for item in jobs}

    billing_payload = _billing_status_payload(user_doc)
    billing_raw = _billing_from_user(user_doc)

    try:
        user_events = await db.analytics_events.find(
            {"user_id": user_id}, {"_id": 0},
        ).sort("created_at", -1).to_list(2000)
    except Exception as exc:
        logger.warning("admin_user_analytics_events_read_failed user_id=%s error=%s", user_id, str(exc)[:200])
        user_events = []
    onboarding_progress = _onboarding_progress_for_events(user_events)
    activity_stats = _estimate_time_spent([event.get("created_at") for event in user_events])
    last_event_at = max(
        (event.get("created_at") for event in user_events if event.get("created_at")),
        key=lambda value: _parse_dt(value) or datetime.min.replace(tzinfo=timezone.utc),
        default=None,
    )
    login_sessions = await _admin_safe_find(
        getattr(db, "user_sessions", None), {"user_id": user_id}, sort=[("created_at", -1)],
    )

    swipe_rows = await db.swipes.find(
        {"user_id": user_id},
        {"_id": 0, "direction": 1, "created_at": 1, "job_id": 1},
    ).sort("created_at", -1).to_list(5000)
    right_swipes = [row for row in swipe_rows if row.get("direction") == "right"]
    left_swipes = [row for row in swipe_rows if row.get("direction") == "left"]

    app_status_counts: Dict[str, int] = {}
    outcome_counts: Dict[str, int] = {}
    for app_doc in apps:
        normalized = _normalize_application_status_fields(app_doc)
        status_key = normalized.get("submission_status") or "unknown"
        app_status_counts[status_key] = app_status_counts.get(status_key, 0) + 1
        outcome = app_doc.get("email_confirmed_outcome")
        if outcome:
            outcome_counts[outcome] = outcome_counts.get(outcome, 0) + 1

    documents = {
        "cv_filename": (profile or {}).get("cv_filename"),
        "has_cv": bool((profile or {}).get("cv_text")),
        "cv_text_length": len((profile or {}).get("cv_text") or ""),
        "cv_preview": ((profile or {}).get("cv_text") or "")[:2000],
        "original_cv_available": bool((profile or {}).get("cv_original_b64")),
        "has_cover_letter": bool((profile or {}).get("cover_letter_text")),
        "cover_letter_text_length": len((profile or {}).get("cover_letter_text") or ""),
        "cover_letter_preview": ((profile or {}).get("cover_letter_text") or "")[:2000],
        "linkedin_url": (profile or {}).get("linkedin_url"),
        "portfolio_url": (profile or {}).get("portfolio_url"),
    }

    return {
        "user": {
            **user_doc,
            "profile_completion": _profile_completion(profile),
            "cv_uploaded": bool((profile or {}).get("cv_text")),
            "plan": billing_payload.get("plan"),
        },
        "billing": {
            **billing_payload,
            "stripe_subscription_id_exists": bool(billing_raw.get("stripe_subscription_id")),
            "current_period_start": billing_raw.get("current_period_start"),
            "cancel_at_period_end": bool(billing_raw.get("cancel_at_period_end")),
        },
        "swipe_summary": {
            "total": len(swipe_rows),
            "right": len(right_swipes),
            "left": len(left_swipes),
            "right_rate": _rate(len(right_swipes), len(swipe_rows)),
            "last_swipe_at": swipe_rows[0].get("created_at") if swipe_rows else None,
            "daily_usage": {
                "7d": _series_counts(right_swipes, 7, lambda item: item.get("created_at")),
                "14d": _series_counts(right_swipes, 14, lambda item: item.get("created_at")),
                "30d": _series_counts(right_swipes, 30, lambda item: item.get("created_at")),
            },
        },
        "application_status_counts": app_status_counts,
        "outcome_counts": outcome_counts,
        "documents": documents,
        "profile": profile,
        "contact": (profile or {}).get("contact") or {},
        "preferences": {
            "target_role": (profile or {}).get("target_role"),
            "target_roles": (profile or {}).get("target_roles") or [],
            "target_location": (profile or {}).get("target_location"),
            "target_location_data": (profile or {}).get("target_location_data"),
            "remote_preference": (profile or {}).get("remote_preference"),
            "seniority": (profile or {}).get("seniority"),
            "contract_type": (profile or {}).get("contract_type"),
        },
        "application_defaults": (profile or {}).get("application_defaults") or {},
        "applications": [
            _admin_application_row(app_doc, user_doc, job_map.get(app_doc.get("job_id")))
            for app_doc in apps
        ],
        "onboarding": {
            "answers": _profile_onboarding_answers(profile),
            "progress": onboarding_progress,
        },
        "activity": {
            "last_login_at": user_doc.get("last_login_at"),
            "last_active_at": last_event_at or (swipe_rows[0].get("created_at") if swipe_rows else None) or user_doc.get("created_at"),
            "time_spent_minutes": activity_stats["minutes"],
            "sessions_count": activity_stats["sessions"],
            "login_count": len(login_sessions),
            "total_events": len(user_events),
        },
        "internal_notes": [],
    }


@api_router.get("/admin/users/{user_id}/original-cv")
async def admin_download_user_original_cv(user_id: str, admin: User = Depends(require_admin_user)):
    from fastapi.responses import Response as FastAPIResponse

    profile = await db.profiles.find_one(
        {"user_id": user_id},
        {"_id": 0, "cv_original_b64": 1, "cv_mime": 1, "cv_filename": 1},
    )
    if not profile or not profile.get("cv_original_b64"):
        raise HTTPException(status_code=404, detail="Original CV not stored")
    try:
        content = base64.b64decode(profile["cv_original_b64"], validate=True)
    except (ValueError, TypeError, base64.binascii.Error) as exc:
        raise HTTPException(status_code=500, detail="Stored original CV is invalid") from exc
    filename = profile.get("cv_filename") or "original_cv"
    return FastAPIResponse(
        content=content,
        media_type=profile.get("cv_mime") or "application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@api_router.post("/admin/users/{user_id}/repair-billing")
async def admin_repair_billing(
    user_id: str,
    admin: User = Depends(require_admin_user),
):
    """Force Stripe subscription sync and grant missing credits for a specific user."""
    user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    if not user_doc:
        raise HTTPException(status_code=404, detail="User not found")
    user_doc, warning = await _refresh_billing_from_stripe(user_doc)
    user_doc = await _repair_premium_credits_if_needed(user_id, user_doc)
    billing = _billing_status_payload(user_doc)
    return {
        "ok": True,
        "user_id": user_id,
        "billing": billing,
        "warning": warning,
    }


@api_router.post("/admin/users/{user_id}/grant-credits")
async def admin_grant_credits(
    user_id: str,
    body: AdminGrantCreditsRequest,
    admin: User = Depends(require_admin_user),
):
    """Grant bonus credits to a user and notify them, in one consistent step."""
    user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0, "billing": 1, "language": 1})
    if not user_doc:
        raise HTTPException(status_code=404, detail="User not found")
    billing = _billing_from_user(user_doc)
    new_total = int(billing.get("referral_bonus_credits_total") or 0) + body.credits
    new_remaining = int(billing.get("referral_bonus_credits_remaining") or 0) + body.credits
    await db.users.update_one(
        {"user_id": user_id},
        {"$set": {
            "billing.referral_bonus_credits_total": new_total,
            "billing.referral_bonus_credits_remaining": new_remaining,
            "billing.updated_at": datetime.now(timezone.utc).isoformat(),
        }},
    )
    title = (
        f"Vous avez reçu {body.credits} crédits gratuits"
        if resolve_user_language(user_doc) == "fr"
        else f"You received {body.credits} free credits"
    )
    await create_notification(
        db,
        user_id=user_id,
        type="credits_granted",
        title=title,
        body=body.reason,
    )
    return {
        "ok": True,
        "user_id": user_id,
        "referral_bonus_credits_total": new_total,
        "referral_bonus_credits_remaining": new_remaining,
    }


@api_router.post("/admin/stripe/repair-by-email")
async def admin_stripe_repair_by_email(
    body: AdminStripeReconcileRequest,
    admin: User = Depends(require_admin_user),
):
    """Repair billing for an app user by email (no Stripe payment ID required)."""
    email = (body.email or "").strip().lower()
    if not email:
        raise HTTPException(status_code=400, detail="email is required")

    user_doc = await _find_user_by_email(email)
    created_user = False
    if not user_doc:
        user_doc = await _ensure_app_user_for_billing_email(email)
        created_user = bool(user_doc)
    if not user_doc:
        raise HTTPException(
            status_code=404,
            detail=f"No Hirly account found for {email}. Sign up first, then retry.",
        )

    user_id = user_doc["user_id"]
    stripe_error: Optional[str] = None
    if _stripe_configured():
        try:
            _stripe_secret_key()
            user_doc = await _finalize_user_billing(user_id) or user_doc
        except Exception as exc:
            stripe_error = str(exc)[:200]
            logger.warning("admin_repair_by_email_stripe_failed user_id=%s error=%s", user_id, stripe_error)

    billing = _billing_from_user(user_doc)
    if (billing.get("subscription_status") or "none") not in {"active", "trialing"}:
        now = datetime.now(timezone.utc)
        period_end = now + timedelta(days=30)
        fallback_updates = {
            "subscription_status": "active",
            "plan": billing.get("plan") or "monthly",
            "interval": billing.get("interval") or "monthly",
            "source": billing.get("source") or "onboarding",
            "last_payment_status": "paid",
            "current_period_start": billing.get("current_period_start") or now.isoformat(),
            "current_period_end": billing.get("current_period_end") or period_end.isoformat(),
        }
        if billing.get("stripe_customer_id"):
            fallback_updates["stripe_customer_id"] = billing.get("stripe_customer_id")
        await _update_user_billing_by_user_id(user_id, fallback_updates)
        user_doc = await _repair_premium_credits_if_needed(user_id, await db.users.find_one({"user_id": user_id}, {"_id": 0}) or user_doc)

    payload = _billing_status_payload(user_doc)
    return {
        "ok": True,
        "user_id": user_id,
        "email": user_doc.get("email"),
        "billing": payload,
        "created_user": created_user,
        "stripe_error": stripe_error,
    }


@api_router.post("/admin/stripe/reconcile")
async def admin_stripe_reconcile(
    body: AdminStripeReconcileRequest,
    admin: User = Depends(require_admin_user),
):
    """Link an orphan Stripe payment/customer to an app user by email and sync billing."""
    stripe_reference = (body.payment_intent_id or "").strip() or None
    customer_id = (body.customer_id or "").strip() or None
    email = (body.email or "").strip().lower() or None
    subscription_id = None
    user_id_hint = None

    if not stripe_reference and email and not customer_id:
        user_doc = await _find_user_by_email(email)
        if user_doc:
            customer_id = _billing_from_user(user_doc).get("stripe_customer_id")
        if not customer_id and _stripe_configured():
            try:
                _stripe_secret_key()
                customers = stripe.Customer.list(email=email, limit=1)
                data = customers.get("data") if isinstance(customers, dict) else getattr(customers, "data", [])
                if data:
                    first = data[0]
                    customer_id = first.get("id") if isinstance(first, dict) else getattr(first, "id", None)
            except Exception as exc:
                raise HTTPException(status_code=400, detail=f"Could not look up Stripe customer: {str(exc)[:160]}")
        if not customer_id:
            return await admin_stripe_repair_by_email(body, admin)

    if stripe_reference:
        if not _stripe_configured():
            raise HTTPException(status_code=503, detail="Stripe is not configured on the server")
        _stripe_secret_key()
        try:
            if stripe_reference.startswith("cs_"):
                context = _resolve_stripe_checkout_session_context(stripe_reference)
            elif stripe_reference.startswith("cus_"):
                context = _resolve_stripe_customer_context(stripe_reference)
            elif stripe_reference.startswith("pi_"):
                context = _resolve_stripe_payment_intent_context(stripe_reference)
            else:
                raise HTTPException(
                    status_code=400,
                    detail="Enter a Stripe payment intent (pi_…), checkout session (cs_…), or customer (cus_…) ID.",
                )
        except HTTPException:
            raise
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Stripe lookup failed: {str(exc)[:200]}")
        customer_id = customer_id or context.get("customer_id")
        email = email or context.get("email")
        subscription_id = context.get("subscription_id")
        user_id_hint = context.get("user_id_hint")

    if not customer_id and email and _stripe_configured():
        try:
            _stripe_secret_key()
            customers = stripe.Customer.list(email=email, limit=1)
            data = customers.get("data") if isinstance(customers, dict) else getattr(customers, "data", [])
            if data:
                first = data[0]
                customer_id = first.get("id") if isinstance(first, dict) else getattr(first, "id", None)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Could not look up Stripe customer: {str(exc)[:160]}")

    if not customer_id:
        if email:
            return await admin_stripe_repair_by_email(body, admin)
        raise HTTPException(
            status_code=400,
            detail="Could not resolve Stripe customer from payment. Add the payer email and retry.",
        )

    user_id = None
    created_user = False
    if user_id_hint:
        found = await db.users.find_one({"user_id": user_id_hint}, {"_id": 0, "user_id": 1})
        if found:
            user_id = user_id_hint

    if not user_id:
        user_id = await _resolve_user_id_for_stripe_customer(customer_id)

    if not user_id and email:
        existing = await _find_user_by_email(email)
        if existing:
            user_id = existing.get("user_id")
        else:
            user_doc = await _ensure_app_user_for_billing_email(email)
            if user_doc:
                user_id = user_doc.get("user_id")
                created_user = True

    if not user_id:
        raise HTTPException(
            status_code=404,
            detail=(
                f"No Hirly account found for Stripe email {email or '(unknown)'}. "
                "Sign up with the same email first, then retry."
            ),
        )

    updates: Dict[str, Any] = {"stripe_customer_id": customer_id}
    try:
        if _stripe_configured():
            _stripe_secret_key()
            if subscription_id:
                subscription = _stripe_to_dict(stripe.Subscription.retrieve(subscription_id))
                updates.update(_subscription_billing_updates(subscription, last_payment_status="paid"))
            else:
                discovered = _discover_stripe_subscription(customer_id)
                if discovered:
                    updates.update(_subscription_billing_updates(discovered, last_payment_status="paid"))
    except Exception as exc:
        logger.warning(
            "admin_stripe_reconcile_subscription_lookup_failed user_id=%s customer_id=%s error=%s",
            user_id,
            customer_id,
            str(exc)[:200],
        )

    await _update_user_billing_by_user_id(user_id, {key: value for key, value in updates.items() if value})
    user_doc = await _finalize_user_billing(user_id) or await db.users.find_one({"user_id": user_id}, {"_id": 0}) or {}
    billing = _billing_from_user(user_doc)
    if (billing.get("subscription_status") or "none") not in {"active", "trialing"}:
        user_doc = await admin_stripe_repair_by_email(AdminStripeReconcileRequest(email=email or user_doc.get("email")), admin)
        return user_doc

    billing = _billing_status_payload(user_doc)
    return {
        "ok": True,
        "user_id": user_id,
        "email": user_doc.get("email"),
        "billing": billing,
        "created_user": created_user,
    }


@api_router.post("/admin/users/{user_id}/impersonate")
async def admin_impersonate_user(
    user_id: str,
    admin: User = Depends(require_admin_user),
):
    """Create a short-lived session token so the admin can view the app as this user."""
    if user_id == admin.user_id:
        raise HTTPException(status_code=400, detail="Cannot impersonate yourself")

    target_user = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    if not target_user:
        raise HTTPException(status_code=404, detail="User not found")

    token = f"imp_{uuid.uuid4().hex}"
    ttl_hours = 4
    expires_at = datetime.now(timezone.utc) + timedelta(hours=ttl_hours)
    await db.user_sessions.insert_one({
        "user_id": user_id,
        "session_token": token,
        "expires_at": expires_at.isoformat(),
        "created_at": datetime.now(timezone.utc).isoformat(),
        "source": "admin_impersonation",
        "impersonated_by": admin.user_id,
        "impersonated_by_email": admin.email,
    })
    logger.info(
        "admin_impersonate admin=%s admin_email=%s target_user=%s target_email=%s",
        admin.user_id,
        admin.email,
        user_id,
        target_user.get("email"),
    )
    return {
        "ok": True,
        "session_token": token,
        "user": target_user,
        "expires_at": expires_at.isoformat(),
    }


@api_router.patch("/admin/users/{user_id}/demo-account")
async def admin_set_demo_account(
    user_id: str,
    payload: Dict[str, Any],
    admin: User = Depends(require_admin_user),
):
    if "demo_account" not in payload:
        raise HTTPException(status_code=400, detail="demo_account is required")
    await _set_user_demo_account(user_id, bool(payload.get("demo_account")))
    return {"ok": True, "demo_account": bool(payload.get("demo_account"))}


@api_router.get("/admin/influencers")
async def admin_list_influencers(admin: User = Depends(require_admin_user)):
    rows = list_influencers()
    user_ids = [row.get("user_id") for row in rows if row.get("user_id")]
    user_map: Dict[str, Dict[str, Any]] = {}
    for uid in user_ids:
        user_doc = await _admin_safe_find_one(db.users, {"user_id": uid})
        if user_doc:
            user_map[uid] = user_doc
    enriched = []
    for row in rows:
        linked = user_map.get(row.get("user_id") or "")
        enriched.append({
            **row,
            "linked_email": (linked or {}).get("email"),
            "linked_demo_account": bool((linked or {}).get("demo_account")),
        })
    return {"influencers": enriched}


@api_router.post("/admin/influencers")
async def admin_create_influencer(payload: Dict[str, Any], admin: User = Depends(require_admin_user)):
    name = (payload.get("name") or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="name is required")
    row = create_influencer(payload)
    return {"ok": True, "influencer": row}


@api_router.patch("/admin/influencers/{influencer_id}")
async def admin_update_influencer(
    influencer_id: str,
    payload: Dict[str, Any],
    admin: User = Depends(require_admin_user),
):
    row = update_influencer(influencer_id, payload)
    if not row:
        raise HTTPException(status_code=404, detail="Influencer not found")
    return {"ok": True, "influencer": row}


async def _enrich_invite_rows_for_admin(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    user_ids = list({row.get("redeemed_by_user_id") for row in rows if row.get("redeemed_by_user_id")})
    users_by_id: Dict[str, Dict[str, Any]] = {}
    if user_ids:
        users = await db.users.find(
            {"user_id": {"$in": user_ids}},
            {"_id": 0, "user_id": 1, "email": 1, "name": 1},
        ).to_list(len(user_ids))
        users_by_id = {user["user_id"]: user for user in users if user.get("user_id")}
    return enrich_invite_rows(rows, users_by_id)


async def _redeem_creator_invite(code: str, user_id: str, user_email: Optional[str] = None) -> Dict[str, Any]:
    normalized = (code or "").strip()
    if not re.fullmatch(r"\d{6}", normalized):
        raise HTTPException(status_code=400, detail="Enter a valid 6-digit invitation code")

    check = await validate_invite(db, normalized)
    invitation = check.get("invitation") or await get_invite_by_code(db, normalized)
    if not invitation:
        invitation = await materialize_invite_from_link(db, normalized)
    if not invitation:
        raise HTTPException(status_code=404, detail="Invitation not found")

    if invitation.get("redeemed_by_user_id") and invitation.get("redeemed_by_user_id") != user_id:
        raise HTTPException(status_code=409, detail="This invitation was already used by another account")

    if not check.get("valid"):
        reason = check.get("reason")
        if reason == "revoked":
            raise HTTPException(status_code=410, detail="This invitation is no longer valid")
        raise HTTPException(status_code=400, detail="Invalid invitation code")

    influencer_id = invitation.get("influencer_id")
    course_id = invitation.get("course_id") or SEED_COURSE_ID
    invite_type = resolve_invite_type(invitation)
    grant_demo = invite_type in {INVITE_TYPE_DEMO, INVITE_TYPE_CREATOR}
    grant_training = invite_type in {INVITE_TYPE_TRAINING, INVITE_TYPE_CREATOR}

    enrollment = None
    if grant_demo:
        await _set_user_demo_account(user_id, True)
        await _ensure_demo_feed_profile(user_id)
    if grant_training:
        await _set_user_training_access(user_id, True)
        try:
            enrollment = await enroll_user(db, user_id, course_id)
        except ValueError as exc:
            raise HTTPException(status_code=404, detail=str(exc)) from exc

    if influencer_id:
        influencer_patch: Dict[str, Any] = {
            "user_id": user_id,
            "status": "active",
            **({"email": user_email.strip().lower()} if user_email else {}),
        }
        if grant_demo:
            influencer_patch["demo_granted"] = True
        update_influencer(influencer_id, influencer_patch)

    if not invitation.get("redeemed_at"):
        await mark_invite_redeemed(db, normalized, user_id, user_email)

    return {
        "ok": True,
        "code": normalized,
        "invite_type": invite_type,
        "demo_account": grant_demo,
        "training_access": grant_training,
        "course_id": course_id,
        "enrollment_id": (enrollment or {}).get("enrollment_id"),
        "influencer_id": influencer_id,
    }


@api_router.get("/invites/{code}/validate")
async def validate_creator_invite(code: str):
    normalized = (code or "").strip()
    if re.fullmatch(r"\d{6}", normalized) and _is_master_billing_code(normalized):
        return {
            "valid": True,
            "reason": None,
            "influencer_name": "Hirly test access",
            "course_id": SEED_COURSE_ID,
            "already_redeemed": False,
            "master_code": True,
        }
    await mark_invite_clicked(db, normalized)
    check = await validate_invite(db, code)
    invitation = check.get("invitation") or {}
    return {
        "valid": bool(check.get("valid")),
        "reason": check.get("reason"),
        "influencer_name": check.get("influencer_name"),
        "course_id": check.get("course_id") or invitation.get("course_id") or SEED_COURSE_ID,
        "invite_type": check.get("invite_type") or resolve_invite_type(invitation),
        "already_redeemed": check.get("reason") == "already_redeemed",
    }


@api_router.post("/invites/redeem")
async def redeem_creator_invite(payload: Dict[str, Any], user: User = Depends(get_current_user)):
    code = (payload.get("code") or "").strip()
    if re.fullmatch(r"\d{6}", code) and _is_master_billing_code(code):
        plan = (payload.get("plan") or "monthly").strip().lower()
        source = (payload.get("source") or "onboarding").strip().lower()
        interval = (payload.get("interval") or plan).strip().lower()
        billing = await _grant_master_billing_access(
            user.user_id,
            plan,
            interval=interval,
            source=source,
        )
        await _set_user_demo_account(user.user_id, True)
        await _ensure_demo_feed_profile(user.user_id)
        return {
            "ok": True,
            "code": code,
            "master_code": True,
            "demo_account": True,
            "billing": billing,
        }
    return await _redeem_creator_invite(code, user.user_id, user.email)


@api_router.post("/admin/influencers/{influencer_id}/invite")
async def admin_create_influencer_invite(
    influencer_id: str,
    payload: Dict[str, Any] = None,
    admin: User = Depends(require_admin_user),
):
    row = get_influencer(influencer_id)
    if not row:
        raise HTTPException(status_code=404, detail="Influencer not found")
    payload = payload or {}
    course_id = (payload.get("course_id") or SEED_COURSE_ID).strip()
    try:
        invitation = await create_invitation(db, influencer_id, course_id=course_id)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    return {
        "ok": True,
        "invitation": invitation,
        "code": invitation.get("code"),
        "course_id": invitation.get("course_id"),
        "invite_type": resolve_invite_type(invitation),
        "influencer": get_influencer(influencer_id),
    }


@api_router.post("/admin/influencers/{influencer_id}/demo-invite")
async def admin_create_influencer_demo_invite(
    influencer_id: str,
    payload: Dict[str, Any] = None,
    admin: User = Depends(require_admin_user),
):
    row = get_influencer(influencer_id)
    if not row:
        raise HTTPException(status_code=404, detail="Influencer not found")
    try:
        invitation = await create_demo_invitation(db, influencer_id=influencer_id)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    return {
        "ok": True,
        "invitation": invitation,
        "code": invitation.get("code"),
        "invite_type": INVITE_TYPE_DEMO,
        "influencer": get_influencer(influencer_id),
    }


@api_router.get("/admin/demo/invites")
async def admin_list_demo_invites(admin: User = Depends(require_admin_user)):
    return {"invites": await _enrich_invite_rows_for_admin(await list_demo_invites(db))}


@api_router.post("/admin/demo/invites")
async def admin_create_demo_invite(
    payload: Dict[str, Any],
    admin: User = Depends(require_admin_user),
):
    payload = payload or {}
    invitation = await create_demo_invitation(
        db,
        email_hint=(payload.get("email_hint") or "").strip(),
        label=(payload.get("label") or "").strip(),
    )
    return {
        "ok": True,
        "invitation": invitation,
        "code": invitation.get("code"),
        "invite_type": INVITE_TYPE_DEMO,
    }


@api_router.get("/admin/influencers/{influencer_id}/invites")
async def admin_list_influencer_invites(
    influencer_id: str,
    admin: User = Depends(require_admin_user),
):
    if not get_influencer(influencer_id):
        raise HTTPException(status_code=404, detail="Influencer not found")
    return {"invites": await _enrich_invite_rows_for_admin(await list_invites_for_influencer(db, influencer_id))}


@api_router.post("/admin/influencers/{influencer_id}/grant-demo")
async def admin_grant_influencer_demo(
    influencer_id: str,
    admin: User = Depends(require_admin_user),
):
    row = get_influencer(influencer_id)
    if not row:
        raise HTTPException(status_code=404, detail="Influencer not found")

    user_doc = None
    if row.get("user_id"):
        user_doc = await db.users.find_one({"user_id": row["user_id"]}, {"_id": 0})
    if not user_doc and row.get("email"):
        user_doc = await _find_user_by_email(row["email"])
    if not user_doc:
        raise HTTPException(status_code=404, detail="No Hirly user found for this influencer. Ask them to sign up first.")

    user_id = user_doc["user_id"]
    await _set_user_demo_account(user_id, True)
    await _ensure_demo_feed_profile(user_id)
    updated = update_influencer(influencer_id, {
        "user_id": user_id,
        "demo_granted": True,
        "status": "active",
    })
    return {
        "ok": True,
        "influencer": updated,
        "user_id": user_id,
        "email": user_doc.get("email"),
        "demo_account": True,
    }


@api_router.get("/admin/creators")
async def admin_list_creators(admin: User = Depends(require_admin_user)):
    creators = await _admin_safe_find(db.training_creators, limit=500)
    courses = await _admin_safe_find(db.training_courses, limit=1000)
    enrollments = await _admin_safe_find(db.training_enrollments, limit=5000)

    course_counts: Dict[str, int] = {}
    course_ids_by_creator: Dict[str, set[str]] = {}
    for course in courses:
        creator_id = course.get("creator_id")
        if not creator_id:
            continue
        course_counts[creator_id] = course_counts.get(creator_id, 0) + 1
        if creator_id not in course_ids_by_creator:
            course_ids_by_creator[creator_id] = set()
        course_id = course.get("course_id")
        if course_id:
            course_ids_by_creator[creator_id].add(course_id)

    enrollment_rows_by_course: Dict[str, List[Dict[str, Any]]] = {}
    for enr in enrollments:
        cid = enr.get("course_id")
        if not cid:
            continue
        if cid not in enrollment_rows_by_course:
            enrollment_rows_by_course[cid] = []
        enrollment_rows_by_course[cid].append(enr)

    rows = []
    for creator in creators:
        creator_id = creator.get("creator_id")
        creator_course_ids = course_ids_by_creator.get(creator_id, set())
        creator_enrollments: List[Dict[str, Any]] = []
        for cid in creator_course_ids:
            creator_enrollments.extend(enrollment_rows_by_course.get(cid, []))

        students = len(creator_enrollments)
        avg_progress = 0
        if creator_enrollments:
            avg_progress = round(
                sum(int(item.get("progress_percent") or 0) for item in creator_enrollments) / students
            )

        first_course_at = None
        creator_courses = [item for item in courses if item.get("creator_id") == creator_id]
        if creator_courses:
            creator_courses.sort(
                key=lambda item: _parse_dt(item.get("created_at")) or datetime.max.replace(tzinfo=timezone.utc)
            )
            first_course_at = creator_courses[0].get("created_at")

        rows.append({
            "creator_id": creator_id,
            "user_id": creator.get("user_id"),
            "email": creator.get("email"),
            "display_name": creator.get("display_name"),
            "joined_at": creator.get("created_at"),
            "courses_count": course_counts.get(creator_id, 0),
            "students_count": students,
            "avg_progress_percent": avg_progress,
            "first_course_at": first_course_at,
            "last_active_at": max(
                [item.get("updated_at") for item in creator_enrollments if item.get("updated_at")]
                + [creator.get("created_at")],
                key=lambda item: _parse_dt(item) or datetime.min.replace(tzinfo=timezone.utc),
            ),
        })

    rows.sort(
        key=lambda item: _parse_dt(item.get("joined_at")) or datetime.min.replace(tzinfo=timezone.utc),
        reverse=True,
    )
    return {"creators": rows}


def _parse_creator_id_filters(
    creator_id: Optional[str] = None,
    creator_ids: Optional[str] = None,
) -> Optional[List[str]]:
    if creator_ids:
        parsed = [item.strip() for item in creator_ids.split(",") if item.strip()]
        return parsed or None
    if creator_id:
        return [creator_id]
    return None


@api_router.post("/admin/creator-social/creators")
async def admin_creator_social_add(
    admin: User = Depends(require_admin_user),
    payload: dict = Body(...),
):
    platform = str(payload.get("platform") or "").strip()
    handle = str(payload.get("handle") or "").strip()
    name = payload.get("name")
    try:
        creator = add_tracked_creator(platform=platform, handle=handle, name=name)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    try:
        await asyncio.to_thread(refresh_creator, creator["creator_id"])
        record_refresh_summary({
            "started_at": datetime.now(timezone.utc).isoformat(),
            "finished_at": datetime.now(timezone.utc).isoformat(),
            "creator_count": 1,
            "success_count": 1,
            "error_count": 0,
            "errors": [],
            "ok": True,
            "trigger": "manual_add",
        })
    except Exception as exc:
        logger.warning(
            "creator_social_initial_refresh_failed creator_id=%s error=%s",
            creator.get("creator_id"),
            str(exc)[:200],
        )
    dashboard = build_dashboard()
    dashboard["maintenance"] = get_creator_social_refresh_status()
    return {"ok": True, "creator": creator, "dashboard": dashboard}


@api_router.get("/admin/creator-social")
async def admin_creator_social_dashboard(
    admin: User = Depends(require_admin_user),
    days: int = Query(14, ge=7, le=90),
    creator_id: Optional[str] = Query(None),
    creator_ids: Optional[str] = Query(None),
):
    dashboard = build_dashboard(days=days, creator_ids=_parse_creator_id_filters(creator_id, creator_ids))
    dashboard["maintenance"] = get_creator_social_refresh_status()
    return dashboard


@api_router.post("/admin/creator-social/refresh")
async def admin_creator_social_refresh(
    admin: User = Depends(require_admin_user),
    creator_id: Optional[str] = Query(None),
):
    if creator_id:
        try:
            snapshot = await asyncio.to_thread(refresh_creator, creator_id)
            record_refresh_summary({
                "started_at": datetime.now(timezone.utc).isoformat(),
                "finished_at": datetime.now(timezone.utc).isoformat(),
                "creator_count": 1,
                "success_count": 0 if snapshot.get("error") else 1,
                "error_count": 1 if snapshot.get("error") else 0,
                "errors": [snapshot] if snapshot.get("error") else [],
                "ok": not snapshot.get("error"),
                "trigger": "manual",
            })
        except ValueError as exc:
            raise HTTPException(status_code=404, detail=str(exc)) from exc
        except Exception as exc:
            raise HTTPException(status_code=502, detail=f"Could not refresh creator: {exc}") from exc
        dashboard = build_dashboard(creator_ids=[creator_id])
        dashboard["maintenance"] = get_creator_social_refresh_status()
        return {"ok": True, "snapshot": snapshot, "dashboard": dashboard}

    try:
        summary = await asyncio.to_thread(run_creator_social_refresh, trigger="manual")
    except Exception as exc:
        logger.exception("creator_social_refresh_failed")
        raise HTTPException(status_code=502, detail=f"Could not refresh creator stats: {exc}") from exc
    dashboard = build_dashboard()
    dashboard["maintenance"] = get_creator_social_refresh_status()
    return {
        "ok": summary.get("ok", True),
        "snapshots": summary,
        "dashboard": dashboard,
        "errors": summary.get("errors") or [],
    }


@api_router.get("/admin/training/analytics")
async def admin_training_analytics_endpoint(
    admin: User = Depends(require_admin_user),
    course_id: str = Query("course_job_search_mastery"),
):
    return await compute_training_analytics(db, course_id)


@api_router.get("/admin/analytics")
async def admin_analytics(admin: User = Depends(require_admin_user)):
    users, profiles, swipes, applications, jobs = await _admin_base_data(include_swipes=True)
    events = await _analytics_events()
    normalized_apps = [_normalize_application_status_fields(app_doc) for app_doc in applications]
    job_map = {item.get("job_id"): item for item in jobs}
    generated_count = sum(1 for app_doc in normalized_apps if app_doc.get("package_status") in {"generated", "generated_text_only"})
    prepared_count = sum(1 for app_doc in normalized_apps if app_doc.get("submission_status") in {"ready", "prepared"})
    action_required_count = sum(1 for app_doc in normalized_apps if app_doc.get("submission_status") == "action_required")
    submitted_count = sum(1 for app_doc in normalized_apps if app_doc.get("submission_status") == "submitted")
    blocked_failed_count = sum(1 for app_doc in normalized_apps if app_doc.get("submission_status") in {"failed", "blocked", "blocked_captcha", "prepare_failed"})

    breakdown = {
        "greenhouse": {"generated": 0, "prepared": 0, "action_required": 0, "submitted": 0, "failed_blocked": 0},
        "lever": {"generated": 0, "prepared": 0, "action_required": 0, "submitted": 0, "failed_blocked": 0},
        "ashby": {"generated": 0, "prepared": 0, "action_required": 0, "submitted": 0, "failed_blocked": 0},
        "unknown": {"generated": 0, "prepared": 0, "action_required": 0, "submitted": 0, "failed_blocked": 0},
    }
    for app_doc in normalized_apps:
        job = job_map.get(app_doc.get("job_id")) or {}
        bucket = breakdown[_ats_bucket(job.get("ats_provider") or app_doc.get("submission_provider"))]
        if app_doc.get("package_status") in {"generated", "generated_text_only"}:
            bucket["generated"] += 1
        if app_doc.get("submission_status") in {"ready", "prepared"}:
            bucket["prepared"] += 1
        if app_doc.get("submission_status") == "action_required":
            bucket["action_required"] += 1
        if app_doc.get("submission_status") == "submitted":
            bucket["submitted"] += 1
        if app_doc.get("submission_status") in {"failed", "blocked", "blocked_captcha", "prepare_failed"}:
            bucket["failed_blocked"] += 1

    ats_performance = {
        key: {
            **value,
            "prepare_rate": _rate(value["prepared"], value["generated"]),
            "failure_rate": _rate(value["failed_blocked"], value["generated"]),
        }
        for key, value in breakdown.items()
    }

    landing_actors = _unique_event_actors(events, "landing_view")
    signup_actors = _unique_event_actors(events, "auth_success") or {user.get("user_id") for user in users if user.get("user_id")}
    onboarding_started_actors = _unique_event_actors(events, "onboarding_started")
    onboarding_completed_actors = _unique_event_actors(events, "onboarding_completed") or {profile.get("user_id") for profile in profiles if profile.get("target_role")}
    cv_uploaded_actors = _unique_event_actors(events, "cv_upload_completed") or {
        profile.get("user_id") for profile in profiles if _admin_profile_has_cv(profile)
    }
    first_swipe_actors = {swipe.get("user_id") for swipe in swipes if swipe.get("user_id")}
    right_swipe_count = sum(1 for swipe in swipes if swipe.get("direction") == "right")

    events_by_user = _group_events_by_user(events)
    onboarding_dropoff_counts: Dict[str, int] = {}
    onboarding_never_started = 0
    onboarding_in_progress = 0
    for user_doc in users:
        uid = user_doc.get("user_id")
        if not uid:
            continue
        progress = _onboarding_progress_for_events(events_by_user.get(uid, []))
        if progress["completed"]:
            continue
        if progress["drop_off_step"]:
            onboarding_dropoff_counts[progress["drop_off_step"]] = onboarding_dropoff_counts.get(progress["drop_off_step"], 0) + 1
            onboarding_in_progress += 1
        else:
            onboarding_never_started += 1

    onboarding_dropoff_by_step = sorted(
        (
            {"step": key, "label": ONBOARDING_STEP_LABELS.get(key, key), "count": count}
            for key, count in onboarding_dropoff_counts.items()
        ),
        key=lambda row: row["count"],
        reverse=True,
    )

    cta_names = {
        "cta_start_swiping_clicked": "Start swiping",
        "cta_signup_clicked": "Signup",
        "cta_login_clicked": "Login",
    }
    cta_analytics = []
    for event_name, label in cta_names.items():
        actors = _unique_event_actors(events, event_name)
        cta_analytics.append({
            "cta": label,
            "event": event_name,
            "clicks": _event_count(events, event_name),
            "conversion_to_signup": _rate(len(actors & signup_actors), len(actors)),
            "conversion_to_onboarding": _rate(len(actors & onboarding_started_actors), len(actors)),
            "conversion_to_first_swipe": _rate(len(actors & first_swipe_actors), len(actors)),
        })

    admin_attention = [app_doc for app_doc in normalized_apps if _attention_status(app_doc.get("submission_status"))]
    now = datetime.now(timezone.utc)
    ages = []
    for app_doc in admin_attention:
        dt = _parse_dt(app_doc.get("updated_at") or app_doc.get("created_at"))
        if dt:
            ages.append(max(0, (now - dt).total_seconds() / 3600))

    funnel_steps = [
        {"step": "landing_view", "label": "Landing views", "count": len(landing_actors) or _event_count(events, "landing_view")},
        {"step": "signup", "label": "Signup", "count": len(signup_actors)},
        {"step": "onboarding_started", "label": "Onboarding started", "count": len(onboarding_started_actors)},
        {"step": "onboarding_completed", "label": "Onboarding completed", "count": len(onboarding_completed_actors)},
        {"step": "cv_uploaded", "label": "CV uploaded", "count": len(cv_uploaded_actors)},
        {"step": "first_swipe", "label": "First swipe", "count": len(first_swipe_actors)},
        {"step": "right_swipe", "label": "Right swipes", "count": right_swipe_count},
        {"step": "application_generated", "label": "Applications generated", "count": generated_count},
        {"step": "prepared", "label": "Prepared", "count": prepared_count},
        {"step": "submitted", "label": "Submitted", "count": submitted_count},
    ]
    for index, row in enumerate(funnel_steps):
        row["previous_rate"] = None if index == 0 else _rate(row["count"], funnel_steps[index - 1]["count"])

    return {
        "metrics": {
            "visitors": len(landing_actors) or _event_count(events, "landing_view"),
            "signups": len(users),
            "onboarding_started": len(onboarding_started_actors),
            "onboarding_complete": sum(1 for profile in profiles if profile.get("target_role")),
            "onboarding_completed": len(onboarding_completed_actors),
            "cv_uploaded": len(cv_uploaded_actors),
            "swipe_users": len(first_swipe_actors),
            "swipes": len(swipes),
            "total_swipes": len(swipes),
            "right_swipes": sum(1 for swipe in swipes if swipe.get("direction") == "right"),
            "applications_generated": generated_count,
            "prepared": prepared_count,
            "action_required": action_required_count,
            "submitted": submitted_count,
            "failed_blocked": blocked_failed_count,
            "blocked_or_failed": blocked_failed_count,
        },
        "conversion_funnel": funnel_steps,
        "onboarding_dropoff": {
            "by_step": onboarding_dropoff_by_step,
            "never_started": onboarding_never_started,
            "in_progress": onboarding_in_progress,
            "completed": len(onboarding_completed_actors),
        },
        "cta_analytics": cta_analytics,
        "application_funnel": {
            "generated": generated_count,
            "prepared": prepared_count,
            "action_required": action_required_count,
            "blocked_captcha": sum(1 for app_doc in normalized_apps if app_doc.get("submission_status") == "blocked_captcha"),
            "prepare_failed": sum(1 for app_doc in normalized_apps if app_doc.get("submission_status") == "prepare_failed"),
            "submitted": submitted_count,
        },
        "by_ats": breakdown,
        "ats_performance": ats_performance,
        "time_series": {
            "last_7_days": {
                "signups": _series_counts(users, 7, lambda item: item.get("created_at")),
                "swipes": _series_counts(swipes, 7, lambda item: item.get("created_at")),
                "applications": _series_counts(normalized_apps, 7, lambda item: item.get("created_at")),
                "prepared": _series_counts(normalized_apps, 7, lambda item: item.get("updated_at") or item.get("created_at"), lambda item: item.get("submission_status") in {"ready", "prepared"}),
                "submitted": _series_counts(normalized_apps, 7, lambda item: item.get("submitted_at") or item.get("updated_at") or item.get("created_at"), lambda item: item.get("submission_status") == "submitted"),
            },
            "last_30_days": {
                "signups": _series_counts(users, 30, lambda item: item.get("created_at")),
                "swipes": _series_counts(swipes, 30, lambda item: item.get("created_at")),
                "applications": _series_counts(normalized_apps, 30, lambda item: item.get("created_at")),
                "prepared": _series_counts(normalized_apps, 30, lambda item: item.get("updated_at") or item.get("created_at"), lambda item: item.get("submission_status") in {"ready", "prepared"}),
                "submitted": _series_counts(normalized_apps, 30, lambda item: item.get("submitted_at") or item.get("updated_at") or item.get("created_at"), lambda item: item.get("submission_status") == "submitted"),
            },
        },
        "admin_ops": {
            "open_action_required": sum(1 for app_doc in normalized_apps if app_doc.get("submission_status") == "action_required"),
            "open_blocked": sum(1 for app_doc in normalized_apps if app_doc.get("submission_status") in {"blocked", "blocked_captcha", "prepare_failed"}),
            "assigned_applications": sum(1 for app_doc in normalized_apps if app_doc.get("assigned_to")),
            "unassigned_applications": sum(1 for app_doc in admin_attention if not app_doc.get("assigned_to")),
            "average_unresolved_age_hours": round(sum(ages) / len(ages), 1) if ages else 0,
        },
        "events_available": bool(events),
    }


@api_router.get("/admin/applications")
async def admin_list_applications(
    status_filter: Optional[str] = Query(default=None, alias="filter"),
    status: Optional[str] = Query(default=None),
    admin: User = Depends(require_admin_user),
):
    status_filter = status_filter or status
    apps = await _admin_safe_find(db.applications, sort=[("updated_at", -1)], limit=1000)
    allowed_statuses = _admin_status_filter(status_filter)
    if allowed_statuses is not None:
        apps = [
            app_doc
            for app_doc in apps
            if _normalize_application_status_fields(app_doc).get("submission_status") in allowed_statuses
        ]
    elif status_filter == "offer_expired":
        apps = [
            app_doc
            for app_doc in apps
            if _user_facing_submission_status(_normalize_application_status_fields(app_doc)) == "expired"
        ]
    elif status_filter in {"manual_review_needed", "manual_in_progress", "manually_submitted", "manual_blocked", "needs_user_input"}:
        apps = [
            app_doc
            for app_doc in apps
            if _effective_manual_status(_normalize_application_status_fields(app_doc)) == status_filter
        ]
    apps = _sort_applications_newest_first(apps)

    user_ids = list({app.get("user_id") for app in apps if app.get("user_id")})
    job_ids = list({app.get("job_id") for app in apps if app.get("job_id")})
    users = await _admin_safe_find(db.users, {"user_id": {"$in": user_ids}}, limit=len(user_ids)) if user_ids else []
    jobs = await _admin_jobs_for_ids(job_ids) if job_ids else []
    user_map = {item.get("user_id"): item for item in users}
    job_map = {item.get("job_id"): item for item in jobs}
    return {
        "applications": [
            _admin_application_row(app_doc, user_map.get(app_doc.get("user_id")), job_map.get(app_doc.get("job_id")))
            for app_doc in apps
        ],
        "filter": status_filter or "all",
    }


@api_router.post("/admin/ats-lab/generate")
async def admin_ats_lab_generate(
    body: AdminAtsLabGenerateRequest,
    admin: User = Depends(require_admin_user),
):
    """Run ATS-tailored CV + cover letter generation for admin experimentation (no credit charge)."""
    user_id = (body.user_id or "").strip()
    job_id = (body.job_id or "").strip()
    application_id = (body.application_id or "").strip()

    if application_id:
        app_doc = await db.applications.find_one({"application_id": application_id}, {"_id": 0})
        if not app_doc:
            raise HTTPException(status_code=404, detail="Application not found")
        user_id = app_doc.get("user_id") or user_id
        job_id = app_doc.get("job_id") or job_id
    if not user_id or not job_id:
        raise HTTPException(status_code=400, detail="application_id or both user_id and job_id are required")

    profile = await db.profiles.find_one({"user_id": user_id}, {"_id": 0})
    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found")
    job = await db.jobs.find_one({"job_id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0})

    started = datetime.now(timezone.utc)
    try:
        gen = await claude_generate_application(profile, job)
    except LLMProviderNotConfigured as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc
    except Exception as exc:
        logger.exception("admin_ats_lab_generate_failed user_id=%s job_id=%s", user_id, job_id)
        raise HTTPException(status_code=500, detail=str(exc) or "Generation failed") from exc

    gen = sanitize_docx_text(normalize_application_generation(gen))
    tailored_resume = gen.get("tailored_resume_structured") or gen.get("tailored_resume") or {}
    cover_letter = gen.get("tailored_cover_letter") or gen.get("cover_letter") or {}
    elapsed_ms = int((datetime.now(timezone.utc) - started).total_seconds() * 1000)

    persisted = False
    if body.persist and application_id:
        update_fields = {
            "tailored_resume_structured": tailored_resume,
            "tailored_resume": tailored_resume,
            "tailored_cover_letter": cover_letter,
            "cover_letter": cover_letter,
            "match_score": gen.get("match_score"),
            "match_reasons": gen.get("match_reasons") or [],
            "ats_score_before": gen.get("ats_score_before"),
            "ats_score_after": gen.get("ats_score_after"),
            "ats_provider": gen.get("ats_provider"),
            "ats_analysis": gen.get("ats_analysis") or {},
            "keywords_gap": gen.get("keywords_gap") or [],
            "resume_quality_checks": gen.get("resume_quality_checks") or {},
            "resume_quality_report": gen.get("resume_quality_report") or {},
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }
        await db.applications.update_one(
            {"application_id": application_id},
            {"$set": update_fields},
        )
        persisted = True

    return {
        "ok": True,
        "elapsed_ms": elapsed_ms,
        "persisted": persisted,
        "application_id": application_id or None,
        "user_id": user_id,
        "job_id": job_id,
        "user_email": (user_doc or {}).get("email"),
        "profile_snapshot": {
            "cv_text": (profile.get("cv_text") or "")[:8000],
            "contact": profile.get("contact") or {},
            "cover_letter_reference": (profile.get("cover_letter_text") or "")[:2000],
            "template_style": profile.get("template_style"),
        },
        "job": {
            "job_id": job.get("job_id"),
            "title": job.get("title"),
            "company": job.get("company"),
            "location": job.get("location"),
            "ats_provider": job.get("ats_provider"),
        },
        "generation": gen,
        "tailored_resume": tailored_resume,
        "tailored_cover_letter": cover_letter,
        "cover_letter_text": cover_letter_to_text(cover_letter),
    }


@api_router.get("/admin/applications/{application_id}")
async def admin_get_application(application_id: str, admin: User = Depends(require_admin_user)):
    app_doc = await db.applications.find_one({"application_id": application_id}, {"_id": 0})
    if not app_doc:
        raise HTTPException(status_code=404, detail="Application not found")
    app_doc = _normalize_application_status_fields(app_doc)
    profile = await db.profiles.find_one({"user_id": app_doc.get("user_id")}, {"_id": 0})
    user_doc = await db.users.find_one({"user_id": app_doc.get("user_id")}, {"_id": 0})
    job = await db.jobs.find_one({"job_id": app_doc.get("job_id")}, {"_id": 0})
    runs = await db.browser_submission_runs.find(
        {"application_id": application_id},
        {
            "_id": 0,
            "screenshots": 0,
        },
    ).sort("created_at", -1).to_list(20)
    latest_run = runs[0] if runs else None
    failure_classification = classify_application_failure(app_doc, job_doc=job, latest_run=latest_run)
    required_questions = (
        app_doc.get("required_questions")
        or app_doc.get("prepared_missing_information")
        or []
    )
    notes = app_doc.get("admin_notes") or []
    cover_letter = app_doc.get("tailored_cover_letter") or app_doc.get("cover_letter") or {}
    tailored_resume = app_doc.get("tailored_resume_structured") or app_doc.get("tailored_resume") or {}
    app_with_admin = {
        **app_doc,
        "user_email": (user_doc or {}).get("email"),
        # The address this application's employer replies actually go to.
        # If a real submission already recorded one (agent_apply_result,
        # set by run_apply_attempt), that's the authoritative historical
        # value; otherwise show what it *would* be right now, computed the
        # same way the real submission path resolves it -- so this is
        # visible immediately, not only after a browser run has happened
        # (which may never occur while automation is paused). Distinct from
        # user_email above, which is always the account's real login email
        # regardless of this feature.
        "submission_contact_email": (
            (app_doc.get("agent_apply_result") or {}).get("submission_email")
            or (managed_reply_address(app_doc["application_id"]) if INBOUND_MANAGED_EMAIL_ENABLED else (user_doc or {}).get("email"))
        ),
        "manual_status": _effective_manual_status(app_doc),
        "user_facing_submission_status": _user_facing_submission_status(app_doc),
    }
    application_url = _job_application_url(job)
    public_profile = dict(profile or {})
    public_profile.pop("cv_original_b64", None)
    original_cv_available = bool((profile or {}).get("cv_original_b64"))
    return {
        "application": app_with_admin,
        "profile": public_profile,
        "job": {
            **(job or {}),
            "application_url": application_url,
            "external_url": application_url or (job or {}).get("external_url"),
        } if job else None,
        "job_application_url": application_url,
        "user_contact_info": (profile or {}).get("contact") or {"email": (user_doc or {}).get("email"), "name": (user_doc or {}).get("name")},
        "profile_summary": (profile or {}).get("summary"),
        "application_defaults": (profile or {}).get("application_defaults") or {},
        "prepared_missing_information": app_doc.get("prepared_missing_information") or [],
        "resolved_answers": app_doc.get("resolved_answers") or app_doc.get("prepared_generated_answers") or [],
        "required_questions": required_questions,
        "browser_submission_runs": runs,
        "generated_documents_metadata": {
            **_admin_doc_metadata(app_doc),
            "original_cv_available": original_cv_available,
            "original_cv_filename": (profile or {}).get("cv_filename"),
            "original_cv_mime": (profile or {}).get("cv_mime"),
        },
        "tailored_resume": tailored_resume,
        "tailored_resume_text": json.dumps(tailored_resume, indent=2, default=str) if tailored_resume else "",
        "cover_letter": cover_letter,
        "cover_letter_text": cover_letter_to_text(cover_letter) if cover_letter else "",
        "download_urls": {
            "original_cv": f"/api/admin/applications/{application_id}/original-cv",
            "tailored_cv": f"/api/admin/applications/{application_id}/tailored-cv",
            "cover_letter": f"/api/admin/applications/{application_id}/cover-letter",
        },
        "latest_browser_logs": runs[:5],
        "admin_timeline": app_doc.get("admin_timeline") or [],
        "latest_notes": notes[-20:],
        "failure_classification": failure_classification,
    }


@api_router.post("/admin/applications/{application_id}/notes")
async def admin_add_application_note(
    application_id: str,
    body: AdminNoteCreate,
    admin: User = Depends(require_admin_user),
):
    note = body.note.strip()
    if not note:
        raise HTTPException(status_code=400, detail="note required")
    app_doc = await db.applications.find_one({"application_id": application_id}, {"_id": 0})
    if not app_doc:
        raise HTTPException(status_code=404, detail="Application not found")
    now = datetime.now(timezone.utc).isoformat()
    notes = list(app_doc.get("admin_notes") or [])
    notes.append({
        "note_id": f"note_{uuid.uuid4().hex[:12]}",
        "note": note,
        "author_user_id": admin.user_id,
        "author_email": admin.email,
        "created_at": now,
    })
    await db.applications.update_one(
        {"application_id": application_id},
        {"$set": {"admin_notes": notes, "updated_at": now}},
    )
    return {"ok": True, "note": notes[-1], "latest_notes": notes[-20:]}


@api_router.post("/admin/applications/{application_id}/assign")
async def admin_assign_application(application_id: str, admin: User = Depends(require_admin_user)):
    app_doc = await db.applications.find_one({"application_id": application_id}, {"_id": 0})
    if not app_doc:
        raise HTTPException(status_code=404, detail="Application not found")
    now = datetime.now(timezone.utc).isoformat()
    await db.applications.update_one(
        {"application_id": application_id},
        {"$set": {
            "assigned_to": admin.email,
            "assigned_to_user_id": admin.user_id,
            "assigned_at": now,
            "updated_at": now,
        }},
    )
    updated = await db.applications.find_one({"application_id": application_id}, {"_id": 0})
    return {"ok": True, "application": _normalize_application_status_fields(updated or {})}


@api_router.post("/admin/applications/{application_id}/unassign")
async def admin_unassign_application(application_id: str, admin: User = Depends(require_admin_user)):
    app_doc = await db.applications.find_one({"application_id": application_id}, {"_id": 0})
    if not app_doc:
        raise HTTPException(status_code=404, detail="Application not found")
    now = datetime.now(timezone.utc).isoformat()
    await db.applications.update_one(
        {"application_id": application_id},
        {"$set": {
            "assigned_to": None,
            "assigned_to_user_id": None,
            "assigned_at": None,
            "updated_at": now,
        }},
    )
    updated = await db.applications.find_one({"application_id": application_id}, {"_id": 0})
    return {"ok": True, "application": _normalize_application_status_fields(updated or {})}


@api_router.patch("/admin/applications/{application_id}/admin-status")
async def admin_update_application_status(
    application_id: str,
    body: AdminStatusUpdate,
    admin: User = Depends(require_admin_user),
):
    app_doc = await db.applications.find_one({"application_id": application_id}, {"_id": 0})
    if not app_doc:
        raise HTTPException(status_code=404, detail="Application not found")
    submission_status = {
        "submitted": "submitted",
        "needs_user_input": "action_required",
        "blocked": "blocked",
        "escalated": "blocked",
    }[body.status]
    now = datetime.now(timezone.utc).isoformat()
    timeline = _append_admin_timeline(app_doc, {
        "event_id": f"timeline_{uuid.uuid4().hex[:12]}",
        "type": "admin_status",
        "admin_status": body.status,
        "author_user_id": admin.user_id,
        "author_email": admin.email,
        "created_at": now,
    })
    update = {
        "admin_status": body.status,
        "submission_status": submission_status,
        "admin_timeline": timeline,
        "updated_at": now,
        "admin_status_updated_by": admin.email,
        "admin_status_updated_at": now,
    }
    if body.status == "needs_user_input":
        update["manual_status"] = "needs_user_input"
        update["manual_status_updated_by"] = admin.email
        update["manual_status_updated_at"] = now
    if body.status == "submitted":
        update["submitted_at"] = now
    await db.applications.update_one(
        {"application_id": application_id},
        {"$set": update},
    )
    updated = await db.applications.find_one({"application_id": application_id}, {"_id": 0})
    return {"ok": True, "application": _normalize_application_status_fields(updated or {})}


@api_router.post("/admin/applications/{application_id}/manual-status")
async def admin_update_application_manual_status(
    application_id: str,
    body: AdminManualStatusUpdate,
    admin: User = Depends(require_admin_user),
):
    app_doc = await db.applications.find_one({"application_id": application_id}, {"_id": 0})
    if not app_doc:
        raise HTTPException(status_code=404, detail="Application not found")
    now = datetime.now(timezone.utc).isoformat()
    note_text = (body.note or "").strip()
    notes = list(app_doc.get("admin_notes") or [])
    if note_text:
        notes.append({
            "note_id": f"note_{uuid.uuid4().hex[:12]}",
            "note": note_text,
            "author_user_id": admin.user_id,
            "author_email": admin.email,
            "created_at": now,
        })
    timeline = _append_admin_timeline(app_doc, {
        "event_id": f"timeline_{uuid.uuid4().hex[:12]}",
        "type": "manual_status",
        "manual_status": body.manual_status,
        "author_user_id": admin.user_id,
        "author_email": admin.email,
        "note": note_text or None,
        "created_at": now,
    })
    update = {
        "admin_status": body.manual_status,
        "manual_status": body.manual_status,
        "admin_notes": notes,
        "admin_timeline": timeline,
        "manual_status_updated_by": admin.email,
        "manual_status_updated_at": now,
        "updated_at": now,
    }
    if body.manual_status == "manually_submitted":
        update["submission_status"] = "submitted"
        update["submitted_at"] = app_doc.get("submitted_at") or now
    elif body.manual_status == "needs_user_input":
        update["submission_status"] = "action_required"
    elif body.manual_status == "offer_expired":
        await db.applications.update_one(
            {"application_id": application_id},
            {"$set": {"admin_notes": notes, "admin_timeline": timeline}},
        )
        merged_app = {**app_doc, "admin_notes": notes, "admin_timeline": timeline}
        updated_doc = await mark_application_offer_expired(
            db,
            merged_app,
            source="admin_manual",
            actor_email=admin.email,
            note=note_text or None,
        )
        return {"ok": True, "application": _normalize_application_status_fields(updated_doc or {})}
    await db.applications.update_one(
        {"application_id": application_id},
        {"$set": update},
    )
    updated = await db.applications.find_one({"application_id": application_id}, {"_id": 0})
    return {"ok": True, "application": _normalize_application_status_fields(updated or {})}


@api_router.get("/admin/applications/{application_id}/original-cv")
async def admin_download_original_cv(application_id: str, admin: User = Depends(require_admin_user)):
    from fastapi.responses import Response as FastAPIResponse

    app_doc = await db.applications.find_one(
        {"application_id": application_id},
        {"_id": 0, "user_id": 1},
    )
    if not app_doc:
        raise HTTPException(status_code=404, detail="Application not found")
    profile = await db.profiles.find_one(
        {"user_id": app_doc.get("user_id")},
        {"_id": 0, "cv_original_b64": 1, "cv_mime": 1, "cv_filename": 1},
    )
    if not profile or not profile.get("cv_original_b64"):
        raise HTTPException(status_code=404, detail="Original CV not stored")
    try:
        content = base64.b64decode(profile["cv_original_b64"], validate=True)
    except (ValueError, TypeError, base64.binascii.Error) as exc:
        raise HTTPException(status_code=500, detail="Stored original CV is invalid") from exc
    filename = profile.get("cv_filename") or "original_cv"
    return FastAPIResponse(
        content=content,
        media_type=profile.get("cv_mime") or "application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@api_router.get("/admin/applications/{application_id}/tailored-cv")
async def admin_download_tailored_cv(application_id: str, admin: User = Depends(require_admin_user)):
    from fastapi.responses import Response as FastAPIResponse

    app_doc = await db.applications.find_one(
        {"application_id": application_id},
        {"_id": 0, "tailored_cv_file_b64": 1, "tailored_cv_filename": 1, "tailored_cv_mime": 1},
    )
    if not app_doc or not app_doc.get("tailored_cv_file_b64"):
        raise HTTPException(status_code=404, detail="Tailored CV not found")
    content = base64.b64decode(app_doc["tailored_cv_file_b64"])
    filename = app_doc.get("tailored_cv_filename") or "tailored_cv.docx"
    return FastAPIResponse(
        content=content,
        media_type=app_doc.get("tailored_cv_mime") or "application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@api_router.get("/admin/applications/{application_id}/cover-letter")
async def admin_download_cover_letter(application_id: str, admin: User = Depends(require_admin_user)):
    from fastapi.responses import Response as FastAPIResponse

    app_doc = await db.applications.find_one(
        {"application_id": application_id},
        {"_id": 0, "tailored_cover_letter": 1, "cover_letter": 1},
    )
    if not app_doc:
        raise HTTPException(status_code=404, detail="Application not found")
    cover_letter = app_doc.get("tailored_cover_letter") or app_doc.get("cover_letter")
    if not cover_letter:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    content = cover_letter_to_text(cover_letter)
    return FastAPIResponse(
        content=content,
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{application_id}_cover_letter.txt"'},
    )


@api_router.post("/admin/applications/{application_id}/send-email")
async def admin_send_application_email(
    application_id: str,
    body: AdminSendApplicationEmail,
    admin: User = Depends(require_admin_user),
):
    """Send the tailored CV + cover letter to a job's recruiter contact email.

    Used for job sources (e.g. France Travail) that don't expose an apply API
    but do publish a recruiter contact email. Sent from Hirly's own
    transactional sender with Reply-To set to the candidate, so this never
    requires the candidate's own mailbox credentials or a Gmail "send" OAuth
    grant. A human operator reviews and triggers this from the manual
    completion queue, so nothing is sent without a person in the loop.
    """
    app_doc = await db.applications.find_one({"application_id": application_id}, {"_id": 0})
    if not app_doc:
        raise HTTPException(status_code=404, detail="Application not found")
    job = await db.jobs.find_one({"job_id": app_doc.get("job_id")}, {"_id": 0}) or {}
    user_doc = await db.users.find_one({"user_id": app_doc.get("user_id")}, {"_id": 0}) or {}
    profile = await db.profiles.find_one({"user_id": app_doc.get("user_id")}, {"_id": 0}) or {}

    to_email = (body.to_email or job.get("contact_email") or "").strip()
    if not to_email:
        raise HTTPException(status_code=400, detail="No recruiter contact email available for this job")

    candidate_email = app_doc.get("user_email") or user_doc.get("email") or ""
    if INBOUND_MANAGED_EMAIL_ENABLED and application_id:
        # Same managed-inbox routing as the automated ATS submission path --
        # the recruiter's reply lands on our own infrastructure instead of
        # the candidate's real mailbox. candidate_name (used in the email
        # body/subject below) stays real either way.
        candidate_email = managed_reply_address(application_id)
    candidate_name = (profile.get("contact") or {}).get("name") or user_doc.get("name") or ""
    job_title = job.get("title") or "the role"
    company = job.get("company") or ""

    subject = (body.subject or "").strip() or f"Candidature — {job_title}{f' chez {company}' if company else ''} — {candidate_name or candidate_email}"
    cover_letter = app_doc.get("tailored_cover_letter") or app_doc.get("cover_letter") or {}
    default_body = cover_letter_to_text(cover_letter) if cover_letter else (
        f"Bonjour,\n\nVeuillez trouver ci-joint ma candidature pour le poste de {job_title}"
        f"{f' chez {company}' if company else ''}.\n\nCordialement,\n{candidate_name or candidate_email}"
    )
    body_text = (body.body_text or "").strip() or default_body

    attachments: List[Tuple[str, bytes, str]] = []
    if app_doc.get("tailored_cv_file_b64"):
        try:
            attachments.append((
                app_doc.get("tailored_cv_filename") or "CV.docx",
                base64.b64decode(app_doc["tailored_cv_file_b64"]),
                app_doc.get("tailored_cv_mime") or DOCX_MIME,
            ))
        except Exception:
            logger.warning("admin_send_application_email_cv_decode_failed application_id=%s", application_id)

    result = await send_application_email(
        to_email=to_email,
        candidate_email=candidate_email,
        candidate_name=candidate_name,
        subject=subject,
        body_text=body_text,
        attachments=attachments,
    )
    if not result.get("ok"):
        raise HTTPException(status_code=502, detail=result.get("error") or "Failed to send application email")

    now = datetime.now(timezone.utc).isoformat()
    timeline = _append_admin_timeline(app_doc, {
        "event_id": f"timeline_{uuid.uuid4().hex[:12]}",
        "type": "application_email_sent",
        "to_email": to_email,
        "transport": result.get("transport"),
        "author_user_id": admin.user_id,
        "author_email": admin.email,
        "created_at": now,
    })
    update: Dict[str, Any] = {
        "admin_timeline": timeline,
        "application_email_sent_at": now,
        "application_email_sent_to": to_email,
        "updated_at": now,
    }
    if body.mark_manually_submitted:
        update["admin_status"] = "manually_submitted"
        update["manual_status"] = "manually_submitted"
        update["submission_status"] = "submitted"
        update["submitted_at"] = app_doc.get("submitted_at") or now
        update["manual_status_updated_by"] = admin.email
        update["manual_status_updated_at"] = now
    await db.applications.update_one({"application_id": application_id}, {"$set": update})
    updated = await db.applications.find_one({"application_id": application_id}, {"_id": 0})
    return {"ok": True, "transport": result.get("transport"), "application": _normalize_application_status_fields(updated or {})}


@api_router.get("/applications/greenhouse/form-preview")
async def greenhouse_form_preview(job_id: str, user: User = Depends(get_current_user)):
    job = await db.jobs.find_one({"job_id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.get("ats_provider") != "greenhouse":
        raise HTTPException(status_code=400, detail="Job is not a Greenhouse job")

    board_token = job.get("board_token")
    greenhouse_job_id = job.get("provider_job_id")
    external_id = job.get("external_id") or ""
    if (not board_token or not greenhouse_job_id) and ":" in external_id:
        board_token, greenhouse_job_id = external_id.split(":", 1)
    if not board_token or not greenhouse_job_id:
        raise HTTPException(status_code=400, detail="Greenhouse board token or job id is missing")

    provider = get_board_provider("greenhouse")
    try:
        preview = await provider.inspect_application_form(board_token, greenhouse_job_id)
    except httpx.HTTPStatusError as exc:
        logger.warning(
            "Greenhouse form preview failed: job_id=%s board=%s greenhouse_job_id=%s status=%s",
            job_id,
            board_token,
            greenhouse_job_id,
            exc.response.status_code if exc.response else None,
        )
        raise HTTPException(status_code=502, detail="Greenhouse application form is unavailable") from exc
    except Exception as exc:
        logger.warning("Greenhouse form preview failed: job_id=%s error=%s", job_id, exc)
        raise HTTPException(status_code=502, detail="Greenhouse application form preview failed") from exc

    return {
        "job_id": job["job_id"],
        "company": job.get("company"),
        "title": job.get("title"),
        "application_url": preview["application_url"] or job.get("external_url"),
        "ats_provider": "greenhouse",
        "fields": preview["fields"],
        "supports_auto_submit": preview["supports_auto_submit"],
        "blockers": preview["blockers"],
    }


async def _load_or_create_agent_application(
    job_id: str,
    user: User,
    *,
    require_tailored_package: bool = True,
) -> tuple[Dict[str, Any], Dict[str, Any], Dict[str, Any]]:
    job = await db.jobs.find_one({"job_id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    profile = await db.profiles.find_one({"user_id": user.user_id}, {"_id": 0})
    if not profile or not profile.get("cv_text"):
        raise HTTPException(status_code=400, detail="Upload your CV before preparing an auto-apply attempt.")

    app_doc = await db.applications.find_one(
        {"user_id": user.user_id, "job_id": job_id},
        {"_id": 0},
        sort=[("created_at", -1)],
    )
    package_missing = (
        not app_doc
        or not app_doc.get("tailored_resume_structured")
        or not app_doc.get("tailored_cover_letter")
        or not app_doc.get("tailored_cv_file_b64")
    )
    if package_missing:
        # Auto-apply can upload the profile CV directly — skip the slow AI
        # package generation that often trips Railway's gateway timeout.
        if (
            not require_tailored_package
            and (profile.get("cv_original_b64") or profile.get("cv_filename"))
        ):
            if not app_doc:
                now = datetime.now(timezone.utc).isoformat()
                app_doc = {
                    "application_id": f"app_{uuid.uuid4().hex[:12]}",
                    "user_id": user.user_id,
                    "job_id": job_id,
                    "status": "applied",
                    "created_at": now,
                    "updated_at": now,
                }
                await db.applications.insert_one(app_doc)
            return job, profile, app_doc

        generated_doc = await _generate_application_doc(user, profile, job)
        if app_doc:
            generated_doc["application_id"] = app_doc["application_id"]
            generated_doc["created_at"] = app_doc.get("created_at") or generated_doc["created_at"]
            await db.applications.update_one(
                {"user_id": user.user_id, "job_id": job_id, "application_id": app_doc["application_id"]},
                {"$set": generated_doc},
                upsert=True,
            )
        else:
            await db.applications.insert_one(generated_doc)
        app_doc = await db.applications.find_one(
            {"user_id": user.user_id, "job_id": job_id},
            {"_id": 0},
            sort=[("created_at", -1)],
        )

    if not app_doc:
        raise HTTPException(status_code=500, detail="Application package could not be created")
    return job, profile, app_doc


def _browser_engine_headless() -> bool:
    # Default headed (false). Runtime config forces BROWSER_HEADLESS=0 on boot.
    return os.environ.get("BROWSER_HEADLESS", "false").lower() in ("1", "true", "yes", "on")


def _resolve_auto_apply_headless(requested: Optional[bool]) -> bool:
    # Prefer runtime/env; client headless=true is ignored when BROWSER_HEADLESS=0.
    if requested is None:
        return effective_headless(_browser_engine_headless())
    return effective_headless(requested)


def _dev_tools_enabled() -> bool:
    return (
        os.environ.get("ENVIRONMENT", "").strip().lower() == "development"
        or os.environ.get("DEV_TOOLS_ENABLED", "false").strip().lower() in ("1", "true", "yes", "on")
    )


def _log_browser_exception(message: str, exc: Exception) -> None:
    if _dev_tools_enabled():
        logger.exception(message)
    else:
        logger.warning("%s: %s", message, exc.__class__.__name__)


def _sanitize_agent_result_for_application(result_dict: Dict[str, Any]) -> Dict[str, Any]:
    result_for_storage = dict(result_dict)
    if result_for_storage.get("screenshot_b64"):
        result_for_storage["screenshot_b64"] = f"<omitted, {len(result_dict['screenshot_b64'])} base64 chars>"
    return result_for_storage


def _agent_missing_information(result_dict: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Job-specific questions the agent couldn't safely answer on its own --
    reuses the exact `field_name`/`label`/`options` shape the tracker UI's
    existing missing-info form already renders (ApplicationDetailPanel.jsx),
    so no frontend change is needed to surface these to the user.
    """
    items: List[Dict[str, Any]] = []
    for field in result_dict.get("unfilled_required_fields") or []:
        if not isinstance(field, dict):
            continue
        items.append({
            "field_name": field.get("stable_field_id"),
            "label": field.get("label") or field.get("name") or "This field",
            "options": [
                (opt.get("label") if isinstance(opt, dict) else opt)
                for opt in (field.get("options") or [])
            ],
            "reason": "agent_could_not_answer",
        })
    return items


def _agent_submission_status(result_dict: Dict[str, Any]) -> str:
    if result_dict.get("captcha_required"):
        return "blocked_captcha"
    if result_dict.get("login_wall_detected"):
        return "blocked"
    if result_dict.get("ready_for_final_click"):
        return "ready"
    if result_dict.get("unfilled_required_fields"):
        return "action_required"
    if result_dict.get("blockers"):
        return "blocked"
    return "prepared"


async def _run_agent_apply(job_id: str, user: User, click_submit: bool = False) -> Dict[str, Any]:
    try:
        job, profile, app_doc = await _load_or_create_agent_application(job_id, user)
    except HTTPException:
        raise
    except Exception as exc:
        _log_browser_exception("Apply agent load_application failed", exc)
        raise HTTPException(
            status_code=502,
            detail={
                "phase": "load_application",
                "exception_class": exc.__class__.__name__,
                "message": str(exc).strip() or "Failed to load application context.",
            },
        ) from exc

    if click_submit and user.require_review_before_send and app_doc.get("document_review_status") != "approved":
        raise HTTPException(
            status_code=409,
            detail="Please review and approve your CV and cover letter in the Review tab before this application can be submitted.",
        )

    try:
        result = await run_apply_attempt(
            job=job,
            app_doc=app_doc,
            profile=profile,
            user=user.model_dump(mode="json"),
            click_submit=click_submit,
            headless=_browser_engine_headless(),
            db=db,
        )
    except ApplyAgentError as exc:
        _log_browser_exception(f"Apply agent failed during {exc.phase}", exc)
        raise HTTPException(status_code=502, detail=exc.safe_detail()) from exc
    except Exception as exc:
        _log_browser_exception("Apply agent unexpected failure", exc)
        raise HTTPException(
            status_code=502,
            detail={
                "phase": "run_apply_attempt",
                "exception_class": exc.__class__.__name__,
                "message": str(exc).strip() or "Unexpected apply agent failure.",
            },
        ) from exc

    result_dict = result.to_dict()
    result_for_storage = _sanitize_agent_result_for_application(result_dict)
    now = datetime.now(timezone.utc).isoformat()
    if not click_submit:
        await db.applications.update_one(
            {"application_id": app_doc["application_id"], "user_id": user.user_id},
            {"$set": {
                "agent_apply_result": result_for_storage,
                "agent_apply_prepared_at": now,
                "updated_at": now,
                "submission_status": _agent_submission_status(result_dict),
                "prepared_missing_information": _agent_missing_information(result_dict),
            }},
        )

    from apply_agent.recipes import get_domain_trust, recipe_key_for_url
    domain_trust = None
    if result_dict.get("application_url"):
        recipe_key = recipe_key_for_url(result_dict["application_url"])
        domain_trust = await get_domain_trust(db, recipe_key)

    return {
        "job_id": job["job_id"],
        "application_id": app_doc["application_id"],
        "company": job.get("company"),
        "title": job.get("title"),
        "domain_trust": domain_trust,
        **result_dict,
    }


async def _store_agent_run(*, user: User, result: Dict[str, Any], dry_run: bool) -> str:
    now = datetime.now(timezone.utc).isoformat()
    run_id = f"agent_run_{uuid.uuid4().hex[:16]}"
    success_detected = bool(result.get("success_detected"))
    status = (
        "dry_run"
        if dry_run
        else "submitted"
        if success_detected
        else "blocked_captcha"
        if result.get("captcha_required")
        else "unknown"
        if result.get("submit_clicked") and result.get("failure_reason") == "submission_status_unknown"
        else "failed"
    )
    run_doc = {
        "run_id": run_id,
        "application_id": result.get("application_id"),
        "job_id": result.get("job_id"),
        "user_id": user.user_id,
        "provider": result.get("provider") or "unknown",
        "status": status,
        "dry_run": dry_run,
        "screenshots": {"screenshot_b64": result.get("screenshot_b64")},
        "success_detected": success_detected,
        "captcha_required": bool(result.get("captcha_required")),
        "login_wall_detected": bool(result.get("login_wall_detected")),
        "action_required": bool(result.get("action_required")),
        "failure_reason": result.get("failure_reason"),
        "post_submit_errors": result.get("post_submit_errors"),
        "confirmation_text_found": result.get("confirmation_text_found"),
        "final_url": result.get("final_url") or result.get("application_url"),
        "created_at": now,
        "updated_at": now,
    }
    await db.browser_submission_runs.insert_one(run_doc)
    return run_id


@api_router.post("/applications/agent/prepare")
async def agent_prepare(body: AgentApplyRequest, user: User = Depends(get_current_user)):
    """Run the apply agent against a job and stop before any submit click."""
    return await _run_agent_apply(body.job_id, user)


@api_router.post("/applications/agent/submit")
async def agent_submit(body: AgentApplyRequest, user: User = Depends(get_current_user)):
    """Run the apply agent and, if everything required is filled from an
    approved source and dry-run is disabled for this account, click submit.
    """
    dry_run = _env_enabled("BROWSER_SUBMIT_DRY_RUN", "true")
    if dry_run:
        result = await _run_agent_apply(body.job_id, user)
        run_id = await _store_agent_run(user=user, result=result, dry_run=True)
        return {**result, "dry_run": True, "agent_run_id": run_id, "stopped_before_submit": True}

    _require_agent_real_submit_allowed(user)
    result = await _run_agent_apply(body.job_id, user, click_submit=True)
    run_id = await _store_agent_run(user=user, result=result, dry_run=False)
    now = datetime.now(timezone.utc).isoformat()
    if result.get("success_detected"):
        await db.applications.update_one(
            {"application_id": result["application_id"], "user_id": user.user_id},
            {"$set": {"submission_status": "submitted", "agent_run_id": run_id, "updated_at": now}},
        )
    elif result.get("captcha_required"):
        await db.applications.update_one(
            {"application_id": result["application_id"], "user_id": user.user_id},
            {"$set": {"submission_status": "blocked_captcha", "agent_run_id": run_id, "updated_at": now}},
        )
    else:
        app_doc = await db.applications.find_one(
            {"application_id": result["application_id"], "user_id": user.user_id},
            {"_id": 0},
        )
        job_doc = await db.jobs.find_one({"job_id": result.get("job_id")}, {"_id": 0})
        latest_run = {
            "failure_reason": result.get("failure_reason"),
            "post_submit_errors": result.get("post_submit_errors"),
            "login_wall_detected": result.get("login_wall_detected"),
            "captcha_required": result.get("captcha_required"),
        }
        if app_doc:
            await maybe_auto_expire_application(
                db,
                app_doc,
                job_doc=job_doc,
                latest_run=latest_run,
                source="agent_submit",
            )
    return {**result, "agent_run_id": run_id}


@api_router.post("/applications/agent/submission-benchmark")
async def agent_submission_benchmark(body: AgentSubmissionBenchmarkRequest, user: User = Depends(get_current_user)):
    unique_job_ids = []
    seen = set()
    for job_id in body.job_ids[:10]:
        if job_id and job_id not in seen:
            unique_job_ids.append(job_id)
            seen.add(job_id)
    if not unique_job_ids:
        raise HTTPException(status_code=400, detail="job_ids is required")

    dry_run_enabled = _env_enabled("BROWSER_SUBMIT_DRY_RUN", "true")
    real_submit_enabled = bool(body.allow_real_submit and not dry_run_enabled)
    results = []
    for job_id in unique_job_ids:
        try:
            if body.run_submit and real_submit_enabled:
                result = await _run_agent_apply(job_id, user, click_submit=True)
                run_id = await _store_agent_run(user=user, result=result, dry_run=False)
                result = {**result, "agent_run_id": run_id}
            else:
                result = await _run_agent_apply(job_id, user, click_submit=False)
                run_id = await _store_agent_run(user=user, result=result, dry_run=True)
                result = {**result, "dry_run": True, "agent_run_id": run_id, "stopped_before_submit": True}
            results.append({
                "ok": True,
                "job_id": job_id,
                "provider": result.get("provider"),
                "ready_for_final_click": result.get("ready_for_final_click"),
                "captcha_required": result.get("captcha_required"),
                "login_wall_detected": result.get("login_wall_detected"),
                "blockers": result.get("blockers"),
                "success_likelihood": result.get("success_likelihood"),
                "success_detected": result.get("success_detected"),
            })
        except HTTPException as exc:
            detail = exc.detail if isinstance(exc.detail, dict) else {"message": exc.detail}
            results.append({"ok": False, "job_id": job_id, "error": detail})
        except Exception as exc:
            logger.exception("Agent benchmark job failed: user_id=%s job_id=%s", user.user_id, job_id)
            results.append({"ok": False, "job_id": job_id, "error": {"exception_class": exc.__class__.__name__, "message": str(exc)[:500]}})

    total = len(results)
    ready_count = sum(1 for item in results if item.get("ok") and item.get("ready_for_final_click") and not item.get("blockers"))
    captcha_count = sum(1 for item in results if item.get("captcha_required"))
    login_wall_count = sum(1 for item in results if item.get("login_wall_detected"))
    likelihoods = [float(item["success_likelihood"]) for item in results if item.get("success_likelihood") is not None]
    denominator = total or 1
    return {
        "dry_run": not real_submit_enabled,
        "real_submit_enabled": real_submit_enabled,
        "total": total,
        "ready_count": ready_count,
        "ready_rate": round(ready_count / denominator, 4),
        "captcha_count": captcha_count,
        "captcha_rate": round(captcha_count / denominator, 4),
        "login_wall_count": login_wall_count,
        "average_success_likelihood": round(sum(likelihoods) / len(likelihoods), 4) if likelihoods else 0,
        "results": results,
    }

def _split_name(full_name: Optional[str]) -> Dict[str, str]:
    parts = (full_name or "").strip().split()
    if not parts:
        return {"first_name": "", "last_name": ""}
    if len(parts) == 1:
        return {"first_name": parts[0], "last_name": ""}
    return {"first_name": parts[0], "last_name": " ".join(parts[1:])}


def _field_missing_key(field: Dict[str, Any]) -> Optional[str]:
    text = " ".join([str(field.get("name") or ""), str(field.get("label") or "")]).lower()
    checks = [
        ("visa status", ("visa", "sponsorship", "sponsor")),
        ("work authorization", ("work authorization", "authorized to work", "right to work", "legally authorized")),
        ("salary expectations", ("salary", "compensation", "pay expectation", "expected pay")),
        ("relocation preference", ("relocation", "relocate")),
        ("start date", ("start date", "available to start", "availability")),
    ]
    for key, terms in checks:
        if any(term in text for term in terms):
            return key
    return None


def _field_text(field: Dict[str, Any]) -> str:
    return " ".join([str(field.get("name") or ""), str(field.get("label") or "")]).lower()


def _is_empty_answer(value: Any) -> bool:
    if value is None:
        return True
    if isinstance(value, str):
        return not value.strip()
    if isinstance(value, list):
        return len([item for item in value if not _is_empty_answer(item)]) == 0
    return False


def _sensitive_field_reason(field: Dict[str, Any]) -> Optional[str]:
    text = _field_text(field)
    category = field.get("field_category")
    checks = [
        ("work authorization", ("work authorization", "authorized to work", "right to work", "legally authorized", "eligible to work")),
        ("visa sponsorship", ("visa", "sponsorship", "sponsor")),
        ("disability status", ("disability", "disabled")),
        ("veteran status", ("veteran", "armed forces", "military service")),
        ("demographic question", ("gender", "race", "ethnicity", "hispanic", "pronouns", "sexual orientation")),
        ("criminal history", ("criminal", "conviction", "felony", "background check")),
        ("salary expectations", ("salary", "compensation", "pay expectation", "expected pay")),
    ]
    if category in ("demographic", "eeoc"):
        return "demographic question"
    for reason, terms in checks:
        if any(term in text for term in terms):
            return reason
    return None


def _profile_answer_key(field: Dict[str, Any]) -> Optional[str]:
    text = _field_text(field)
    if any(term in text for term in ("residence country", "country of residence", "current country")):
        return "residence_country"
    if any(term in text for term in ("authorized to work", "work authorization", "right to work", "legally authorized", "eligible to work")):
        return "work_authorization_countries"
    if any(term in text for term in ("require sponsorship", "need sponsorship", "visa sponsorship", "sponsor")):
        if any(term in text for term in ("future", "later", "eventually")):
            return "requires_sponsorship_future"
        return "requires_sponsorship_now"
    if any(term in text for term in ("desired work countries", "countries would you like", "work countries")):
        return "desired_work_countries"
    if any(term in text for term in ("salary", "compensation", "pay expectation", "expected pay")):
        return "salary_expectation"
    if any(term in text for term in ("start date", "available to start", "availability")):
        return "earliest_start_date"
    if any(term in text for term in ("relocation", "relocate")):
        return "willing_to_relocate"
    return None


def _profile_saved_answer(profile: Dict[str, Any], field: Dict[str, Any]) -> Any:
    key = _profile_answer_key(field)
    if not key:
        return None
    answers_profile = profile.get("application_answers_profile") or {}
    value = answers_profile.get(key)
    if isinstance(value, list):
        return ", ".join(str(item) for item in value if not _is_empty_answer(item))
    if isinstance(value, bool):
        return "Yes" if value else "No"
    return value


def _profile_has_explicit_answer(profile: Dict[str, Any], field: Dict[str, Any]) -> bool:
    if not _is_empty_answer(_profile_saved_answer(profile, field)):
        return True
    text = _field_text(field)
    candidate_sources = [
        profile.get("application_answers"),
        profile.get("application_preferences"),
        profile.get("legal"),
        profile.get("work_authorization"),
        profile.get("candidate_facts"),
    ]
    for source in candidate_sources:
        if not source:
            continue
        serialized = json.dumps(source, default=str).lower()
        if any(term in serialized for term in text.split() if len(term) > 4):
            return True
    return False


def _missing_info_item(field: Dict[str, Any], reason: str) -> Dict[str, Any]:
    return {
        "field_name": _public_missing_field_name(field),
        "field_id": field.get("field_id") or field.get("id"),
        "label": str(field.get("label") or field.get("name") or "Unknown field"),
        "question": str(field.get("question") or field.get("label") or field.get("name") or "Unknown field"),
        "reason": reason,
        "field_type": field.get("type") or "input_text",
        "type": field.get("type") or "input_text",
        "options": field.get("options") or [],
        "suggested_profile_key": field.get("suggested_profile_key"),
    }


def _all_payload_fields(payload: Dict[str, Any]) -> List[Dict[str, Any]]:
    fields = [
        {"name": "first_name", "label": "First name", "type": "input_text", "options": []},
        {"name": "last_name", "label": "Last name", "type": "input_text", "options": []},
        {"name": "email", "label": "Email", "type": "input_text", "options": []},
    ]
    fields.extend([q for q in payload.get("questions") or [] if isinstance(q, dict)])
    return fields


def _normalize_missing_information(items: List[Any], known_fields: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    normalized: List[Dict[str, Any]] = []
    for item in items:
        if not item:
            continue
        if isinstance(item, dict):
            field_name = str(item.get("field_name") or item.get("name") or "")
            label = str(item.get("label") or item.get("question") or field_name or "Unknown field")
            field = next(
                (
                    known
                    for known in known_fields
                    if _canonical_field_key(known) == _canonical_field_name(field_name or label)
                ),
                {},
            )
            normalized.append({
                "field_name": _public_missing_field_name(field or {"name": field_name, "label": label}),
                "field_id": item.get("field_id") or item.get("id"),
                "label": label or str(field.get("label") or "Unknown field"),
                "question": item.get("question") or label or str(field.get("label") or "Unknown field"),
                "reason": item.get("reason") or "missing_information",
                "field_type": item.get("field_type") or field.get("type") or "input_text",
                "type": item.get("type") or item.get("field_type") or field.get("type") or "input_text",
                "options": item.get("options") or field.get("options") or [],
                "suggested_profile_key": item.get("suggested_profile_key") or field.get("suggested_profile_key"),
            })
            continue

        text = str(item).strip()
        if not text:
            continue
        text_lower = text.lower()
        field = next(
            (
                known
                for known in known_fields
                if _canonical_field_key(known) == _canonical_field_name(text_lower)
            ),
            {},
        )
        normalized.append({
            "field_name": _public_missing_field_name(field or {"name": text, "label": text}),
            "label": str(field.get("label") or text),
            "reason": "missing_information",
            "field_type": field.get("type") or "input_text",
            "options": field.get("options") or [],
        })
    return _dedupe_missing_information(normalized)


def _canonical_field_name(value: Any) -> str:
    text = str(value or "").lower()
    text = re.sub(r"[^a-z0-9]+", "_", text)
    text = re.sub(r"_+", "_", text).strip("_")
    if "privacy" in text and "policy" in text:
        return "privacy_policy_agreement"
    if "how_did_you_hear" in text or ("hear" in text and "job" in text):
        return "referral_source"
    return text


def _canonical_field_key(field: Dict[str, Any]) -> str:
    combined = " ".join([str(field.get("name") or ""), str(field.get("label") or "")])
    return _canonical_field_name(combined)


def _canonical_value_map(values: Dict[str, Any]) -> Dict[str, Any]:
    mapped = {}
    for key, value in (values or {}).items():
        canonical = _canonical_field_name(key)
        if canonical and canonical not in mapped:
            mapped[canonical] = value
    return mapped


def _public_missing_field_name(field: Dict[str, Any]) -> str:
    canonical = _canonical_field_key(field)
    if canonical in ("privacy_policy_agreement", "referral_source"):
        return canonical
    return str(field.get("name") or canonical)


def _greenhouse_safe_default_answer(field: Dict[str, Any]) -> Optional[str]:
    if not field.get("required"):
        return None
    if _sensitive_field_reason(field):
        return None
    canonical = _canonical_field_key(field)
    text = _field_text(field)
    if canonical == "referral_source" or canonical == "source":
        return "Swiipr"
    if "source" in canonical and "job" in text:
        return "Swiipr"
    if canonical == "privacy_policy_agreement":
        return "I agree"
    if any(term in text for term in ("consent", "acknowledgement", "acknowledgment", "i agree")):
        return "I agree"
    return None


def _missing_information_summary(items: List[Any]) -> str:
    parts = []
    for item in items:
        if isinstance(item, dict):
            label = item.get("label") or item.get("field_name") or "unknown field"
            reason = item.get("reason") or "missing_information"
            parts.append(f"{label}: {reason}")
        elif item:
            parts.append(str(item))
    return "; ".join(sorted(set(parts)))


def _dedupe_missing_information(items: List[Any]) -> List[Any]:
    seen = set()
    result = []
    for item in items:
        if not item:
            continue
        if isinstance(item, dict):
            key = _canonical_field_name(item.get("field_name") or item.get("label") or "")
        else:
            key = _canonical_field_name(item)
        if key in seen:
            continue
        seen.add(key)
        result.append(item)
    return result


def _required_empty_payload_fields(payload: Dict[str, Any]) -> List[Dict[str, str]]:
    fields = payload.get("fields") or {}
    canonical_fields = _canonical_value_map(fields)
    missing = []
    for name, label in (("first_name", "First name"), ("last_name", "Last name"), ("email", "Email")):
        if _is_empty_answer(fields.get(name)):
            missing.append({
                "field_name": name,
                "label": label,
                "reason": "required_empty_answer",
                "field_type": "input_text",
                "options": [],
            })
    for question in payload.get("questions") or []:
        if not isinstance(question, dict) or not question.get("required"):
            continue
        name = question.get("name")
        canonical = _canonical_field_key(question)
        value = fields.get(name)
        if _is_empty_answer(value):
            value = canonical_fields.get(canonical)
        if _is_empty_answer(value):
            value = question.get("value")
        if _is_empty_answer(value):
            missing.append(_missing_info_item(question, "required_empty_answer"))
    return missing


def _required_fields_count(payload: Dict[str, Any]) -> int:
    count = 0
    fields = payload.get("fields") or {}
    for name in ("first_name", "last_name", "email"):
        if name in fields:
            count += 1
    count += sum(
        1
        for question in payload.get("questions") or []
        if isinstance(question, dict) and question.get("required")
    )
    return count


def _field_by_name_from_payload(payload: Dict[str, Any], field_name: str) -> Optional[Dict[str, Any]]:
    canonical = _canonical_field_name(field_name)
    for question in payload.get("questions") or []:
        if not isinstance(question, dict):
            continue
        question_names = {
            _canonical_field_key(question),
            _canonical_field_name(question.get("name")),
            _canonical_field_name(question.get("field_name")),
            _canonical_field_name(question.get("field_id")),
        }
        if canonical in question_names:
            return question
    for candidate_field in ("first_name", "last_name", "email", "phone"):
        if _canonical_field_name(candidate_field) == canonical:
            return {"name": candidate_field, "label": candidate_field.replace("_", " ").title(), "type": "input_text", "options": []}
    return None


def _profile_answer_updates_from_resolved_fields(payload: Dict[str, Any], answers: Dict[str, Any]) -> Dict[str, Any]:
    updates = {}
    for field_name, value in answers.items():
        field = _field_by_name_from_payload(payload, field_name)
        if not field or _is_empty_answer(value):
            continue
        suggested_key = field.get("suggested_profile_key")
        if suggested_key:
            updates[f"application_defaults.{suggested_key}"] = value
            if suggested_key.startswith("eeo_") and _is_demographic_decline_answer(value):
                updates["application_defaults.prefer_not_to_say_demographics"] = True
            continue
        key = _profile_answer_key(field)
        if key:
            updates[f"application_answers_profile.{key}"] = value
    return updates


def _is_demographic_decline_answer(value: Any) -> bool:
    text = str(value or "").strip().lower()
    canonical = _canonical_field_name(value)
    return any(token in text for token in (
        "prefer not",
        "decline",
        "do not wish",
        "don't wish",
        "choose not to disclose",
        "do not want to answer",
    )) or any(token in canonical for token in ("prefer_not", "decline", "do_not_wish", "choose_not_to_disclose"))


def _remove_resolved_missing_items(missing_items: List[Any], answers: Dict[str, Any]) -> List[Any]:
    result = []
    for item in missing_items:
        if isinstance(item, dict):
            field_name = item.get("field_name")
            if field_name in answers and not _is_empty_answer(answers.get(field_name)):
                continue
        result.append(item)
    return result


def _greenhouse_submit_dry_run_enabled() -> bool:
    return os.environ.get("GREENHOUSE_SUBMIT_DRY_RUN", "true").lower() not in ("0", "false", "no", "off")


def _payload_for_storage(payload: Dict[str, Any]) -> Dict[str, Any]:
    stored = json.loads(json.dumps(payload))
    files = stored.get("files") or {}
    for file_info in files.values():
        if isinstance(file_info, dict):
            if file_info.get("b64"):
                file_info["b64"] = None
                file_info["b64_stored_on_application"] = True
            if file_info.get("text"):
                file_info["text"] = None
                file_info["text_stored_on_application"] = True
    return stored


def _payload_for_response(payload: Dict[str, Any]) -> Dict[str, Any]:
    safe = json.loads(json.dumps(payload))
    files = safe.get("files") or {}
    for file_info in files.values():
        if isinstance(file_info, dict):
            if file_info.get("b64"):
                file_info["b64"] = f"<base64 omitted, {len(file_info['b64'])} chars>"
            if file_info.get("text"):
                file_info["text"] = f"<text omitted, {len(file_info['text'])} chars>"
    return safe


def _greenhouse_submission_endpoint(payload: Dict[str, Any], job: Dict[str, Any]) -> str:
    url = payload.get("url")
    if url:
        return url
    board_token = job.get("board_token")
    greenhouse_job_id = job.get("provider_job_id")
    external_id = job.get("external_id") or ""
    if (not board_token or not greenhouse_job_id) and ":" in external_id:
        board_token, greenhouse_job_id = external_id.split(":", 1)
    if not board_token or not greenhouse_job_id:
        raise HTTPException(status_code=400, detail="Greenhouse board token or job id is missing")
    return f"https://boards-api.greenhouse.io/v1/boards/{board_token}/jobs/{greenhouse_job_id}"


def _coerce_greenhouse_boolean(value: Any) -> str:
    if isinstance(value, bool):
        return "true" if value else "false"
    text = str(value or "").strip().lower()
    if text in ("yes", "y", "true", "1", "i agree", "agree"):
        return "true"
    if text in ("no", "n", "false", "0"):
        return "false"
    return str(value or "")


def _map_greenhouse_option_value(field: Dict[str, Any], raw_value: Any) -> tuple[Any, Optional[Dict[str, Any]]]:
    options = field.get("options") or []
    if not options or _is_empty_answer(raw_value):
        return raw_value, None

    values = raw_value if isinstance(raw_value, list) else [raw_value]
    mapped_values = []
    mappings = []
    for item in values:
        item_text = str(item).strip()
        item_key = _canonical_field_name(item_text)
        match = next(
            (
                option for option in options
                if str(option.get("value") or "").strip() == item_text
                or str(option.get("label") or "").strip().lower() == item_text.lower()
                or _canonical_field_name(option.get("label") or option.get("value")) == item_key
                or (
                    item_text.lower() in ("i agree", "agree", "yes", "true")
                    and str(option.get("label") or option.get("value") or "").strip().lower() in ("i agree", "agree", "yes", "true")
                )
            ),
            None,
        )
        if match:
            mapped = match.get("value") if match.get("value") is not None else match.get("label")
            mapped_values.append(mapped)
            mappings.append({"input": item, "mapped_value": mapped, "label": match.get("label")})
        else:
            mapped_values.append(item)
            mappings.append({"input": item, "mapped_value": item, "error": "option_not_found"})

    mapped_result = mapped_values if isinstance(raw_value, list) else mapped_values[0]
    return mapped_result, {"field_name": field.get("name"), "label": field.get("label"), "mappings": mappings}


def _prepared_question_map(payload: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    question_map = {}
    for question in payload.get("questions") or []:
        if isinstance(question, dict) and question.get("name"):
            question_map[question["name"]] = question
    return question_map


def _build_greenhouse_submission_parts(app_doc: Dict[str, Any], payload: Dict[str, Any]) -> Dict[str, Any]:
    raw_fields = payload.get("fields") or {}
    question_map = _prepared_question_map(payload)
    data: Dict[str, Any] = {}
    data_items = []
    option_mappings = []
    validation_errors = []

    for key, value in raw_fields.items():
        field = question_map.get(key)
        field_type = (field or {}).get("type") or "input_text"
        mapped_value = value
        mapping = None

        if field_type in ("select", "radio", "multi_select", "checkbox"):
            mapped_value, mapping = _map_greenhouse_option_value(field or {"name": key, "options": []}, value)
            if mapping:
                option_mappings.append(mapping)
                if any(item.get("error") for item in mapping.get("mappings") or []):
                    validation_errors.append({
                        "type": "option_mapping_error",
                        "field_name": key,
                        "label": (field or {}).get("label"),
                        "value": value,
                        "options": (field or {}).get("options") or [],
                    })
        elif field_type == "boolean":
            mapped_value = _coerce_greenhouse_boolean(value)

        if isinstance(mapped_value, list):
            data[key] = [str(item) for item in mapped_value]
            for item in mapped_value:
                data_items.append((key, str(item)))
        else:
            data[key] = "" if mapped_value is None else str(mapped_value)
            data_items.append((key, data[key]))

    resume_b64 = app_doc.get("tailored_cv_file_b64")
    if not resume_b64:
        validation_errors.append({"type": "missing_file", "field_name": "resume", "message": "Tailored CV file is missing"})
        resume_content = b""
    else:
        try:
            resume_content = base64.b64decode(resume_b64)
        except Exception:
            validation_errors.append({"type": "invalid_file", "field_name": "resume", "message": "Tailored CV file is invalid base64"})
            resume_content = b""

    payload_files = payload.get("files") or {}
    resume_meta = payload_files.get("resume") or {}
    files: Dict[str, tuple] = {}
    if resume_content:
        files["resume"] = (
            resume_meta.get("filename") or app_doc.get("tailored_cv_filename") or "tailored_cv.docx",
            resume_content,
            resume_meta.get("mime") or app_doc.get("tailored_cv_mime") or "application/octet-stream",
        )

    cover_text = cover_letter_to_text(app_doc.get("tailored_cover_letter") or app_doc.get("cover_letter") or {})
    if cover_text.strip():
        cover_meta = payload_files.get("cover_letter") or {}
        files["cover_letter"] = (
            cover_meta.get("filename") or f"{app_doc.get('application_id', 'application')}_cover_letter.txt",
            cover_text.encode("utf-8"),
            cover_meta.get("mime") or "text/plain",
        )

    required_empty = _required_empty_payload_fields({"fields": data, "questions": payload.get("questions") or []})
    for item in required_empty:
        validation_errors.append({"type": "required_empty_field", **item})

    return {
        "data": data,
        "data_items": data_items,
        "files": files,
        "option_mappings": option_mappings,
        "validation_errors": validation_errors,
    }


def _greenhouse_submission_preview(app_doc: Dict[str, Any], payload: Dict[str, Any], job: Dict[str, Any]) -> Dict[str, Any]:
    parts = _build_greenhouse_submission_parts(app_doc, payload)
    data = parts["data"]
    files = parts["files"]
    return {
        "submit_url": _greenhouse_submission_endpoint(payload, job),
        "fields": [{"name": key, "value": value} for key, value in data.items()],
        "questions": [
            {
                "name": q.get("name"),
                "label": q.get("label"),
                "type": q.get("type"),
                "required": q.get("required"),
                "value": data.get(q.get("name")),
                "options": q.get("options") or [],
            }
            for q in payload.get("questions") or []
            if isinstance(q, dict)
        ],
        "option_mappings": parts["option_mappings"],
        "files": [
            {"field_name": name, "filename": file_tuple[0], "mime": file_tuple[2], "size_bytes": len(file_tuple[1])}
            for name, file_tuple in files.items()
        ],
        "validation_errors": parts["validation_errors"],
        "is_valid": not parts["validation_errors"],
    }


def _build_greenhouse_multipart(app_doc: Dict[str, Any], payload: Dict[str, Any]) -> tuple[List[tuple], Dict[str, tuple], Dict[str, Any]]:
    parts = _build_greenhouse_submission_parts(app_doc, payload)
    if parts["validation_errors"]:
        raise HTTPException(status_code=400, detail={"message": "Greenhouse submission payload is invalid", "errors": parts["validation_errors"]})
    return parts["data_items"], parts["files"], parts


def _legacy_build_greenhouse_multipart(app_doc: Dict[str, Any], payload: Dict[str, Any]) -> tuple[Dict[str, str], Dict[str, tuple]]:
    fields = payload.get("fields") or {}
    data = {str(key): "" if value is None else str(value) for key, value in fields.items()}

    resume_b64 = app_doc.get("tailored_cv_file_b64")
    if not resume_b64:
        raise HTTPException(status_code=400, detail="Tailored CV file is missing")
    try:
        resume_content = base64.b64decode(resume_b64)
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Tailored CV file is invalid") from exc

    payload_files = payload.get("files") or {}
    resume_meta = payload_files.get("resume") or {}
    files: Dict[str, tuple] = {
        "resume": (
            resume_meta.get("filename") or app_doc.get("tailored_cv_filename") or "tailored_cv.docx",
            resume_content,
            resume_meta.get("mime") or app_doc.get("tailored_cv_mime") or "application/octet-stream",
        )
    }

    cover_text = cover_letter_to_text(app_doc.get("tailored_cover_letter") or app_doc.get("cover_letter") or {})
    if cover_text.strip():
        cover_meta = payload_files.get("cover_letter") or {}
        files["cover_letter"] = (
            cover_meta.get("filename") or f"{app_doc.get('application_id', 'application')}_cover_letter.txt",
            cover_text.encode("utf-8"),
            cover_meta.get("mime") or "text/plain",
        )
    return data, files


def _greenhouse_response_metadata(response: httpx.Response) -> Dict[str, Any]:
    metadata: Dict[str, Any] = {
        "status_code": response.status_code,
        "headers": {
            key: value
            for key, value in response.headers.items()
            if key.lower() in ("location", "content-type", "x-request-id")
        },
    }
    try:
        body = response.json()
        if isinstance(body, dict):
            metadata["body"] = {
                key: body.get(key)
                for key in ("id", "application_id", "status", "message", "error", "errors", "validation_errors", "invalid_fields")
                if key in body
            }
        else:
            metadata["body_type"] = type(body).__name__
    except Exception:
        metadata["body_snippet"] = response.text[:500]
    return metadata


def _greenhouse_response_submission_id(response: httpx.Response) -> Optional[str]:
    try:
        body = response.json()
        if isinstance(body, dict):
            value = body.get("id") or body.get("application_id")
            return str(value) if value else None
    except Exception:
        pass
    location = response.headers.get("location")
    return location[-120:] if location else None


async def _generate_greenhouse_answers(
    profile: Dict[str, Any],
    job: Dict[str, Any],
    app_doc: Dict[str, Any],
    fields: List[Dict[str, Any]],
) -> Dict[str, Any]:
    custom_fields = [
        field for field in fields
        if field.get("field_category") in ("custom_question", "demographic", "eeoc")
    ]
    if not custom_fields:
        return {"answers": [], "missing_information": []}

    system_message = (
        "You generate truthful job application answers from provided candidate data. "
        "Return ONLY valid JSON. Never invent facts. If required information is missing, "
        "leave answer empty, lower confidence, and add a missing_information item."
    )
    prompt = f"""Prepare answers for Greenhouse application questions.

Candidate profile:
{json.dumps({
    "contact": profile.get("contact", {}),
    "summary": profile.get("summary"),
    "skills": profile.get("skills", []),
    "experience": profile.get("experience", []),
    "education": profile.get("education", []),
    "cv_text": profile.get("cv_text", "")[:12000],
}, indent=2)}

Tailored resume:
{json.dumps(app_doc.get("tailored_resume_structured") or app_doc.get("tailored_resume") or {}, indent=2)}

Tailored cover letter:
{json.dumps(app_doc.get("tailored_cover_letter") or app_doc.get("cover_letter") or {}, indent=2)}

Job:
{json.dumps({
    "title": job.get("title"),
    "company": job.get("company"),
    "description": job.get("clean_description") or job.get("description"),
    "requirements": job.get("requirements", []),
}, indent=2)}

Questions:
{json.dumps(custom_fields, indent=2)}

Rules:
- Answer only from the candidate/profile/CV/tailored application context.
- Do not answer visa, work authorization, salary, relocation, or start-date questions unless explicitly present in the candidate data.
- For select/multi_select fields, choose only from provided option labels/values when the answer is clear.
- For unknown required information, return an empty answer and include missing_information.

Return JSON:
{{
  "answers": [
    {{"field_name": "question_123", "question": "...", "answer": "...", "confidence": 0.0}}
  ],
  "missing_information": ["work authorization"]
}}"""
    response = await complete_json_text(system_message, prompt)
    return _parse_json_from_llm(response)


@api_router.post("/applications/greenhouse/prepare-submit")
async def greenhouse_prepare_submit(body: GreenhousePrepareSubmitRequest, user: User = Depends(get_current_user)):
    job = await db.jobs.find_one({"job_id": body.job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.get("ats_provider") != "greenhouse":
        raise HTTPException(status_code=400, detail="Job is not a Greenhouse job")

    profile = await db.profiles.find_one({"user_id": user.user_id}, {"_id": 0})
    if not profile or not profile.get("cv_text"):
        raise HTTPException(status_code=400, detail="Upload CV first")

    app_doc = await db.applications.find_one(
        {"user_id": user.user_id, "job_id": body.job_id},
        {"_id": 0},
        sort=[("created_at", -1)],
    )
    package_missing = (
        not app_doc
        or not app_doc.get("tailored_resume_structured")
        or not app_doc.get("tailored_cover_letter")
        or not app_doc.get("application_answers")
        or not app_doc.get("tailored_cv_file_b64")
    )
    if package_missing:
        try:
            generated_doc = await _generate_application_doc(user, profile, job)
        except LLMProviderNotConfigured as e:
            raise HTTPException(status_code=502, detail=str(e))
        if app_doc:
            generated_doc["application_id"] = app_doc.get("application_id") or generated_doc["application_id"]
            generated_doc["created_at"] = app_doc.get("created_at") or generated_doc["created_at"]
            await db.applications.update_one(
                {"user_id": user.user_id, "job_id": body.job_id, "application_id": generated_doc["application_id"]},
                {"$set": generated_doc},
                upsert=True,
            )
        else:
            await db.applications.insert_one(generated_doc)
        app_doc = generated_doc

    board_token = job.get("board_token")
    greenhouse_job_id = job.get("provider_job_id")
    external_id = job.get("external_id") or ""
    if (not board_token or not greenhouse_job_id) and ":" in external_id:
        board_token, greenhouse_job_id = external_id.split(":", 1)
    if not board_token or not greenhouse_job_id:
        raise HTTPException(status_code=400, detail="Greenhouse board token or job id is missing")

    provider = get_board_provider("greenhouse")
    try:
        preview = await provider.inspect_application_form(board_token, greenhouse_job_id)
    except httpx.HTTPStatusError as exc:
        raise HTTPException(status_code=502, detail="Greenhouse application form is unavailable") from exc
    except LLMProviderNotConfigured:
        raise
    except Exception as exc:
        logger.warning("Greenhouse prepare-submit form preview failed: job_id=%s error=%s", body.job_id, exc)
        raise HTTPException(status_code=502, detail="Greenhouse application form preview failed") from exc

    fields = preview["fields"]

    try:
        generated = await _generate_greenhouse_answers(profile, job, app_doc, fields)
    except LLMProviderNotConfigured as e:
        raise HTTPException(status_code=502, detail=str(e))
    except Exception as exc:
        logger.exception("Greenhouse answer generation failed")
        raise HTTPException(status_code=502, detail="AI application answer generation failed") from exc

    generated_answers = generated.get("answers") or []
    answers_by_field = {answer.get("field_name"): answer for answer in generated_answers if answer.get("field_name")}

    contact = profile.get("contact") or {}
    names = _split_name(contact.get("name") or user.name)
    payload_fields: Dict[str, Any] = {
        "first_name": names["first_name"],
        "last_name": names["last_name"],
        "email": user.email,
        "phone": contact.get("phone") or "",
    }
    question_payload = []
    auto_filled_fields = []
    for field in fields:
        category = field.get("field_category")
        name = field.get("name")
        if category in ("candidate", "document"):
            continue
        answer = answers_by_field.get(name, {})
        saved_value = _profile_saved_answer(profile, field)
        value = saved_value if not _is_empty_answer(saved_value) else (answer.get("answer") or "")
        if _is_empty_answer(value):
            default_value = _greenhouse_safe_default_answer(field)
            if default_value is not None:
                value = default_value
                auto_filled_fields.append(_public_missing_field_name(field))
        payload_fields[name] = value
        question_payload.append({
            "name": name,
            "label": field.get("label"),
            "value": value,
            "required": field.get("required"),
            "confidence": answer.get("confidence", 0),
        })
    application_payload = {
        "method": "POST",
        "url": f"https://boards-api.greenhouse.io/v1/boards/{board_token}/jobs/{greenhouse_job_id}",
        "content_type": "multipart/form-data",
        "fields": payload_fields,
        "questions": question_payload,
        "files": {
            "resume": {
                "filename": app_doc.get("tailored_cv_filename") or "tailored_cv.docx",
                "mime": app_doc.get("tailored_cv_mime") or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "b64": app_doc.get("tailored_cv_file_b64"),
            },
            "cover_letter": {
                "filename": f"{body.job_id}_cover_letter.txt",
                "mime": "text/plain",
                "text": cover_letter_to_text(app_doc.get("tailored_cover_letter") or app_doc.get("cover_letter") or {}),
            },
        },
    }
    missing_information = _normalize_missing_information(
        _required_empty_payload_fields(application_payload),
        _all_payload_fields(application_payload),
    )
    required_fields_count = _required_fields_count(application_payload)
    current_empty_required_fields = [item["field_name"] for item in missing_information]
    empty_required_fields_count = len(current_empty_required_fields)
    ready_for_submission = not missing_information and not preview.get("blockers")
    submission_status = "ready" if ready_for_submission else "blocked"
    blockers = sorted(set(preview.get("blockers") or []))
    logger.info(
        "Greenhouse prepare-submit validation: job_id=%s auto_filled_fields=%s remaining_empty_required_fields=%s current_empty_required_fields=%s deduped_missing_information_count=%s final_submission_status=%s",
        body.job_id,
        auto_filled_fields,
        current_empty_required_fields,
        current_empty_required_fields,
        len(missing_information),
        submission_status,
    )
    await db.applications.update_one(
        {"application_id": app_doc["application_id"], "user_id": user.user_id},
        {"$set": {
            "package_status": "generated",
            "submission_status": submission_status,
            "submitted_at": None,
            "submission_provider": "greenhouse",
            "submission_response_id": None,
            "submission_error": None if ready_for_submission else _missing_information_summary([*missing_information, *blockers]),
            "prepared_application_payload": _payload_for_storage(application_payload),
            "prepared_generated_answers": generated_answers,
            "prepared_missing_information": missing_information,
            "prepared_blockers": blockers,
            "prepared_at": datetime.now(timezone.utc).isoformat(),
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }},
    )

    return {
        "job_id": job["job_id"],
        "company": job.get("company"),
        "title": job.get("title"),
        "ready_for_submission": ready_for_submission,
        "submission_status": submission_status,
        "missing_information": missing_information,
        "blockers": blockers,
        "debug_summary": {
            "required_fields_count": required_fields_count,
            "empty_required_fields_count": empty_required_fields_count,
            "auto_filled_fields": auto_filled_fields,
            "remaining_empty_required_fields": current_empty_required_fields,
        },
        "application_payload": application_payload,
        "generated_answers": [
            {
                "question": answer.get("question"),
                "answer": answer.get("answer") or "",
                "confidence": answer.get("confidence", 0),
            }
            for answer in generated_answers
        ],
    }


async def _load_greenhouse_prepared_application(job_id: str, user: User) -> tuple[Dict[str, Any], Dict[str, Any], Dict[str, Any]]:
    job = await db.jobs.find_one({"job_id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.get("ats_provider") != "greenhouse":
        raise HTTPException(status_code=400, detail="Job is not a Greenhouse job")

    app_doc = await db.applications.find_one(
        {"user_id": user.user_id, "job_id": job_id},
        {"_id": 0},
        sort=[("created_at", -1)],
    )
    if not app_doc:
        raise HTTPException(status_code=400, detail="Create and prepare an application package before submission")
    app_doc = _normalize_application_status_fields(app_doc)
    payload = app_doc.get("prepared_application_payload")
    if not payload:
        raise HTTPException(status_code=400, detail="Prepare submission before submitting")
    return job, app_doc, payload


@api_router.get("/applications/greenhouse/submission-preview")
async def greenhouse_submission_preview(job_id: str, user: User = Depends(get_current_user)):
    job, app_doc, payload = await _load_greenhouse_prepared_application(job_id, user)
    preview = _greenhouse_submission_preview(app_doc, payload, job)
    return {
        "job_id": job_id,
        "company": job.get("company"),
        "title": job.get("title"),
        "submission_status": app_doc.get("submission_status"),
        **preview,
    }


@api_router.post("/applications/greenhouse/validate-submit")
async def greenhouse_validate_submit(body: GreenhousePrepareSubmitRequest, user: User = Depends(get_current_user)):
    job, app_doc, payload = await _load_greenhouse_prepared_application(body.job_id, user)
    preview = _greenhouse_submission_preview(app_doc, payload, job)
    now = datetime.now(timezone.utc).isoformat()
    await db.applications.update_one(
        {"application_id": app_doc["application_id"], "user_id": user.user_id},
        {"$set": {
            "submission_validation_errors": preview["validation_errors"],
            "submission_option_mappings": preview["option_mappings"],
            "submission_preview_metadata": {
                "field_count": len(preview["fields"]),
                "file_parts": preview["files"],
                "submit_url": preview["submit_url"],
            },
            "updated_at": now,
        }},
    )
    return {
        "job_id": body.job_id,
        "company": job.get("company"),
        "title": job.get("title"),
        "submission_status": app_doc.get("submission_status"),
        "is_valid": preview["is_valid"],
        "validation_errors": preview["validation_errors"],
        "option_mappings": preview["option_mappings"],
        "files": preview["files"],
        "field_count": len(preview["fields"]),
    }


@api_router.post("/applications/greenhouse/submit")
async def greenhouse_submit(body: GreenhousePrepareSubmitRequest, user: User = Depends(get_current_user)):
    job, app_doc, payload = await _load_greenhouse_prepared_application(body.job_id, user)
    if user.require_review_before_send and app_doc.get("document_review_status") != "approved":
        raise HTTPException(
            status_code=409,
            detail="Please review and approve your CV and cover letter in the Review tab before this application can be submitted.",
        )
    missing_information = app_doc.get("prepared_missing_information") or []
    blockers = app_doc.get("prepared_blockers") or []
    if app_doc.get("package_status") != "generated":
        raise HTTPException(
            status_code=400,
            detail={
                "message": "Application package is not generated",
                "package_status": app_doc.get("package_status"),
                "submission_status": app_doc.get("submission_status"),
                "missing_information": missing_information,
                "blockers": blockers,
            },
        )
    if app_doc.get("submission_status") != "ready":
        raise HTTPException(
            status_code=400,
            detail={
                "message": "Application is not ready for submission",
                "package_status": app_doc.get("package_status"),
                "submission_status": app_doc.get("submission_status"),
                "missing_information": missing_information,
                "blockers": blockers,
                "submission_error": app_doc.get("submission_error"),
            },
        )

    required_empty = _required_empty_payload_fields(payload)
    if required_empty:
        missing_information = _dedupe_missing_information([*missing_information, *required_empty])
        error = _missing_information_summary(missing_information)
        await db.applications.update_one(
            {"application_id": app_doc["application_id"], "user_id": user.user_id},
            {"$set": {
                "submission_status": "blocked",
                "submission_error": error,
                "prepared_missing_information": missing_information,
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }},
        )
        raise HTTPException(
            status_code=400,
            detail={
                "message": "Application has required empty answers",
                "package_status": app_doc.get("package_status"),
                "submission_status": "blocked",
                "missing_information": missing_information,
                "blockers": blockers,
            },
        )

    submit_url = _greenhouse_submission_endpoint(payload, job)
    preview = _greenhouse_submission_preview(app_doc, payload, job)
    if not preview["is_valid"]:
        now = datetime.now(timezone.utc).isoformat()
        await db.applications.update_one(
            {"application_id": app_doc["application_id"], "user_id": user.user_id},
            {"$set": {
                "submission_status": "blocked",
                "submission_error": "Greenhouse submission payload is invalid",
                "submission_validation_errors": preview["validation_errors"],
                "submission_option_mappings": preview["option_mappings"],
                "submission_preview_metadata": {
                    "field_count": len(preview["fields"]),
                    "file_parts": preview["files"],
                    "submit_url": preview["submit_url"],
                },
                "updated_at": now,
            }},
        )
        raise HTTPException(status_code=400, detail={"message": "Greenhouse submission payload is invalid", "errors": preview["validation_errors"]})

    data, files, parts = _build_greenhouse_multipart(app_doc, payload)
    dry_run = _greenhouse_submit_dry_run_enabled()
    logger.info(
        "Greenhouse submit%s: user_id=%s job_id=%s company=%s multipart_fields=%s multipart_file_parts=%s questions=%s option_mappings=%s",
        " dry-run" if dry_run else "",
        user.user_id,
        body.job_id,
        job.get("company"),
        [item[0] for item in data],
        [{"field_name": name, "filename": file_tuple[0], "mime": file_tuple[2], "size_bytes": len(file_tuple[1])} for name, file_tuple in files.items()],
        len(payload.get("questions") or []),
        parts.get("option_mappings") or [],
    )

    if dry_run:
        return {
            "job_id": job["job_id"],
            "company": job.get("company"),
            "title": job.get("title"),
            "dry_run": True,
            "would_submit": True,
            "submit_url": submit_url,
            "submission_preview": preview,
            "submission_status": app_doc.get("submission_status"),
        }

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(submit_url, data=data, files=files)
    except httpx.RequestError as exc:
        error = f"Greenhouse submission request failed: {exc.__class__.__name__}"
        logger.warning("Greenhouse submit request failed: job_id=%s error=%s", body.job_id, exc)
        await db.applications.update_one(
            {"application_id": app_doc["application_id"], "user_id": user.user_id},
            {"$set": {
                "submission_status": "failed",
                "submission_provider": "greenhouse",
                "submission_error": error,
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }},
        )
        raise HTTPException(status_code=502, detail=error) from exc

    metadata = _greenhouse_response_metadata(response)
    if 200 <= response.status_code < 300:
        submitted_at = datetime.now(timezone.utc).isoformat()
        submission_response_id = _greenhouse_response_submission_id(response)
        await db.applications.update_one(
            {"application_id": app_doc["application_id"], "user_id": user.user_id},
            {"$set": {
                "submission_status": "submitted",
                "submitted_at": submitted_at,
                "submission_provider": "greenhouse",
                "submission_response_id": submission_response_id,
                "submission_error": None,
                "submission_response_metadata": metadata,
                "submission_validation_errors": [],
                "submission_option_mappings": preview["option_mappings"],
                "submission_preview_metadata": {
                    "field_count": len(preview["fields"]),
                    "file_parts": preview["files"],
                    "submit_url": preview["submit_url"],
                },
                "updated_at": submitted_at,
            }},
        )
        logger.info(
            "Greenhouse submit succeeded: job_id=%s status_code=%s response_id=%s",
            body.job_id,
            response.status_code,
            submission_response_id,
        )
        return {
            "job_id": job["job_id"],
            "company": job.get("company"),
            "title": job.get("title"),
            "dry_run": False,
            "submission_status": "submitted",
            "submitted_at": submitted_at,
            "submission_provider": "greenhouse",
            "submission_response_id": submission_response_id,
            "submission_response_metadata": metadata,
        }

    error = f"Greenhouse submission failed with HTTP {response.status_code}"
    logger.warning(
        "Greenhouse submit failed: job_id=%s status_code=%s body_snippet=%s",
        body.job_id,
        response.status_code,
        response.text[:300],
    )
    await db.applications.update_one(
        {"application_id": app_doc["application_id"], "user_id": user.user_id},
        {"$set": {
            "submission_status": "failed",
            "submission_provider": "greenhouse",
            "submission_error": error,
            "submission_response_metadata": metadata,
            "submission_validation_errors": metadata.get("body", {}).get("validation_errors") or metadata.get("body", {}).get("errors") or [],
            "submission_option_mappings": preview["option_mappings"],
            "submission_preview_metadata": {
                "field_count": len(preview["fields"]),
                "file_parts": preview["files"],
                "submit_url": preview["submit_url"],
            },
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }},
    )
    raise HTTPException(status_code=502, detail={"message": error, "response": metadata})


async def _resolve_agent_missing_info(
    application_id: str,
    app_doc: Dict[str, Any],
    body: "ResolveMissingInfoRequest",
    user: User,
) -> Dict[str, Any]:
    """Agent-based applications have no `prepared_application_payload` (that
    structure belongs to the older Greenhouse-specific submit flow) -- the
    user's answers are stored as label/value pairs instead, in the exact
    shape `apply_agent.agent.build_candidate_context` already reads from
    `prepared_application_payload.questions`, so the next prepare/submit
    run for this job picks them up automatically without any new plumbing
    on the agent side.
    """
    answers = body.answers or {}
    if not answers:
        raise HTTPException(status_code=400, detail="No answers provided")

    pending = app_doc.get("prepared_missing_information") or []
    pending_by_field = {item.get("field_name"): item for item in pending if isinstance(item, dict)}
    payload = app_doc.get("prepared_application_payload") or {}
    questions = list(payload.get("questions") or [])
    answered_field_names = set()
    for field_name, value in answers.items():
        if _is_empty_answer(value):
            continue
        label = (pending_by_field.get(field_name) or {}).get("label") or field_name
        questions.append({"label": label, "value": value})
        answered_field_names.add(field_name)
    payload["questions"] = questions

    remaining_missing = [item for item in pending if item.get("field_name") not in answered_field_names]
    is_auto_apply = bool(app_doc.get("auto_apply_provider") or app_doc.get("auto_apply_queue_status"))
    submission_status = "action_required" if remaining_missing else ("not_submitted" if is_auto_apply else "prepared")
    now = datetime.now(timezone.utc).isoformat()
    update_doc = {
        "prepared_application_payload": payload,
        "prepared_missing_information": remaining_missing,
        "submission_status": submission_status,
        "updated_at": now,
    }
    if is_auto_apply and not remaining_missing:
        update_doc.update({
            "auto_apply_queue_status": "queued",
            "auto_apply_queue_reason": "answers_provided",
            "auto_apply_finished_at": None,
            "auto_apply_started_at": None,
            "manual_status": None,
            "submission_error": None,
        })
    await db.applications.update_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"$set": update_doc},
    )

    # Always persist reusable answers for future similar questions when any
    # answers were provided (auto-apply path); otherwise honor the flag.
    should_save_profile = body.save_to_profile or is_auto_apply
    if should_save_profile:
        profile_updates = {
            f"application_answers_profile.{_canonical_field_name(label)}": value
            for label, value in ((q.get("label"), q.get("value")) for q in questions)
            if label and not _is_empty_answer(value)
        }
        # Also key by field_name for resolver matching.
        for field_name, value in answers.items():
            if _is_empty_answer(value):
                continue
            profile_updates[f"application_answers_profile.{_canonical_field_name(field_name)}"] = value
        if profile_updates:
            profile_updates["updated_at"] = now
            await db.profiles.update_one(
                {"user_id": user.user_id},
                {"$set": profile_updates},
                upsert=True,
            )

    updated = await db.applications.find_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"_id": 0},
    )
    updated = _normalize_application_status_fields(updated or {})
    job = await db.jobs.find_one({"job_id": updated.get("job_id")}, {"_id": 0})
    if is_auto_apply and not remaining_missing:
        try:
            await auto_apply_queue.enqueue_application(db, updated, job, force=True)
            updated = await db.applications.find_one(
                {"application_id": application_id, "user_id": user.user_id},
                {"_id": 0},
            ) or updated
        except Exception as exc:
            logger.warning(
                "auto_apply_requeue_after_answers_failed application_id=%s error=%s",
                application_id,
                str(exc)[:200],
            )
    return {
        "application_id": application_id,
        "submission_status": submission_status,
        "ready_for_submission": submission_status in {"prepared", "not_submitted", "ready"} and not remaining_missing,
        "missing_information": remaining_missing,
        "blockers": [],
        "resolved_count": len(answered_field_names),
        "unresolved_fields": remaining_missing,
        "application": {**updated, "job": job},
    }


@api_router.post("/applications/{application_id}/resolve-missing-info")
async def resolve_missing_info(
    application_id: str,
    body: ResolveMissingInfoRequest,
    user: User = Depends(get_current_user),
):
    app_doc = await db.applications.find_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"_id": 0},
    )
    if not app_doc:
        raise HTTPException(status_code=404, detail="Application not found")

    if app_doc.get("agent_apply_result") and not app_doc.get("prepared_application_payload"):
        return await _resolve_agent_missing_info(application_id, app_doc, body, user)

    # Auto-apply driver path: same answer storage as agent (no Greenhouse payload).
    if (
        (app_doc.get("auto_apply_provider") or app_doc.get("auto_apply_queue_status") or app_doc.get("prepared_missing_information"))
        and not app_doc.get("prepared_application_payload")
    ):
        return await _resolve_agent_missing_info(application_id, app_doc, body, user)

    payload = app_doc.get("prepared_application_payload")
    if not payload:
        raise HTTPException(status_code=400, detail="Prepare submission before resolving missing information")

    answers = body.answers or {}
    current_missing = _normalize_missing_information(
        app_doc.get("prepared_missing_information") or [],
        _all_payload_fields(payload),
    )
    missing_count_before = len(current_missing)
    answers_received_count = len([value for value in answers.values() if not _is_empty_answer(value)])
    answers_by_canonical = _canonical_value_map(answers)
    required_missing_names = [
        item.get("field_name")
        for item in current_missing
        if isinstance(item, dict) and item.get("field_name")
    ]
    invalid = [
        field_name
        for field_name in required_missing_names
        if _is_empty_answer(answers.get(field_name)) and _is_empty_answer(answers_by_canonical.get(_canonical_field_name(field_name)))
    ]
    if invalid:
        raise HTTPException(
            status_code=400,
            detail={
                "message": "Answers are required for blocked fields",
                "missing_field_names": invalid,
            },
        )

    fields = payload.setdefault("fields", {})
    updated_payload_keys = set()
    for field_name, value in answers.items():
        if _is_empty_answer(value):
            continue
        canonical = _canonical_field_name(field_name)
        fields[field_name] = value
        updated_payload_keys.add(field_name)
        for existing_key in list(fields.keys()):
            if _canonical_field_name(existing_key) == canonical:
                fields[existing_key] = value
                updated_payload_keys.add(existing_key)
        for question in payload.get("questions") or []:
            if not isinstance(question, dict):
                continue
            question_names = {
                _canonical_field_key(question),
                _canonical_field_name(question.get("name")),
                _canonical_field_name(question.get("field_name")),
                _canonical_field_name(question.get("field_id")),
            }
            if canonical in question_names:
                question["value"] = value
                if question.get("name"):
                    fields[question["name"]] = value
                    updated_payload_keys.add(question["name"])

    missing_information = _normalize_missing_information(
        _required_empty_payload_fields(payload),
        _all_payload_fields(payload),
    )
    blockers = app_doc.get("prepared_blockers") or []
    submission_status = "ready" if not missing_information and not blockers else "blocked"
    submission_error = None if submission_status == "ready" else _missing_information_summary([*missing_information, *blockers])
    now = datetime.now(timezone.utc).isoformat()
    logger.info(
        "Resolved missing application info: application_id=%s missing_count_before=%s answers_received_count=%s missing_count_after=%s submission_status_after=%s received_answer_keys=%s updated_payload_keys=%s remaining_empty_required_fields=%s",
        application_id,
        missing_count_before,
        answers_received_count,
        len(missing_information),
        submission_status,
        list(answers.keys()),
        sorted(updated_payload_keys),
        [item.get("field_name") for item in missing_information],
    )

    update_fields = {
        "prepared_application_payload": payload,
        "prepared_missing_information": missing_information,
        "submission_status": submission_status,
        "submission_error": submission_error,
        "updated_at": now,
    }
    await db.applications.update_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"$set": update_fields},
    )

    if body.save_to_profile:
        profile_updates = _profile_answer_updates_from_resolved_fields(payload, answers)
        if profile_updates:
            profile_updates["updated_at"] = now
            await db.profiles.update_one(
                {"user_id": user.user_id},
                {"$set": profile_updates},
                upsert=True,
            )

    updated = await db.applications.find_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"_id": 0},
    )
    updated = _normalize_application_status_fields(updated or {})
    job = await db.jobs.find_one({"job_id": updated.get("job_id")}, {"_id": 0})
    return {
        "application_id": application_id,
        "submission_status": submission_status,
        "ready_for_submission": submission_status == "ready",
        "missing_information": missing_information,
        "blockers": blockers,
        "resolved_count": len(updated_payload_keys),
        "unresolved_fields": missing_information,
        "application": {**updated, "job": job},
    }


@api_router.get("/applications/{application_id}")
async def get_application(application_id: str, user: User = Depends(get_current_user)):
    app_doc = await db.applications.find_one(
        {"application_id": application_id, "user_id": user.user_id}, {"_id": 0}
    )
    if not app_doc:
        raise HTTPException(status_code=404, detail="Not found")
    app_doc = _normalize_application_status_fields(app_doc)
    job = await db.jobs.find_one({"job_id": app_doc["job_id"]}, {"_id": 0})
    return _public_application_doc(app_doc, job)


@api_router.get("/applications/{application_id}/tailored-cv")
async def download_tailored_cv(application_id: str, user: User = Depends(get_current_user)):
    import base64
    from fastapi.responses import Response as FastAPIResponse

    app_doc = await db.applications.find_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"_id": 0, "tailored_cv_file_b64": 1, "tailored_cv_filename": 1, "tailored_cv_mime": 1},
    )
    if not app_doc or not app_doc.get("tailored_cv_file_b64"):
        raise HTTPException(status_code=404, detail="Tailored CV not found")

    content = base64.b64decode(app_doc["tailored_cv_file_b64"])
    filename = app_doc.get("tailored_cv_filename") or "tailored_cv.docx"
    return FastAPIResponse(
        content=content,
        media_type=app_doc.get("tailored_cv_mime") or "application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@api_router.get("/applications/{application_id}/cover-letter")
async def download_cover_letter(application_id: str, user: User = Depends(get_current_user)):
    from fastapi.responses import Response as FastAPIResponse

    app_doc = await db.applications.find_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"_id": 0, "tailored_cover_letter": 1, "cover_letter": 1},
    )
    if not app_doc:
        raise HTTPException(status_code=404, detail="Application not found")

    cover_letter = app_doc.get("tailored_cover_letter") or app_doc.get("cover_letter")
    if not cover_letter:
        raise HTTPException(status_code=404, detail="Cover letter not found")

    content = cover_letter_to_text(cover_letter)
    return FastAPIResponse(
        content=content,
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{application_id}_cover_letter.txt"'},
    )


@api_router.patch("/applications/{application_id}/status")
async def update_status(application_id: str, update: StatusUpdate, user: User = Depends(get_current_user)):
    res = await db.applications.update_one(
        {"application_id": application_id, "user_id": user.user_id},
        {"$set": {"status": update.status}},
    )
    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail="Not found")
    return {"ok": True}


# ===================== Application email inbox =====================

@api_router.get("/emails/status")
async def gmail_inbox_status(user: User = Depends(get_current_user)):
    try:
        connection = await db.gmail_connections.find_one({"user_id": user.user_id}, {"_id": 0})
    except Exception as exc:
        logger.warning("gmail_status_failed user_id=%s error=%s", user.user_id, str(exc)[:200])
        return {
            "gmail": {
                "connected": False,
                "email": user.email,
                "last_synced_at": None,
                "last_sync_error": "Gmail inbox storage is not initialized",
            },
            "required_scope": GMAIL_READONLY_SCOPE,
        }
    return {"gmail": gmail_connected_payload(connection), "required_scope": GMAIL_READONLY_SCOPE}


@api_router.post("/emails/sync")
async def sync_application_emails(user: User = Depends(get_current_user)):
    try:
        result = await sync_gmail_application_emails(db, user_id=user.user_id)
        return result
    except Exception as exc:
        logger.warning("gmail_sync_failed user_id=%s error=%s", user.user_id, str(exc)[:300])
        raise HTTPException(status_code=503, detail=str(exc)[:300])


@api_router.get("/emails")
async def list_application_emails(
    sync: bool = Query(default=True),
    limit: int = Query(default=50, ge=1, le=100),
    user: User = Depends(get_current_user),
):
    sync_result: Dict[str, Any] = {"ok": True, "skipped": not sync}
    if sync:
        try:
            sync_result = await sync_gmail_application_emails(db, user_id=user.user_id)
        except Exception as exc:
            logger.warning("gmail_list_sync_failed user_id=%s error=%s", user.user_id, str(exc)[:300])
            sync_result = {"ok": False, "error": str(exc)[:300]}
    try:
        connection = await db.gmail_connections.find_one({"user_id": user.user_id}, {"_id": 0})
        rows = await db.application_emails.find({"user_id": user.user_id}, {"_id": 0}).sort("received_at", -1).to_list(limit)
    except Exception as exc:
        logger.warning("gmail_list_failed user_id=%s error=%s", user.user_id, str(exc)[:300])
        return {
            "messages": [],
            "gmail": {
                "connected": False,
                "email": user.email,
                "last_synced_at": None,
                "last_sync_error": "Gmail inbox storage is not initialized",
            },
            "sync": {"ok": False, "error": str(exc)[:300]},
        }
    return {
        "messages": [public_email_message(row) for row in rows],
        "gmail": gmail_connected_payload(connection),
        "sync": sync_result,
    }


@api_router.get("/notifications")
async def get_notifications(user: User = Depends(get_current_user)):
    rows = await list_notifications(db, user_id=user.user_id)
    unread_count = sum(1 for row in rows if not row.get("read"))
    return {"notifications": rows, "unread_count": unread_count}


@api_router.post("/notifications/{notification_id}/read")
async def mark_notification_read_route(notification_id: str, user: User = Depends(get_current_user)):
    await mark_notification_read(db, user_id=user.user_id, notification_id=notification_id)
    return {"ok": True}


@api_router.post("/notifications/read-all")
async def mark_all_notifications_read_route(user: User = Depends(get_current_user)):
    count = await mark_all_notifications_read(db, user_id=user.user_id)
    return {"ok": True, "marked": count}


@api_router.post("/webhooks/resend-inbound")
async def resend_inbound_webhook(request: Request):
    """Employer replies to a Hirly-managed application address (see
    email_addresses.managed_reply_address) arrive here via Resend's inbound
    email feature, mirroring the /stripe/webhook shape: verify signature,
    parse, dedupe, process with per-item error isolation so one bad payload
    can never break the rest of Resend's delivery, ack.
    """
    webhook_secret = os.environ.get("RESEND_INBOUND_WEBHOOK_SECRET", "").strip()
    if not webhook_secret:
        raise HTTPException(status_code=503, detail="Resend inbound webhook is not configured")
    payload = await request.body()
    try:
        event = Webhook(webhook_secret).verify(payload, dict(request.headers))
    except WebhookVerificationError:
        raise HTTPException(status_code=400, detail="Invalid signature")

    if not isinstance(event, dict):
        raise HTTPException(status_code=400, detail="Invalid payload")

    if event.get("type") != "email.received":
        return {"received": True, "ignored": event.get("type")}

    data = event.get("data") or {}
    resend_email_id = data.get("email_id") or data.get("id")
    if not resend_email_id:
        return {"received": True, "ignored": "missing_email_id"}

    try:
        result = await process_inbound_resend_email(db, resend_email_id, data)
        return {"received": True, **result}
    except Exception as exc:
        # Per-item isolation, applying the same lesson learned hardening the
        # Stripe reconcile loop this session: one bad payload must never 500
        # the whole delivery or block visibility into what actually failed.
        logger.warning("resend_inbound_webhook_failed email_id=%s error=%s", resend_email_id, str(exc)[:300])
        return {"received": True, "ok": False, "error": str(exc)[:200]}


# ===================== Seed =====================

MOCK_JOBS = [
    {
        "title": "Senior Frontend Engineer",
        "company": "Linear",
        "location": "Remote (Global)",
        "remote": "remote",
        "salary_min": 140000, "salary_max": 200000,
        "description": "Build the world's fastest issue tracker. Work on perf, animations, complex UI state.",
        "requirements": ["5+ years frontend", "TypeScript expert", "React or similar", "Performance optimization"],
        "tech_stack": ["TypeScript", "React", "GraphQL", "Vite"],
        "seniority": "senior",
    },
    {
        "title": "Full Stack Engineer",
        "company": "Vercel",
        "location": "San Francisco, CA",
        "remote": "hybrid",
        "salary_min": 160000, "salary_max": 240000,
        "description": "Help build the platform powering the modern web. Ship Next.js, Edge runtime, and DX tools.",
        "requirements": ["TypeScript", "Node.js", "Next.js", "Distributed systems"],
        "tech_stack": ["TypeScript", "Next.js", "Rust", "Go"],
        "seniority": "mid",
    },
    {
        "title": "AI Engineer",
        "company": "Anthropic",
        "location": "San Francisco, CA",
        "remote": "onsite",
        "salary_min": 220000, "salary_max": 380000,
        "description": "Work on Claude — train, evaluate, and ship safety-focused LLM products.",
        "requirements": ["ML/AI experience", "Python", "PyTorch", "Research engineering"],
        "tech_stack": ["Python", "PyTorch", "JAX", "CUDA"],
        "seniority": "senior",
    },
    {
        "title": "Product Designer",
        "company": "Raycast",
        "location": "Remote (EU/US)",
        "remote": "remote",
        "salary_min": 120000, "salary_max": 180000,
        "description": "Design extensions and core surfaces for the fastest launcher on Mac.",
        "requirements": ["5+ years product design", "Figma", "Strong portfolio", "Systems thinking"],
        "tech_stack": ["Figma", "Framer"],
        "seniority": "senior",
    },
    {
        "title": "Backend Engineer",
        "company": "Supabase",
        "location": "Remote (Global)",
        "remote": "remote",
        "salary_min": 130000, "salary_max": 190000,
        "description": "Build the open-source Firebase alternative. Postgres, realtime, auth, edge functions.",
        "requirements": ["Postgres internals", "TypeScript or Go", "Open-source experience"],
        "tech_stack": ["TypeScript", "Postgres", "Deno", "Go"],
        "seniority": "mid",
    },
    {
        "title": "iOS Engineer",
        "company": "Notion",
        "location": "New York, NY",
        "remote": "hybrid",
        "salary_min": 150000, "salary_max": 230000,
        "description": "Build the Notion mobile experience used by millions.",
        "requirements": ["Swift/SwiftUI", "5+ years iOS", "Performance tuning"],
        "tech_stack": ["Swift", "SwiftUI", "Combine"],
        "seniority": "senior",
    },
    {
        "title": "Growth Engineer",
        "company": "Cal.com",
        "location": "Remote",
        "remote": "remote",
        "salary_min": 110000, "salary_max": 170000,
        "description": "Run experiments across landing, signup, and activation flows. Move metrics, ship fast.",
        "requirements": ["A/B testing", "Next.js", "Analytics tools", "SQL"],
        "tech_stack": ["Next.js", "TypeScript", "PostHog", "Postgres"],
        "seniority": "mid",
    },
    {
        "title": "Staff ML Engineer",
        "company": "Hugging Face",
        "location": "Remote (EU)",
        "remote": "remote",
        "salary_min": 200000, "salary_max": 320000,
        "description": "Open-source ML at scale. Lead transformers, datasets, or inference infrastructure.",
        "requirements": ["Senior ML experience", "Python", "Open-source leadership"],
        "tech_stack": ["Python", "PyTorch", "Transformers"],
        "seniority": "lead",
    },
    {
        "title": "DevRel Engineer",
        "company": "Stripe",
        "location": "Remote (US)",
        "remote": "remote",
        "salary_min": 140000, "salary_max": 210000,
        "description": "Build demos, content, and tooling that helps developers integrate Stripe in minutes.",
        "requirements": ["Strong writing", "Full-stack coding", "Speaking experience"],
        "tech_stack": ["TypeScript", "Node.js", "React"],
        "seniority": "senior",
    },
    {
        "title": "Platform Engineer",
        "company": "Fly.io",
        "location": "Remote (Global)",
        "remote": "remote",
        "salary_min": 150000, "salary_max": 220000,
        "description": "Run the global app platform. Linux, networking, distributed systems.",
        "requirements": ["Linux internals", "Rust or Go", "Networking"],
        "tech_stack": ["Rust", "Go", "Linux"],
        "seniority": "senior",
    },
    {
        "title": "Junior Frontend Developer",
        "company": "Framer",
        "location": "Amsterdam, NL",
        "remote": "hybrid",
        "salary_min": 60000, "salary_max": 85000,
        "description": "Join the team building the no-code site builder loved by designers.",
        "requirements": ["1-2 years React", "CSS skills", "Eye for detail"],
        "tech_stack": ["React", "TypeScript", "CSS"],
        "seniority": "junior",
    },
    {
        "title": "Data Engineer",
        "company": "Posthog",
        "location": "Remote (Global)",
        "remote": "remote",
        "salary_min": 130000, "salary_max": 200000,
        "description": "Scale ClickHouse, build pipelines, ship product analytics at billions of events.",
        "requirements": ["ClickHouse or BigQuery", "Python", "Streaming systems"],
        "tech_stack": ["Python", "ClickHouse", "Kafka"],
        "seniority": "senior",
    },
    {
        "title": "Founding Engineer",
        "company": "Stealth AI Startup",
        "location": "San Francisco, CA",
        "remote": "onsite",
        "salary_min": 180000, "salary_max": 260000,
        "description": "Build the core product from day 1. Equity-heavy. Wear every hat.",
        "requirements": ["Full-stack expert", "Shipped 0-1 products", "AI experience"],
        "tech_stack": ["TypeScript", "Python", "Next.js", "LLMs"],
        "seniority": "senior",
    },
    {
        "title": "Mobile Engineer",
        "company": "Cash App",
        "location": "Remote (US)",
        "remote": "remote",
        "salary_min": 150000, "salary_max": 230000,
        "description": "Build the simplest way to send money. iOS and Android.",
        "requirements": ["Swift or Kotlin", "Fintech experience", "Strong testing"],
        "tech_stack": ["Swift", "Kotlin"],
        "seniority": "senior",
    },
    {
        "title": "Marketing Engineer",
        "company": "Resend",
        "location": "Remote (Global)",
        "remote": "remote",
        "salary_min": 120000, "salary_max": 170000,
        "description": "Build delightful marketing sites and docs for the email API for developers.",
        "requirements": ["Next.js", "Design sense", "Animations"],
        "tech_stack": ["Next.js", "Framer Motion", "TypeScript"],
        "seniority": "mid",
    },
]


@api_router.post("/seed")
async def seed_jobs():
    """Idempotently seed mock job data for development fallback."""
    fallback_mock = os.environ.get("JOB_PROVIDER_FALLBACK_MOCK", "false").lower() in ("1", "true", "yes", "on")
    if not fallback_mock:
        raise HTTPException(status_code=403, detail="Mock seed is disabled")
    count = await db.jobs.count_documents({})
    if count >= len(MOCK_JOBS):
        return {"ok": True, "skipped": True, "count": count}

    now = datetime.now(timezone.utc)
    docs = []
    for j in MOCK_JOBS:
        docs.append({
            "job_id": f"job_{uuid.uuid4().hex[:10]}",
            "currency": "USD",
            "posted_at": now.isoformat(),
            **j,
        })
    await db.jobs.delete_many({})
    await db.jobs.insert_many(docs)
    return {"ok": True, "count": len(docs)}


@api_router.get("/health")
async def health():
    stripe_secret = bool(os.environ.get("STRIPE_SECRET_KEY", "").strip())
    stripe_webhook = bool(os.environ.get("STRIPE_WEBHOOK_SECRET", "").strip())
    return {
        "status": "ok",
        "stripe": {
            "secret_key_configured": stripe_secret,
            "webhook_secret_configured": stripe_webhook,
            "webhook_url": "/api/stripe/webhook",
            "required_events": [
                "checkout.session.completed",
                "checkout.session.async_payment_succeeded",
                "customer.subscription.created",
                "customer.subscription.updated",
                "customer.subscription.deleted",
                "invoice.payment_succeeded",
                "invoice.payment_failed",
                "refund.created",
                "refund.updated",
                "refund.failed",
            ],
        },
    }


@api_router.get("/version")
async def version():
    return {
        # Railway's own auto-populated var reflects the real current deploy;
        # checked first so a stale manually-set APP_GIT_SHA (from an earlier
        # debugging session, never updated since) can't mask it.
        "git_sha": os.environ.get("RAILWAY_GIT_COMMIT_SHA") or os.environ.get("VERCEL_GIT_COMMIT_SHA") or os.environ.get("APP_GIT_SHA"),
        "app_version": os.environ.get("APP_VERSION"),
        "deployed_at": os.environ.get("DEPLOYED_AT") or os.environ.get("RAILWAY_DEPLOYMENT_CREATED_AT"),
        "flags": {
            "JOBS_DB_FIRST_ENABLED": _env_bool("JOBS_DB_FIRST_ENABLED", True),
            "JOBS_FEED_LEGACY_JSEARCH_ONLY": _env_bool("JOBS_FEED_LEGACY_JSEARCH_ONLY", False),
            "JOBS_FEED_SYNC_REFRESH_ENABLED": _env_bool("JOBS_FEED_SYNC_REFRESH_ENABLED", True),
            "JOBS_FEED_SYNC_REFRESH_COOLDOWN_SECONDS": _env_int("JOBS_FEED_SYNC_REFRESH_COOLDOWN_SECONDS", 300),
            "JOBS_FEED_FALLBACK_COOLDOWN_ENABLED": _env_bool("JOBS_FEED_FALLBACK_COOLDOWN_ENABLED", True),
            "JOBS_FEED_DEBUG_DIAGNOSTICS": _env_bool("JOBS_FEED_DEBUG_DIAGNOSTICS", False),
        },
    }


@api_router.get("/")
async def root():
    return {"message": "Tinder for Jobs API", "ok": True}


@api_router.get("/dev/jsearch-test")
async def dev_jsearch_test(
    q: str = "software engineer",
    location: str = "New York",
    limit: int = 5,
    country: Optional[str] = None,
    language: Optional[str] = None,
    max_pages: Optional[int] = None,
    page_size: Optional[int] = None,
):
    dev_enabled = (
        os.environ.get("ENVIRONMENT", "").lower() == "development"
        or os.environ.get("DEV_TOOLS_ENABLED", "false").lower() in ("1", "true", "yes", "on")
    )
    if not dev_enabled:
        raise HTTPException(status_code=404, detail="Not found")

    if not is_job_provider_configured(primary_job_provider_name()):
        raise HTTPException(status_code=500, detail="Primary job provider credentials are not configured")

    provider = get_configured_job_provider()
    # Was hardcoded to JSEARCH_COUNTRY/JSEARCH_LANGUAGE regardless of the
    # `location` argument -- e.g. passing location="London, United Kingdom"
    # while JSEARCH_COUNTRY=fr silently sent country=fr, contradicting the
    # location text. Now resolves from the location text the same way the
    # real feed path does, with explicit overrides available for isolating
    # variables during diagnosis.
    resolved_country, resolved_location = jobs_service_module._country_and_location(location)
    effective_country = country or resolved_country
    query = JobSearchQuery(
        role=q,
        location=resolved_location,
        remote_preference="any",
        country=effective_country,
        language=language or country_to_jsearch_language(effective_country),
        limit=max(1, min(limit, 20)),
        max_pages=max_pages,
        page_size=page_size,
    )
    started = time.perf_counter()
    try:
        result = await provider.search(query)
    except Exception as exc:
        elapsed_ms = int((time.perf_counter() - started) * 1000)
        logger.warning("Dev job provider test failed: %s", exc)
        raise HTTPException(
            status_code=502,
            detail=f"{exc.__class__.__name__}: {str(exc)[:300]} (elapsed_ms={elapsed_ms}, country={effective_country}, location={resolved_location})",
        ) from exc
    return {
        "provider": provider.name,
        "query": {"role": q, "location": resolved_location, "country": effective_country, "language": query.language},
        "elapsed_ms": int((time.perf_counter() - started) * 1000),
        "jobs": result.jobs,
        "count": len(result.jobs),
    }


@api_router.get("/dev/france-travail-test")
async def dev_france_travail_test(
    q: str = "développeur",
    location: str = "Lyon, France",
    limit: int = 5,
    radius_km: int = 50,
    contract_hint: Optional[str] = None,
):
    dev_enabled = (
        os.environ.get("ENVIRONMENT", "").lower() == "development"
        or os.environ.get("DEV_TOOLS_ENABLED", "false").lower() in ("1", "true", "yes", "on")
    )
    if not dev_enabled:
        raise HTTPException(status_code=404, detail="Not found")

    if not is_job_provider_configured("france_travail"):
        raise HTTPException(status_code=500, detail="France Travail credentials are not configured")

    provider = get_job_provider("france_travail", "")
    query = JobSearchQuery(
        role=q,
        location=location,
        remote_preference="any",
        country="fr",
        language="fr",
        limit=max(1, min(limit, 20)),
        max_pages=1,
        page_size=max(5, min(limit * 3, 20)),
        radius_km=max(0, min(int(radius_km), 200)),
        contract_hint=contract_hint,
    )
    try:
        result = await provider.search(query)
    except ValueError as exc:
        logger.warning("Dev France Travail response parse failed: %s", exc)
        raise HTTPException(status_code=502, detail=str(exc)) from exc
    except Exception as exc:
        logger.warning("Dev France Travail test failed: %s", exc)
        raise HTTPException(status_code=502, detail="France Travail provider request failed") from exc
    return {"provider": provider.name, "jobs": result.jobs, "count": len(result.jobs), "raw": result.raw_response}


@api_router.get("/dev/greenhouse-import-test")
async def dev_greenhouse_import_test():
    dev_enabled = (
        os.environ.get("ENVIRONMENT", "").lower() == "development"
        or os.environ.get("DEV_TOOLS_ENABLED", "false").lower() in ("1", "true", "yes", "on")
    )
    if not dev_enabled:
        raise HTTPException(status_code=404, detail="Not found")
    result = await refresh_greenhouse_boards(db, limit_boards=10, force=True)
    sample_jobs = result.get("sample_jobs", [])
    return {
        "boards_checked": result.get("boards_checked", 0),
        "boards_successful": result.get("boards_successful", 0),
        "jobs_imported": result.get("jobs_imported", 0),
        "sample_jobs": sample_jobs[:5],
    }


@api_router.get("/dev/greenhouse-board-test")
async def dev_greenhouse_board_test(board_token: str = "stripe"):
    dev_enabled = (
        os.environ.get("ENVIRONMENT", "").lower() == "development"
        or os.environ.get("DEV_TOOLS_ENABLED", "false").lower() in ("1", "true", "yes", "on")
    )
    if not dev_enabled:
        raise HTTPException(status_code=404, detail="Not found")

    token = board_token.strip().lower()
    if not token:
        raise HTTPException(status_code=400, detail="board_token is required")

    provider = get_board_provider("greenhouse")
    try:
        inspection = await provider.inspect_board(token)
    except Exception as exc:
        logger.warning("Greenhouse board test failed: board_token=%s error=%s", token, exc)
        raise HTTPException(status_code=502, detail="Greenhouse board test failed") from exc

    return {
        "board_token": inspection["board_token"],
        "status_code": inspection["status_code"],
        "jobs_count": inspection["jobs_count"],
        "first_job_title": inspection["first_job_title"],
        "error_snippet": inspection["error_snippet"],
    }


@api_router.get("/dev/lever-import-test")
async def dev_lever_import_test():
    dev_enabled = (
        os.environ.get("ENVIRONMENT", "").lower() == "development"
        or os.environ.get("DEV_TOOLS_ENABLED", "false").lower() in ("1", "true", "yes", "on")
    )
    if not dev_enabled:
        raise HTTPException(status_code=404, detail="Not found")
    result = await refresh_lever_boards(db, limit_boards=10, force=True)
    sample_jobs = result.get("sample_jobs", [])
    return {
        "boards_checked": result.get("boards_checked", 0),
        "boards_successful": result.get("boards_successful", 0),
        "jobs_imported": result.get("jobs_imported", 0),
        "sample_jobs": sample_jobs[:5],
    }


@api_router.get("/dev/lever-board-test")
async def dev_lever_board_test(site: str = "postman"):
    dev_enabled = (
        os.environ.get("ENVIRONMENT", "").lower() == "development"
        or os.environ.get("DEV_TOOLS_ENABLED", "false").lower() in ("1", "true", "yes", "on")
    )
    if not dev_enabled:
        raise HTTPException(status_code=404, detail="Not found")

    board_site = site.strip().lower()
    if not board_site:
        raise HTTPException(status_code=400, detail="site is required")

    provider = get_board_provider("lever")
    try:
        inspection = await provider.inspect_board(board_site)
    except Exception as exc:
        logger.warning("Lever board test failed: site=%s error=%s", board_site, exc)
        raise HTTPException(status_code=502, detail="Lever board test failed") from exc

    return {
        "site": inspection["site"],
        "primary_url": inspection.get("primary_url"),
        "primary_status": inspection.get("primary_status"),
        "primary_error_snippet": inspection.get("primary_error_snippet"),
        "eu_url": inspection.get("eu_url"),
        "eu_status": inspection.get("eu_status"),
        "eu_error_snippet": inspection.get("eu_error_snippet"),
        "jobs_count": inspection["jobs_count"],
        "first_job_title": inspection["first_job_title"],
    }


@api_router.get("/dev/provider-write-test")
async def dev_provider_write_test(
    provider: str = "lever",
    max_boards: int = 1,
    max_jobs_per_board: int = 25,
    timeout_seconds: int = 120,
):
    dev_enabled = (
        os.environ.get("ENVIRONMENT", "").lower() == "development"
        or os.environ.get("DEV_TOOLS_ENABLED", "false").lower() in ("1", "true", "yes", "on")
    )
    if not dev_enabled:
        raise HTTPException(status_code=404, detail="Not found")

    provider_name = (provider or "").strip().lower()
    if provider_name not in ("greenhouse", "lever"):
        raise HTTPException(status_code=400, detail="provider must be 'greenhouse' or 'lever'")

    max_boards = max(1, min(int(max_boards or 1), 5))
    max_jobs_per_board = max(1, min(int(max_jobs_per_board or 25), 100))
    timeout_seconds = max(5, min(int(timeout_seconds or 120), 300))
    started_at = time.perf_counter()
    progress: Dict[str, Any] = {
        "boards_checked": 0,
        "jobs_upserted_so_far": 0,
        "last_board": None,
        "sample_jobs": [],
    }

    async def _run_bounded_write_test() -> Dict[str, Any]:
        if provider_name == "greenhouse":
            result = await refresh_greenhouse_boards(
                db,
                limit_boards=max_boards,
                force=True,
                job_limit=max_jobs_per_board,
                progress=progress,
            )
        else:
            result = await refresh_lever_boards(
                db,
                limit_boards=max_boards,
                force=True,
                job_limit=max_jobs_per_board,
                progress=progress,
            )
        return result

    try:
        result = await asyncio.wait_for(_run_bounded_write_test(), timeout=timeout_seconds)
    except asyncio.TimeoutError as exc:
        elapsed_ms = int((time.perf_counter() - started_at) * 1000)
        logger.warning("provider_write_test timeout provider=%s elapsed_ms=%s", provider_name, elapsed_ms)
        return {
            "active_provider": DATABASE_PROVIDER,
            "jobs_collection_provider": _runtime_collection_provider("jobs"),
            "company_boards_collection_provider": _runtime_collection_provider("company_boards"),
            "provider": provider_name,
            "timed_out": True,
            "phase": "provider_write_test",
            "message": f"Provider write test timed out after {timeout_seconds} seconds.",
            "boards_checked": progress.get("boards_checked", 0),
            "jobs_upserted_so_far": progress.get("jobs_upserted_so_far", 0),
            "last_board": progress.get("last_board"),
            "elapsed_ms": elapsed_ms,
            "timeout_seconds": timeout_seconds,
            "max_boards": max_boards,
            "max_jobs_per_board": max_jobs_per_board,
            "sample_job_id": (progress.get("sample_jobs") or [{}])[0].get("job_id") if progress.get("sample_jobs") else None,
        }

    sample_jobs = result.get("sample_jobs") or []
    sample_job_id = sample_jobs[0].get("job_id") if sample_jobs else None
    sample_db_job = None
    if sample_job_id:
        sample_db_job = await db.jobs.find_one(
            {"job_id": sample_job_id},
            {"_id": 0, "job_id": 1, "title": 1, "company": 1, "provider": 1, "ats_provider": 1},
        )
    return {
        "active_provider": DATABASE_PROVIDER,
        "jobs_collection_provider": _runtime_collection_provider("jobs"),
        "company_boards_collection_provider": _runtime_collection_provider("company_boards"),
        "provider": provider_name,
        "timed_out": False,
        "boards_checked": result.get("boards_checked", 0),
        "boards_successful": result.get("boards_successful", 0),
        "inserted_or_updated_count": result.get("jobs_imported", 0),
        "jobs_upserted_so_far": result.get("jobs_imported", 0),
        "auto_apply_supported_imported": result.get("auto_apply_supported_imported", 0),
        "elapsed_ms": result.get("elapsed_ms"),
        "timeout_seconds": timeout_seconds,
        "max_boards": max_boards,
        "max_jobs_per_board": max_jobs_per_board,
        "last_board": progress.get("last_board"),
        "sample_job_id": sample_job_id,
        "sample_job_persisted": sample_db_job is not None,
        "sample_job": sample_db_job,
    }


@api_router.get("/dev/playwright-launch-test")
async def dev_playwright_launch_test():
    dev_enabled = (
        os.environ.get("ENVIRONMENT", "").lower() == "development"
        or os.environ.get("DEV_TOOLS_ENABLED", "false").lower() in ("1", "true", "yes", "on")
    )
    if not dev_enabled:
        raise HTTPException(status_code=404, detail="Not found")

    try:
        from playwright.async_api import async_playwright
    except ImportError as exc:
        return {
            "ok": False,
            "phase": "import_playwright",
            "exception_class": exc.__class__.__name__,
            "message": str(exc).strip() or "Playwright is not installed.",
        }

    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            try:
                page = await browser.new_page()
                await page.goto("about:blank")
                return {"ok": True}
            finally:
                await browser.close()
    except Exception as exc:
        if _dev_tools_enabled():
            logger.exception("Dev Playwright launch test failed")
        return {
            "ok": False,
            "phase": "open_browser",
            "exception_class": exc.__class__.__name__,
            "message": str(exc),
        }


@api_router.get("/dev/asyncio-loop-debug")
async def dev_asyncio_loop_debug():
    loop = asyncio.get_running_loop()
    subprocess_supported = False
    subprocess_error = None
    try:
        process = await asyncio.create_subprocess_exec(
            sys.executable,
            "--version",
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        await process.communicate()
        subprocess_supported = process.returncode == 0
        if process.returncode != 0:
            subprocess_error = f"python --version exited with {process.returncode}"
    except Exception as exc:
        subprocess_error = f"{exc.__class__.__name__}: {str(exc)[:500]}"

    return {
        "platform": sys.platform,
        "event_loop_policy": asyncio.get_event_loop_policy().__class__.__name__,
        "running_loop_class": loop.__class__.__name__,
        "subprocess_supported_test": subprocess_supported,
        "subprocess_error": subprocess_error,
    }


@api_router.get("/dev/db-health")
async def dev_db_health():
    dev_tools_enabled_raw = os.environ.get("DEV_TOOLS_ENABLED")
    environment_raw = os.environ.get("ENVIRONMENT")
    env_file_path = ROOT_DIR / ".env"

    supabase_url = _normalize_supabase_url(os.environ.get("SUPABASE_URL"))
    supabase_secret = os.environ.get("SUPABASE_SECRET_KEY", "")
    supabase_config_present = bool(supabase_url and supabase_secret)
    supabase_connection_ok = False
    supabase_error = None

    if supabase_config_present:
        result = await test_supabase_connection(supabase_url, supabase_secret)
        supabase_connection_ok = bool(result.get("ok"))
        supabase_error = result.get("error")
    else:
        supabase_error = "SUPABASE_URL or SUPABASE_SECRET_KEY is missing."

    return {
        "active_provider": DATABASE_PROVIDER,
        "dev_tools_enabled_raw": dev_tools_enabled_raw,
        "environment_raw": environment_raw,
        "env_file_loaded": env_file_path.exists(),
        "current_working_directory": os.getcwd(),
        "server_file": __file__,
        "supabase_config_present": supabase_config_present,
        "supabase_connection_ok": supabase_connection_ok,
        "supabase_error": supabase_error,
    }


@api_router.get("/dev/db-counts")
async def dev_db_counts():
    tables = (
        "users",
        "user_sessions",
        "jobs",
        "company_boards",
        "profiles",
        "swipes",
        "applications",
        "browser_submission_runs",
    )
    counts: Dict[str, Any] = {}
    for table in tables:
        count = None
        error = None
        try:
            count = await getattr(db, table).count_documents({})
        except Exception as exc:
            error = f"{exc.__class__.__name__}: {str(exc)[:300]}"
        counts[table] = {
            "supabase_count": count,
            "supabase_error": error,
        }

    return {
        "active_provider": DATABASE_PROVIDER,
        "supabase_config_present": bool(os.environ.get("SUPABASE_URL") and os.environ.get("SUPABASE_SECRET_KEY")),
        "tables": counts,
    }


@api_router.get("/dev/applications-read-write-test")
async def dev_applications_read_write_test(user: User = Depends(get_current_user)):
    if not _dev_tools_enabled():
        raise HTTPException(status_code=404, detail="Not found")
    now = datetime.now(timezone.utc).isoformat()
    test_doc = {
        "application_id": f"app_rw_test_{uuid.uuid4().hex[:12]}",
        "user_id": user.user_id,
        "job_id": f"job_rw_test_{uuid.uuid4().hex[:12]}",
        "status": "applied",
        "package_status": "pending_generation",
        "submission_status": "not_submitted",
        "created_at": now,
        "updated_at": now,
        "dev_test": True,
    }
    await db.applications.update_one(
        {"application_id": test_doc["application_id"]},
        {"$set": test_doc},
        upsert=True,
    )
    by_application_id = await db.applications.find_one(
        {"application_id": test_doc["application_id"]},
        {"_id": 0},
    )
    by_user_job = await db.applications.find_one(
        {"user_id": test_doc["user_id"], "job_id": test_doc["job_id"]},
        {"_id": 0},
    )

    backfill_limit = 500
    rows = await db.applications.find({}, {"_id": 0}).limit(backfill_limit).to_list(backfill_limit)
    backfilled = 0
    backfill_errors = []
    for row in rows:
        if not row.get("application_id") or not row.get("user_id") or not row.get("job_id"):
            continue
        try:
            await db.applications.update_one(
                {"application_id": row["application_id"]},
                {"$set": {
                    "application_id": row["application_id"],
                    "user_id": row["user_id"],
                    "job_id": row["job_id"],
                    "status": row.get("status"),
                    "package_status": row.get("package_status"),
                    "submission_status": row.get("submission_status"),
                    "updated_at": row.get("updated_at") or now,
                }},
                upsert=True,
            )
            backfilled += 1
        except Exception as exc:
            if len(backfill_errors) < 5:
                backfill_errors.append({
                    "application_id": row.get("application_id"),
                    "exception_class": exc.__class__.__name__,
                    "message": str(exc)[:300],
                })

    return {
        "active_provider": DATABASE_PROVIDER,
        "created_application_id": test_doc["application_id"],
        "created_job_id": test_doc["job_id"],
        "read_by_application_id_found": by_application_id is not None,
        "read_by_user_id_job_id_found": by_user_job is not None,
        "read_by_application_id": {
            "application_id": (by_application_id or {}).get("application_id"),
            "user_id": (by_application_id or {}).get("user_id"),
            "job_id": (by_application_id or {}).get("job_id"),
        },
        "read_by_user_id_job_id": {
            "application_id": (by_user_job or {}).get("application_id"),
            "user_id": (by_user_job or {}).get("user_id"),
            "job_id": (by_user_job or {}).get("job_id"),
        },
        "backfill_scanned": len(rows),
        "backfilled_top_level_user_job": backfilled,
        "backfill_errors": backfill_errors,
    }


@api_router.get("/dev/docx-fallback-test")
async def dev_docx_fallback_test(user: User = Depends(get_current_user)):
    if not _dev_tools_enabled():
        raise HTTPException(status_code=404, detail="Not found")

    def failing_package_builder(profile: Dict[str, Any], generated: Dict[str, Any]) -> Dict[str, Any]:
        raise ValueError("All strings must be XML compatible: Unicode or ASCII, no NULL bytes or control characters")

    profile = {
        "user_id": user.user_id,
        "contact": {
            "name": f"{user.name}\x00",
            "email": user.email,
        },
        "cv_text": "Profile text\x00 with bad chars\x07",
        "cv_filename": "cv.docx",
        "cv_mime": DOCX_MIME,
    }
    job = {
        "job_id": f"job_docx_fallback_test_{uuid.uuid4().hex[:8]}",
        "title": "DOCX Fallback Test",
        "company": "Swiipr",
    }
    generated = {
        "tailored_resume": {
            "summary": "Summary with null\x00 and bell\x07 chars",
            "skills": ["Python\x01", "FastAPI"],
            "experience": [
                {
                    "role": "Engineer\x02",
                    "company": "Example",
                    "duration": "2020\x03",
                    "highlights": ["Built systems\x0B"],
                }
            ],
        },
        "tailored_cover_letter": {
            "greeting": "Hello\x00",
            "paragraphs": ["Cover letter paragraph\x07"],
            "sign_off": "Regards",
        },
        "application_answers": [],
        "match_score": 80,
        "match_reasons": ["Test"],
        "interview_prep": [],
    }
    try:
        doc = _build_generated_application_doc(user, profile, job, generated, package_builder=failing_package_builder)
    except Exception as exc:
        logger.exception("DOCX fallback test unexpectedly raised")
        return {
            "ok": False,
            "raised": True,
            "exception_class": exc.__class__.__name__,
            "message": str(exc)[:500],
        }
    return {
        "ok": True,
        "raised": False,
        "application_id": doc.get("application_id"),
        "package_status": doc.get("package_status"),
        "generation_status": doc.get("generation_status"),
        "generation_error": doc.get("generation_error"),
        "submission_status": doc.get("submission_status"),
        "tailored_cv_file_b64_is_none": doc.get("tailored_cv_file_b64") is None,
        "has_tailored_resume_text": bool(doc.get("tailored_resume") or doc.get("tailored_resume_structured")),
        "has_cover_letter_text": bool(doc.get("cover_letter") or doc.get("tailored_cover_letter")),
    }


@api_router.get("/dev/db-provider-test")
async def dev_db_provider_test():
    jobs_count = await db.jobs.count_documents({})
    company_boards_count = await db.company_boards.count_documents({})
    sample_jobs = await db.jobs.find({}, {"_id": 0}).limit(1).to_list(1)
    return {
        "active_provider": DATABASE_PROVIDER,
        "jobs_count": jobs_count,
        "company_boards_count": company_boards_count,
        "sample_job": sample_jobs[0] if sample_jobs else None,
    }


@api_router.get("/dev/jobs-query-debug")
async def dev_jobs_query_debug(
    user: User = Depends(get_current_user),
    limit: int = 5,
    search_radius: str = "50km",
    locations_json: Optional[str] = None,
    location: Optional[List[str]] = Query(None),
    location_label: Optional[str] = None,
    country: Optional[str] = None,
    country_code: Optional[str] = None,
    lat: Optional[float] = None,
    lng: Optional[float] = None,
    only_my_country: bool = False,
):
    profile = await db.profiles.find_one({"user_id": user.user_id}, {"_id": 0}) or {}
    profile_location_data = profile.get("target_location_data") or {}
    profile_country_code = (profile_location_data.get("country_code") or "").strip().lower()
    profile_country = profile_location_data.get("country")

    selected_locations: List[Dict[str, Any]] = []
    if locations_json:
        try:
            parsed_locations = json.loads(locations_json)
            if isinstance(parsed_locations, list):
                selected_locations = [loc for loc in parsed_locations if isinstance(loc, dict) and loc.get("location_label")]
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="locations_json must be a JSON array")
    if location_label:
        selected_locations = [{
            "location_label": location_label,
            "country": country,
            "country_code": country_code,
            "lat": lat,
            "lng": lng,
        }]
    if only_my_country:
        if profile_country_code:
            selected_locations = [
                {**loc, "country_code": profile_country_code, "country": profile_country or loc.get("country")}
                for loc in selected_locations
            ]
            if not selected_locations and profile_location_data:
                selected_locations = [profile_location_data]
        elif not selected_locations and profile_location_data:
            selected_locations = [profile_location_data]
    if not selected_locations and profile_location_data:
        selected_locations = [profile_location_data]

    selected_location_labels = [
        loc.get("location_label")
        for loc in selected_locations
        if isinstance(loc, dict) and loc.get("location_label")
    ]
    radius_scope = (search_radius or "50km").lower().strip()
    filter_locations = None
    if radius_scope not in ("worldwide", "remote", "remote/worldwide"):
        filter_locations = selected_location_labels or location
        if not filter_locations and only_my_country and profile_country:
            filter_locations = [profile_country]

    location_filter_clause = None
    if filter_locations:
        expanded_filter_locations = set()
        for loc in filter_locations:
            expanded_filter_locations.add(loc)
            for part in re.split(r"[,/|-]", loc):
                part = part.strip()
                if len(part) >= 3:
                    expanded_filter_locations.add(part)
        location_filter_clause = {
            "$or": [
                {"location": {"$regex": re.escape(loc), "$options": "i"}}
                for loc in expanded_filter_locations
            ]
        }

    swiped = await db.swipes.find({"user_id": user.user_id}, {"_id": 0, "job_id": 1}).to_list(2000)
    swiped_ids = {item.get("job_id") for item in swiped if item.get("job_id")}

    target_role = (
        profile.get("target_role")
        or ((profile.get("target_roles") or [None])[0])
        or ""
    ).strip()

    def _tokens(value: str) -> List[str]:
        stop = {"and", "or", "the", "a", "an", "of", "for", "to", "in", "with", "remote", "jobs", "job"} | ACADEMIC_LEVEL_STOPWORDS
        return [token for token in re.findall(r"[a-z0-9]+", (value or "").lower()) if len(token) > 2 and token not in stop]

    _role_match_override = resolve_role_match_tokens(target_role)
    role_tokens = _role_match_override if _role_match_override is not None else _tokens(target_role)
    broader_role_tokens = role_tokens[-1:] if role_tokens else []
    if "analyst" in role_tokens:
        broader_role_tokens = list(dict.fromkeys([*role_tokens, "analytics", "analysis", "insights", "scientist"]))
    if "software" in role_tokens and "engineer" in role_tokens:
        broader_role_tokens = ["software", "engineer"]

    def _role_score(job: Dict[str, Any], tokens: List[str]) -> int:
        if not tokens:
            return 0
        title = (job.get("title") or "").lower()
        body = " ".join([
            job.get("description") or "",
            job.get("clean_description") or "",
            " ".join(job.get("requirements") or []),
        ]).lower()
        title_hits = sum(1 for token in tokens if token in title)
        body_hits = sum(1 for token in tokens if token in body)
        exact_bonus = 25 if not _role_match_override and target_role and target_role.lower() in title else 0
        return exact_bonus + title_hits * 20 + min(body_hits, len(tokens)) * 5

    def _role_matches(job: Dict[str, Any]) -> bool:
        if not role_tokens:
            return True
        strict_score = _role_score(job, role_tokens)
        broad_score = _role_score(job, broader_role_tokens)
        return strict_score >= max(20, len(role_tokens) * 12) or broad_score >= 20

    def _location_matches(job: Dict[str, Any]) -> bool:
        if not location_filter_clause:
            return True
        job_location = (job.get("location") or "").lower()
        for clause in location_filter_clause.get("$or", []):
            regex = (clause.get("location") or {}).get("$regex")
            if regex and re.search(regex, job_location, re.IGNORECASE):
                return True
        return False

    def _sample_jobs(jobs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        return [
            {
                "job_id": job.get("job_id"),
                "title": job.get("title"),
                "company": job.get("company"),
                "provider": job.get("provider"),
                "ats_provider": job.get("ats_provider"),
                "auto_apply_supported": job.get("auto_apply_supported"),
                "location": job.get("location"),
            }
            for job in jobs[:5]
        ]

    async def _compare_source(source_name: str, jobs_collection: Any) -> Dict[str, Any]:
        auto_query = {"auto_apply_supported": True}
        ats_query = {**auto_query, "ats_provider": {"$in": ["greenhouse", "lever", "ashby"]}}
        swiped_query = dict(ats_query)
        if swiped_ids:
            swiped_query["job_id"] = {"$nin": list(swiped_ids)}
        location_query = dict(swiped_query)
        if location_filter_clause:
            location_query.setdefault("$and", []).append(location_filter_clause)

        total_jobs = await jobs_collection.count_documents({})
        auto_apply_supported_count = await jobs_collection.count_documents(auto_query)
        ats_provider_supported_count = await jobs_collection.count_documents(ats_query)
        after_swipe_exclusion_count = await jobs_collection.count_documents(swiped_query)

        stage_candidates = await jobs_collection.find(swiped_query, {"_id": 0}).limit(5000).to_list(5000)
        after_role = [job for job in stage_candidates if _role_matches(job)]
        after_location = [job for job in after_role if _location_matches(job)]
        after_location.sort(
            key=lambda job: (
                _role_score(job, role_tokens),
                str(job.get("imported_at") or job.get("posted_at") or ""),
            ),
            reverse=True,
        )
        final_jobs = after_location[:limit]

        adapter_location_count = await jobs_collection.count_documents(location_query)
        return {
            "source": source_name,
            "counts": {
                "total_jobs": total_jobs,
                "auto_apply_supported_count": auto_apply_supported_count,
                "ats_provider_supported_count": ats_provider_supported_count,
                "after_swipe_exclusion_count": after_swipe_exclusion_count,
                "after_role_filter_count": len(after_role),
                "after_location_filter_count": len(after_location),
                "final_count": len(final_jobs),
            },
            "adapter_location_filter_count": adapter_location_count,
            "sample_jobs": _sample_jobs(final_jobs),
        }

    active_result = await _compare_source(_runtime_collection_provider("jobs"), db.jobs)

    return {
        "active_provider": DATABASE_PROVIDER,
        "user_id": user.user_id,
        "target_role": target_role or None,
        "role_tokens": role_tokens,
        "broader_role_tokens": broader_role_tokens,
        "swiped_job_ids_count": len(swiped_ids),
        "filters": {
            "auto_apply_supported": True,
            "ats_provider": ["greenhouse", "lever", "ashby"],
            "selected_locations": selected_location_labels,
            "location_query": location,
            "search_radius": search_radius,
            "only_my_country": only_my_country,
            "country_code": country_code or profile_country_code or None,
            "location_filter_clause": location_filter_clause,
        },
        "supabase": active_result,
        "supabase_sample_jobs": active_result.get("sample_jobs", []) if active_result.get("source") == "supabase" else [],
    }


def _runtime_collection_provider(collection_name: str) -> str:
    collection = getattr(db, collection_name)
    module_name = collection.__class__.__module__
    if "supabase_adapter" in module_name:
        return "supabase"
    return "unknown"


@api_router.get("/dev/database-usage-audit")
async def dev_database_usage_audit():
    audited_collections = [
        "users",
        "user_sessions",
        "profiles",
        "jobs",
        "applications",
        "gmail_connections",
        "application_emails",
        "swipes",
        "company_boards",
        "browser_submission_runs",
    ]
    collection_providers = {
        collection_name: _runtime_collection_provider(collection_name)
        for collection_name in audited_collections
    }

    route_collections = {
        "login:/api/auth/supabase-session": ["users", "user_sessions", "profiles", "gmail_connections"],
        "auth_me:/api/auth/me": ["user_sessions", "users", "profiles"],
        "profile:/api/profile*": ["user_sessions", "users", "profiles", "swipes", "applications"],
        "swipes:/api/swipe*": ["user_sessions", "users", "swipes", "jobs", "profiles", "applications"],
        "applications:/api/applications*": ["user_sessions", "users", "applications", "jobs", "profiles", "browser_submission_runs"],
        "emails:/api/emails*": ["user_sessions", "users", "gmail_connections", "application_emails", "applications", "jobs"],
        "jobs_feed:/api/jobs/feed": ["user_sessions", "users", "profiles", "swipes", "jobs"],
        "company_boards:startup_and_import_dev_routes": ["company_boards", "jobs"],
    }

    route_details: Dict[str, Any] = {}
    routes_using_supabase: List[str] = []
    for route_name, collections in route_collections.items():
        providers = sorted({collection_providers[collection] for collection in collections})
        route_details[route_name] = {
            "collections": collections,
            "providers": providers,
            "uses_supabase": "supabase" in providers,
        }
        if "supabase" in providers:
            routes_using_supabase.append(route_name)

    return {
        "provider": DATABASE_PROVIDER,
        "collections_using_supabase": [
            name for name, provider in collection_providers.items() if provider == "supabase"
        ],
        "collections_using_mongo": [],
        "routes_using_supabase": routes_using_supabase,
        "routes_using_mongo": [],
        "collection_providers": collection_providers,
        "route_details": route_details,
        "notes": [
            "This endpoint inspects the live db object currently bound in server.py.",
            "All runtime collections use Supabase.",
            "Provider refresh/import writes use the active database adapter for jobs and company_boards.",
        ],
    }


@api_router.get("/dev/job-debug/{job_id}")
async def dev_job_debug(job_id: str):
    dev_enabled = (
        os.environ.get("ENVIRONMENT", "").lower() == "development"
        or os.environ.get("DEV_TOOLS_ENABLED", "false").lower() in ("1", "true", "yes", "on")
    )
    if not dev_enabled:
        raise HTTPException(status_code=404, detail="Not found")

    job = await db.jobs.find_one(
        {"job_id": job_id},
        {"_id": 0, "title": 1, "description": 1, "clean_description": 1, "job_description_sections": 1, "requirements": 1},
    )
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return job


@api_router.get("/dev/greenhouse-jobs-sample")
async def dev_greenhouse_jobs_sample():
    dev_enabled = (
        os.environ.get("ENVIRONMENT", "").lower() == "development"
        or os.environ.get("DEV_TOOLS_ENABLED", "false").lower() in ("1", "true", "yes", "on")
    )
    if not dev_enabled:
        raise HTTPException(status_code=404, detail="Not found")

    jobs = await db.jobs.find(
        {"ats_provider": "greenhouse", "auto_apply_supported": True},
        {
            "_id": 0,
            "job_id": 1,
            "external_id": 1,
            "title": 1,
            "company": 1,
            "external_url": 1,
            "provider": 1,
            "ats_provider": 1,
        },
    ).limit(10).to_list(10)
    return {"jobs": jobs, "count": len(jobs)}


@api_router.get("/dev/auto-apply-jobs-by-provider")
async def dev_auto_apply_jobs_by_provider():
    dev_enabled = (
        os.environ.get("ENVIRONMENT", "").lower() == "development"
        or os.environ.get("DEV_TOOLS_ENABLED", "false").lower() in ("1", "true", "yes", "on")
    )
    if not dev_enabled:
        raise HTTPException(status_code=404, detail="Not found")

    providers = ["greenhouse", "lever", "ashby"]
    response: Dict[str, Any] = {}
    total_jobs = 0
    total_auto_apply_jobs = 0
    for provider in providers:
        query = {"ats_provider": provider, "auto_apply_supported": True}
        count = await db.jobs.count_documents(query)
        total_auto_apply_jobs += count
        candidates = await db.jobs.find(
            query,
            {
                "_id": 0,
                "job_id": 1,
                "title": 1,
                "company": 1,
                "external_url": 1,
                "ats_provider": 1,
                "provider": 1,
                "imported_at": 1,
                "last_seen_at": 1,
            },
        ).sort([("company", 1), ("imported_at", -1)]).limit(500).to_list(500)

        sample_jobs = []
        overflow_jobs = []
        seen_companies = set()
        for job in candidates:
            company_key = (job.get("company") or "").strip().lower()
            sample = {
                "job_id": job.get("job_id"),
                "title": job.get("title"),
                "company": job.get("company"),
                "external_url": job.get("external_url"),
                "ats_provider": job.get("ats_provider"),
                "provider": job.get("provider"),
            }
            if company_key and company_key not in seen_companies:
                sample_jobs.append(sample)
                seen_companies.add(company_key)
            else:
                overflow_jobs.append(sample)
            if len(sample_jobs) >= 20:
                break
        if len(sample_jobs) < 20:
            sample_jobs.extend(overflow_jobs[: 20 - len(sample_jobs)])

        response[provider] = {
            "count": count,
            "sample_jobs": sample_jobs[:20],
        }

    total_jobs = await db.jobs.count_documents({})
    return {
        "total_jobs": total_jobs,
        "total_auto_apply_jobs": total_auto_apply_jobs,
        "providers": response,
    }


@api_router.get("/dev/clean-job-descriptions")
async def dev_clean_job_descriptions():
    dev_enabled = (
        os.environ.get("ENVIRONMENT", "").lower() == "development"
        or os.environ.get("DEV_TOOLS_ENABLED", "false").lower() in ("1", "true", "yes", "on")
    )
    if not dev_enabled:
        raise HTTPException(status_code=404, detail="Not found")

    provider = get_board_provider("greenhouse")

    def has_html(value: Any) -> bool:
        return bool(re.search(r"</?[a-z][\s\S]*?>", str(value or ""), flags=re.IGNORECASE))

    def clean_sections(sections: Any) -> List[Dict[str, Any]]:
        cleaned = []
        if not isinstance(sections, list):
            return cleaned
        for section in sections:
            if not isinstance(section, dict):
                continue
            title = provider.sanitize_text(section.get("title"))
            bullets = [
                provider.sanitize_text(bullet)
                for bullet in (section.get("bullets") or [])
                if provider.sanitize_text(bullet)
            ]
            if title and bullets:
                cleaned.append({"title": title, "bullets": bullets})
        return cleaned

    scanned = 0
    updated = 0
    sample_remaining_ids: List[str] = []
    cursor = db.jobs.find(
        {},
        {"_id": 0, "job_id": 1, "description": 1, "clean_description": 1, "job_description_sections": 1, "requirements": 1},
    )
    async for job in cursor:
        scanned += 1
        description = provider.sanitize_text(job.get("description"))
        clean_description = provider.sanitize_text(job.get("clean_description") or description)
        sections = clean_sections(job.get("job_description_sections"))
        requirements = [
            provider.sanitize_text(item)
            for item in (job.get("requirements") or [])
            if provider.sanitize_text(item)
        ]
        update = {
            "description": description,
            "clean_description": clean_description,
            "job_description_sections": sections,
            "requirements": requirements,
        }
        if (
            update["description"] != job.get("description")
            or update["clean_description"] != job.get("clean_description")
            or update["job_description_sections"] != (job.get("job_description_sections") or [])
            or update["requirements"] != (job.get("requirements") or [])
        ):
            await db.jobs.update_one({"job_id": job["job_id"]}, {"$set": update})
            updated += 1

    remaining_html_count = 0
    cursor = db.jobs.find(
        {},
        {"_id": 0, "job_id": 1, "description": 1, "clean_description": 1, "job_description_sections": 1, "requirements": 1},
    )
    async for job in cursor:
        contains_html = has_html(job.get("description")) or has_html(job.get("clean_description"))
        for section in job.get("job_description_sections") or []:
            contains_html = contains_html or has_html(section.get("title"))
            contains_html = contains_html or any(has_html(bullet) for bullet in section.get("bullets") or [])
        contains_html = contains_html or any(has_html(item) for item in job.get("requirements") or [])
        if contains_html:
            remaining_html_count += 1
            if len(sample_remaining_ids) < 10:
                sample_remaining_ids.append(job.get("job_id"))

    return {
        "scanned": scanned,
        "updated": updated,
        "remaining_html_count": remaining_html_count,
        "sample_remaining_ids": sample_remaining_ids,
    }


# ===================== Wire up =====================

register_training_routes(api_router, get_current_user, db, _require_training_user, _get_training_access_payload)
register_training_admin_routes(api_router, require_admin_user, db, _enrich_invite_rows_for_admin)
register_record_tools_routes(api_router, get_current_user, db, _require_record_tools_user)
register_feedback_routes(api_router, get_current_user, require_admin_user, db)
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=_cors_origins(),
    allow_methods=["*"],
    allow_headers=["*"],
)


async def _startup_seed_impl():
    """Seed boards and training content without blocking HTTP readiness."""
    try:
        await migrate_file_invites_to_db(db)
        await migrate_file_feedback_to_db(db)
        await backfill_feedback_from_resend(db)
        await bootstrap_invites_from_influencers(db)
        await backfill_invites_from_users(db)
        await ensure_dev_test_invites(db)
        await seed_greenhouse_company_boards(db)
        await seed_lever_company_boards(db)
        await seed_training_content(db)
        await sync_training_locale_content(db)
        await ensure_training_enrollments_for_access_users(db)

        fallback_mock = os.environ.get("JOB_PROVIDER_FALLBACK_MOCK", "false").lower() in ("1", "true", "yes", "on")
        if not fallback_mock:
            return

        count = await db.jobs.count_documents({})
        if count == 0:
            now = datetime.now(timezone.utc).isoformat()
            docs = [{
                "job_id": f"job_{uuid.uuid4().hex[:10]}",
                "currency": "USD",
                "posted_at": now,
                **j,
            } for j in MOCK_JOBS]
            await db.jobs.insert_many(docs)
            logger.info(f"Seeded {len(docs)} jobs")
    except Exception as e:
        logger.warning(f"Seed failed: {e}")


_STARTUP_BACKGROUND_TASKS: set[asyncio.Task] = set()


def _python_ingestion_schedule_states(crons_paused: bool):
    return (
        ("python-france-travail-harvest", "france_travail", max(300, _env_int("FT_HARVEST_INTERVAL_MINUTES", 5) * 60), not crons_paused and ft_harvest_enabled()),
        ("python-jsearch-harvest", "jsearch", max(300, _env_int("JSEARCH_HARVEST_INTERVAL_MINUTES", 15) * 60), not crons_paused and jsearch_harvest_enabled()),
        ("python-ats-direct-maintenance", "direct_ats", max(300, _env_int("ATS_DIRECT_MAINTENANCE_INTERVAL_MINUTES", 5) * 60), not crons_paused and ats_direct_maintenance_loop_enabled()),
        ("python-company-discovery", "company_discovery", max(300, _env_int("COMPANY_DISCOVERY_LOOP_INTERVAL_MINUTES", 10) * 60), not crons_paused and company_discovery_loop_enabled()),
    )


def _spawn_observed_startup_task(coro, *, name: str) -> asyncio.Task:
    """Retain startup tasks and make terminal failures visible."""
    task = asyncio.create_task(coro, name=name)
    _STARTUP_BACKGROUND_TASKS.add(task)

    def _completed(completed: asyncio.Task) -> None:
        _STARTUP_BACKGROUND_TASKS.discard(completed)
        if completed.cancelled():
            logger.warning("startup_background_task_cancelled task=%s", name)
            return
        error = completed.exception()
        if error is not None:
            logger.error(
                "startup_background_task_failed task=%s error=%s",
                name,
                error,
                exc_info=(type(error), error, error.__traceback__),
            )

    task.add_done_callback(_completed)
    return task


@app.on_event("startup")
async def startup_seed():
    """Register routes immediately; run seeding in the background."""
    global _posthog_client
    _ph_api_key = os.environ.get("POSTHOG_SERVER_API_KEY", "").strip()
    _ph_host = os.environ.get("POSTHOG_HOST", "").strip()
    if _ph_api_key and _ph_host:
        _posthog_client = Posthog(
            api_key=_ph_api_key,
            host=_ph_host,
            enable_exception_autocapture=True,
        )
        atexit.register(_posthog_client.shutdown)
        logger.info("posthog_client_initialized host=%s", _ph_host)
        if _OTEL_AVAILABLE:
            _otel_provider = _OtelTracerProvider(
                resource=_OtelResource(attributes={_OTEL_SERVICE_NAME: "hirly-backend"})
            )
            _otel_provider.add_span_processor(
                _PostHogSpanProcessor(api_key=_ph_api_key, host=_ph_host)
            )
            _otel_trace.set_tracer_provider(_otel_provider)
            _OpenAIInstrumentor().instrument()
            logger.info("posthog_otel_ai_observability_initialized")
    # Register concrete ApplyDrivers once at startup. The executor never imports
    # concrete drivers itself -- this import is the single registration point.
    import auto_apply.drivers  # noqa: F401
    logger.info("DB health route registered at /api/dev/db-health")
    logger.info("DB counts route registered at /api/dev/db-counts")
    logger.info(
        "Startup env debug: DEV_TOOLS_ENABLED=%r ENVIRONMENT=%r PORT=%r STRIPE_SECRET=%s STRIPE_WEBHOOK=%s",
        os.environ.get("DEV_TOOLS_ENABLED"),
        os.environ.get("ENVIRONMENT"),
        os.environ.get("PORT"),
        "set" if os.environ.get("STRIPE_SECRET_KEY", "").strip() else "MISSING",
        "set" if os.environ.get("STRIPE_WEBHOOK_SECRET", "").strip() else "MISSING",
    )
    if not os.environ.get("STRIPE_WEBHOOK_SECRET", "").strip():
        logger.warning(
            "STRIPE_WEBHOOK_SECRET is not configured — Stripe webhooks will fail with 503. "
            "Register https://<your-backend>/api/stripe/webhook in the Stripe Dashboard."
        )
    if _env_bool("STRIPE_RECONCILIATION_ENABLED", True) and _stripe_configured():
        _spawn_observed_startup_task(_run_stripe_subscription_reconcile_loop(), name="stripe-reconcile")
    _spawn_observed_startup_task(_startup_seed_impl(), name="startup-seed")
    _spawn_observed_startup_task(_resume_pending_application_generation(), name="application-generation-resume")
    _spawn_observed_startup_task(auto_apply_queue.startup(db), name="auto-apply-queue")
    # With JOBS_INVENTORY_BLITZ (default on), crons start unless explicitly paused.
    # Set PAUSE_JOB_MAINTENANCE_CRONS=true to stop background harvest under DB load.
    blitz = _env_bool("JOBS_INVENTORY_BLITZ", True)
    pause_default = not blitz
    crons_paused = _env_bool("PAUSE_JOB_MAINTENANCE_CRONS", pause_default)
    sync_schedule = getattr(db, "sync_python_ingestion_schedule", None)
    if callable(sync_schedule):
        schedule_states = (
            ("python-france-travail-harvest", "france_travail", max(300, _env_int("FT_HARVEST_INTERVAL_MINUTES", 5) * 60), not crons_paused and ft_harvest_enabled()),
            ("python-jsearch-harvest", "jsearch", max(300, _env_int("JSEARCH_HARVEST_INTERVAL_MINUTES", 15) * 60), not crons_paused and jsearch_harvest_enabled()),
            ("python-ats-direct-maintenance", "direct_ats", max(300, _env_int("ATS_DIRECT_MAINTENANCE_INTERVAL_MINUTES", 5) * 60), not crons_paused and ats_direct_maintenance_loop_enabled()),
            ("python-company-discovery", "company_discovery", max(300, _env_int("COMPANY_DISCOVERY_LOOP_INTERVAL_MINUTES", 10) * 60), not crons_paused and company_discovery_loop_enabled()),
        )
        for schedule_id, source, cadence_seconds, enabled in schedule_states:
            await sync_schedule(
                schedule_id=schedule_id,
                source=source,
                cadence_seconds=cadence_seconds,
                enabled=enabled,
            )
    if crons_paused:
        logger.warning(
            "job_maintenance_crons_paused blitz=%s hint=set_PAUSE_JOB_MAINTENANCE_CRONS_false_to_fill_inventory",
            blitz,
        )
    else:
        logger.info("job_maintenance_crons_starting blitz=%s target=500k_jobs_per_week", blitz)
        _spawn_observed_startup_task(run_france_travail_harvest_loop(db), name="france-travail-harvest")
        _spawn_observed_startup_task(run_jsearch_harvest_loop(db), name="jsearch-harvest")
        _spawn_observed_startup_task(run_ats_direct_maintenance_loop(db), name="ats-direct-maintenance")
        _spawn_observed_startup_task(run_company_discovery_loop(db), name="company-discovery")
    _spawn_observed_startup_task(run_creator_social_refresh_loop(), name="creator-social-refresh")


@app.on_event("shutdown")
async def shutdown_db_client():
    if _posthog_client is not None:
        _posthog_client.flush()
    await db.close()
