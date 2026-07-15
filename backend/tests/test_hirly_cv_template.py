import io

import docx as docx_lib
from PIL import Image

import application_documents as ad
import server


def _fake_image_bytes(width, height, color=(180, 60, 60), fmt="JPEG"):
    buffer = io.BytesIO()
    Image.new("RGB", (width, height), color).save(buffer, format=fmt)
    return buffer.getvalue()


# ---------------------------------------------------------------- template

def test_template_name_always_hirly_default_regardless_of_input():
    assert ad._template_name({}) == "hirly_default"
    assert ad._template_name({"template_recommendation": "modern_pro"}) == "hirly_default"
    assert ad._template_name({"template_recommendation": "does_not_exist"}) == "hirly_default"


def test_build_application_package_always_uses_hirly_builder(monkeypatch):
    calls = {"count": 0}

    def fake_builder(profile, tailored, job_title="", photo_bytes=None):
        calls["count"] += 1
        calls["job_title"] = job_title
        calls["photo_bytes"] = photo_bytes
        return b"FAKE_DOCX_BYTES"

    monkeypatch.setattr(ad, "_build_hirly_docx", fake_builder)

    for mime, filename in (
        (ad.DOCX_MIME, "cv.docx"),
        ("application/pdf", "cv.pdf"),
        ("image/png", "cv.png"),
    ):
        calls["count"] = 0
        profile = {"cv_mime": mime, "cv_filename": filename, "contact": {"name": "Jane"}}
        generated = {"tailored_resume_structured": {"experience": []}, "job_title": "Backend Engineer"}
        result = ad.build_application_package(profile, generated)
        assert calls["count"] == 1
        assert calls["job_title"] == "Backend Engineer"
        assert result["tailored_cv_file_b64"] is not None
        assert result["template_used"] == "hirly_default"


# ------------------------------------------------------------- docx builder

def test_build_hirly_docx_produces_valid_document_with_expected_content():
    photo_bytes = _fake_image_bytes(200, 200)
    profile = {"contact": {"name": "Jane Doe", "phone": "0600000000", "email": "jane@example.com", "location": "Paris"}}
    tailored = {
        "contact": {},
        "experience": [{
            "role": "Software Engineer",
            "company": "Acme",
            "location": "Paris",
            "duration": "01/2020 - Present",
            "highlights": ["Shipped a thing"],
        }],
        "education": [{"degree": "MSc Computer Science", "school": "Sorbonne", "year": "2019"}],
        "languages": ["English - Fluent", "French - Native", "Spanish - Basic"],
    }

    out = ad._build_hirly_docx(profile, tailored, job_title="Lead Backend Engineer", photo_bytes=photo_bytes)
    assert isinstance(out, bytes)
    assert len(out) > 0

    document = docx_lib.Document(io.BytesIO(out))
    all_text = "\n".join(p.text for p in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text

    assert "Jane Doe" in all_text
    assert "Lead Backend Engineer" in all_text
    assert "Software Engineer" in all_text
    assert "Acme" in all_text
    assert "MSc Computer Science" in all_text
    assert "English" in all_text
    assert any(document.inline_shapes), "expected the photo to be embedded"


def test_build_hirly_docx_without_photo_has_no_inline_shapes():
    profile = {"contact": {"name": "Jane Doe"}}
    tailored = {"contact": {}, "experience": [], "education": [], "languages": []}

    out = ad._build_hirly_docx(profile, tailored, job_title="", photo_bytes=None)
    document = docx_lib.Document(io.BytesIO(out))
    assert not any(document.inline_shapes)


# --------------------------------------------------------- photo extraction

def test_looks_like_headshot_accepts_reasonable_portrait():
    assert server._looks_like_headshot(_fake_image_bytes(300, 300)) is True
    assert server._looks_like_headshot(_fake_image_bytes(250, 320)) is True


def test_looks_like_headshot_rejects_small_or_wide_images():
    assert server._looks_like_headshot(_fake_image_bytes(60, 60)) is False
    assert server._looks_like_headshot(_fake_image_bytes(800, 100)) is False


def test_extract_cv_photo_pdf_picks_first_valid_candidate(monkeypatch):
    small_logo = _fake_image_bytes(40, 40)
    good_photo = _fake_image_bytes(300, 300)
    monkeypatch.setattr(server, "_extract_pdf_image_bytes", lambda content: [small_logo, good_photo])

    photo_bytes, mime = server._extract_cv_photo("pdf", b"irrelevant")

    assert photo_bytes is not None
    assert mime == "image/jpeg"


def test_extract_cv_photo_docx_uses_inline_images(monkeypatch):
    good_photo = _fake_image_bytes(300, 300)
    monkeypatch.setattr(server, "_extract_docx_inline_images", lambda content: [good_photo])

    photo_bytes, mime = server._extract_cv_photo("docx", b"irrelevant")

    assert photo_bytes is not None
    assert mime == "image/jpeg"


def test_extract_cv_photo_returns_none_for_unsupported_format():
    photo_bytes, mime = server._extract_cv_photo("txt", b"irrelevant")
    assert photo_bytes is None
    assert mime is None


def test_extract_cv_photo_returns_none_when_no_candidate_matches(monkeypatch):
    monkeypatch.setattr(server, "_extract_pdf_image_bytes", lambda content: [_fake_image_bytes(40, 40)])

    photo_bytes, mime = server._extract_cv_photo("pdf", b"irrelevant")

    assert photo_bytes is None
    assert mime is None
