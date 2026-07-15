import base64
import io

import docx

from application_documents import build_application_package
from application_documents import SUPPORTED_GENERATED_TEMPLATES
from cv_quality import clean_cv_text, normalize_application_generation, validate_resume_quality


def test_clean_cv_text_removes_ai_artifacts_without_breaking_urls():
    text = "Built AI // ML tooling — see https://example.com/profile **now**"

    cleaned = clean_cv_text(text)

    assert "//" not in cleaned.replace("https://", "")
    assert "https://example.com/profile" in cleaned
    assert "**" not in cleaned
    assert "—" not in cleaned


def test_normalize_application_generation_cleans_resume_and_reports_quality():
    generated = {
        "tailored_resume_structured": {
            "template_recommendation": "luxe_minimal",
            "headline": "Frontend Engineer for React Platforms",
            "contact": {"name": "Jane Doe", "email": "jane@example.com"},
            "summary": "Results-driven engineer with React // TypeScript experience.",
            "role_keywords": ["React", "React", "Design Systems", "Unsupported extra"],
            "skills": ["React", "React", "TypeScript", "GraphQL"],
            "languages": ["French - Native", "French - Native", "English - Full professional"],
            "experience": [{
                "role": "Frontend Engineer",
                "company": "Acme",
                "duration": "2021-2024",
                "location": "Paris",
                "highlights": ["Built design system // components **used by teams**."],
                "source_evidence": "Frontend Engineer at Acme",
            }],
            "education": [{"degree": "BS CS", "school": "MIT", "year": "2019"}],
            "evidence_notes": ["React -> Acme frontend role"],
            "unsupported_requirements": ["Kubernetes"],
        },
        "tailored_cover_letter": {
            "greeting": "Dear team,",
            "paragraphs": ["I am excited to apply."],
            "sign_off": "Best,",
        },
    }

    normalized = normalize_application_generation(generated)
    resume = normalized["tailored_resume_structured"]

    assert resume["skills"] == ["React", "TypeScript", "GraphQL"]
    assert resume["template_recommendation"] == "luxe_minimal"
    assert resume["headline"] == "Frontend Engineer for React Platforms"
    assert resume["role_keywords"] == ["React", "Design Systems", "Unsupported extra"]
    assert resume["languages"] == ["French - Native", "English - Full professional"]
    assert "//" not in resume["summary"]
    assert "**" not in resume["experience"][0]["highlights"][0]
    assert resume["experience"][0]["source_evidence"] == "Frontend Engineer at Acme"
    assert resume["evidence_notes"] == ["React -> Acme frontend role"]
    assert resume["unsupported_requirements"] == ["Kubernetes"]
    assert normalized["resume_quality_report"]["status"] in {"pass", "needs_review"}
    assert "ats_score" in normalized["resume_quality_report"]
    assert "recruiter_score" in normalized["resume_quality_report"]
    assert "tailored_resume" in normalized


def test_validate_resume_quality_flags_missing_sections_and_ai_phrases():
    report = validate_resume_quality({
        "summary": "As an AI language model, I am a dynamic professional.",
        "skills": [],
        "experience": [],
    })

    assert report["status"] == "needs_review"
    assert "missing_skills" in report["issues"]
    assert "missing_experience" in report["issues"]
    assert "contains_ai_phrase" in report["issues"]


def test_normalize_application_generation_preserves_ats_analysis():
    generated = {
        "tailored_resume_structured": {
            "template_recommendation": "ats_classic",
            "contact": {"name": "Jane Doe"},
            "summary": "Frontend engineer focused on React applications.",
            "skills": ["React"],
            "experience": [{
                "role": "Frontend Engineer",
                "company": "Acme",
                "duration": "2021-2024",
                "highlights": ["Built React interfaces for customer workflows."],
            }],
        },
        "ats_analysis": {
            "score_current": 52,
            "score_after_optimization": 78,
            "critical_keywords": [{
                "keyword": "React",
                "importance": "elevee",
                "present_in_cv": "oui",
                "section": "Experience",
                "integration_recommendation": "Keep React in skills and experience.",
            }],
            "final_checklist": ["Use standard headings."],
        },
    }

    normalized = normalize_application_generation(generated)

    assert normalized["ats_analysis"]["score_current"] == 52
    assert normalized["ats_analysis"]["critical_keywords"][0]["keyword"] == "React"
    assert normalized["ats_analysis"]["final_checklist"] == ["Use standard headings."]


def _docx_text(document):
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    return text


def test_build_application_package_outputs_hirly_default_docx():
    profile = {
        "contact": {
            "name": "Jane Doe",
            "email": "jane@example.com",
            "phone": "+33 1 23 45 67 89",
            "location": "Paris",
            "linkedin": "linkedin.com/in/janedoe",
        },
        "cv_filename": "jane.pdf",
        "cv_mime": "application/pdf",
        "cv_original_b64": base64.b64encode(b"%PDF-1.4 fake").decode("ascii"),
    }
    generated = {
        "job_title": "Frontend Engineer",
        "tailored_resume_structured": {
            "template_recommendation": "ats_classic",
            "headline": "Frontend Engineer",
            "contact": {"name": "Jane Doe", "email": "jane@example.com", "location": "Paris"},
            "summary": "Frontend engineer focused on React and TypeScript applications.",
            "role_keywords": ["React", "TypeScript", "Accessibility"],
            "skills": ["React", "TypeScript", "Accessibility"],
            "languages": ["English - Full professional"],
            "experience": [{
                "role": "Frontend Engineer",
                "company": "Acme",
                "duration": "2021-2024",
                "location": "Paris",
                "highlights": ["Built accessible React components for customer-facing workflows."],
            }],
            "education": [{"degree": "BS Computer Science", "school": "MIT", "year": "2019"}],
        },
        "tailored_cover_letter": {"greeting": "Dear team,", "paragraphs": ["Relevant fit."], "sign_off": "Best,"},
    }

    package = build_application_package(profile, generated)

    assert package["tailored_cv_mime"].endswith("document")
    # Hard-locked to the single Hirly template regardless of template_recommendation.
    assert package["template_used"] == "hirly_default"
    assert set(package["available_templates"]) == SUPPORTED_GENERATED_TEMPLATES
    assert package["resume_quality_report"]["status"] == "pass"
    document = docx.Document(io.BytesIO(base64.b64decode(package["tailored_cv_file_b64"])))
    text = _docx_text(document)
    assert "EXPERIENCE" in text
    assert "EDUCATION" in text
    assert "LANGUAGES" in text
    assert "Jane Doe" in text
    assert "Frontend Engineer" in text


def test_luxe_minimal_recommendation_is_overridden_to_hirly_default():
    profile = {
        "contact": {
            "name": "Jane Doe",
            "email": "jane@example.com",
            "phone": "+33 1 23 45 67 89",
            "location": "Paris",
        },
        "cv_filename": "jane.pdf",
        "cv_mime": "application/pdf",
        "cv_original_b64": base64.b64encode(b"%PDF-1.4 fake").decode("ascii"),
    }
    generated = {
        "tailored_resume_structured": {
            "template_recommendation": "luxe_minimal",
            "headline": "Business Development Representative",
            "contact": {"name": "Jane Doe", "email": "jane@example.com", "location": "Paris"},
            "summary": "Business development representative experienced in outbound prospecting and CRM workflows.",
            "skills": ["Outbound Prospecting", "CRM", "Pipeline Management"],
            "experience": [{
                "role": "Business Development Representative",
                "company": "Acme",
                "duration": "2022-2024",
                "location": "Paris",
                "highlights": ["Qualified enterprise prospects and maintained CRM records for account executives."],
            }],
            "education": [],
        },
    }

    package = build_application_package(profile, generated)
    document = docx.Document(io.BytesIO(base64.b64decode(package["tailored_cv_file_b64"])))
    text = _docx_text(document)

    assert package["template_used"] == "hirly_default"
    assert "Business Development Representative" in text
    assert "EXPERIENCE" in text


def test_all_template_recommendations_are_overridden_to_hirly_default():
    profile = {
        "contact": {
            "name": "Jane Doe",
            "email": "jane@example.com",
            "phone": "+33 1 23 45 67 89",
            "location": "Paris",
        },
        "cv_filename": "jane.pdf",
        "cv_mime": "application/pdf",
        "cv_original_b64": base64.b64encode(b"%PDF-1.4 fake").decode("ascii"),
    }

    for template_name in SUPPORTED_GENERATED_TEMPLATES:
        generated = {
            "tailored_resume_structured": {
                "template_recommendation": template_name,
                "headline": "Account Executive",
                "contact": {"name": "Jane Doe", "email": "jane@example.com", "location": "Paris"},
                "summary": "Account executive experienced in B2B sales, CRM workflows, and pipeline management.",
                "skills": ["B2B Sales", "CRM", "Pipeline Management"],
                "experience": [{
                    "role": "Account Executive",
                    "company": "Acme",
                    "duration": "2021-2024",
                    "location": "Paris",
                    "highlights": ["Managed qualified opportunities and maintained accurate CRM records for revenue reporting."],
                }],
                "education": [],
            },
        }

        package = build_application_package(profile, generated)
        document = docx.Document(io.BytesIO(base64.b64decode(package["tailored_cv_file_b64"])))
        text = _docx_text(document)

        assert package["template_used"] == "hirly_default"
        assert "Jane Doe" in text
        assert "EXPERIENCE" in text
        assert "Account Executive" in text
