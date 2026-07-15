"""Application document generation helpers.

V1 prioritizes preserving DOCX structure where possible. PDF preservation is
reported as approximate because PDFs are not reliably editable as semantic
resume templates.
"""

import base64
import io
import logging
import re
from copy import deepcopy
from typing import Any, Dict, List, Tuple

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Inches, Pt

from cv_quality import normalize_application_generation, normalize_resume_structured, validate_resume_quality

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XML_INVALID_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]")
SUPPORTED_GENERATED_TEMPLATES = {
    "ats_classic",
    "modern_pro",
    "executive_compact",
    "luxe_minimal",
    "studio_slate",
    "blue_split",
    "hirly_default",
}
TEMPLATE_SPECS = {
    "hirly_default": {
        "font": "Arial",
        "body_size": 10,
        "name_size": 20,
        "heading_size": 11,
        "heading_color": RGBColor(0, 0, 0),
        "accent_color": RGBColor(107, 114, 128),
        "top_margin": 0.55,
        "bottom_margin": 0.55,
        "left_margin": 0.6,
        "right_margin": 0.6,
        "space_after": 3,
    },
    "ats_classic": {
        "font": "Arial",
        "body_size": 10,
        "name_size": 18,
        "heading_size": 10.5,
        "heading_color": RGBColor(0, 0, 0),
        "accent_color": RGBColor(0, 0, 0),
        "top_margin": 0.55,
        "bottom_margin": 0.55,
        "left_margin": 0.65,
        "right_margin": 0.65,
        "space_after": 3,
    },
    "luxe_minimal": {
        "font": "Arial",
        "body_size": 10,
        "name_size": 20,
        "heading_size": 11,
        "heading_color": RGBColor(38, 47, 56),
        "accent_color": RGBColor(111, 92, 74),
        "top_margin": 0.6,
        "bottom_margin": 0.6,
        "left_margin": 0.72,
        "right_margin": 0.72,
        "space_after": 4,
    },
    "modern_pro": {
        "font": "Arial",
        "body_size": 10,
        "name_size": 19,
        "heading_size": 11,
        "heading_color": RGBColor(31, 78, 121),
        "accent_color": RGBColor(31, 78, 121),
        "top_margin": 0.58,
        "bottom_margin": 0.58,
        "left_margin": 0.68,
        "right_margin": 0.68,
        "space_after": 4,
    },
    "executive_compact": {
        "font": "Arial",
        "body_size": 9.7,
        "name_size": 18,
        "heading_size": 10.5,
        "heading_color": RGBColor(30, 30, 30),
        "accent_color": RGBColor(90, 90, 90),
        "top_margin": 0.48,
        "bottom_margin": 0.48,
        "left_margin": 0.58,
        "right_margin": 0.58,
        "space_after": 2,
    },
    "studio_slate": {
        "font": "Arial",
        "body_size": 10,
        "name_size": 19,
        "heading_size": 11,
        "heading_color": RGBColor(49, 63, 76),
        "accent_color": RGBColor(49, 63, 76),
        "top_margin": 0.62,
        "bottom_margin": 0.62,
        "left_margin": 0.7,
        "right_margin": 0.7,
        "space_after": 4,
    },
    "blue_split": {
        "font": "Arial",
        "body_size": 10,
        "name_size": 19,
        "heading_size": 11,
        "heading_color": RGBColor(22, 70, 122),
        "accent_color": RGBColor(22, 70, 122),
        "top_margin": 0.58,
        "bottom_margin": 0.58,
        "left_margin": 0.68,
        "right_margin": 0.68,
        "space_after": 4,
    },
}
logger = logging.getLogger(__name__)


def sanitize_docx_text(value: Any) -> Any:
    if isinstance(value, str):
        cleaned, removed = XML_INVALID_CONTROL_CHARS.subn("", value)
        if removed:
            logger.warning("DOCX_SANITIZE_REMOVED_CHARS count=%s", removed)
        return cleaned
    if isinstance(value, list):
        return [sanitize_docx_text(item) for item in value]
    if isinstance(value, tuple):
        return tuple(sanitize_docx_text(item) for item in value)
    if isinstance(value, dict):
        return {key: sanitize_docx_text(item) for key, item in value.items()}
    return value


def _safe_text(value: Any) -> str:
    return str(sanitize_docx_text("" if value is None else value))


def _safe_add_paragraph(document: Any, text: Any = "", *args: Any, **kwargs: Any) -> Any:
    return document.add_paragraph(_safe_text(text), *args, **kwargs)


def _safe_set_run_text(run: Any, text: Any) -> None:
    run.text = _safe_text(text)


def build_application_package(profile: Dict[str, Any], generated: Dict[str, Any]) -> Dict[str, Any]:
    profile = sanitize_docx_text(profile)
    generated = sanitize_docx_text(normalize_application_generation(generated))
    original_filename = profile.get("cv_filename") or "cv"
    tailored = normalize_resume_structured(generated.get("tailored_resume_structured") or generated.get("tailored_resume") or {})
    quality_report = generated.get("resume_quality_report") or validate_resume_quality(tailored)
    template_used = _template_name(tailored)

    photo_b64 = profile.get("cv_photo_b64")
    photo_bytes = base64.b64decode(photo_b64) if photo_b64 else None
    job_title = generated.get("job_title") or ""
    file_bytes = _build_hirly_docx(profile, tailored, job_title=job_title, photo_bytes=photo_bytes)
    filename = _tailored_filename(original_filename, ".docx")
    mime = DOCX_MIME
    status = "generated"
    notes = "Generated using the Hirly CV template."

    logger.info("DOCX_BUILD_SUCCESS filename=%s status=%s", filename, status)
    return {
        "tailored_resume_structured": tailored,
        "tailored_cover_letter": generated.get("tailored_cover_letter") or generated.get("cover_letter") or {},
        "application_answers": generated.get("application_answers") or [],
        "tailored_cv_file_b64": base64.b64encode(file_bytes).decode("ascii"),
        "tailored_cv_filename": filename,
        "tailored_cv_mime": mime,
        "template_preservation_status": status,
        "template_preservation_notes": notes,
        "template_used": template_used,
        "available_templates": sorted(SUPPORTED_GENERATED_TEMPLATES),
        "resume_quality_report": quality_report,
    }


def cover_letter_to_text(cover_letter: Dict[str, Any]) -> str:
    cover_letter = sanitize_docx_text(cover_letter)
    if isinstance(cover_letter, str):
        return cover_letter
    parts = []
    if cover_letter.get("template") == "french_formal" or cover_letter.get("subject"):
        sender_lines = [
            cover_letter.get("sender_name"),
            cover_letter.get("sender_address"),
            cover_letter.get("sender_phone"),
            cover_letter.get("sender_email"),
        ]
        sender_block = "\n".join(str(line) for line in sender_lines if line)
        if sender_block:
            parts.append(sender_block)
        recipient_lines = [
            cover_letter.get("recipient_attention"),
            cover_letter.get("recipient_company"),
            cover_letter.get("recipient_address"),
        ]
        recipient_block = "\n".join(str(line) for line in recipient_lines if line)
        if recipient_block:
            parts.append(recipient_block)
        if cover_letter.get("date_line"):
            parts.append(str(cover_letter.get("date_line")))
        if cover_letter.get("subject"):
            parts.append(f"Objet : {cover_letter.get('subject')}")
    greeting = cover_letter.get("greeting")
    if greeting:
        parts.append(str(greeting))
    for paragraph in cover_letter.get("paragraphs") or []:
        parts.append(str(paragraph))
    sign_off = cover_letter.get("sign_off")
    if sign_off:
        parts.append(str(sign_off))
    signature = cover_letter.get("signature_name")
    if signature:
        parts.append(str(signature))
    return "\n\n".join(parts).strip()


def _tailored_filename(original_filename: str, suffix: str) -> str:
    stem = original_filename.rsplit(".", 1)[0] if "." in original_filename else original_filename
    return f"{stem}_tailored{suffix}"


def _circular_photo_png(photo_bytes: bytes, size_px: int = 300) -> bytes:
    """Center-crop to square and mask to a circle via alpha channel -- python-docx
    has no native circular-clip, so the transparency itself is what reads as
    circular once embedded on the page's white background."""
    from PIL import Image, ImageDraw

    with Image.open(io.BytesIO(photo_bytes)) as image:
        image = image.convert("RGB")
        width, height = image.size
        side = min(width, height)
        left = (width - side) // 2
        top = (height - side) // 2
        image = image.crop((left, top, left + side, top + side)).resize((size_px, size_px))
        mask = Image.new("L", (size_px, size_px), 0)
        ImageDraw.Draw(mask).ellipse((0, 0, size_px, size_px), fill=255)
        circular = Image.new("RGBA", (size_px, size_px))
        circular.paste(image, (0, 0), mask)
        buffer = io.BytesIO()
        circular.save(buffer, format="PNG")
        return buffer.getvalue()


def _clear_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def _set_dotted_bottom_border(paragraph: Any) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "dotted")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "999999")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _parse_language_entry(entry: Any) -> Tuple[str, str]:
    raw = _safe_text(entry).strip()
    if not raw:
        return "", ""
    for sep in (" - ", " – ", " — "):
        if sep in raw:
            name, _, level = raw.partition(sep)
            return name.strip(), level.strip()
    if ":" in raw:
        name, _, level = raw.partition(":")
        return name.strip(), level.strip()
    return raw, ""


def _add_hirly_section_heading(document: Any, text: str, template: Dict[str, Any]) -> Any:
    paragraph = _safe_add_paragraph(document)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(_safe_text(text).upper())
    run.bold = True
    run.font.name = template["font"]
    run.font.size = Pt(template["heading_size"])
    run.font.color.rgb = template["heading_color"]
    _set_dotted_bottom_border(paragraph)
    return paragraph


def _add_title_dates_line(document: Any, title: str, dates: Any, template: Dict[str, Any]) -> Any:
    """Bold title on the left, dates right-aligned on the same line via a
    right tab stop at the text area's right edge -- python-docx has no
    side-by-side paragraph layout, but tab stops give the same visual result
    without needing a table for this part."""
    paragraph = _safe_add_paragraph(document)
    section = document.sections[0]
    text_width = section.page_width - section.left_margin - section.right_margin
    paragraph.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)
    run = paragraph.add_run(_safe_text(title))
    run.bold = True
    run.font.name = template["font"]
    run.font.size = Pt(template["body_size"])
    if dates:
        date_run = paragraph.add_run(f"\t{_safe_text(dates)}")
        date_run.font.name = template["font"]
        date_run.font.size = Pt(template["body_size"])
        date_run.font.color.rgb = template["accent_color"]
    return paragraph


def _write_hirly_header(
    document: Any,
    contact: Dict[str, Any],
    name: str,
    job_title: str,
    photo_bytes: bytes | None,
    template: Dict[str, Any],
) -> None:
    table = document.add_table(rows=1, cols=2)
    _clear_table_borders(table)
    left_cell, right_cell = table.rows[0].cells
    left_cell.width = Inches(2.3)
    right_cell.width = Inches(3.7)

    left_para = left_cell.paragraphs[0]
    if photo_bytes:
        try:
            circular = _circular_photo_png(photo_bytes)
            run = left_para.add_run()
            run.add_picture(io.BytesIO(circular), width=Inches(0.9), height=Inches(0.9))
        except Exception:
            logger.warning("hirly_docx_photo_embed_failed", exc_info=True)

    name_para = left_cell.add_paragraph()
    name_run = name_para.add_run(_safe_text(name))
    name_run.bold = True
    name_run.font.name = template["font"]
    name_run.font.size = Pt(template["name_size"])

    if job_title:
        title_para = left_cell.add_paragraph()
        title_run = title_para.add_run(_safe_text(job_title))
        title_run.italic = True
        title_run.font.name = template["font"]
        title_run.font.size = Pt(template["body_size"])
        title_run.font.color.rgb = template["accent_color"]

    contact_lines = [
        ("☎", contact.get("phone")),
        ("✉", contact.get("email")),
        ("\U0001F517", contact.get("linkedin")),
        ("\U0001F4CD", contact.get("location")),
    ]
    first = True
    for icon, value in contact_lines:
        if not value:
            continue
        para = right_cell.paragraphs[0] if first else right_cell.add_paragraph()
        first = False
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run(f"{icon}  {_safe_text(value)}")
        run.font.name = template["font"]
        run.font.size = Pt(template["body_size"])


def _write_hirly_languages(document: Any, languages: List[Any], template: Dict[str, Any]) -> None:
    parsed = [_parse_language_entry(entry) for entry in languages]
    parsed = [item for item in parsed if item[0]]
    if not parsed:
        return
    half = (len(parsed) + 1) // 2
    columns = (parsed[:half], parsed[half:])
    rows = max(len(columns[0]), len(columns[1]))
    table = document.add_table(rows=rows, cols=2)
    _clear_table_borders(table)
    for row_index in range(rows):
        for col_index, column in enumerate(columns):
            if row_index >= len(column):
                continue
            name, level = column[row_index]
            paragraph = table.cell(row_index, col_index).paragraphs[0]
            name_run = paragraph.add_run(_safe_text(name))
            name_run.bold = True
            name_run.font.name = template["font"]
            name_run.font.size = Pt(template["body_size"])
            if level:
                level_run = paragraph.add_run(f"   {_safe_text(level)}")
                level_run.font.name = template["font"]
                level_run.font.size = Pt(template["body_size"])
                level_run.font.color.rgb = template["accent_color"]


def _build_hirly_docx(
    profile: Dict[str, Any],
    tailored: Dict[str, Any],
    job_title: str = "",
    photo_bytes: bytes | None = None,
) -> bytes:
    document = docx.Document()
    template = TEMPLATE_SPECS["hirly_default"]
    _configure_ats_safe_document(document, template)

    contact = sanitize_docx_text(deepcopy(profile.get("contact") or {}))
    tailored_contact = tailored.get("contact") or {}
    merged_contact = {**contact, **{key: value for key, value in tailored_contact.items() if value}}
    name = merged_contact.get("name") or "Candidate"

    _write_hirly_header(document, merged_contact, name, job_title, photo_bytes, template)

    experience = tailored.get("experience") or []
    if experience:
        _add_hirly_section_heading(document, "Experience", template)
        for item in experience:
            _add_title_dates_line(document, item.get("role") or "", item.get("duration"), template)
            meta_line = " — ".join(
                _safe_text(item.get(key)) for key in ("company", "location") if item.get(key)
            )
            if meta_line:
                paragraph = _safe_add_paragraph(document, meta_line)
                for run in paragraph.runs:
                    run.font.color.rgb = template["accent_color"]
                    run.font.name = template["font"]
                    run.font.size = Pt(template["body_size"])
            for highlight in item.get("highlights") or []:
                _safe_add_paragraph(document, highlight, style="List Bullet")

    education = tailored.get("education") or []
    if education:
        _add_hirly_section_heading(document, "Education", template)
        for item in education:
            _add_title_dates_line(document, item.get("degree") or "", item.get("year"), template)
            if item.get("school"):
                paragraph = _safe_add_paragraph(document, item.get("school"))
                for run in paragraph.runs:
                    run.font.color.rgb = template["accent_color"]
                    run.font.name = template["font"]
                    run.font.size = Pt(template["body_size"])

    languages = tailored.get("languages") or []
    if languages:
        _add_hirly_section_heading(document, "Languages", template)
        _write_hirly_languages(document, languages, template)

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def _build_preserved_docx(original_bytes: bytes, profile: Dict[str, Any], tailored: Dict[str, Any]) -> Tuple[bytes, str, str]:
    try:
        document = docx.Document(io.BytesIO(original_bytes))
    except Exception:
        return (
            _build_clean_docx(profile, tailored),
            "approximate",
            "Could not open original DOCX. Generated a clean tailored DOCX instead.",
        )

    replacements = _replacement_pairs(profile, tailored)
    replacements_made = 0
    for old, new in replacements:
        if not old or not new or old == new:
            continue
        replacements_made += _replace_text(document, old, new)

    if replacements_made == 0:
        _append_tailored_section(document, tailored)
        status = "approximate"
        notes = (
            "Original DOCX structure was preserved, but exact source text could not be matched. "
            "Added a tailored content section to the existing document."
        )
    else:
        status = "preserved"
        notes = (
            "Original DOCX structure was reused and matching text was replaced in-place where possible. "
            "Fonts, section layout, margins, tables, and existing styles are preserved by the original document."
        )

    out = io.BytesIO()
    document.save(out)
    return out.getvalue(), status, notes


def _replace_text(document, old: str, new: str) -> int:
    count = 0
    for paragraph in list(document.paragraphs) + _table_paragraphs(document):
        if old in paragraph.text:
            _set_paragraph_text(paragraph, paragraph.text.replace(old, new))
            count += 1
    return count


def _table_paragraphs(document) -> List[Any]:
    paragraphs = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def _set_paragraph_text(paragraph, text: str) -> None:
    first_run = paragraph.runs[0] if paragraph.runs else None
    for run in list(paragraph.runs):
        _safe_set_run_text(run, "")
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    _safe_set_run_text(run, text)
    if first_run is not None and run is not first_run:
        run.bold = first_run.bold
        run.italic = first_run.italic
        run.underline = first_run.underline
        run.style = first_run.style


def _replacement_pairs(profile: Dict[str, Any], tailored: Dict[str, Any]) -> List[Tuple[str, str]]:
    pairs: List[Tuple[str, str]] = []
    if profile.get("summary") and tailored.get("summary"):
        pairs.append((_safe_text(profile["summary"]), _safe_text(tailored["summary"])))

    original_skills = profile.get("skills") or []
    tailored_skills = tailored.get("skills") or []
    if original_skills and tailored_skills:
        pairs.append((", ".join(map(_safe_text, original_skills)), ", ".join(map(_safe_text, tailored_skills))))

    original_experience = profile.get("experience") or []
    tailored_experience = tailored.get("experience") or []
    for original, tailored_item in zip(original_experience, tailored_experience):
        original_highlights = original.get("highlights") or []
        tailored_highlights = tailored_item.get("highlights") or []
        for old, new in zip(original_highlights, tailored_highlights):
            pairs.append((_safe_text(old), _safe_text(new)))
    return pairs


def _append_tailored_section(document, tailored: Dict[str, Any]) -> None:
    document.add_page_break()
    document.add_heading("Tailored Resume Content", level=1)
    _write_tailored_content(document, tailored)


def _build_clean_docx(profile: Dict[str, Any], tailored: Dict[str, Any]) -> bytes:
    document = docx.Document()
    template = _template_spec(tailored)
    _configure_ats_safe_document(document, template)
    contact = sanitize_docx_text(deepcopy(profile.get("contact") or {}))
    tailored_contact = tailored.get("contact") or {}
    name = contact.get("name") or "Candidate"
    if tailored_contact.get("name"):
        name = tailored_contact.get("name")

    header = _safe_add_paragraph(document)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(_safe_text(name))
    run.bold = True
    run.font.name = template["font"]
    run.font.size = Pt(template["name_size"])

    contact_line = " | ".join(
        _safe_text(tailored_contact.get(key) or contact.get(key))
        for key in ("email", "phone", "location", "linkedin", "website")
        if tailored_contact.get(key) or contact.get(key)
    )
    if contact_line:
        paragraph = _safe_add_paragraph(document, contact_line)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.style = document.styles["Normal"]

    headline = _safe_text(tailored.get("headline"))
    if headline:
        paragraph = _safe_add_paragraph(document, headline)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.style = document.styles["Normal"]
        for run in paragraph.runs:
            run.italic = True

    _write_tailored_content(document, tailored, template)
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def _template_name(tailored: Dict[str, Any]) -> str:
    # Hard-locked to the single Hirly-branded template for every user --
    # see the "New default Hirly CV template" plan for context. Ignores
    # tailored.template_recommendation entirely (kept as a field for now in
    # case per-user template variety returns later).
    return "hirly_default"


def _template_spec(tailored: Dict[str, Any]) -> Dict[str, Any]:
    return TEMPLATE_SPECS[_template_name(tailored)]


def _configure_ats_safe_document(document: Any, template: Dict[str, Any]) -> None:
    section = document.sections[0]
    section.top_margin = Inches(template["top_margin"])
    section.bottom_margin = Inches(template["bottom_margin"])
    section.left_margin = Inches(template["left_margin"])
    section.right_margin = Inches(template["right_margin"])

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = template["font"]
    normal.font.size = Pt(template["body_size"])
    normal.paragraph_format.space_after = Pt(template["space_after"])
    normal.paragraph_format.line_spacing = 1.0

    for style_name in ("Heading 1", "Heading 2"):
        style = styles[style_name]
        style.font.name = template["font"]
        style.font.bold = True
        style.font.size = Pt(template["heading_size"] if style_name == "Heading 1" else template["body_size"])
        style.font.color.rgb = template["heading_color"]
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(3)


def _write_tailored_content(document, tailored: Dict[str, Any], template: Dict[str, Any] | None = None) -> None:
    template = template or TEMPLATE_SPECS["ats_classic"]
    if tailored.get("summary"):
        _add_section_heading(document, "Professional Summary", template)
        _safe_add_paragraph(document, tailored["summary"])

    role_keywords = tailored.get("role_keywords") or []
    if role_keywords:
        _add_section_heading(document, "Relevant Focus", template)
        _safe_add_paragraph(document, " | ".join(map(_safe_text, role_keywords)))

    skills = tailored.get("skills") or []
    if skills:
        _add_section_heading(document, "Core Skills", template)
        _safe_add_paragraph(document, " | ".join(map(_safe_text, skills)))

    languages = tailored.get("languages") or []
    if languages:
        _add_section_heading(document, "Languages", template)
        _safe_add_paragraph(document, " | ".join(map(_safe_text, languages)))

    experience = tailored.get("experience") or []
    if experience:
        _add_section_heading(document, "Professional Experience", template)
        for item in experience:
            title = " - ".join(str(item.get(key)) for key in ("role", "company") if item.get(key))
            if title:
                paragraph = _safe_add_paragraph(document)
                run = paragraph.add_run(_safe_text(title))
                run.bold = True
            meta = " | ".join(_safe_text(item.get(key)) for key in ("duration", "location") if item.get(key))
            if meta:
                _safe_add_paragraph(document, meta)
            for highlight in item.get("highlights") or []:
                _safe_add_paragraph(document, highlight, style="List Bullet")

    education = tailored.get("education") or []
    if education:
        _add_section_heading(document, "Education", template)
        for item in education:
            line = " - ".join(_safe_text(item.get(key)) for key in ("degree", "school", "year") if item.get(key))
            if line:
                _safe_add_paragraph(document, line, style="List Bullet")


def _add_section_heading(document: Any, text: str, template: Dict[str, Any]) -> Any:
    paragraph = document.add_heading("", level=1)
    run = paragraph.add_run(_safe_text(text))
    run.bold = True
    run.font.name = template["font"]
    run.font.size = Pt(template["heading_size"])
    run.font.color.rgb = template["heading_color"]
    return paragraph
